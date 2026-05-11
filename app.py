"""
dpc Advance Sheet App — Flask Backend
Run: python app.py  (after running init_db.py first)
"""
import os
import sqlite3
import json
import math
import shutil
import logging
import logging.handlers
import atexit
import subprocess
import gzip
import threading
import secrets
import re
import socket
import time
import uuid
import html as _html_mod
from datetime import datetime, date, timedelta
from functools import wraps
from io import BytesIO

import db_adapter
from db_adapter import DBIntegrityError
import s3_storage
import pdf_layouts

from flask import (Flask, render_template, request, redirect, url_for,
                   flash, session, jsonify, make_response, abort, send_file)
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename


def _safe_content_disposition(filename):
    """Build a safe Content-Disposition header value, stripping injection chars."""
    safe = secure_filename(filename) or 'download'
    return f'inline; filename="{safe}"'


# Allowed HTML tags/attributes for user-supplied rich text (message body)
_MSG_ALLOWED_TAGS  = frozenset(['p','br','b','strong','i','em','a','ul','ol','li','span','h3','h4','hr','div'])
_MSG_ALLOWED_ATTRS = {'a': ['href', 'title']}

def _sanitize_html(raw):
    """Strip disallowed HTML tags, keeping a safe subset. Prevents stored XSS."""
    from html.parser import HTMLParser
    class _S(HTMLParser):
        def __init__(self):
            super().__init__()
            self.out = []
        def handle_starttag(self, tag, attrs):
            t = tag.lower()
            if t not in _MSG_ALLOWED_TAGS:
                return
            allowed = _MSG_ALLOWED_ATTRS.get(t, [])
            safe_attrs = ''
            for k, v in attrs:
                k = k.lower()
                if k not in allowed or not v:
                    continue
                # Reject javascript: and data: URIs
                if re.match(r'\s*(javascript|data|vbscript)\s*:', v, re.I):
                    continue
                safe_attrs += f' {_html_mod.escape(k)}="{_html_mod.escape(v)}"'
            self.out.append(f'<{t}{safe_attrs}>')
        def handle_endtag(self, tag):
            t = tag.lower()
            if t in _MSG_ALLOWED_TAGS:
                self.out.append(f'</{t}>')
        def handle_data(self, data):
            self.out.append(_html_mod.escape(data))
    s = _S()
    s.feed(raw or '')
    return ''.join(s.out)

app = Flask(__name__)

# ── SECRET_KEY — generate and persist if not provided via environment ─────────
_secret = os.environ.get('SECRET_KEY', '')
if not _secret:
    _env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '.env')
    # Try to read from .env file
    if os.path.exists(_env_path):
        with open(_env_path) as _ef:
            for _line in _ef:
                if _line.strip().startswith('SECRET_KEY='):
                    _secret = _line.strip().split('=', 1)[1]
                    break
    # Auto-generate if still missing (first run / dev mode)
    if not _secret:
        import secrets as _secrets_mod
        _secret = _secrets_mod.token_hex(32)
        app.logger.warning(
            'SECRET_KEY not set — generated an ephemeral key. '
            'Run install.sh or set SECRET_KEY in .env for persistent sessions.'
        )
app.secret_key = _secret

# ── Session cookie security ───────────────────────────────────────────────────
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
# Default: omit Secure flag so the LAN/HTTP deployment still works. HTTPS
# deployments should set SESSION_COOKIE_SECURE=1 in the env to upgrade the
# cookie to Secure-only — the sid is now a bearer token (DB-backed sessions),
# so leaking it over plaintext is more dangerous than the old signed cookie.
if os.environ.get('SESSION_COOKIE_SECURE', '').lower() in ('1', 'true', 'yes'):
    app.config['SESSION_COOKIE_SECURE'] = True
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=12)


# ── DB-backed sessions ────────────────────────────────────────────────────────
# Replaces Flask's signed-cookie sessions with a server-side store. The cookie
# carries only an opaque random session ID; all the data lives in the
# `app_sessions` table. When two apps point at the same PostgreSQL DB (sharing
# the `shared` schema) they automatically share login state — sign in once
# and you're authenticated in both.
#
# Set DISABLE_DB_SESSIONS=1 in the environment to fall back to Flask's default
# signed-cookie sessions (e.g. for troubleshooting).
import secrets as _secrets_mod
from flask.sessions import SessionInterface as _FlaskSessionInterface, SessionMixin as _FlaskSessionMixin
from werkzeug.datastructures import CallbackDict as _CallbackDict

_SID_RE = re.compile(r'^[A-Za-z0-9_-]{20,128}$')


class _DBSession(_CallbackDict, _FlaskSessionMixin):
    def __init__(self, initial=None, sid=None, new=False):
        def _on_update(_self):
            _self.modified = True
        _CallbackDict.__init__(self, initial, _on_update)
        self.sid = sid
        self.new = new
        self.modified = False


class _DBSessionInterface(_FlaskSessionInterface):
    """Server-side session backend that stores session data in the
    `app_sessions` table (lives in the shared schema on PostgreSQL).
    The cookie carries only a 256-bit random session ID."""

    def _new_sid(self):
        return _secrets_mod.token_urlsafe(32)

    def _load(self, sid):
        """Return (data_dict, ok). ok=False if the row is missing or expired."""
        try:
            db = get_db()
        except Exception:
            return {}, False
        try:
            row = db.execute(
                "SELECT data, expires_at FROM app_sessions WHERE sid = ?", (sid,)
            ).fetchone()
        except Exception:
            try: db.close()
            except Exception: pass
            return {}, False
        if not row:
            try: db.close()
            except Exception: pass
            return {}, False
        expires = row['expires_at']
        if isinstance(expires, str):
            try:
                expires_dt = datetime.fromisoformat(expires.split('.')[0].replace('Z', ''))
            except Exception:
                expires_dt = datetime.utcnow() - timedelta(seconds=1)  # treat as expired
        else:
            expires_dt = expires
        if expires_dt < datetime.utcnow():
            try:
                db.execute("DELETE FROM app_sessions WHERE sid = ?", (sid,))
                db.commit()
            except Exception:
                pass
            try: db.close()
            except Exception: pass
            return {}, False
        try:
            data = json.loads(row['data']) if row['data'] else {}
            if not isinstance(data, dict):
                data = {}
        except (json.JSONDecodeError, TypeError, ValueError):
            data = {}
        try: db.close()
        except Exception: pass
        return data, True

    def open_session(self, app, request):
        cookie_name = app.config.get('SESSION_COOKIE_NAME', 'session')
        sid = request.cookies.get(cookie_name)
        if not sid or not _SID_RE.match(sid):
            return _DBSession(sid=self._new_sid(), new=True)
        data, ok = self._load(sid)
        if not ok:
            # Never adopt a client-supplied sid that doesn't exist in our store —
            # mint a fresh one. Otherwise an attacker who plants a known sid via
            # an open-redirect / XSS-adjacent vector could pre-fixate the victim's
            # session before login.
            return _DBSession(sid=self._new_sid(), new=True)
        return _DBSession(data, sid=sid, new=False)

    def save_session(self, app, session, response):
        domain = self.get_cookie_domain(app)
        path = self.get_cookie_path(app)
        cookie_name = app.config.get('SESSION_COOKIE_NAME', 'session')

        # Session was cleared — drop the row and the cookie
        if not session:
            if session.modified:
                try:
                    db = get_db()
                    try:
                        db.execute("DELETE FROM app_sessions WHERE sid = ?", (session.sid,))
                        db.commit()
                    finally:
                        db.close()
                except Exception:
                    pass
                response.delete_cookie(cookie_name, domain=domain, path=path)
            return

        # Nothing to persist
        if not session.modified and not session.new:
            return

        lifetime = app.permanent_session_lifetime
        expires_dt = datetime.utcnow() + lifetime
        try:
            data_json = json.dumps(dict(session), default=str)
        except (TypeError, ValueError):
            data_json = '{}'
        user_id = session.get('user_id')
        try:
            db = get_db()
            try:
                db.execute(
                    "INSERT OR REPLACE INTO app_sessions "
                    "(sid, user_id, data, last_seen, expires_at) "
                    "VALUES (?, ?, ?, CURRENT_TIMESTAMP, ?)",
                    (session.sid, user_id, data_json, expires_dt)
                )
                db.commit()
            finally:
                db.close()
        except Exception as e:
            app.logger.warning(f'DB session save failed: {e}')
            return

        response.set_cookie(
            cookie_name,
            session.sid,
            expires=expires_dt,
            httponly=self.get_cookie_httponly(app),
            domain=domain,
            path=path,
            secure=self.get_cookie_secure(app),
            samesite=self.get_cookie_samesite(app),
        )


if os.environ.get('DISABLE_DB_SESSIONS', '').lower() not in ('1', 'true', 'yes'):
    app.session_interface = _DBSessionInterface()


@app.after_request
def _set_security_headers(response):
    """Add defense-in-depth security headers to every response."""
    response.headers.setdefault('X-Content-Type-Options', 'nosniff')
    response.headers.setdefault('X-Frame-Options', 'SAMEORIGIN')
    response.headers.setdefault('Referrer-Policy', 'strict-origin-when-cross-origin')
    return response


# ── CSRF Protection ───────────────────────────────────────────────────────────
# For AJAX: require X-Requested-With header (cannot be set cross-origin without CORS)
# For form POSTs: validate Origin/Referer header matches our host
_CSRF_SAFE_METHODS = frozenset(('GET', 'HEAD', 'OPTIONS'))
_CSRF_EXEMPT_ENDPOINTS = frozenset(('login', 'static'))


@app.before_request
def _csrf_protect():
    """Block cross-site state-changing requests."""
    if request.method in _CSRF_SAFE_METHODS:
        return
    if request.endpoint in _CSRF_EXEMPT_ENDPOINTS:
        return
    if not session.get('user_id'):
        return  # Not logged in — auth decorators will handle it

    # AJAX requests: require X-Requested-With header
    # (Browsers block cross-origin custom headers without CORS preflight)
    if request.is_json or request.content_type == 'application/json':
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return  # Valid AJAX request
        # Also accept if Origin matches (for fetch() without custom header)
        if _origin_matches():
            return
        abort(403)

    # Form POSTs: validate Origin or Referer
    if _origin_matches():
        return

    app.logger.warning(
        f'CSRF blocked: endpoint={request.endpoint} '
        f'origin={request.headers.get("Origin")} '
        f'referer={request.headers.get("Referer")} '
        f'user={session.get("username")}'
    )
    abort(403)


def _origin_matches():
    """Check that Origin or Referer header matches our server."""
    from urllib.parse import urlparse
    # Check Origin header first (most reliable)
    origin = request.headers.get('Origin')
    if origin:
        parsed = urlparse(origin)
        return parsed.hostname == request.host.split(':')[0]
    # Fall back to Referer
    referer = request.headers.get('Referer')
    if referer:
        parsed = urlparse(referer)
        return parsed.hostname == request.host.split(':')[0]
    # No Origin or Referer — could be a direct form submission from same host
    # (some privacy extensions strip Referer). Allow only if SameSite=Lax is set.
    return True


@app.context_processor
def inject_version():
    return {'app_version': APP_VERSION}


@app.template_filter('pretty_json')
def pretty_json_filter(value):
    """Pretty-print a JSON string in templates."""
    try:
        return json.dumps(json.loads(value), indent=2, ensure_ascii=False)
    except Exception:
        return value or ''


@app.template_filter('multi')
def multi_filter(value, sep=', '):
    """Render a multi-select value cleanly. Accepts either a JSON-encoded
    list (e.g. '["01", "02"]' from multi-checkbox fields) or a plain string;
    returns a human-readable comma-separated string. Empty/N-A values
    collapse to ''."""
    if value is None:
        return ''
    s = value if isinstance(value, str) else str(value)
    s = s.strip()
    if not s or s in ('-', '—', 'None', 'none', '[]'):
        return ''
    # When this filter receives output from a Jinja macro, HTML entities
    # have already been escaped (e.g. '"' → '&#34;'/'&quot;'). Unescape
    # so JSON.loads can still recognize the list.
    if '&' in s:
        try:
            import html as _html
            s = _html.unescape(s)
        except Exception:
            pass
    if s.startswith('[') and s.endswith(']'):
        try:
            parsed = json.loads(s)
            if isinstance(parsed, list):
                return sep.join(str(x).strip() for x in parsed if str(x).strip())
        except (json.JSONDecodeError, ValueError):
            pass
    return s


@app.template_filter('hhmm')
def hhmm_filter(value):
    """Normalize a time string to HH:MM for display (e.g. '1900' → '19:00')."""
    s = (value or '').strip() if isinstance(value, str) else str(value or '').strip()
    if not s:
        return ''
    if ':' in s:
        parts = s.split(':', 1)
        try:
            h = int(parts[0]); m = int(parts[1])
        except ValueError:
            return s
    elif s.isdigit() and len(s) in (3, 4):
        h = int(s[:-2]); m = int(s[-2:])
    else:
        return s
    if 0 <= h <= 23 and 0 <= m <= 59:
        return f'{h:02d}:{m:02d}'
    return s


DATABASE = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'advance.db')
BACKUP_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'backups')

# ── Application Version ───────────────────────────────────────────────────────
# Format: MAJOR.MINOR.PATCH
#   MAJOR — breaking schema or architectural changes
#   MINOR — new feature sets (e.g. asset manager, user enhancements)
#   PATCH — bug fixes, small improvements, security patches
APP_VERSION = '2.11.0'

# Flask-Limiter for login rate limiting
try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address
    limiter = Limiter(
        app=app,
        key_func=get_remote_address,
        default_limits=[],
        storage_uri="memory://",
    )
    _limiter_available = True
except ImportError:
    limiter = None
    _limiter_available = False
    import warnings
    warnings.warn(
        'flask-limiter is not installed — login rate limiting is DISABLED. '
        'Install it: pip install flask-limiter',
        stacklevel=1,
    )


def _get_upload_max():
    """Read upload_max_mb from app_settings, default 20."""
    if not os.path.exists(DATABASE):
        return 20 * 1024 * 1024
    try:
        db = get_db()
        row = db.execute("SELECT value FROM app_settings WHERE key='upload_max_mb'").fetchone()
        db.close()
        mb = int(row['value']) if row and row['value'] else 20
        return mb * 1024 * 1024
    except Exception:
        return 20 * 1024 * 1024

DEPARTMENTS = ['Production', 'Programming', 'Event Manager', 'Education Team',
               'Hospitality', 'Guest Services', 'Security', 'Runners']

# ─── Syslog ───────────────────────────────────────────────────────────────────

syslog_logger = logging.getLogger('showadvance')
syslog_logger.setLevel(logging.INFO)
syslog_logger.addHandler(logging.NullHandler())
_syslog_handler = None


def reload_syslog_handler():
    """Read syslog settings from DB and reconfigure the handler."""
    global _syslog_handler
    if not os.path.exists(DATABASE):
        return
    try:
        db = get_db()
        rows = db.execute(
            "SELECT key, value FROM app_settings WHERE key LIKE 'syslog_%'"
        ).fetchall()
        db.close()
    except Exception:
        return
    settings = {r['key']: r['value'] for r in rows}

    if _syslog_handler:
        syslog_logger.removeHandler(_syslog_handler)
        _syslog_handler.close()
        _syslog_handler = None

    if settings.get('syslog_enabled') != '1':
        return

    host = settings.get('syslog_host', '127.0.0.1')
    try:
        port = int(settings.get('syslog_port', 514))
    except ValueError:
        port = 514
    facility_name = settings.get('syslog_facility', 'LOG_LOCAL0')
    facility = getattr(logging.handlers.SysLogHandler, facility_name,
                       logging.handlers.SysLogHandler.LOG_LOCAL0)
    try:
        _syslog_handler = logging.handlers.SysLogHandler(
            address=(host, port), facility=facility
        )
        _syslog_handler.setFormatter(
            logging.Formatter('showadvance: %(levelname)s %(message)s')
        )
        syslog_logger.addHandler(_syslog_handler)
    except Exception as e:
        app.logger.error(f'Failed to configure syslog: {e}')


# ─── App Settings Helper ──────────────────────────────────────────────────────

def get_app_setting(key, default=''):
    """
    Fetch a single app_setting value.
    Always reads from the SQLite bootstrap file so this is safe to call
    at startup before the active DB connection type is resolved.
    """
    if not os.path.exists(DATABASE):
        return default
    try:
        _conn = sqlite3.connect(DATABASE)
        _conn.row_factory = db_adapter._row_factory
        row = _conn.execute('SELECT value FROM app_settings WHERE key=?', (key,)).fetchone()
        _conn.close()
        return row['value'] if row and row['value'] is not None else default
    except Exception:
        return default


# ─── Database ─────────────────────────────────────────────────────────────────

def get_db():
    """Return a normalized DB connection (SQLite or PostgreSQL based on settings)."""
    settings = db_adapter.read_db_settings(DATABASE)
    return db_adapter.connect(DATABASE, settings)


# ─── Cluster Heartbeat (multi-server leader election) ────────────────────────
#
# PROTOCOL
# --------
# Each running app instance (each Gunicorn worker process counts as its own
# instance) generates a UUID once at process start (`_CLUSTER_INSTANCE_ID`) and
# writes a row to the shared `cluster_instances` table every
# `cluster_heartbeat_interval_sec` (default 10 s) with last_seen=NOW().
#
# To find live peers, any code SELECTs rows whose last_seen is within
# `cluster_peer_timeout_sec` (default 30 s) of now. Stale rows from crashed
# instances are simply filtered out by that WHERE clause — no garbage
# collector required.
#
# LEADER ELECTION
# ---------------
# The leader is the live peer with the lowest IPv4 address (compared as a
# 4-tuple of ints, so 192.168.1.9 sorts before 192.168.1.10), tiebroken by
# `instance_id`. Every instance arrives at the same answer independently —
# no election message is exchanged. If the leader crashes or shuts down,
# its row goes stale (or is DELETEd by atexit on graceful exit) and the
# next-lowest IP becomes leader on the very next read.
#
# Three settings affect leader behaviour:
#   cluster_heartbeat_enabled  ('1'/'0', default '1')
#       Master toggle. When '0', the heartbeat thread is not started, no row
#       is written for this instance, and `am_i_leader()` returns True
#       unconditionally (single-server fallback).
#   cluster_force_leader  ('auto' | 'always' | 'never', default 'auto')
#       Operational escape hatch. 'always' = this instance is leader no
#       matter what; 'never' = this instance never claims leadership.
#   cluster_heartbeat_interval_sec / cluster_peer_timeout_sec
#       Tune the heartbeat cadence and how long a missing peer takes to be
#       considered dead.
#
# WRITING A NEW SCHEDULED TASK THAT MUST RUN ONCE CLUSTER-WIDE
# ------------------------------------------------------------
# Add the job to `start_scheduler()` at the bottom of this module like:
#
#     scheduler.add_job(my_periodic_task, 'interval',
#                       hours=1, id='my_periodic_task')
#
# Then make the FIRST line of the task body:
#
#     def my_periodic_task():
#         if not am_i_leader():
#             app.logger.info('my_periodic_task skipped — not cluster leader')
#             return
#         ...real work here...
#
# That is the only change required. With Gunicorn running 4 workers per
# server and 2 servers, your job will fire 8x per tick at the APScheduler
# level, but only the single global leader executes it — the other 7
# return immediately. No locks, no DB rows, no Redis required.
#
# WHEN NOT TO USE am_i_leader()
# -----------------------------
# Some scheduled tasks intentionally run on every instance — e.g. local
# backups (`run_hourly_backup`, `run_daily_backup`) write to a per-server
# /backups/ directory and we want both servers to keep their own copy in
# case one dies. Those jobs deliberately do NOT call am_i_leader().
#
# Rule of thumb:
#   - Sends external side-effects (email, SMS, webhooks)?       → gate on am_i_leader()
#   - Modifies shared DB rows that should change exactly once?  → gate on am_i_leader()
#   - Writes only to instance-local state (logs, local files)?  → do NOT gate
#
# The leader check is cached for 3 s, so calling it at the start of a job
# that fires every minute is essentially free.

_CLUSTER_INSTANCE_ID = uuid.uuid4().hex
_CLUSTER_STARTED_AT = datetime.utcnow()
_cluster_thread = None
_cluster_stop_event = threading.Event()
_cluster_lock = threading.RLock()
_leader_cache = {'at': 0.0, 'is_leader': True, 'leader_id': None, 'leader_ip': None}
_LEADER_CACHE_TTL = 3.0  # seconds


def _get_local_ip():
    """Discover the outbound interface IP without sending traffic."""
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        s.connect(('8.8.8.8', 80))
        return s.getsockname()[0]
    except Exception:
        return '127.0.0.1'
    finally:
        try:
            s.close()
        except Exception:
            pass


def _ip_sort_key(ip):
    """Convert '192.168.1.10' to (192, 168, 1, 10) for correct numeric ordering."""
    try:
        return tuple(int(p) for p in (ip or '').split('.'))
    except Exception:
        return (999, 999, 999, 999)


def _cluster_heartbeat_iteration():
    """Upsert this instance's row. Called on a timer; failures are logged."""
    try:
        ip = _get_local_ip()
        hostname = socket.gethostname()
        port_str = get_app_setting('app_port', '5400')
        try:
            port = int(port_str)
        except ValueError:
            port = 5400
        version = get_app_setting('app_version', '') or ''
        now = datetime.utcnow()
        db = get_db()
        try:
            db.execute("""
                INSERT OR REPLACE INTO cluster_instances
                    (instance_id, ip, hostname, port, app_version, started_at, last_seen)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (_CLUSTER_INSTANCE_ID, ip, hostname, port, version,
                  _CLUSTER_STARTED_AT, now))
            db.commit()
        finally:
            db.close()
    except Exception as e:
        app.logger.warning(f'cluster heartbeat failed: {e}')


def _cluster_heartbeat_loop():
    """Daemon thread loop. Heartbeats until stop_event is set."""
    while not _cluster_stop_event.is_set():
        if get_app_setting('cluster_heartbeat_enabled', '1') in ('1', 'true'):
            _cluster_heartbeat_iteration()
        try:
            interval = int(get_app_setting('cluster_heartbeat_interval_sec', '10'))
        except ValueError:
            interval = 10
        if interval < 2:
            interval = 2
        # Use Event.wait so a stop signal interrupts immediately
        if _cluster_stop_event.wait(timeout=interval):
            break


def _query_live_peers():
    """Return list of dicts for instances seen within the timeout window."""
    try:
        timeout = int(get_app_setting('cluster_peer_timeout_sec', '30'))
    except ValueError:
        timeout = 30
    cutoff = datetime.utcnow() - timedelta(seconds=timeout)
    db = get_db()
    try:
        rows = db.execute(
            "SELECT instance_id, ip, hostname, port, app_version, "
            "started_at, last_seen FROM cluster_instances "
            "WHERE last_seen > ? ORDER BY ip, instance_id",
            (cutoff,)
        ).fetchall()
        return [dict(r) for r in rows]
    finally:
        db.close()


def _compute_leader(peers):
    """Return the leader peer dict (lowest IP, tiebreak by instance_id) or None."""
    if not peers:
        return None
    return min(peers, key=lambda p: (_ip_sort_key(p.get('ip', '')),
                                     p.get('instance_id', '')))


def get_cluster_status():
    """Build the dict consumed by /api/cluster/peers.

    The protocol writes one row per Gunicorn worker (each worker process gets
    its own UUID) so leader election can pick exactly one worker to run
    scheduled jobs. For UI display, however, we aggregate rows by
    (ip, hostname) and report a worker_count so users see one entry per
    physical server instead of N entries per server.
    """
    enabled = get_app_setting('cluster_heartbeat_enabled', '1') in ('1', 'true')
    force = get_app_setting('cluster_force_leader', 'auto')
    self_ip = _get_local_ip()
    self_hostname = socket.gethostname()
    raw_peers = _query_live_peers() if enabled else []
    leader = _compute_leader(raw_peers)

    if force == 'always':
        leader_id = _CLUSTER_INSTANCE_ID
        leader_ip = self_ip
        is_leader = True
    elif force == 'never':
        leader_id = leader.get('instance_id') if leader else None
        leader_ip = leader.get('ip') if leader else None
        is_leader = False
    elif not enabled or not raw_peers:
        # Single-instance / disabled fallback: self is leader
        leader_id = _CLUSTER_INSTANCE_ID
        leader_ip = self_ip
        is_leader = True
    else:
        leader_id = leader.get('instance_id')
        leader_ip = leader.get('ip')
        is_leader = (leader_id == _CLUSTER_INSTANCE_ID)

    # Aggregate raw worker rows by (ip, hostname) into per-server rows.
    groups = {}
    for p in raw_peers:
        key = (p.get('ip', ''), p.get('hostname', '') or '')
        if key not in groups:
            groups[key] = {
                'ip':            p.get('ip'),
                'hostname':      p.get('hostname'),
                'port':          p.get('port'),
                'app_version':   p.get('app_version') or '',
                'started_at':    p.get('started_at'),
                'last_seen':     p.get('last_seen'),
                'worker_count':  0,
                '_instance_ids': set(),
            }
        g = groups[key]
        g['worker_count'] += 1
        iid = p.get('instance_id')
        if iid:
            g['_instance_ids'].add(iid)
        # Earliest started_at = whenever the first worker on this server came up
        sa = p.get('started_at')
        if sa and (not g['started_at'] or sa < g['started_at']):
            g['started_at'] = sa
        # Latest last_seen across the workers on this server
        ls = p.get('last_seen')
        if ls and (not g['last_seen'] or ls > g['last_seen']):
            g['last_seen'] = ls

    aggregated = []
    for g in groups.values():
        ids = g.pop('_instance_ids')
        for k in ('started_at', 'last_seen'):
            v = g.get(k)
            if hasattr(v, 'isoformat'):
                g[k] = v.isoformat()
        g['is_self']   = (_CLUSTER_INSTANCE_ID in ids)
        g['is_leader'] = (leader_id in ids) if leader_id else False
        aggregated.append(g)

    aggregated.sort(key=lambda x: _ip_sort_key(x.get('ip', '')))

    # Server-level leader flag: True if the elected leader worker lives on the
    # same server as this worker (same IP). Used for the "This Instance" UI
    # badge so that all 4 Gunicorn workers on the leader server show LEADER,
    # even though only one of them is actually firing scheduled jobs.
    is_self_server_leader = bool(leader_ip and leader_ip == self_ip)

    return {
        'self_id':                _CLUSTER_INSTANCE_ID,
        'self_ip':                self_ip,
        'self_hostname':          self_hostname,
        'self_started_at':        _CLUSTER_STARTED_AT.isoformat(),
        'leader_id':              leader_id,
        'leader_ip':              leader_ip,
        'is_leader':              is_leader,             # per-worker (drives am_i_leader)
        'is_self_server_leader':  is_self_server_leader, # per-server (drives UI badge)
        'enabled':                enabled,
        'force_leader':           force,
        'peers':                  aggregated,
    }


def am_i_leader():
    """Return True if this instance should run scheduled jobs.

    Cached for ~3 s to avoid hammering the DB if multiple callers exist.
    Single-instance / disabled fallback returns True so existing single-server
    installs keep firing scheduled emails with no config changes.
    """
    now = time.monotonic()
    with _cluster_lock:
        if (now - _leader_cache['at']) < _LEADER_CACHE_TTL:
            return _leader_cache['is_leader']
    status = get_cluster_status()
    with _cluster_lock:
        _leader_cache['at'] = now
        _leader_cache['is_leader'] = status['is_leader']
        _leader_cache['leader_id'] = status['leader_id']
        _leader_cache['leader_ip'] = status['leader_ip']
    return status['is_leader']


def _cluster_cleanup_on_exit():
    """Best-effort delete of this instance's row so failover is instant."""
    try:
        db = get_db()
        try:
            db.execute('DELETE FROM cluster_instances WHERE instance_id=?',
                       (_CLUSTER_INSTANCE_ID,))
            db.commit()
        finally:
            db.close()
    except Exception:
        pass


def start_cluster_heartbeat():
    """Spawn the heartbeat daemon thread. Idempotent."""
    global _cluster_thread
    if _cluster_thread and _cluster_thread.is_alive():
        return _cluster_thread
    _cluster_stop_event.clear()
    _cluster_thread = threading.Thread(
        target=_cluster_heartbeat_loop,
        name='cluster-heartbeat',
        daemon=True,
    )
    _cluster_thread.start()
    atexit.register(_cluster_cleanup_on_exit)
    return _cluster_thread


def stop_cluster_heartbeat():
    """Signal the heartbeat thread to stop; best-effort. Used for restarts."""
    _cluster_stop_event.set()


# ─── Auth Decorators ──────────────────────────────────────────────────────────

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login', next=request.path))
        return f(*args, **kwargs)
    return decorated


def readonly_blocked(f):
    """Block read-only users from mutating actions."""
    @wraps(f)
    def decorated(*args, **kwargs):
        if session.get('is_readonly'):
            if request.is_json:
                return jsonify({'error': 'Read-only access'}), 403
            abort(403)
        return f(*args, **kwargs)
    return decorated


def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        if session.get('user_role') != 'admin':
            abort(403)
        return f(*args, **kwargs)
    return decorated


def content_admin_required(f):
    """Allow system admins AND users in an 'admin_group' type group."""
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        if session.get('user_role') == 'admin':
            return f(*args, **kwargs)
        if session.get('is_content_admin'):
            return f(*args, **kwargs)
        abort(403)
    return decorated


def staff_or_admin_required(f):
    """Allow staff and admin roles (but not plain users) to access a route."""
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        if session.get('user_role') not in ('admin', 'staff'):
            abort(403)
        return f(*args, **kwargs)
    return decorated


def _can_schedule_labor():
    """Mirror of scheduler_required's permission check, usable inline.
    Returns True if the current session may toggle scheduled / assign techs."""
    if 'user_id' not in session:
        return False
    if session.get('user_role') in ('admin', 'staff'):
        return True
    if session.get('is_labor_scheduler') or session.get('is_content_admin'):
        return True
    if session.get('is_scheduler'):
        return True
    return False


def scheduler_required(f):
    """Allow admins, staff, scheduler_group members, or users with is_scheduler flag."""
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        if _can_schedule_labor():
            return f(*args, **kwargs)
        abort(403)
    return decorated


def asset_manager_required(f):
    """Allow admins, content-admins, or users with is_asset_manager flag."""
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        if session.get('user_role') == 'admin':
            return f(*args, **kwargs)
        if session.get('is_content_admin'):
            return f(*args, **kwargs)
        if session.get('is_asset_manager'):
            return f(*args, **kwargs)
        abort(403)
    return decorated


def show_advance_editor_required(f):
    """Allow any user with access to the show to edit its advance section
    (add / edit / remove show_assets and external rentals).

    Blocks anonymous, read-only, restricted users, and users without show
    access. Admins and content admins continue to qualify since they always
    pass can_access_show.
    """
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        show_id = kwargs.get('show_id')
        if show_id is None:
            abort(403)
        if session.get('is_readonly') or session.get('is_restricted'):
            if request.is_json:
                return jsonify({'error': 'Read-only access'}), 403
            abort(403)
        if not can_access_show(session['user_id'], show_id):
            if request.is_json:
                return jsonify({'error': 'Access denied'}), 403
            abort(403)
        return f(*args, **kwargs)
    return decorated


def _populate_session_from_user(user, ts=None):
    """Copy role / permission flags from a users row into the Flask session.
    Used by both the login route (initial population) and the periodic role
    refresh below."""
    uid = user['id']
    session['user_id']            = uid
    if user.get('username') is not None:
        session['username']       = user['username']
    session['display_name']       = user['display_name'] or session.get('username', '')
    session['user_role']          = user['role']
    if user.get('theme') is not None:
        session['theme']          = user['theme'] or 'dark'
    session['is_restricted']      = is_restricted_user(uid)
    session['is_content_admin']   = is_content_admin(uid)
    session['is_labor_scheduler'] = is_labor_scheduler(uid)
    session['is_readonly']        = bool(user.get('is_readonly', 0))
    session['is_scheduler']       = bool(user.get('is_scheduler', 0))
    session['is_asset_manager']   = bool(user.get('is_asset_manager', 0))
    session['is_document_viewer'] = bool(user.get('is_document_viewer', 0))
    session['viewer_venues']      = _decode_json_list(user.get('viewer_venues'))
    session['viewer_doc_types']   = _decode_json_list(user.get('viewer_doc_types'))
    session['_role_checked_at']   = ts if ts is not None else datetime.utcnow().timestamp()


@app.before_request
def _refresh_session_roles():
    """Re-check user role/permissions from DB every 5 minutes to catch demotions."""
    if 'user_id' not in session:
        return
    last_check = session.get('_role_checked_at', 0)
    now = datetime.utcnow().timestamp()
    if now - last_check < 300:  # 5 minutes
        return
    try:
        db = get_db()
        user = db.execute(
            'SELECT id, role, display_name, is_readonly, is_scheduler, is_asset_manager, '
            '       is_document_viewer, viewer_venues, viewer_doc_types '
            'FROM users WHERE id=?',
            (session['user_id'],)
        ).fetchone()
        db.close()
        if not user:
            session.clear()
            return redirect(url_for('login'))
        _populate_session_from_user(user, ts=now)
    except Exception:
        pass


def get_current_user():
    if 'user_id' in session:
        return {
            'id': session['user_id'],
            'username': session['username'],
            'display_name': session.get('display_name', session['username']),
            'role': session.get('user_role', 'user'),
            'theme': session.get('theme', 'dark'),
            'is_restricted': session.get('is_restricted', False),
            'is_content_admin': session.get('is_content_admin', False),
            'is_labor_scheduler': session.get('is_labor_scheduler', False),
            'is_scheduler': session.get('is_scheduler', False),
            'is_asset_manager': session.get('is_asset_manager', False),
        }
    return None


# ─── Access Control Helpers ───────────────────────────────────────────────────

def _get_user_group_types(user_id):
    """Returns a list of group_type strings for the user's groups."""
    db = get_db()
    rows = db.execute("""
        SELECT ug.group_type FROM user_groups ug
        JOIN user_group_members ugm ON ug.id = ugm.group_id
        WHERE ugm.user_id = ?
    """, (user_id,)).fetchall()
    db.close()
    return [r['group_type'] for r in rows]


def is_content_admin(user_id):
    """True if the user is a system admin, a staff user, or is in an 'admin_group' type group."""
    db = get_db()
    user = db.execute('SELECT role FROM users WHERE id=?', (user_id,)).fetchone()
    db.close()
    if not user:
        return False
    if user['role'] in ('admin', 'staff'):
        return True
    return 'admin_group' in _get_user_group_types(user_id)


def get_accessible_shows(user_id):
    """Returns None (all shows) or a list of accessible show IDs."""
    db = get_db()
    user = db.execute('SELECT role FROM users WHERE id=?', (user_id,)).fetchone()
    if not user or user['role'] == 'admin':
        db.close()
        return None

    group_types = _get_user_group_types(user_id)

    # admin_group and all_access both get unrestricted show access
    if not group_types or any(t in ('all_access', 'admin_group') for t in group_types):
        return None

    rows = db.execute("""
        SELECT DISTINCT sga.show_id
        FROM show_group_access sga
        JOIN user_group_members ugm ON sga.group_id = ugm.group_id
        WHERE ugm.user_id = ?
    """, (user_id,)).fetchall()
    db.close()
    return [r['show_id'] for r in rows]


def can_access_show(user_id, show_id):
    accessible = get_accessible_shows(user_id)
    if accessible is None:
        return True
    return show_id in accessible


def is_restricted_user(user_id):
    """True if the user is ONLY in restricted groups (read-only, assigned shows only)."""
    db = get_db()
    user = db.execute('SELECT role FROM users WHERE id=?', (user_id,)).fetchone()
    if not user or user['role'] == 'admin':
        db.close()
        return False
    group_types = _get_user_group_types(user_id)
    if not group_types:
        return False
    # Restricted only if ALL groups are 'restricted' (no all_access or admin_group)
    return all(t == 'restricted' for t in group_types)


def is_labor_scheduler(user_id):
    """True if the user can access the Labor Scheduler page.

    System admins and staff always qualify; additionally any user in a group
    with type 'scheduler_group' or 'admin_group' qualifies.
    """
    db = get_db()
    user = db.execute('SELECT role FROM users WHERE id=?', (user_id,)).fetchone()
    db.close()
    if not user:
        return False
    if user['role'] in ('admin', 'staff'):
        return True
    group_types = _get_user_group_types(user_id)
    return any(t in ('scheduler_group', 'admin_group') for t in group_types)


# ─── Document Viewer ─────────────────────────────────────────────────────────
# A document viewer is a stricter read-only role: they're bounced to /viewer
# on login and can only see PDFs / read-only document pages whose VENUE and
# DOCUMENT TYPE match their per-user allow-lists. Empty list = "all".
DOCUMENT_TYPES = ('advance', 'schedule', 'postnotes')
DOCUMENT_TYPE_LABELS = {
    'advance':   'Advance',
    'schedule':  'Production Schedule',
    'postnotes': 'Post-show Notes',
}


def _decode_json_list(raw):
    """Parse a JSON list column. None / '' / invalid → []."""
    if not raw:
        return []
    if isinstance(raw, list):
        return [str(x).strip() for x in raw if str(x).strip()]
    try:
        v = json.loads(raw)
        if isinstance(v, list):
            return [str(x).strip() for x in v if str(x).strip()]
    except (json.JSONDecodeError, TypeError, ValueError):
        pass
    return []


def get_document_viewer_settings(user_id):
    """Return (is_viewer: bool, venues: list[str], doc_types: list[str]).
    Empty lists mean 'no filter' (allow everything of that dimension)."""
    db = get_db()
    try:
        user = db.execute(
            'SELECT is_document_viewer, viewer_venues, viewer_doc_types '
            'FROM users WHERE id=?', (user_id,)
        ).fetchone()
    except Exception:
        user = None
    db.close()
    if not user:
        return False, [], []
    return (
        bool(user['is_document_viewer'] or 0),
        _decode_json_list(user['viewer_venues']),
        _decode_json_list(user['viewer_doc_types']),
    )


def viewer_can_see_doc_type(doc_type):
    """Session-aware check. Non-viewers always return True."""
    if not session.get('is_document_viewer'):
        return True
    allowed = session.get('viewer_doc_types') or []
    return (not allowed) or (doc_type in allowed)


def viewer_can_see_venue(venue):
    if not session.get('is_document_viewer'):
        return True
    allowed = session.get('viewer_venues') or []
    if not allowed:
        return True
    return (venue or '').strip() in allowed


def viewer_accessible_shows(user_id):
    """Return list of show_ids the viewer can see (intersection of their
    group-show-ACL and venue filter). None means no restriction (admins).
    A viewer with no venue filter still respects the group ACL."""
    base = get_accessible_shows(user_id)
    is_viewer, venues, _ = get_document_viewer_settings(user_id)
    if not is_viewer:
        return base
    db = get_db()
    if base is None:
        # Viewer with all_access group — still apply venue filter if any
        if not venues:
            rows = db.execute(
                "SELECT id FROM shows WHERE COALESCE(status, 'active') != 'archived'"
            ).fetchall()
        else:
            ph = ','.join(['?'] * len(venues))
            rows = db.execute(
                f"SELECT id FROM shows WHERE COALESCE(status, 'active') != 'archived' "
                f"AND venue IN ({ph})", venues
            ).fetchall()
        db.close()
        return [r['id'] for r in rows]
    if not base:
        db.close()
        return []
    if not venues:
        db.close()
        return list(base)
    ph_ids = ','.join(['?'] * len(base))
    ph_v   = ','.join(['?'] * len(venues))
    rows = db.execute(
        f"SELECT id FROM shows WHERE id IN ({ph_ids}) AND venue IN ({ph_v})",
        list(base) + venues
    ).fetchall()
    db.close()
    return [r['id'] for r in rows]


# Whitelist of endpoint names a document viewer is allowed to hit. Anything
# else triggers a redirect to /viewer in _viewer_gate() below.
_VIEWER_ALLOWED_ENDPOINTS = frozenset({
    'viewer_home', 'viewer_show', 'viewer_export',
    'login', 'logout', 'static',
    'force_change_password',
    # Theme + password change so the user can still toggle dark/light + reset
    'set_theme', 'change_own_password',
})


@app.before_request
def _viewer_gate():
    """Redirect document viewers away from any non-viewer endpoint."""
    if not session.get('is_document_viewer'):
        return None
    ep = request.endpoint or ''
    if ep in _VIEWER_ALLOWED_ENDPOINTS:
        return None
    if ep.startswith('static'):
        return None
    # AJAX / API calls: respond 403 JSON instead of redirecting
    if request.is_json or request.path.startswith('/api/'):
        return jsonify({'error': 'Document viewers are restricted to /viewer.'}), 403
    return redirect(url_for('viewer_home'))


# ─── Form Fields Helper ───────────────────────────────────────────────────────


def _normalize_row_dates(d):
    """Convert any datetime.date / datetime objects in a dict to ISO-format strings.

    psycopg2 returns Python date objects for DATE columns; Flask's default JSON
    encoder serialises those as HTTP-date ("Tue, 14 Apr 2026 00:00:00 GMT")
    which breaks ``<input type="date">`` (requires yyyy-MM-dd).  Call this on
    every dict(row) before passing to templates or jsonify.
    """
    for k, v in d.items():
        if isinstance(v, datetime):
            d[k] = v.isoformat()
        elif isinstance(v, date):
            d[k] = v.isoformat()
    return d

def get_form_fields_for_template():
    """Returns ordered list of sections, each with a .fields list."""
    db = get_db()
    sections = db.execute(
        'SELECT * FROM form_sections ORDER BY sort_order'
    ).fetchall()
    fields = db.execute(
        'SELECT * FROM form_fields ORDER BY section_id, sort_order'
    ).fetchall()
    db.close()

    # Keys already hardcoded in the template (load-in/out section)
    _hardcoded_keys = {'load_in_date', 'load_in_time', 'load_out_date', 'load_out_time'}
    field_map = {}
    for f in fields:
        fd = dict(f)
        if fd['field_key'] in _hardcoded_keys:
            continue  # skip — these are rendered by the hardcoded load-in/out section
        if fd['options_json']:
            try:
                fd['options'] = json.loads(fd['options_json'])
            except (json.JSONDecodeError, TypeError):
                fd['options'] = []
        else:
            fd['options'] = []
        field_map.setdefault(fd['section_id'], []).append(fd)

    result = []
    for s in sections:
        sd = dict(s)
        sd['fields'] = field_map.get(s['id'], [])
        result.append(sd)
    return result


def get_schedule_meta_fields():
    """Returns ordered list of schedule meta field templates."""
    db = get_db()
    rows = db.execute(
        'SELECT * FROM schedule_meta_fields ORDER BY sort_order, id'
    ).fetchall()
    db.close()
    return [dict(r) for r in rows]


# ─── Backup Scheduler ─────────────────────────────────────────────────────────

def _ensure_backup_dirs():
    os.makedirs(os.path.join(BACKUP_DIR, 'hourly'), exist_ok=True)
    os.makedirs(os.path.join(BACKUP_DIR, 'daily'), exist_ok=True)


def _run_pg_dump(dest_path, settings):
    """Run pg_dump and write the compressed SQL dump to dest_path (.sql.gz)."""
    env = os.environ.copy()
    env['PGPASSWORD'] = settings.get('pg_password', '')
    cmd = [
        'pg_dump',
        '-h', settings.get('pg_host', 'localhost'),
        '-p', str(settings.get('pg_port', '5432')),
        '-U', settings.get('pg_user', ''),
        '-d', settings.get('pg_dbname', '321theater'),
    ]
    result = subprocess.run(cmd, capture_output=True, env=env, timeout=300)
    if result.returncode != 0:
        raise RuntimeError(f"pg_dump failed: {result.stderr.decode('utf-8', errors='replace')}")
    with gzip.open(dest_path, 'wb') as f:
        f.write(result.stdout)


def _backup_file_ext():
    """Return the expected backup file extension for the active database type."""
    settings = db_adapter.read_db_settings(DATABASE)
    return '.sql.gz' if settings.get('db_type') == 'postgres' else '.db'


def run_hourly_backup():
    _ensure_backup_dirs()
    ts = datetime.now().strftime('%Y%m%d_%H%M')
    hourly_dir = os.path.join(BACKUP_DIR, 'hourly')
    settings = db_adapter.read_db_settings(DATABASE)
    if settings.get('db_type') == 'postgres':
        dest = os.path.join(hourly_dir, f'advance_{ts}.sql.gz')
        _run_pg_dump(dest, settings)
        ext = '.sql.gz'
    else:
        dest = os.path.join(hourly_dir, f'advance_{ts}.db')
        shutil.copy2(DATABASE, dest)
        ext = '.db'
    syslog_logger.info(f'BACKUP_CREATED type=hourly file={dest}')
    files = sorted(
        [f for f in os.listdir(hourly_dir) if f.endswith(ext)],
        reverse=True
    )
    for old in files[24:]:
        os.remove(os.path.join(hourly_dir, old))


def run_daily_backup():
    _ensure_backup_dirs()
    ts = datetime.now().strftime('%Y%m%d')
    daily_dir = os.path.join(BACKUP_DIR, 'daily')
    settings = db_adapter.read_db_settings(DATABASE)
    if settings.get('db_type') == 'postgres':
        dest = os.path.join(daily_dir, f'advance_{ts}.sql.gz')
        _run_pg_dump(dest, settings)
        ext = '.sql.gz'
    else:
        dest = os.path.join(daily_dir, f'advance_{ts}.db')
        shutil.copy2(DATABASE, dest)
        ext = '.db'
    syslog_logger.info(f'BACKUP_CREATED type=daily file={dest}')
    files = sorted(
        [f for f in os.listdir(daily_dir) if f.endswith(ext)],
        reverse=True
    )
    for old in files[30:]:
        os.remove(os.path.join(daily_dir, old))


def _get_smtp_settings():
    """Return a dict of SMTP config from app_settings."""
    keys = ('smtp_host', 'smtp_port', 'smtp_user', 'smtp_pass',
            'smtp_from', 'smtp_tls')
    return {k: get_app_setting(k, '') for k in keys}


def _log_email_error(recipients, subject, error, *,
                     pdf_type=None, show_id=None,
                     triggered_by=None, smtp_code=None):
    """Record one row per failed recipient in email_send_errors.

    `recipients` may be a list, a single string, or None (treated as one
    'unknown' row so the failure isn't lost). Never raises — best-effort
    logging only."""
    if recipients is None:
        recipients = ['(unknown)']
    elif isinstance(recipients, str):
        recipients = [recipients]
    if not recipients:
        return
    err_str  = str(error)[:1000]
    code_str = str(smtp_code or '')[:50]
    by_str   = (triggered_by or session.get('username') or 'system')[:100]
    subj_str = (subject or '')[:200]
    pdf_str  = (pdf_type or '')[:50]
    try:
        db = get_db()
        for r in recipients:
            try:
                db.execute("""
                    INSERT INTO email_send_errors
                        (recipient, subject, error_msg, smtp_code,
                         pdf_type, show_id, triggered_by)
                    VALUES (?,?,?,?,?,?,?)
                """, (str(r)[:200], subj_str, err_str, code_str,
                      pdf_str, show_id, by_str))
            except Exception:
                pass
        db.commit()
        db.close()
    except Exception as e:
        app.logger.warning(f'_log_email_error: could not write error row: {e}')


def _build_mime_message(subject, from_addr, recipients, body_text=None,
                        body_html=None, attachments=None, use_bcc=True):
    """
    Build a MIME email message.

    Args:
        subject (str): Email subject
        from_addr (str): Sender address
        recipients (list[str]): Recipient addresses
        body_text (str|None): Plain text body
        body_html (str|None): HTML body
        attachments (list[dict]|None): Each dict: {'filename', 'data' (bytes), 'mimetype'}
        use_bcc (bool): When True (default), recipients are NOT listed in the
            visible To: header — delivery is via the SMTP envelope only so
            recipients can't see each other. To: shows the from address.

    Returns:
        email.mime.multipart.MIMEMultipart
    """
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders as email_encoders

    # Sanitize email headers to prevent header injection
    def _clean_header(v):
        return v.replace('\r', '').replace('\n', '') if v else v

    msg = MIMEMultipart()
    msg['From'] = _clean_header(from_addr)
    if use_bcc:
        # Most mail clients want a non-empty To: header; show the from address
        # rather than "Undisclosed Recipients" so the message looks normal.
        msg['To'] = _clean_header(from_addr)
    else:
        msg['To'] = _clean_header(', '.join(recipients))
    msg['Subject'] = _clean_header(subject)

    if body_text and body_html:
        alt = MIMEMultipart('alternative')
        alt.attach(MIMEText(body_text, 'plain'))
        alt.attach(MIMEText(body_html, 'html'))
        msg.attach(alt)
    elif body_html:
        msg.attach(MIMEText(body_html, 'html'))
    elif body_text:
        msg.attach(MIMEText(body_text, 'plain'))

    for att in (attachments or []):
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(att['data'])
        email_encoders.encode_base64(part)
        mime = att.get('mimetype', 'application/octet-stream')
        part.add_header('Content-Type', mime)
        part.add_header('Content-Disposition',
                        f'attachment; filename="{att["filename"]}"')
        msg.attach(part)

    return msg


def _send_email_smtp(subject, recipients, body_text=None, body_html=None,
                     attachments=None, from_address=None, error_context=None):
    """Send email via configured SMTP relay. Returns (success, message).

    `error_context` is an optional dict of {pdf_type, show_id, triggered_by}
    that's attached to any rows written into email_send_errors so the admin
    can trace which feature triggered the failure."""
    import smtplib

    smtp_cfg = _get_smtp_settings()
    if not smtp_cfg.get('smtp_host'):
        _log_email_error(recipients, subject, 'SMTP not configured.',
                         **(error_context or {}))
        return False, 'SMTP not configured.'

    from_addr = from_address or smtp_cfg.get('smtp_from') or smtp_cfg.get('smtp_user', '')
    msg = _build_mime_message(subject, from_addr, recipients, body_text,
                              body_html, attachments)

    try:
        port = int(smtp_cfg.get('smtp_port') or 587)
        use_tls = smtp_cfg.get('smtp_tls', '1') not in ('0', 'false', 'False', '')
        if use_tls:
            server = smtplib.SMTP(smtp_cfg['smtp_host'], port, timeout=15)
            server.ehlo()
            server.starttls()
            server.ehlo()
        else:
            server = smtplib.SMTP_SSL(smtp_cfg['smtp_host'], port, timeout=15)
        if smtp_cfg.get('smtp_user') and smtp_cfg.get('smtp_pass'):
            server.login(smtp_cfg['smtp_user'], smtp_cfg['smtp_pass'])
        try:
            # sendmail returns a dict of refused recipients for partial failures
            refused = server.sendmail(from_addr, recipients, msg.as_string())
            if refused:
                # SMTP relay refused some addresses but accepted others
                for addr, (code, why) in refused.items():
                    _log_email_error(addr, subject, f'{code} {why!r}',
                                     smtp_code=str(code), **(error_context or {}))
        finally:
            try: server.quit()
            except Exception: pass
    except smtplib.SMTPRecipientsRefused as e:
        # Every recipient was refused — log each with its specific code
        app.logger.error(f'SMTP all recipients refused: {e.recipients}')
        for addr, (code, why) in (e.recipients or {}).items():
            _log_email_error(addr, subject, f'{code} {why!r}',
                             smtp_code=str(code), **(error_context or {}))
        return False, f'SMTP error: all recipients refused.'
    except smtplib.SMTPResponseException as e:
        app.logger.error(f'SMTP send failed: {e.smtp_code} {e.smtp_error}')
        _log_email_error(recipients, subject, str(e.smtp_error),
                         smtp_code=str(e.smtp_code), **(error_context or {}))
        return False, f'SMTP error: {e.smtp_code} {e.smtp_error}'
    except Exception as e:
        app.logger.error(f'SMTP send failed: {e}')
        _log_email_error(recipients, subject, str(e),
                         **(error_context or {}))
        return False, f'SMTP error: {e}'

    return True, f'Sent to {len(recipients)} recipient(s).'


def _send_email_direct(subject, recipients, body_text=None, body_html=None,
                       attachments=None, from_address=None, error_context=None):
    """Send email directly via MX lookup (no relay). Returns (success, message)."""
    import smtplib
    import dns.resolver

    smtp_cfg = _get_smtp_settings()
    from_addr = from_address or smtp_cfg.get('smtp_from') or 'noreply@localhost'
    ehlo_hostname = get_app_setting('direct_ehlo_hostname', '').strip() or None
    display_name = get_app_setting('direct_display_name', '').strip()

    # Wrap from address with display name if configured
    if display_name:
        from email.utils import formataddr
        from_addr_header = formataddr((display_name, from_addr))
    else:
        from_addr_header = from_addr

    msg = _build_mime_message(subject, from_addr_header, recipients, body_text,
                              body_html, attachments)
    msg_str = msg.as_string()

    # Group recipients by domain
    from collections import defaultdict
    by_domain = defaultdict(list)
    for addr in recipients:
        if '@' in addr:
            by_domain[addr.split('@')[1].lower()].append(addr)

    errors = []
    sent_count = 0

    for domain, addrs in by_domain.items():
        # Resolve MX records
        try:
            mx_records = dns.resolver.resolve(domain, 'MX')
            mx_hosts = sorted(mx_records, key=lambda r: r.preference)
        except Exception as e:
            err_msg = f'MX lookup failed for {domain}: {e}'
            errors.append(err_msg)
            _log_email_error(addrs, subject, err_msg, **(error_context or {}))
            continue

        # Try each MX host in priority order
        delivered = False
        last_error = None
        for mx in mx_hosts:
            mx_host = str(mx.exchange).rstrip('.')
            try:
                server = smtplib.SMTP(mx_host, 25, timeout=15,
                                      local_hostname=ehlo_hostname)
                server.ehlo(ehlo_hostname)
                try:
                    server.starttls()
                    server.ehlo(ehlo_hostname)
                except smtplib.SMTPNotSupportedError:
                    pass  # Server doesn't support STARTTLS, continue unencrypted
                refused = server.sendmail(from_addr, addrs, msg_str)
                try: server.quit()
                except Exception: pass
                # Partial refusal — server accepted the connection but rejected
                # specific addresses; log each so the admin can act on bad
                # addresses (e.g. typo'd recipient).
                if refused:
                    for addr, (code, why) in refused.items():
                        _log_email_error(addr, subject, f'{code} {why!r}',
                                         smtp_code=str(code), **(error_context or {}))
                sent_count += len(addrs) - len(refused or {})
                delivered = True
                break
            except Exception as e:
                last_error = f'{mx_host}: {e}'
                app.logger.warning(f'Direct send to MX {mx_host} for {domain} failed: {e}')
                continue

        if not delivered:
            detail = f' ({last_error})' if last_error else ''
            err_msg = f'All MX hosts failed for {domain}{detail}'
            errors.append(err_msg)
            _log_email_error(addrs, subject, err_msg, **(error_context or {}))

    if errors and sent_count == 0:
        return False, '; '.join(errors)
    elif errors:
        return True, f'Sent to {sent_count} recipient(s). Failures: {"; ".join(errors)}'
    return True, f'Sent to {sent_count} recipient(s).'


def _send_email(subject, recipients, body_text=None, body_html=None,
                attachments=None, from_address=None, error_context=None):
    """
    General-purpose email sender. Dispatches to SMTP relay or direct MX
    based on the email_provider setting.

    Args:
        subject (str): Email subject
        recipients (list[str]): Recipient email addresses
        body_text (str|None): Plain text body
        body_html (str|None): HTML body
        attachments (list[dict]|None): Each dict: {'filename', 'data' (bytes), 'mimetype'}
        from_address (str|None): Override the configured from address
        error_context (dict|None): {pdf_type, show_id, triggered_by} stamped
            onto any rows the helpers write into email_send_errors.

    Returns:
        (bool, str): (success, message)
    """
    provider = get_app_setting('email_provider', 'smtp')
    if provider == 'direct':
        return _send_email_direct(subject, recipients, body_text, body_html,
                                  attachments, from_address,
                                  error_context=error_context)
    return _send_email_smtp(subject, recipients, body_text, body_html,
                            attachments, from_address,
                            error_context=error_context)


def _send_pdf_email(show_id, pdf_type, triggered_by, exported_by_id=None, days_before=None):
    """
    Build a PDF (advance or schedule) and email it to all report recipients.

    triggered_by : display string for the email body ('username' or 'system')
    exported_by_id : user.id to log in export_log (None for scheduled sends)
    days_before : int used for dedup key in email_send_log

    Returns (success: bool, message: str, recipient_count: int)
    """
    provider = get_app_setting('email_provider', 'smtp')
    if provider == 'smtp':
        smtp_cfg = _get_smtp_settings()
        if not smtp_cfg.get('smtp_host'):
            return False, 'SMTP not configured.', 0

    # Fetch recipients based on pdf_type
    db = get_db()
    if pdf_type == 'advance':
        recip_col = 'advance_recipient'
    elif pdf_type == 'schedule':
        recip_col = 'production_recipient'
    elif pdf_type == 'postnotes':
        recip_col = 'postnotes_recipient'
    else:
        recip_col = 'report_recipient'
    recipients = [
        r['email'] for r in
        db.execute(
            f"SELECT email FROM contacts WHERE ({recip_col}=1 OR report_recipient=1) AND email != '' ORDER BY name"
        ).fetchall()
    ]
    db.close()

    if not recipients:
        return False, 'No report recipients configured.', 0

    # Build PDF bytes — run inside app context, no request context needed
    try:
        with app.app_context():
            if pdf_type == 'advance':
                _, pdf_version, show_dict, pdf_bytes, pdf_log_id = _build_advance_pdf(
                    show_id, exported_by_id=exported_by_id, base_url='/'
                )
            elif pdf_type == 'postnotes':
                _, pdf_version, show_dict, pdf_bytes, pdf_log_id = _build_postnotes_pdf(
                    show_id, exported_by_id=exported_by_id, base_url='/'
                )
            else:
                _, pdf_version, show_dict, pdf_bytes, pdf_log_id = _build_schedule_pdf(
                    show_id, exported_by_id=exported_by_id, base_url='/'
                )
    except Exception as e:
        app.logger.error(f'PDF build failed for email show={show_id} type={pdf_type}: {e}')
        return False, f'PDF generation failed: {e}', 0

    # Push PDF to S3 for archival (synchronous — no user waiting on this path)
    if pdf_bytes and pdf_log_id and s3_storage.is_configured():
        try:
            s3_key = f"exports/{show_id}/{pdf_type}/v{pdf_version}.pdf"
            s3_storage.upload_file(s3_key, pdf_bytes, 'application/pdf')
            _db_s3 = get_db()
            _db_s3.execute('UPDATE export_log SET s3_key=? WHERE id=?', (s3_key, pdf_log_id))
            _db_s3.commit()
            _db_s3.close()
        except Exception as e:
            app.logger.error(f"S3 push failed for email PDF show={show_id} type={pdf_type}: {e}")
            syslog_logger.error(f"S3_PUSH_FAILED context=email_pdf show_id={show_id} type={pdf_type} error={e}")

    if not pdf_bytes:
        return False, 'PDF generation produced no output.', 0

    # Build email subject
    show_name  = show_dict.get('name', 'Show')
    show_date  = show_dict.get('show_date', '')
    venue      = show_dict.get('venue', '')
    pm_name    = show_dict.get('production_manager', '')
    if not pm_name:
        # Try pulling from advance_data
        try:
            _db2 = get_db()
            _row = _db2.execute(
                "SELECT field_value FROM advance_data WHERE show_id=? AND field_key='production_manager'",
                (show_id,)
            ).fetchone()
            _db2.close()
            pm_name = _row['field_value'] if _row else ''
        except Exception:
            pm_name = ''

    type_label = ({
        'advance': 'Advance Sheet',
        'schedule': 'Production Schedule',
        'postnotes': 'Post-Show Report',
    }).get(pdf_type, 'Production Schedule')
    # PostgreSQL returns DATE columns as datetime.date — coerce every part to
    # string before join() so we never hit
    # "sequence item N: expected str instance, datetime.date found".
    def _s(v):
        if v is None: return ''
        if isinstance(v, (datetime, date)): return v.isoformat()[:10]
        return str(v)
    subject_parts = ['3·2·1→Theater', type_label, _s(show_name)]
    if show_date:
        subject_parts.append(_s(show_date))
    if venue:
        subject_parts.append(_s(venue))
    if pm_name:
        subject_parts.append(f'PM: {_s(pm_name)}')
    subject = ' | '.join(subject_parts)

    # Email body
    if triggered_by == 'system':
        body_line = (f'This {type_label} was automatically generated and sent by 3·2·1→Theater '
                     f'on {datetime.now().strftime("%B %d, %Y at %I:%M %p")}.')
    else:
        body_line = (f'This {type_label} was generated and sent by {triggered_by} '
                     f'on {datetime.now().strftime("%B %d, %Y at %I:%M %p")}.')

    safe_show = show_name.replace(' ', '_').replace('/', '-')
    filename   = f"{type_label.replace(' ','_')}_{safe_show}_{show_date}.pdf"

    # Send email via configured provider (SMTP relay or direct MX)
    attachments = [{'filename': filename, 'data': pdf_bytes, 'mimetype': 'application/pdf'}]
    success, send_message = _send_email(
        subject=subject, recipients=recipients,
        body_text=body_line, attachments=attachments,
        error_context={
            'pdf_type':     pdf_type,
            'show_id':      show_id,
            'triggered_by': triggered_by,
        },
    )
    if not success:
        app.logger.error(f'Email send failed show={show_id} type={pdf_type}: {send_message}')
        return False, send_message, 0

    # Log the send
    try:
        _db3 = get_db()
        _db3.execute("""
            INSERT INTO email_send_log
              (show_id, pdf_type, trigger_type, days_before, sent_by, recipient_count)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (show_id, pdf_type,
              'scheduled' if triggered_by == 'system' else 'manual',
              days_before, triggered_by, len(recipients)))
        _db3.commit()
        _db3.close()
    except Exception as e:
        app.logger.warning(f'email_send_log write failed: {e}')

    syslog_logger.info(
        f"PDF_EMAIL show_id={show_id} type={pdf_type} "
        f"recipients={len(recipients)} by={triggered_by}"
    )
    return True, f'Sent to {len(recipients)} recipient(s).', len(recipients)


def run_scheduled_pdf_emails():
    """
    APScheduler job: hourly tick.  Sends PDFs when the configured send hour
    matches the current hour and no send has been recorded for this show/type/day.
    """
    if not am_i_leader():
        app.logger.info('PDF email check skipped — not cluster leader')
        return
    send_hour = int(get_app_setting('pdf_email_send_hour', '6'))
    if datetime.now().hour != send_hour:
        return

    today = date.today()
    today_str = today.isoformat()

    db = get_db()
    shows = db.execute(
        "SELECT id, name, show_date FROM shows WHERE status='active'"
    ).fetchall()

    # Pre-load first perf date per show
    perfs = db.execute(
        "SELECT show_id, MIN(perf_date) as first_perf FROM show_performances GROUP BY show_id"
    ).fetchall()
    first_perf = {r['show_id']: r['first_perf'] for r in perfs}

    # Already sent today (to avoid duplicate sends within the same day)
    sent_today = set()
    for r in db.execute(
        "SELECT show_id, pdf_type, days_before FROM email_send_log "
        "WHERE trigger_type='scheduled' AND DATE(sent_at)=?", (today_str,)
    ).fetchall():
        sent_today.add((r['show_id'], r['pdf_type'], r['days_before']))

    db.close()

    configs = [
        ('advance',  'advance_email_enabled',     'advance_email_days_before',  None),
        ('schedule', 'schedule_email_enabled_1',  'schedule_email_days_1',      None),
        ('schedule', 'schedule_email_enabled_2',  'schedule_email_days_2',      None),
    ]

    for show_row in shows:
        show_id   = show_row['id']
        perf_date_str = first_perf.get(show_id) or show_row['show_date']
        if not perf_date_str:
            continue
        try:
            perf_date = date.fromisoformat(perf_date_str)
        except ValueError:
            continue
        days_until = (perf_date - today).days

        for pdf_type, enabled_key, days_key, _ in configs:
            if get_app_setting(enabled_key, '0') not in ('1', 'true'):
                continue
            try:
                trigger_days = int(get_app_setting(days_key, '0'))
            except ValueError:
                continue
            if trigger_days <= 0:
                continue
            if days_until != trigger_days:
                continue
            if (show_id, pdf_type, trigger_days) in sent_today:
                continue

            ok, msg, _ = _send_pdf_email(
                show_id, pdf_type, 'system', days_before=trigger_days
            )
            app.logger.info(
                f'Scheduled PDF email show={show_id} type={pdf_type} '
                f'days_before={trigger_days}: {msg}'
            )


def start_scheduler():
    """Register all background jobs.

    IMPORTANT — multi-server safety:
    Under Gunicorn each worker process loads this module independently, so
    `start_scheduler()` runs once per worker. With the default 4 workers on
    a 2-server cluster, every job below is fired 8 times per tick. Most
    jobs must therefore guard their entry point with `am_i_leader()` so
    only the single global leader does the work. See the Cluster Heartbeat
    section near the top of this file for the convention and exceptions
    (backups intentionally run everywhere for redundancy).

    To add a new job:
      1. Add `scheduler.add_job(my_func, ..., id='my_func')` below.
      2. Inside `my_func`, decide whether it should be leader-gated:
         - Yes (almost always): start with
             if not am_i_leader():
                 app.logger.info('my_func skipped — not cluster leader')
                 return
         - No (only for instance-local effects like local backups): no gate.
    """
    try:
        from apscheduler.schedulers.background import BackgroundScheduler
        scheduler = BackgroundScheduler()
        # Backups: NOT leader-gated — every instance keeps its own copy on
        # local disk so a failing primary doesn't take its backup history
        # with it. The duplication is intentional.
        scheduler.add_job(run_hourly_backup, 'interval', hours=1, id='hourly_backup')
        scheduler.add_job(run_daily_backup, 'cron', hour=0, minute=0, id='daily_backup')
        # PDF emails: leader-gated inside run_scheduled_pdf_emails() so
        # recipients never receive a duplicate when multiple instances run.
        scheduler.add_job(run_scheduled_pdf_emails, 'interval', hours=1, id='pdf_email_check')
        scheduler.start()
        return scheduler
    except ImportError:
        app.logger.warning('APScheduler not installed — backups disabled.')
        return None


# ─── General Helpers ──────────────────────────────────────────────────────────

def auto_archive_past_shows():
    """Move shows whose last performance date has passed into 'archived' status."""
    db = get_db()
    today = date.today().isoformat()
    db.execute("""
        UPDATE shows SET status = 'archived'
        WHERE status = 'active'
          AND (
            -- Has performances: archive only when ALL have passed
            (id IN (SELECT DISTINCT show_id FROM show_performances)
             AND id NOT IN (
               SELECT DISTINCT show_id FROM show_performances
               WHERE perf_date IS NULL OR perf_date >= ?
             ))
            OR
            -- No performances: use legacy show_date field
            (id NOT IN (SELECT DISTINCT show_id FROM show_performances)
             AND show_date IS NOT NULL
             AND show_date < ?)
          )
    """, (today, today))
    db.commit()
    db.close()


def _sync_show_primary_date(db, show_id):
    """Keep shows.show_date/show_time in sync with the earliest performance."""
    first = db.execute("""
        SELECT perf_date, perf_time FROM show_performances
        WHERE show_id = ?
        ORDER BY CASE WHEN perf_date IS NULL THEN 1 ELSE 0 END, perf_date, id
        LIMIT 1
    """, (show_id,)).fetchone()
    if first:
        db.execute("""
            UPDATE shows SET show_date=?, show_time=?, updated_at=CURRENT_TIMESTAMP WHERE id=?
        """, (first['perf_date'], first['perf_time'], show_id))
        for key, val in [('show_date', first['perf_date'] or ''),
                         ('show_time', first['perf_time'] or '')]:
            db.execute("""
                INSERT OR REPLACE INTO advance_data (show_id, field_key, field_value, updated_at)
                VALUES (?, ?, ?, CURRENT_TIMESTAMP)
            """, (show_id, key, val))
    else:
        db.execute("""
            UPDATE shows SET show_date=NULL, show_time='', updated_at=CURRENT_TIMESTAMP WHERE id=?
        """, (show_id,))


def get_show_or_404(show_id):
    db = get_db()
    show = db.execute('SELECT * FROM shows WHERE id = ?', (show_id,)).fetchone()
    db.close()
    if not show:
        abort(404)
    return show


def get_contacts_by_dept():
    db = get_db()
    contacts = db.execute('SELECT * FROM contacts ORDER BY department, name').fetchall()
    db.close()
    by_dept = {}
    for c in contacts:
        dept = c['department'] or 'Other'
        by_dept.setdefault(dept, []).append(dict(c))
    return by_dept


def _snapshot_form_history(db, show_id, form_type, snapshot_data):
    """Insert a history snapshot and prune to 50 entries."""
    db.execute("""
        INSERT INTO form_history (show_id, form_type, saved_by, snapshot_json)
        VALUES (?, ?, ?, ?)
    """, (show_id, form_type, session.get('user_id'), json.dumps(snapshot_data)))
    db.execute("""
        DELETE FROM form_history
        WHERE show_id = ? AND form_type = ?
          AND id NOT IN (
            SELECT id FROM form_history
            WHERE show_id = ? AND form_type = ?
            ORDER BY saved_at DESC LIMIT 50
          )
    """, (show_id, form_type, show_id, form_type))


def log_audit(db, action, entity_type, entity_id=None, show_id=None,
              before=None, after=None, detail=None):
    """Write one row to audit_log. Never raises — audit failure must not block normal flow."""
    try:
        db.execute("""
            INSERT INTO audit_log
              (user_id, username, action, entity_type, entity_id,
               show_id, before_json, after_json, ip_address, detail)
            VALUES (?,?,?,?,?,?,?,?,?,?)
        """, (
            session.get('user_id'),
            session.get('username', ''),
            action,
            entity_type,
            str(entity_id) if entity_id is not None else None,
            show_id,
            json.dumps(before) if before is not None else None,
            json.dumps(after)  if after  is not None else None,
            request.remote_addr,
            detail,
        ))
    except Exception:
        pass


# ─── Audit Undo Infrastructure ────────────────────────────────────────────────
#
# Maps audit_log.entity_type values to the SQL table that stores the entity,
# for routes that need to look up or reverse a mutation. Only leaf tables with
# integer primary keys named `id` are listed — these are the entities we can
# undo. Entities not in this map (form, setting, system, attachment, etc.)
# either use a separate restore flow or are not reversibly logged.

UNDO_TABLE_MAP = {
    'contact':                'contacts',
    'form_field':             'form_fields',
    'form_section':           'form_sections',
    'schedule_template':      'schedule_templates',
    'position_category':      'position_categories',
    'job_position':           'job_positions',
    'labor_request':          'labor_requests',
    'pay_rate_level':         'pay_rate_levels',
    'crew_member':            'crew_members',
    'warehouse_location':     'warehouse_locations',
    'asset_category':         'asset_categories',
    'arts_group':             'arts_groups',
    'asset_type':             'asset_types',
    'asset_item':             'asset_items',
    'show_asset':             'show_assets',
    'show_external_rental':   'show_external_rentals',
    'site_message':           'site_messages',
    'group':                  'user_groups',
    'schedule_meta_field':    'schedule_meta_fields',
    'overhead_project':       'overhead_projects',
    'overhead_labor_group':   'overhead_labor_groups',
    'overhead_labor_request': 'overhead_labor_requests',
    'overhead_labor_template':'overhead_labor_templates',
}

# Action suffixes we know how to reverse. Maps suffix → operation kind.
_UNDO_VERB = {
    # Creates — reverse with DELETE FROM <table> WHERE id = entity_id
    'ADD':     'create',
    'CREATE':  'create',
    'POST':    'create',
    'UPLOAD':  'create',
    # Updates — reverse by UPDATE …SET (cols from before_json) WHERE id = entity_id
    'EDIT':    'update',
    'UPDATE':  'update',
    'RENAME':  'update',
    # Deletes — reverse by INSERT (cols from before_json) INTO <table>
    'DELETE':  'delete',
    'REMOVE':  'delete',
    'RETIRE':  'delete',
}


def _snapshot_row(db, table, row_id):
    """Return a dict of all columns for a single row, or None if not found."""
    try:
        row = db.execute(f'SELECT * FROM {table} WHERE id = ?', (row_id,)).fetchone()
        return dict(row) if row else None
    except Exception:
        return None


def log_audit_change(db, action, entity_type, entity_id, *, show_id=None, detail=None,
                     before=None, after=None, table=None):
    """Write an audit row with before/after state captured for undo.

    If `table` is provided and `before`/`after` are None, the wrapper fetches the
    current row snapshot automatically — useful right *before* a delete or right
    *after* an insert/update. Most callers should pass explicit `before`/`after`
    dicts when they already have the data in hand.
    """
    if table and before is None and after is None and entity_id is not None:
        # Default: snapshot current row as `after` (fits create/update patterns)
        after = _snapshot_row(db, table, entity_id)
    log_audit(db, action, entity_type, entity_id, show_id=show_id,
              before=before, after=after, detail=detail)


def _classify_undo_action(action):
    """Return one of 'create' / 'update' / 'delete' / None for an audit action."""
    if not action:
        return None
    # Last word-chunk after the final underscore drives the verb (ASSET_ITEM_ADD -> ADD)
    suffix = action.rsplit('_', 1)[-1].upper()
    return _UNDO_VERB.get(suffix)


def _can_undo_audit_row(row):
    """Return (ok, reason) — True if this audit row has the data needed to reverse it."""
    if not row:
        return False, 'Audit row not found'
    if row['entity_type'] not in UNDO_TABLE_MAP:
        return False, f"Entity type '{row['entity_type']}' is not undoable"
    kind = _classify_undo_action(row['action'])
    if kind is None:
        return False, f"Action '{row['action']}' has no known reverse"
    if kind == 'create' and not row['entity_id']:
        return False, 'No entity_id recorded — cannot target the created row'
    if kind == 'update' and not row['before_json']:
        return False, 'No before-state recorded — cannot restore prior values'
    if kind == 'delete' and not row['before_json']:
        return False, 'No before-state recorded — cannot re-create deleted row'
    if row.get('undone_at'):
        return False, 'Already undone'
    return True, kind


# ─── Auth Routes ──────────────────────────────────────────────────────────────

def _login_route():
    if 'user_id' in session:
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        db = get_db()
        user = db.execute('SELECT * FROM users WHERE username = ?', (username,)).fetchone()
        if user and check_password_hash(user['password_hash'], password):
            # Update last_login timestamp
            try:
                db.execute('UPDATE users SET last_login=CURRENT_TIMESTAMP WHERE id=?', (user['id'],))
                db.commit()
            except Exception:
                pass
            # Regenerate session to prevent session fixation. session.clear()
            # empties the dict but does NOT rotate the sid on the DB-backed
            # session, so we explicitly mint a new sid here. The old DB row
            # (if any) becomes orphaned and is harvested on expiry.
            next_url = request.form.get('next') or url_for('dashboard')
            session.clear()
            try:
                session.sid = _secrets_mod.token_urlsafe(32)
                session.new = True
            except AttributeError:
                pass  # Non-DB session backend in use (DISABLE_DB_SESSIONS=1)
            _populate_session_from_user(user)
            log_audit(db, 'LOGIN', 'user', user['id'], detail=username)
            db.commit()
            db.close()
            session.permanent = True
            syslog_logger.info(f"LOGIN user={username} ip={request.remote_addr}")
            # Prevent open redirect — only allow relative paths
            if not next_url or not next_url.startswith('/') or next_url.startswith('//'):
                next_url = url_for('dashboard')
            # Force password change if still using default
            try:
                _must_change = user['must_change_password']
            except (KeyError, IndexError):
                _must_change = False
            if _must_change:
                session['must_change_password'] = True
                return redirect(url_for('force_change_password'))
            # Document viewers always land on their restricted page.
            if session.get('is_document_viewer'):
                return redirect(url_for('viewer_home'))
            return redirect(next_url)
        else:
            # Constant-time failure: always hash something to prevent user enumeration
            if not user:
                check_password_hash(
                    'scrypt:32768:8:1$dummy$0000000000000000000000000000000000000000000000000000000000000000',
                    password,
                )
        db.close()
        flash('Invalid username or password.', 'error')

    return render_template('login.html', next=request.args.get('next', ''))


if _limiter_available and limiter:
    @app.route('/login', methods=['GET', 'POST'])
    @limiter.limit("15 per minute", methods=["POST"])
    def login():
        return _login_route()
else:
    @app.route('/login', methods=['GET', 'POST'])
    def login():
        return _login_route()


@app.route('/logout')
def logout():
    syslog_logger.info(f"LOGOUT user={session.get('username')}")
    if session.get('user_id'):
        db = get_db()
        log_audit(db, 'LOGOUT', 'user', session['user_id'])
        db.commit()
        db.close()
    session.clear()
    return redirect(url_for('login'))


@app.route('/admin/view-as', methods=['POST'])
@login_required
def admin_view_as():
    """Admin-only: temporarily view the site as a different role."""
    if session.get('user_role') != 'admin' and not session.get('_real_role'):
        return jsonify({'error': 'Forbidden'}), 403
    # If already in view-as mode, check the saved real role
    real_role = session.get('_real_role', session.get('user_role'))
    if real_role != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    data = request.get_json() or {}
    view_as = data.get('role', '')
    if view_as not in ('staff', 'user', 'readonly'):
        return jsonify({'error': 'Invalid role'}), 400
    # Save real values if not already saved
    if '_real_role' not in session:
        session['_real_role'] = session.get('user_role')
        session['_real_is_readonly'] = session.get('is_readonly', False)
        session['_real_is_content_admin'] = session.get('is_content_admin', False)
        session['_real_is_labor_scheduler'] = session.get('is_labor_scheduler', False)
        session['_real_is_scheduler'] = session.get('is_scheduler', False)
        session['_real_is_asset_manager'] = session.get('is_asset_manager', False)
    session['_view_as'] = view_as
    syslog_logger.info(f"ADMIN_VIEW_AS view_as={view_as} by={session.get('username')}")
    if view_as == 'readonly':
        session['user_role'] = 'user'
        session['is_readonly'] = True
        session['is_content_admin'] = False
        session['is_labor_scheduler'] = False
        session['is_scheduler'] = False
        session['is_asset_manager'] = False
    elif view_as == 'user':
        session['user_role'] = 'user'
        session['is_readonly'] = False
        session['is_content_admin'] = False
        session['is_labor_scheduler'] = False
        session['is_scheduler'] = False
        session['is_asset_manager'] = False
    elif view_as == 'staff':
        session['user_role'] = 'staff'
        session['is_readonly'] = False
        session['is_content_admin'] = False
        session['is_labor_scheduler'] = False
        session['is_scheduler'] = False
        session['is_asset_manager'] = False
    return jsonify({'success': True, 'view_as': view_as})


@app.route('/admin/view-as/reset', methods=['POST'])
@login_required
def admin_view_as_reset():
    """Restore the admin's real role after view-as preview."""
    if '_real_role' not in session:
        return jsonify({'success': True})
    session['user_role'] = session.pop('_real_role')
    session['is_readonly'] = session.pop('_real_is_readonly', False)
    session['is_content_admin'] = session.pop('_real_is_content_admin', False)
    session['is_labor_scheduler'] = session.pop('_real_is_labor_scheduler', False)
    session['is_scheduler'] = session.pop('_real_is_scheduler', False)
    session['is_asset_manager'] = session.pop('_real_is_asset_manager', False)
    session.pop('_view_as', None)
    syslog_logger.info(f"ADMIN_VIEW_AS_RESET by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/change-password', methods=['GET', 'POST'])
@login_required
def force_change_password():
    """Force password change screen (shown after login when must_change_password is set)."""
    if not session.get('must_change_password'):
        return redirect(url_for('dashboard'))
    if request.method == 'POST':
        new_pw = request.form.get('new_password', '')
        confirm = request.form.get('confirm_password', '')
        if new_pw != confirm:
            flash('Passwords do not match.', 'error')
            return render_template('force_change_password.html', user=get_current_user())
        pw_err = _validate_password(new_pw)
        if pw_err:
            flash(pw_err, 'error')
            return render_template('force_change_password.html', user=get_current_user())
        db = get_db()
        db.execute('UPDATE users SET password_hash=?, must_change_password=0 WHERE id=?',
                   (generate_password_hash(new_pw), session['user_id']))
        db.commit()
        db.close()
        session.pop('must_change_password', None)
        syslog_logger.info(f"FORCED_PASSWORD_CHANGE user_id={session['user_id']}")
        flash('Password changed successfully.', 'success')
        return redirect(url_for('dashboard'))
    return render_template('force_change_password.html', user=get_current_user())


@app.before_request
def _enforce_password_change():
    """Block all routes except logout/change-password if must_change_password is set."""
    if session.get('must_change_password'):
        allowed = ('force_change_password', 'logout', 'static')
        if request.endpoint not in allowed:
            return redirect(url_for('force_change_password'))


# ─── Dashboard ────────────────────────────────────────────────────────────────

@app.route('/')
@login_required
def index():
    return redirect(url_for('dashboard'))


@app.route('/dashboard')
@login_required
def dashboard():
    auto_archive_past_shows()
    accessible = get_accessible_shows(session['user_id'])
    db = get_db()

    _eff_date = """COALESCE(s.show_date,
        (SELECT MIN(perf_date) FROM show_performances
         WHERE show_id=s.id AND perf_date IS NOT NULL))"""

    if accessible is None:
        active = db.execute(f"""
            SELECT s.*, u.display_name as creator,
              (SELECT COUNT(*) FROM show_performances WHERE show_id=s.id) as perf_count,
              {_eff_date} as show_date
            FROM shows s LEFT JOIN users u ON s.created_by = u.id
            WHERE s.status = 'active'
            ORDER BY {_eff_date} ASC NULLS LAST
        """).fetchall()
        archived = db.execute(f"""
            SELECT s.*, u.display_name as creator,
              (SELECT COUNT(*) FROM show_performances WHERE show_id=s.id) as perf_count,
              {_eff_date} as show_date
            FROM shows s LEFT JOIN users u ON s.created_by = u.id
            WHERE s.status = 'archived'
            ORDER BY {_eff_date} DESC
            LIMIT 30
        """).fetchall()
    else:
        if accessible:
            placeholders = ','.join('?' * len(accessible))
            active = db.execute(f"""
                SELECT s.*, u.display_name as creator,
                  (SELECT COUNT(*) FROM show_performances WHERE show_id=s.id) as perf_count,
                  {_eff_date} as show_date
                FROM shows s LEFT JOIN users u ON s.created_by = u.id
                WHERE s.status = 'active' AND s.id IN ({placeholders})
                ORDER BY {_eff_date} ASC NULLS LAST
            """, accessible).fetchall()
            archived = db.execute(f"""
                SELECT s.*, u.display_name as creator,
                  (SELECT COUNT(*) FROM show_performances WHERE show_id=s.id) as perf_count,
                  {_eff_date} as show_date
                FROM shows s LEFT JOIN users u ON s.created_by = u.id
                WHERE s.status = 'archived' AND s.id IN ({placeholders})
                ORDER BY {_eff_date} DESC LIMIT 30
            """, accessible).fetchall()
        else:
            active = []
            archived = []

    # Attach full performance list per show (for multi-date display on card)
    def _attach_perfs(rows):
        if not rows:
            return []
        ids = [r['id'] for r in rows]
        ph = ','.join('?' * len(ids))
        perfs = db.execute(
            f"""SELECT show_id, perf_date, perf_time FROM show_performances
                WHERE show_id IN ({ph})
                ORDER BY CASE WHEN perf_date IS NULL THEN 1 ELSE 0 END,
                         perf_date, perf_time, id""",
            ids
        ).fetchall()
        by_show = {}
        for p in perfs:
            pd = p['perf_date']
            if pd is not None and not isinstance(pd, str):
                try: pd = pd.strftime('%Y-%m-%d')
                except AttributeError: pd = str(pd)
            by_show.setdefault(p['show_id'], []).append(
                {'perf_date': pd, 'perf_time': p['perf_time']}
            )
        out = []
        for r in rows:
            d = dict(r)
            d['performances'] = by_show.get(r['id'], [])
            out.append(d)
        return out

    active = _attach_perfs(active)
    archived = _attach_perfs(archived)

    db.close()
    restricted = session.get('is_restricted', False)

    # Group active shows by venue for column layout
    _venue_map = {}
    for s in active:
        v = (s['venue'] or '').strip() or 'Unassigned'
        _venue_map.setdefault(v, []).append(s)
    _names = sorted([v for v in _venue_map if v != 'Unassigned'], key=str.lower)
    if 'Unassigned' in _venue_map:
        _names.append('Unassigned')
    for v in _names:
        _venue_map[v].sort(key=lambda s: (s['show_date'] is None, s['show_date'] or ''))
    venue_groups = [(v, _venue_map[v]) for v in _names]

    return render_template('dashboard.html',
                           active_shows=active,
                           archived_shows=archived,
                           venue_groups=venue_groups,
                           restricted=restricted,
                           motd_messages=get_active_messages(session.get('user_id'), 'motd'),
                           user=get_current_user())


# ─── New Show ─────────────────────────────────────────────────────────────────

@app.route('/shows/new', methods=['GET', 'POST'])
@login_required
def new_show():
    if session.get('is_restricted'):
        abort(403)

    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        if not name:
            flash('Show name is required.', 'error')
            return render_template('new_show.html', user=get_current_user())

        show_date = request.form.get('show_date') or None
        show_time = request.form.get('show_time', '')
        venue     = request.form.get('venue', "Judson's Live")

        db = get_db()
        cur = db.execute("""
            INSERT INTO shows (name, show_date, show_time, venue, created_by)
            VALUES (?, ?, ?, ?, ?)
        """, (name, show_date, show_time, venue, session['user_id']))
        show_id = cur.lastrowid

        for key, val in [('show_name', name), ('show_date', show_date or ''),
                         ('show_time', show_time), ('venue', venue)]:
            if val:
                db.execute("""
                    INSERT OR REPLACE INTO advance_data (show_id, field_key, field_value)
                    VALUES (?, ?, ?)
                """, (show_id, key, val))

        if show_date:
            db.execute("""
                INSERT INTO show_performances (show_id, perf_date, perf_time, sort_order)
                VALUES (?, ?, ?, 0)
            """, (show_id, show_date, show_time))

        log_audit(db, 'SHOW_CREATE', 'show', show_id, show_id=show_id,
                  after={'name': name, 'show_date': show_date, 'venue': venue})
        db.commit()
        db.close()
        syslog_logger.info(f"SHOW_CREATE show_id={show_id} name={name} by={session.get('username')}")
        return redirect(url_for('show_page', show_id=show_id, tab='advance'))

    return render_template('new_show.html', user=get_current_user())


# ─── Show Page (all tabs) ─────────────────────────────────────────────────────

@app.route('/shows/<int:show_id>')
@login_required
def show_page(show_id):
    if not can_access_show(session['user_id'], show_id):
        abort(403)

    tab = request.args.get('tab', 'advance')
    db = get_db()
    show = db.execute('SELECT * FROM shows WHERE id = ?', (show_id,)).fetchone()
    if not show:
        abort(404)

    # Last-saved-by info
    last_saved_display_name = None
    last_saved_at = None
    try:
        _last_saved_by_id = show['last_saved_by']
    except (IndexError, KeyError):
        _last_saved_by_id = None
    if _last_saved_by_id:
        saver = db.execute(
            'SELECT display_name, username FROM users WHERE id=?',
            (_last_saved_by_id,)
        ).fetchone()
        if saver:
            last_saved_display_name = saver['display_name'] or saver['username']
        try:
            last_saved_at = show['last_saved_at']
        except (IndexError, KeyError):
            last_saved_at = None

    # Advance data
    adv_rows = db.execute(
        'SELECT field_key, field_value FROM advance_data WHERE show_id = ?', (show_id,)
    ).fetchall()
    advance_data = {r['field_key']: r['field_value'] for r in adv_rows}

    # Production schedule
    sched_rows = db.execute("""
        SELECT * FROM schedule_rows WHERE show_id = ?
        ORDER BY sort_order, id
    """, (show_id,)).fetchall()
    meta_rows = db.execute(
        'SELECT field_key, field_value FROM schedule_meta WHERE show_id = ?', (show_id,)
    ).fetchall()
    schedule_meta = {r['field_key']: r['field_value'] for r in meta_rows}

    # Post-show notes
    note_rows = db.execute(
        'SELECT field_key, field_value FROM post_show_notes WHERE show_id = ?', (show_id,)
    ).fetchall()
    notes_data = {r['field_key']: r['field_value'] for r in note_rows}

    # Performances
    performances = db.execute("""
        SELECT * FROM show_performances WHERE show_id = ?
        ORDER BY CASE WHEN perf_date IS NULL THEN 1 ELSE 0 END, perf_date, perf_time, id
    """, (show_id,)).fetchall()
    performances = [dict(p) for p in performances]
    for _p in performances:
        _pd = _p.get('perf_date')
        if _pd is not None and not isinstance(_pd, str):
            try: _p['perf_date'] = _pd.strftime('%Y-%m-%d')
            except AttributeError: _p['perf_date'] = str(_pd)

    # Group performances by date, then extend with every date in the
    # load-in → load-out range so the schedule has a day tab for each
    # calendar day of the production.
    schedule_days = []
    _date_to_idx = {}
    for _p in performances:
        key = _p.get('perf_date') or f"__null_{_p['id']}"
        idx = _date_to_idx.get(key)
        if idx is None:
            idx = len(schedule_days)
            _date_to_idx[key] = idx
            schedule_days.append({
                'perf_date': _p.get('perf_date'),
                'date_key': _p.get('perf_date') or key,
                'perfs': [],
                'perf_ids': [],
            })
        schedule_days[idx]['perfs'].append(_p)
        schedule_days[idx]['perf_ids'].append(_p['id'])

    # Add any load-in → load-out calendar dates not already covered by a perf
    def _iso(v):
        s = str(v) if v is not None else ''
        return s[:10] if s else ''
    _li = _iso(show['load_in_date'])
    _lo = _iso(show['load_out_date'])
    if _li and _lo:
        try:
            from datetime import date as _date, timedelta as _td
            _d1 = _date.fromisoformat(_li)
            _d2 = _date.fromisoformat(_lo)
            if _d2 >= _d1:
                _cur = _d1
                while _cur <= _d2:
                    _k = _cur.isoformat()
                    if _k not in _date_to_idx:
                        _date_to_idx[_k] = len(schedule_days)
                        schedule_days.append({
                            'perf_date': _k,
                            'date_key': _k,
                            'perfs': [],
                            'perf_ids': [],
                        })
                    _cur = _cur + _td(days=1)
        except (ValueError, TypeError):
            pass

    # Sort days by calendar date (unscheduled/null keys sort last)
    def _sort_key(d):
        pd = d.get('perf_date')
        return (0, pd) if pd else (1, d.get('date_key') or '')
    schedule_days.sort(key=_sort_key)

    for d in schedule_days:
        d['primary_perf_id'] = d['perf_ids'][0] if d['perf_ids'] else None
        if not d.get('date_key'):
            d['date_key'] = d.get('perf_date') or f"__null_{d['primary_perf_id']}"

    # Export log
    exports = db.execute("""
        SELECT e.*, u.display_name as exporter
        FROM export_log e LEFT JOIN users u ON e.exported_by = u.id
        WHERE e.show_id = ?
        ORDER BY e.exported_at DESC
        LIMIT 10
    """, (show_id,)).fetchall()

    # Contacts
    contacts = db.execute('SELECT * FROM contacts ORDER BY department, name').fetchall()
    contacts_by_dept = {}
    for c in contacts:
        dept = c['department'] or 'Other'
        contacts_by_dept.setdefault(dept, []).append(dict(c))

    # All users — for @mention autocomplete in comments
    all_users_rows = db.execute(
        'SELECT id, username, display_name FROM users ORDER BY display_name'
    ).fetchall()
    all_users = [{'id': u['id'], 'username': u['username'],
                  'display_name': u['display_name'] or u['username']} for u in all_users_rows]

    arts_groups = [dict(r) for r in db.execute(
        'SELECT id, name FROM arts_groups ORDER BY sort_order, name'
    ).fetchall()]

    db.close()

    form_sections = get_form_fields_for_template()
    restricted = session.get('is_restricted', False)
    can_edit_advance = (
        not session.get('is_readonly')
        and not restricted
        and can_access_show(session['user_id'], show_id)
    )

    # Global WiFi for schedule display (no longer a per-show editable field)
    global_wifi_network  = get_app_setting('wifi_network', '')
    global_wifi_password = get_app_setting('wifi_password', '')

    # Schedule templates for the schedule tab
    db2 = get_db()
    sched_templates = [dict(r) for r in db2.execute(
        'SELECT id, name FROM schedule_templates ORDER BY sort_order, name'
    ).fetchall()]

    # Labor requests for this show
    labor_rows = db2.execute("""
        SELECT lr.*, jp.name as position_name,
               cm.name as scheduled_crew_name
        FROM labor_requests lr
        LEFT JOIN job_positions jp ON lr.position_id = jp.id
        LEFT JOIN crew_members cm ON lr.scheduled_crew_member_id = cm.id
        WHERE lr.show_id = ?
        ORDER BY lr.sort_order, lr.id
    """, (show_id,)).fetchall()
    labor_requests_data = [_normalize_row_dates(dict(r)) for r in labor_rows]

    # Asset categories (for the Assets tab)
    asset_cats = db2.execute('SELECT * FROM asset_categories ORDER BY sort_order, name').fetchall()
    asset_categories_for_tab = [dict(c) for c in asset_cats]
    db2.close()

    return render_template('show.html',
                           show=show,
                           tab=tab,
                           advance_data=advance_data,
                           performances=performances,
                           schedule_days=schedule_days,
                           schedule_rows=[dict(r) for r in sched_rows],
                           schedule_meta=schedule_meta,
                           sched_meta_fields=get_schedule_meta_fields(),
                           notes_data=notes_data,
                           exports=exports,
                           contacts_by_dept=contacts_by_dept,
                           arts_groups=arts_groups,
                           departments=DEPARTMENTS,
                           form_sections=form_sections,
                           last_saved_display_name=last_saved_display_name,
                           last_saved_at=last_saved_at,
                           restricted=restricted,
                           all_users=all_users,
                           global_wifi_network=global_wifi_network,
                           global_wifi_password=global_wifi_password,
                           sched_templates=sched_templates,
                           labor_requests_data=labor_requests_data,
                           asset_categories=asset_categories_for_tab,
                           is_content_admin_user=session.get('is_content_admin', False),
                           can_edit_advance=can_edit_advance,
                           ollama_enabled=get_app_setting('ollama_enabled', '0') == '1',
                           user=get_current_user())


# ─── Save Endpoints (AJAX) ────────────────────────────────────────────────────

@app.route('/shows/<int:show_id>/save/advance', methods=['POST'])
@login_required
def save_advance(show_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    if session.get('is_restricted'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403

    get_show_or_404(show_id)
    data = request.get_json(force=True) or {}
    db = get_db()

    # Field keys whose type is arts_group_dropdown — any new value gets
    # auto-created in the global arts_groups table on save.
    arts_group_keys = {
        r['field_key'] for r in db.execute(
            "SELECT field_key FROM form_fields WHERE field_type='arts_group_dropdown'"
        ).fetchall()
    }

    for key, value in data.items():
        db.execute("""
            INSERT OR REPLACE INTO advance_data (show_id, field_key, field_value, updated_at)
            VALUES (?, ?, ?, CURRENT_TIMESTAMP)
        """, (show_id, key, str(value) if value is not None else ''))
        if key in arts_group_keys:
            name = (str(value) if value is not None else '').strip()
            if name:
                try:
                    max_order = db.execute(
                        'SELECT MAX(sort_order) FROM arts_groups'
                    ).fetchone()[0] or 0
                    db.execute(
                        'INSERT OR IGNORE INTO arts_groups (name, sort_order) VALUES (?, ?)',
                        (name, max_order + 10)
                    )
                except Exception as e:
                    app.logger.warning(f'arts_groups upsert failed for {name!r}: {e}')

    # Sync core show fields
    if 'show_name' in data and data['show_name']:
        db.execute('UPDATE shows SET name=?, updated_at=CURRENT_TIMESTAMP WHERE id=?',
                   (data['show_name'], show_id))
    if 'show_date' in data:
        db.execute('UPDATE shows SET show_date=?, updated_at=CURRENT_TIMESTAMP WHERE id=?',
                   (data['show_date'] or None, show_id))
    if 'show_time' in data:
        db.execute('UPDATE shows SET show_time=?, updated_at=CURRENT_TIMESTAMP WHERE id=?',
                   (data['show_time'], show_id))
    if 'venue' in data:
        db.execute('UPDATE shows SET venue=?, updated_at=CURRENT_TIMESTAMP WHERE id=?',
                   (data['venue'], show_id))
    if 'load_in_date' in data:
        val = data['load_in_date'].strip() if data['load_in_date'] else None
        db.execute('UPDATE shows SET load_in_date=?, updated_at=CURRENT_TIMESTAMP WHERE id=?',
                   (val or None, show_id))
    if 'load_in_time' in data:
        db.execute('UPDATE shows SET load_in_time=?, updated_at=CURRENT_TIMESTAMP WHERE id=?',
                   (data['load_in_time'].strip() if data['load_in_time'] else '', show_id))
    if 'load_out_date' in data:
        val = data['load_out_date'].strip() if data['load_out_date'] else None
        db.execute('UPDATE shows SET load_out_date=?, updated_at=CURRENT_TIMESTAMP WHERE id=?',
                   (val or None, show_id))
    if 'load_out_time' in data:
        db.execute('UPDATE shows SET load_out_time=?, updated_at=CURRENT_TIMESTAMP WHERE id=?',
                   (data['load_out_time'].strip() if data['load_out_time'] else '', show_id))

    # Track last saved
    db.execute("""
        UPDATE shows SET last_saved_by=?, last_saved_at=CURRENT_TIMESTAMP WHERE id=?
    """, (session['user_id'], show_id))

    # Version snapshot
    _snapshot_form_history(db, show_id, 'advance', {'advance_data': data})
    log_audit(db, 'FORM_SAVE', 'form', show_id, show_id=show_id, detail='type=advance')

    db.commit()
    db.close()
    syslog_logger.info(f"FORM_SAVE show_id={show_id} type=advance by={session.get('username')}")
    return jsonify({'success': True})


# ─── Performances (multiple dates/times per show) ─────────────────────────────

def _normalize_perf_time(raw):
    """Accept '19:00', '1900', '7:00', '' — return canonical 'HH:MM' or ''."""
    s = (raw or '').strip()
    if not s:
        return ''
    if ':' in s:
        parts = s.split(':', 1)
        try:
            h = int(parts[0]); m = int(parts[1])
        except ValueError:
            return s
    elif s.isdigit() and len(s) in (3, 4):
        h = int(s[:-2]); m = int(s[-2:])
    else:
        return s
    if 0 <= h <= 23 and 0 <= m <= 59:
        return f'{h:02d}:{m:02d}'
    return s


def _perf_to_json(row):
    """Serialize a performance row with perf_date as YYYY-MM-DD string."""
    d = dict(row)
    pd = d.get('perf_date')
    if pd is not None and not isinstance(pd, str):
        try:
            d['perf_date'] = pd.strftime('%Y-%m-%d')
        except AttributeError:
            d['perf_date'] = str(pd)
    return d


@app.route('/shows/<int:show_id>/performances', methods=['POST'])
@login_required
def add_performance(show_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    if session.get('is_restricted'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403
    get_show_or_404(show_id)
    data = request.get_json(force=True) or {}
    perf_date = data.get('perf_date') or None
    perf_time = _normalize_perf_time(data.get('perf_time'))
    db = get_db()
    cur = db.execute("""
        INSERT INTO show_performances (show_id, perf_date, perf_time, sort_order)
        VALUES (?, ?, ?,
          (SELECT COALESCE(MAX(sort_order)+1, 0) FROM show_performances WHERE show_id=?))
    """, (show_id, perf_date, perf_time, show_id))
    perf_id = cur.lastrowid
    _sync_show_primary_date(db, show_id)
    db.commit()
    perf = db.execute('SELECT * FROM show_performances WHERE id=?', (perf_id,)).fetchone()
    db.close()
    return jsonify({'success': True, 'performance': _perf_to_json(perf)})


@app.route('/shows/<int:show_id>/performances/<int:perf_id>', methods=['PUT'])
@login_required
def update_performance(show_id, perf_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    if session.get('is_restricted'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403
    db = get_db()
    perf = db.execute(
        'SELECT * FROM show_performances WHERE id=? AND show_id=?', (perf_id, show_id)
    ).fetchone()
    if not perf:
        db.close()
        return jsonify({'success': False, 'error': 'Not found.'}), 404
    data = request.get_json(force=True) or {}
    db.execute("""
        UPDATE show_performances SET perf_date=?, perf_time=? WHERE id=?
    """, (data.get('perf_date') or None, _normalize_perf_time(data.get('perf_time')), perf_id))
    _sync_show_primary_date(db, show_id)
    db.commit()
    perf = db.execute('SELECT * FROM show_performances WHERE id=?', (perf_id,)).fetchone()
    db.close()
    return jsonify({'success': True, 'performance': _perf_to_json(perf)})


@app.route('/shows/<int:show_id>/performances/<int:perf_id>', methods=['DELETE'])
@login_required
def delete_performance(show_id, perf_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    if session.get('is_restricted'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403
    db = get_db()
    perf = db.execute(
        'SELECT * FROM show_performances WHERE id=? AND show_id=?', (perf_id, show_id)
    ).fetchone()
    if not perf:
        db.close()
        return jsonify({'success': False, 'error': 'Not found.'}), 404
    db.execute('DELETE FROM show_performances WHERE id=?', (perf_id,))
    _sync_show_primary_date(db, show_id)
    db.commit()
    db.close()
    return jsonify({'success': True})


@app.route('/shows/<int:show_id>/save/schedule', methods=['POST'])
@login_required
def save_schedule(show_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    if session.get('is_restricted'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403

    get_show_or_404(show_id)
    data = request.get_json(force=True) or {}
    db = get_db()
    if 'meta' in data:
        for key, val in data['meta'].items():
            db.execute("""
                INSERT OR REPLACE INTO schedule_meta (show_id, field_key, field_value)
                VALUES (?, ?, ?)
            """, (show_id, key, val or ''))
    if 'rows' in data:
        db.execute('DELETE FROM schedule_rows WHERE show_id = ?', (show_id,))
        for i, row in enumerate(data['rows']):
            perf_id = row.get('perf_id')  # None for single-day / first day
            day_date = row.get('day_date') or None
            if isinstance(day_date, str):
                day_date = day_date.strip() or None
            db.execute("""
                INSERT INTO schedule_rows (show_id, perf_id, day_date, sort_order, start_time, end_time, description, notes)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (show_id, perf_id, day_date, i,
                  row.get('start_time', ''), row.get('end_time', ''),
                  row.get('description', ''), row.get('notes', '')))

    db.execute('UPDATE shows SET updated_at=CURRENT_TIMESTAMP WHERE id=?', (show_id,))
    db.execute("""
        UPDATE shows SET last_saved_by=?, last_saved_at=CURRENT_TIMESTAMP WHERE id=?
    """, (session['user_id'], show_id))

    _snapshot_form_history(db, show_id, 'schedule', data)
    log_audit(db, 'FORM_SAVE', 'form', show_id, show_id=show_id, detail='type=schedule')

    db.commit()
    db.close()
    syslog_logger.info(f"FORM_SAVE show_id={show_id} type=schedule by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/shows/<int:show_id>/save/postnotes', methods=['POST'])
@login_required
def save_postnotes(show_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    if session.get('is_restricted'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403

    get_show_or_404(show_id)
    data = request.get_json(force=True) or {}
    db = get_db()
    for key, val in data.items():
        db.execute("""
            INSERT OR REPLACE INTO post_show_notes (show_id, field_key, field_value)
            VALUES (?, ?, ?)
        """, (show_id, key, val or ''))
    db.execute('UPDATE shows SET updated_at=CURRENT_TIMESTAMP WHERE id=?', (show_id,))
    db.execute("""
        UPDATE shows SET last_saved_by=?, last_saved_at=CURRENT_TIMESTAMP WHERE id=?
    """, (session['user_id'], show_id))

    _snapshot_form_history(db, show_id, 'postnotes', {'notes_data': data})
    log_audit(db, 'FORM_SAVE', 'form', show_id, show_id=show_id, detail='type=postnotes')

    db.commit()
    db.close()
    syslog_logger.info(f"FORM_SAVE show_id={show_id} type=postnotes by={session.get('username')}")
    return jsonify({'success': True})


# ─── Version History ──────────────────────────────────────────────────────────

@app.route('/shows/<int:show_id>/history/<form_type>')
@login_required
def form_history_list(show_id, form_type):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    db = get_db()
    entries = db.execute("""
        SELECT fh.id, fh.saved_at, fh.form_type,
               u.display_name as saved_by_name, u.username as saved_by_username
        FROM form_history fh
        LEFT JOIN users u ON fh.saved_by = u.id
        WHERE fh.show_id = ? AND fh.form_type = ?
        ORDER BY fh.saved_at DESC
        LIMIT 50
    """, (show_id, form_type)).fetchall()
    db.close()
    return jsonify([{
        'id': e['id'],
        'saved_at': e['saved_at'],
        'form_type': e['form_type'],
        'saved_by_name': e['saved_by_name'] or e['saved_by_username'] or 'Unknown',
    } for e in entries])


@app.route('/shows/<int:show_id>/history/<int:hist_id>/snapshot')
@login_required
def history_snapshot(show_id, hist_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    db = get_db()
    entry = db.execute(
        'SELECT * FROM form_history WHERE id=? AND show_id=?', (hist_id, show_id)
    ).fetchone()
    db.close()
    if not entry:
        return jsonify({'success': False, 'error': 'Snapshot not found.'}), 404
    return jsonify({'id': entry['id'], 'form_type': entry['form_type'],
                    'saved_at': entry['saved_at'],
                    'data': json.loads(entry['snapshot_json'])})


@app.route('/shows/<int:show_id>/history/<int:hist_id>/restore', methods=['POST'])
@login_required
def restore_history(show_id, hist_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    if session.get('is_restricted'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403

    db = get_db()
    entry = db.execute(
        'SELECT * FROM form_history WHERE id=? AND show_id=?', (hist_id, show_id)
    ).fetchone()
    if not entry:
        db.close()
        return jsonify({'success': False, 'error': 'Snapshot not found.'}), 404

    snapshot = json.loads(entry['snapshot_json'])
    form_type = entry['form_type']

    # Check for newer snapshots — warn the user before overwriting
    force = request.args.get('force') == '1'
    if not force:
        newer = db.execute("""
            SELECT fh.id, fh.saved_at, u.username, u.display_name, fh.snapshot_json
            FROM form_history fh
            LEFT JOIN users u ON fh.saved_by = u.id
            WHERE fh.show_id=? AND fh.form_type=? AND fh.id > ?
            ORDER BY fh.saved_at DESC LIMIT 1
        """, (show_id, form_type, hist_id)).fetchone()
        if newer:
            db.close()
            return jsonify({
                'conflict':          True,
                'newer_saved_at':    newer['saved_at'],
                'newer_saved_by':    newer['display_name'] or newer['username'] or 'Unknown',
                'restoring_snapshot': snapshot,
                'current_snapshot':  json.loads(newer['snapshot_json']),
            }), 409

    if form_type == 'advance':
        adv = snapshot.get('advance_data', {})
        for key, val in adv.items():
            db.execute("""
                INSERT OR REPLACE INTO advance_data (show_id, field_key, field_value, updated_at)
                VALUES (?, ?, ?, CURRENT_TIMESTAMP)
            """, (show_id, key, str(val) if val is not None else ''))

    elif form_type == 'schedule':
        if 'meta' in snapshot:
            for key, val in snapshot['meta'].items():
                db.execute("""
                    INSERT OR REPLACE INTO schedule_meta (show_id, field_key, field_value)
                    VALUES (?, ?, ?)
                """, (show_id, key, val or ''))
        if 'rows' in snapshot:
            db.execute('DELETE FROM schedule_rows WHERE show_id=?', (show_id,))
            for i, row in enumerate(snapshot['rows']):
                db.execute("""
                    INSERT INTO schedule_rows (show_id, sort_order, start_time, end_time, description, notes)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (show_id, i, row.get('start_time',''), row.get('end_time',''),
                      row.get('description',''), row.get('notes','')))

    elif form_type == 'postnotes':
        notes = snapshot.get('notes_data', {})
        for key, val in notes.items():
            db.execute("""
                INSERT OR REPLACE INTO post_show_notes (show_id, field_key, field_value)
                VALUES (?, ?, ?)
            """, (show_id, key, val or ''))

    db.execute("""
        UPDATE shows SET last_saved_by=?, last_saved_at=CURRENT_TIMESTAMP WHERE id=?
    """, (session['user_id'], show_id))
    log_audit(db, 'HISTORY_RESTORE', 'form', hist_id, show_id=show_id,
              detail=f'type={form_type}')
    db.commit()
    db.close()
    syslog_logger.info(
        f"HISTORY_RESTORE show_id={show_id} hist_id={hist_id} type={form_type} by={session.get('username')}"
    )
    return jsonify({'success': True})


# ─── Comments ─────────────────────────────────────────────────────────────────

@app.route('/shows/<int:show_id>/comments', methods=['GET'])
@login_required
def get_comments(show_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    is_admin = session.get('user_role') == 'admin'
    db = get_db()
    rows = db.execute("""
        SELECT sc.id, sc.body, sc.created_at, sc.edited_at, sc.deleted_at,
               u.display_name, u.username, u.id as uid,
               du.display_name as deleted_by_name, du.username as deleted_by_username
        FROM show_comments sc
        JOIN users u ON sc.user_id = u.id
        LEFT JOIN users du ON sc.deleted_by = du.id
        WHERE sc.show_id = ? AND (sc.deleted_at IS NULL OR ?)
        ORDER BY sc.created_at ASC
    """, (show_id, is_admin)).fetchall()
    db.close()
    result = []
    for r in rows:
        author = r['display_name'] or r['username']
        entry = {
            'id':          r['id'],
            'body':        r['body'],
            'created_at':  r['created_at'],
            'edited_at':   r['edited_at'],
            'deleted_at':  r['deleted_at'],
            'author':      author,
            'author_id':   r['uid'],
            'initials':    ''.join(w[0].upper() for w in author.split()[:2]),
            'is_own':      r['uid'] == session['user_id'],
        }
        if is_admin and r['deleted_at']:
            entry['deleted_by'] = r['deleted_by_name'] or r['deleted_by_username'] or 'Unknown'
        result.append(entry)
    return jsonify(result)


@app.route('/shows/<int:show_id>/comments', methods=['POST'])
@login_required
def post_comment(show_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    if session.get('is_restricted'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403
    data = request.get_json(force=True) or {}
    body = data.get('body', '').strip()
    if not body:
        return jsonify({'success': False, 'error': 'Comment cannot be empty.'}), 400
    if len(body) > 2000:
        return jsonify({'success': False, 'error': 'Comment too long (max 2000 chars).'}), 400
    db = get_db()
    cur = db.execute(
        'INSERT INTO show_comments (show_id, user_id, body) VALUES (?, ?, ?)',
        (show_id, session['user_id'], body)
    )
    cid = cur.lastrowid
    log_audit(db, 'COMMENT_POST', 'comment', cid, show_id=show_id,
              after={'body': body})
    db.commit()
    row = db.execute("""
        SELECT sc.id, sc.body, sc.created_at,
               u.display_name, u.username, u.id as uid
        FROM show_comments sc JOIN users u ON sc.user_id = u.id
        WHERE sc.id = ?
    """, (cid,)).fetchone()
    db.close()
    syslog_logger.info(f"COMMENT_POST show_id={show_id} by={session.get('username')}")
    return jsonify({
        'success': True,
        'comment': {
            'id':        row['id'],
            'body':      row['body'],
            'created_at': row['created_at'],
            'author':    row['display_name'] or row['username'],
            'author_id': row['uid'],
            'initials':  ''.join(w[0].upper() for w in (row['display_name'] or row['username']).split()[:2]),
            'is_own':    True,
        }
    })


@app.route('/shows/<int:show_id>/comments/<int:cid>/delete', methods=['POST'])
@login_required
def delete_comment(show_id, cid):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    db = get_db()
    comment = db.execute(
        'SELECT * FROM show_comments WHERE id=? AND show_id=? AND deleted_at IS NULL',
        (cid, show_id)
    ).fetchone()
    if not comment:
        db.close()
        return jsonify({'success': False, 'error': 'Comment not found.'}), 404
    if comment['user_id'] != session['user_id'] and session.get('user_role') != 'admin':
        db.close()
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    log_audit(db, 'COMMENT_DELETE', 'comment', cid, show_id=show_id,
              before={'body': comment['body']})
    db.execute(
        'UPDATE show_comments SET deleted_at=CURRENT_TIMESTAMP, deleted_by=? WHERE id=?',
        (session['user_id'], cid)
    )
    db.commit()
    db.close()
    syslog_logger.info(f"COMMENT_DELETE show_id={show_id} comment_id={cid} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/shows/<int:show_id>/comments/<int:cid>', methods=['PUT'])
@login_required
def edit_comment(show_id, cid):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    if session.get('is_restricted'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403
    data = request.get_json(force=True) or {}
    new_body = data.get('body', '').strip()
    if not new_body:
        return jsonify({'success': False, 'error': 'Comment cannot be empty.'}), 400
    if len(new_body) > 2000:
        return jsonify({'success': False, 'error': 'Comment too long (max 2000 chars).'}), 400
    db = get_db()
    comment = db.execute(
        'SELECT * FROM show_comments WHERE id=? AND show_id=? AND deleted_at IS NULL',
        (cid, show_id)
    ).fetchone()
    if not comment:
        db.close()
        return jsonify({'success': False, 'error': 'Comment not found.'}), 404
    if comment['user_id'] != session['user_id'] and session.get('user_role') != 'admin':
        db.close()
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    old_body = comment['body']
    # Save previous version
    db.execute(
        'INSERT INTO comment_versions (comment_id, body, edited_by) VALUES (?,?,?)',
        (cid, old_body, session['user_id'])
    )
    db.execute(
        'UPDATE show_comments SET body=?, edited_at=CURRENT_TIMESTAMP WHERE id=?',
        (new_body, cid)
    )
    log_audit(db, 'COMMENT_EDIT', 'comment', cid, show_id=show_id,
              before={'body': old_body}, after={'body': new_body})
    db.commit()
    db.close()
    syslog_logger.info(f"COMMENT_EDIT show_id={show_id} comment_id={cid} by={session.get('username')}")
    return jsonify({'success': True, 'body': new_body})


@app.route('/shows/<int:show_id>/comments/<int:cid>/restore', methods=['POST'])
@admin_required
def restore_comment(show_id, cid):
    db = get_db()
    comment = db.execute(
        'SELECT * FROM show_comments WHERE id=? AND show_id=? AND deleted_at IS NOT NULL',
        (cid, show_id)
    ).fetchone()
    if not comment:
        db.close()
        return jsonify({'success': False, 'error': 'Comment not found.'}), 404
    db.execute(
        'UPDATE show_comments SET deleted_at=NULL, deleted_by=NULL WHERE id=?', (cid,)
    )
    log_audit(db, 'COMMENT_RESTORE', 'comment', cid, show_id=show_id)
    db.commit()
    db.close()
    syslog_logger.info(f"COMMENT_RESTORE show_id={show_id} comment_id={cid} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/shows/<int:show_id>/comments/<int:cid>/versions', methods=['GET'])
@admin_required
def comment_versions_list(show_id, cid):
    db = get_db()
    rows = db.execute("""
        SELECT cv.id, cv.body, cv.edited_at,
               u.display_name, u.username
        FROM comment_versions cv
        LEFT JOIN users u ON cv.edited_by = u.id
        WHERE cv.comment_id = ?
        ORDER BY cv.edited_at DESC
    """, (cid,)).fetchall()
    db.close()
    return jsonify([{
        'id':        r['id'],
        'body':      r['body'],
        'edited_at': r['edited_at'],
        'edited_by': r['display_name'] or r['username'] or 'Unknown',
    } for r in rows])


@app.route('/shows/<int:show_id>/comments/<int:cid>/versions/<int:vid>/restore', methods=['POST'])
@admin_required
def comment_version_restore(show_id, cid, vid):
    db = get_db()
    version = db.execute(
        'SELECT * FROM comment_versions WHERE id=? AND comment_id=?', (vid, cid)
    ).fetchone()
    if not version:
        db.close()
        return jsonify({'success': False, 'error': 'Version not found.'}), 404
    comment = db.execute(
        'SELECT body FROM show_comments WHERE id=? AND show_id=?', (cid, show_id)
    ).fetchone()
    if not comment:
        db.close()
        return jsonify({'success': False, 'error': 'Comment not found.'}), 404
    old_body = comment['body']
    # Save current as a version before restoring
    db.execute(
        'INSERT INTO comment_versions (comment_id, body, edited_by) VALUES (?,?,?)',
        (cid, old_body, session['user_id'])
    )
    db.execute(
        'UPDATE show_comments SET body=?, edited_at=CURRENT_TIMESTAMP WHERE id=?',
        (version['body'], cid)
    )
    log_audit(db, 'COMMENT_VERSION_RESTORE', 'comment', cid, show_id=show_id,
              before={'body': old_body}, after={'body': version['body']},
              detail=f'version_id={vid}')
    db.commit()
    db.close()
    syslog_logger.info(f"COMMENT_VERSION_RESTORE show_id={show_id} comment_id={cid} version_id={vid} by={session.get('username')}")
    return jsonify({'success': True, 'body': version['body']})


# ─── File Attachments ──────────────────────────────────────────────────────────

@app.route('/shows/<int:show_id>/attachments', methods=['GET'])
@login_required
def get_attachments(show_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    db = get_db()
    field_key = request.args.get('field_key')
    if field_key:
        rows = db.execute("""
            SELECT sa.id, sa.filename, sa.mime_type, sa.file_size, sa.created_at,
                   sa.field_key, sa.description,
                   u.display_name, u.username
            FROM show_attachments sa
            LEFT JOIN users u ON sa.uploaded_by = u.id
            WHERE sa.show_id = ? AND sa.field_key = ?
            ORDER BY sa.created_at ASC
        """, (show_id, field_key)).fetchall()
    else:
        rows = db.execute("""
            SELECT sa.id, sa.filename, sa.mime_type, sa.file_size, sa.created_at,
                   sa.field_key, sa.description,
                   u.display_name, u.username
            FROM show_attachments sa
            LEFT JOIN users u ON sa.uploaded_by = u.id
            WHERE sa.show_id = ?
            ORDER BY sa.created_at ASC
        """, (show_id,)).fetchall()
    db.close()
    return jsonify([{
        'id':         r['id'],
        'filename':   r['filename'],
        'mime_type':  r['mime_type'],
        'file_size':  r['file_size'],
        'created_at': r['created_at'],
        'field_key':  r['field_key'],
        'description': r['description'] or '',
        'uploader':   r['display_name'] or r['username'] or 'Unknown',
    } for r in rows])


@app.route('/shows/<int:show_id>/attachments', methods=['POST'])
@login_required
def upload_attachment(show_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    if session.get('is_restricted'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403
    f = request.files.get('file')
    if not f or not f.filename:
        return jsonify({'success': False, 'error': 'No file provided.'}), 400
    data = f.read()
    max_size = _get_upload_max()
    max_mb = max_size // (1024 * 1024)
    if len(data) > max_size:
        return jsonify({'success': False, 'error': f'File too large (max {max_mb} MB).'}), 413
    filename  = secure_filename(f.filename) or 'file'
    mime_type = f.content_type or 'application/octet-stream'
    # Optional per-form-field association (NULL = general show attachment).
    field_key = (request.form.get('field_key') or '').strip() or None
    description = (request.form.get('description') or '').strip()
    db = get_db()
    # Insert row first (without file data) to get the auto-assigned id
    cur = db.execute("""
        INSERT INTO show_attachments (show_id, uploaded_by, filename, mime_type, file_data, file_size, field_key, description)
        VALUES (?, ?, ?, ?, NULL, ?, ?, ?)
    """, (show_id, session['user_id'], filename, mime_type, len(data), field_key, description))
    aid = cur.lastrowid
    # Upload to S3; fall back to DB storage if S3 is unavailable
    if s3_storage.is_configured():
        try:
            s3_key = f"attachments/{show_id}/{aid}/{filename}"
            s3_storage.upload_file(s3_key, data, mime_type)
            db.execute('UPDATE show_attachments SET s3_key=? WHERE id=?', (s3_key, aid))
        except Exception as e:
            app.logger.warning(f"S3 upload failed for attachment {aid}, falling back to DB: {e}")
            syslog_logger.warning(f"S3_UPLOAD_FAILED table=show_attachments id={aid} show_id={show_id} error={e}")
            db.execute('UPDATE show_attachments SET file_data=? WHERE id=?', (data, aid))
    else:
        db.execute('UPDATE show_attachments SET file_data=? WHERE id=?', (data, aid))
    log_audit(db, 'FILE_UPLOAD', 'attachment', aid, show_id=show_id, detail=filename)
    db.commit()
    row = db.execute("""
        SELECT sa.id, sa.filename, sa.mime_type, sa.file_size, sa.created_at,
               u.display_name, u.username
        FROM show_attachments sa LEFT JOIN users u ON sa.uploaded_by = u.id
        WHERE sa.id = ?
    """, (aid,)).fetchone()
    db.close()
    syslog_logger.info(f"FILE_UPLOAD show_id={show_id} filename={filename} field_key={field_key} by={session.get('username')}")
    return jsonify({
        'success': True,
        'attachment': {
            'id':         row['id'],
            'filename':   row['filename'],
            'mime_type':  row['mime_type'],
            'file_size':  row['file_size'],
            'created_at': row['created_at'],
            'field_key':  field_key,
            'description': description,
            'uploader':   row['display_name'] or row['username'] or 'Unknown',
        }
    })


@app.route('/shows/<int:show_id>/attachments/<int:aid>/download')
@login_required
def download_attachment(show_id, aid):
    if not can_access_show(session['user_id'], show_id):
        abort(403)
    db = get_db()
    row = db.execute(
        'SELECT * FROM show_attachments WHERE id=? AND show_id=?', (aid, show_id)
    ).fetchone()
    db.close()
    if not row:
        abort(404)
    if row['s3_key']:
        try:
            data = s3_storage.download_file(row['s3_key'])
        except Exception as e:
            app.logger.error(f"S3 download failed for attachment {aid}: {e}")
            abort(503)
    elif row['file_data']:
        data = bytes(row['file_data'])
    else:
        abort(404)
    resp = make_response(data)
    resp.headers['Content-Type'] = row['mime_type']
    resp.headers['Content-Disposition'] = _safe_content_disposition(row['filename'])
    return resp


@app.route('/shows/<int:show_id>/attachments/<int:aid>/delete', methods=['POST'])
@login_required
def delete_attachment(show_id, aid):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    db = get_db()
    row = db.execute(
        'SELECT * FROM show_attachments WHERE id=? AND show_id=?', (aid, show_id)
    ).fetchone()
    if not row:
        db.close()
        return jsonify({'success': False, 'error': 'File not found.'}), 404
    if row['uploaded_by'] != session['user_id'] and session.get('user_role') != 'admin':
        db.close()
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    if row['s3_key']:
        try:
            s3_storage.delete_file(row['s3_key'])
        except Exception as e:
            app.logger.error(f"S3 delete failed for attachment {aid} key={row['s3_key']}: {e}")
            syslog_logger.error(f"S3_DELETE_FAILED table=show_attachments id={aid} show_id={show_id} error={e}")
    log_audit(db, 'FILE_DELETE', 'attachment', aid, show_id=show_id,
              detail=row['filename'] if row else str(aid))
    db.execute('DELETE FROM show_attachments WHERE id=?', (aid,))
    db.commit()
    db.close()
    syslog_logger.info(f"FILE_DELETE show_id={show_id} aid={aid} by={session.get('username')}")
    return jsonify({'success': True})


# ─── Read Receipts ─────────────────────────────────────────────────────────────

@app.route('/shows/<int:show_id>/read', methods=['POST'])
@login_required
def mark_advance_read(show_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    db = get_db()
    show = db.execute('SELECT advance_version FROM shows WHERE id=?', (show_id,)).fetchone()
    if not show:
        db.close()
        return jsonify({'success': False, 'error': 'Show not found.'}), 404
    version = show['advance_version'] or 0
    db.execute("""
        INSERT INTO advance_reads (show_id, user_id, version_read, read_at)
        VALUES (?, ?, ?, CURRENT_TIMESTAMP)
        ON CONFLICT(show_id, user_id) DO UPDATE SET
            version_read = excluded.version_read,
            read_at      = excluded.read_at
    """, (show_id, session['user_id'], version))
    db.commit()
    db.close()
    return jsonify({'success': True})


@app.route('/shows/<int:show_id>/reads')
@login_required
def get_advance_reads(show_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    db = get_db()
    rows = db.execute("""
        SELECT ar.version_read, ar.read_at,
               u.display_name, u.username, u.id as uid
        FROM advance_reads ar
        JOIN users u ON ar.user_id = u.id
        WHERE ar.show_id = ?
        ORDER BY ar.read_at DESC
    """, (show_id,)).fetchall()
    db.close()
    return jsonify([{
        'version_read':     r['version_read'],
        'read_at':          r['read_at'],
        'author':           r['display_name'] or r['username'],
        'initials':         ''.join(w[0].upper() for w in (r['display_name'] or r['username']).split()[:2]),
        'is_current_user':  r['uid'] == session['user_id'],
    } for r in rows])


# ─── Real-time Sync ───────────────────────────────────────────────────────────

def _upsert_active_session(db, user_id, show_id, tab, focused_field=None):
    """Record that a user is actively on a show page and prune stale sessions."""
    db.execute("""
        INSERT INTO active_sessions (user_id, show_id, tab, focused_field, last_seen)
        VALUES (?, ?, ?, ?, CURRENT_TIMESTAMP)
        ON CONFLICT(user_id, show_id) DO UPDATE SET
            tab=excluded.tab,
            focused_field=excluded.focused_field,
            last_seen=excluded.last_seen
    """, (user_id, show_id, tab, focused_field or None))
    # Prune sessions idle > 60 s
    db.execute("DELETE FROM active_sessions WHERE last_seen < datetime('now', '-60 seconds')")


def _get_other_active_users(db, user_id, show_id):
    """Return list of other users active on this show in the last 45 s."""
    rows = db.execute("""
        SELECT u.display_name, u.username, acs.tab, acs.focused_field
        FROM active_sessions acs
        JOIN users u ON acs.user_id = u.id
        WHERE acs.show_id = ?
          AND acs.user_id != ?
          AND acs.last_seen > datetime('now', '-45 seconds')
        ORDER BY acs.last_seen DESC
    """, (show_id, user_id)).fetchall()
    return [{
        'name':          r['display_name'] or r['username'],
        'initials':      ''.join(w[0].upper() for w in (r['display_name'] or r['username']).split()[:2]),
        'tab':           r['tab'],
        'focused_field': r['focused_field'],
    } for r in rows]


@app.route('/shows/<int:show_id>/sync/advance')
@login_required
def sync_advance(show_id):
    """
    Lightweight poll endpoint for real-time field sync.
    Query param:  since=<YYYY-MM-DD HH:MM:SS>  — last sync timestamp
                  tab=<advance|schedule|postnotes>  — caller's current tab
    Returns changed advance_data fields + active-user presence list.
    """
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403

    since         = request.args.get('since', '')
    tab           = request.args.get('tab', 'advance')
    focused_field = request.args.get('field') or None

    db = get_db()

    # Fields changed since last poll (exclude the current user's own saves so
    # we don't echo back what they just wrote)
    if since:
        changed_rows = db.execute("""
            SELECT ad.field_key, ad.field_value
            FROM advance_data ad
            WHERE ad.show_id = ?
              AND ad.updated_at > ?
              AND (
                SELECT last_saved_by FROM shows WHERE id = ad.show_id
              ) != ?
        """, (show_id, since, session['user_id'])).fetchall()
    else:
        changed_rows = []

    # New "since" cursor = latest updated_at across the whole show's advance data
    ts_row = db.execute(
        "SELECT MAX(updated_at) FROM advance_data WHERE show_id = ?", (show_id,)
    ).fetchone()
    new_since = ts_row[0] if ts_row and ts_row[0] else since

    # Update presence (including which field is focused) and get other active users
    _upsert_active_session(db, session['user_id'], show_id, tab, focused_field)
    others = _get_other_active_users(db, session['user_id'], show_id)

    db.commit()
    db.close()

    return jsonify({
        'since':        new_since,
        'fields':       {r['field_key']: r['field_value'] for r in changed_rows},
        'active_users': others,
    })


@app.route('/shows/<int:show_id>/heartbeat', methods=['POST'])
@login_required
def show_heartbeat(show_id):
    """
    Thin presence-only update for schedule/postnotes tabs where the advance
    sync poll isn't running. Also detects when another user saved non-advance
    data so the client can show a "someone else saved" notice.
    """
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403

    data          = request.get_json(force=True) or {}
    tab           = data.get('tab', 'advance')
    focused_field = data.get('focused_field') or None

    db = get_db()
    _upsert_active_session(db, session['user_id'], show_id, tab, focused_field)
    others = _get_other_active_users(db, session['user_id'], show_id)

    # For schedule / postnotes: tell the client if someone else saved recently
    # so it can show a "reload?" banner without fetching the full dataset.
    show = db.execute('SELECT last_saved_by, last_saved_at FROM shows WHERE id=?',
                      (show_id,)).fetchone()
    other_saved = (show and show['last_saved_by'] and
                   show['last_saved_by'] != session['user_id'])

    db.commit()
    db.close()

    return jsonify({
        'active_users': others,
        'other_saved':  other_saved,
        'last_saved_at': show['last_saved_at'] if show else None,
    })


# ─── PDF Export ───────────────────────────────────────────────────────────────

def _build_advance_pdf(show_id, exported_by_id=None, base_url=None):
    """
    Build the advance PDF.  Works both in a request context (base_url=None reads
    from request.url_root) and in a background context (pass base_url='/').
    exported_by_id defaults to session['user_id'] when not supplied.
    """
    if exported_by_id is None:
        exported_by_id = session.get('user_id')
    if base_url is None:
        base_url = request.url_root

    db = get_db()
    show = db.execute('SELECT * FROM shows WHERE id = ?', (show_id,)).fetchone()
    adv_rows = db.execute(
        'SELECT field_key, field_value FROM advance_data WHERE show_id = ?', (show_id,)
    ).fetchall()
    advance_data = {r['field_key']: r['field_value'] for r in adv_rows}
    contacts = db.execute('SELECT * FROM contacts ORDER BY name').fetchall()
    contact_map = {c['id']: dict(c) for c in contacts}

    logo_data = get_app_setting('logo_data', '')

    new_v = (show['advance_version'] or 0) + 1
    db.execute('UPDATE shows SET advance_version=? WHERE id=?', (new_v, show_id))
    log_cur = db.execute("""INSERT INTO export_log (show_id, export_type, version, exported_by)
                  VALUES (?, 'advance', ?, ?)""", (show_id, new_v, exported_by_id))
    log_id = log_cur.lastrowid
    db.commit()
    db.close()

    # Fetch form sections; if fetch fails, PDF falls back to generic rendering
    try:
        form_sections = get_form_fields_for_template()
    except Exception as e:
        app.logger.warning(f'Could not load form_sections for PDF: {e}')
        form_sections = []

    # Rented assets, grouped by asset category, for sections that have a
    # linked category. Skip hidden rentals; sort within each category by
    # type name for stable PDF output.
    assets_by_category = {}
    try:
        db2 = get_db()
        rental_rows = db2.execute("""
            SELECT sa.quantity, sa.notes, sa.rental_start, sa.rental_end,
                   at.name AS type_name, at.manufacturer, at.model,
                   ac.id AS category_id, ac.name AS category_name
            FROM show_assets sa
            JOIN asset_types at ON at.id = sa.asset_type_id
            JOIN asset_categories ac ON ac.id = at.category_id
            WHERE sa.show_id = ? AND sa.is_hidden = 0
            ORDER BY ac.sort_order, ac.name, at.name
        """, (show_id,)).fetchall()
        db2.close()
        for r in rental_rows:
            assets_by_category.setdefault(r['category_id'], []).append(dict(r))
    except Exception as e:
        app.logger.warning(f'Could not load rented assets for advance PDF: {e}')

    # Map section.id -> list of rentals for sections that link a category.
    assets_by_section = {
        s['id']: assets_by_category.get(s.get('asset_category_id'), [])
        for s in form_sections
        if s.get('asset_category_id')
    }

    layout = pdf_layouts.PdfLayout('advance', get_app_setting, form_sections=form_sections)

    try:
        html = render_template('pdf/advance_pdf.html',
                               show=show, advance_data=advance_data,
                               contact_map=contact_map,
                               form_sections=form_sections,
                               assets_by_section=assets_by_section,
                               logo_data=logo_data,
                               version=new_v,
                               layout=layout,
                               export_date=datetime.now().strftime('%B %d, %Y at %I:%M %p'))
    except Exception as e:
        app.logger.error(f'advance_pdf template error for show {show_id}: {e}')
        html = render_template('pdf/advance_pdf.html',
                               show=show, advance_data=advance_data,
                               contact_map=contact_map,
                               form_sections=[],
                               assets_by_section={},
                               logo_data=logo_data,
                               version=new_v,
                               layout=layout,
                               export_date=datetime.now().strftime('%B %d, %Y at %I:%M %p'))

    # Generate PDF bytes (S3 push is handled by the caller)
    try:
        from weasyprint import HTML as WP_HTML
        pdf_bytes = WP_HTML(string=html, base_url=base_url).write_pdf()
    except Exception as e:
        app.logger.error(f"PDF_GENERATION_FAILED show_id={show_id} type=advance error={e}")
        pdf_bytes = None

    # Append per-field uploaded files (file_upload form fields). PDFs merge as-is;
    # images and Word docs are converted to PDF wrapper pages first. The
    # watermark below brands each appended page so a printed copy is
    # traceable back to this advance sheet.
    if pdf_bytes:
        try:
            extras = _collect_advance_field_attachments(show_id, base_url)
            if extras:
                wm_text = (
                    f"Attached to ADVANCE SHEET v{new_v}  ·  {show['name']}  ·  "
                    f"{show['show_date'] or '—'}  ·  Exported "
                    f"{datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
                )
                pdf_bytes = _merge_pdfs(pdf_bytes, extras, extras_watermark=wm_text)
        except Exception as e:
            app.logger.error(f"PDF append failed for show {show_id}: {e}")

    syslog_logger.info(
        f"PDF_EXPORT show_id={show_id} type=advance v={new_v} by={exported_by_id}"
    )
    return html, new_v, dict(show), pdf_bytes, log_id


def _collect_advance_field_attachments(show_id, base_url):
    """Fetch all show_attachments tied to file_upload advance-form fields, in
    advance-form section/field order. Convert each to PDF bytes (wrapping
    images and Word docs in a generated cover page). Returns a list of PDF
    byte-strings ready to append to the advance PDF."""
    db = get_db()
    rows = db.execute("""
        SELECT sa.id, sa.filename, sa.mime_type, sa.file_data, sa.s3_key,
               sa.field_key, sa.description, sa.created_at,
               ff.label AS field_label, fs.label AS section_label,
               fs.sort_order AS section_order, ff.sort_order AS field_order
        FROM show_attachments sa
        LEFT JOIN form_fields   ff ON ff.field_key = sa.field_key
        LEFT JOIN form_sections fs ON fs.id = ff.section_id
        WHERE sa.show_id = ?
          AND sa.field_key IS NOT NULL AND sa.field_key != ''
          AND ff.field_type = 'file_upload'
        ORDER BY fs.sort_order, ff.sort_order, sa.created_at
    """, (show_id,)).fetchall()
    db.close()

    extras = []
    for r in rows:
        try:
            data = None
            if r['s3_key']:
                try:
                    data = s3_storage.download_file(r['s3_key'])
                except Exception as e:
                    app.logger.warning(f"S3 fetch failed for attachment {r['id']}: {e}")
            elif r['file_data']:
                data = bytes(r['file_data'])
            if not data:
                continue
            mime = (r['mime_type'] or '').lower()
            fname = (r['filename'] or '').lower()
            if 'pdf' in mime or fname.endswith('.pdf'):
                extras.append(data)
                continue
            wrapper = _render_attachment_wrapper_pdf(
                data, mime, r['filename'],
                section_label=r['section_label'],
                field_label=r['field_label'],
                description=r['description'] or '',
                base_url=base_url
            )
            if wrapper:
                extras.append(wrapper)
        except Exception as e:
            app.logger.warning(f"Could not append attachment {r['id']} to PDF: {e}")
    return extras


def _render_attachment_wrapper_pdf(data, mime, filename, section_label,
                                   field_label, description, base_url):
    """Build a single-section HTML page for a non-PDF attachment and render
    it via WeasyPrint. Supports image/* (embeds the image) and DOCX (extracts
    text). Other types render a placeholder page noting the file is attached
    separately. Returns PDF bytes or None on failure."""
    try:
        import base64
        from io import BytesIO
        section = (section_label or '').strip()
        flabel  = (field_label or filename or '').strip()
        desc    = (description or '').strip()
        body_html = ''

        if mime.startswith('image/'):
            b64 = base64.b64encode(data).decode('ascii')
            body_html = f'<img src="data:{mime};base64,{b64}" style="max-width:100%;max-height:9in;display:block;margin:0 auto">'
        elif filename.lower().endswith('.docx') or 'wordprocessingml' in mime:
            try:
                import docx as _docx
                document = _docx.Document(BytesIO(data))
                paras = [p.text for p in document.paragraphs if p.text.strip()]
                from markupsafe import escape as _e
                body_html = '<div style="font-size:10pt;line-height:1.4;white-space:pre-wrap">' + \
                            '<br><br>'.join(str(_e(p)) for p in paras) + '</div>'
                if not paras:
                    body_html = '<p style="color:#666">[Word document had no extractable text — original file is attached to the show.]</p>'
            except Exception as e:
                app.logger.warning(f"DOCX text extract failed: {e}")
                body_html = f'<p style="color:#666">[Word document <strong>{filename}</strong> could not be embedded — original file is attached to the show.]</p>'
        else:
            body_html = f'<p style="color:#666">[File <strong>{filename}</strong> ({mime or "unknown type"}) is attached to the show but cannot be rendered inline. Download it from the show\'s Files tab.]</p>'

        html = f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>
            @page {{ size: letter; margin: 0.6in 0.55in; }}
            body {{ font-family: Arial, Helvetica, sans-serif; font-size: 9pt; color: #1a1a1a; }}
            .head {{ border-bottom: 2px solid #000; padding-bottom: 6px; margin-bottom: 14px; }}
            .head .section {{ font-size: 8pt; font-weight: bold; letter-spacing: 0.08em; text-transform: uppercase; color: #B8840A; }}
            .head .title {{ font-size: 14pt; font-weight: 700; margin-top: 2px; }}
            .head .meta {{ font-size: 8pt; color: #555; margin-top: 4px; }}
            .desc {{ font-size: 9pt; color: #333; margin-bottom: 12px; padding: 6px 10px; background: #f5f5f5; border-left: 3px solid #B8840A; }}
        </style></head><body>
            <div class="head">
                {f'<div class="section">{section}</div>' if section else ''}
                <div class="title">{flabel}</div>
                <div class="meta">Attached file: {filename}</div>
            </div>
            {f'<div class="desc">{desc}</div>' if desc else ''}
            {body_html}
        </body></html>"""

        from weasyprint import HTML as WP_HTML
        return WP_HTML(string=html, base_url=base_url).write_pdf()
    except Exception as e:
        app.logger.warning(f"Wrapper PDF render failed for {filename}: {e}")
        return None


def _build_schedule_pdf(show_id, exported_by_id=None, base_url=None):
    """
    Build the schedule PDF.  Works both in a request context (base_url=None reads
    from request.url_root) and in a background context (pass base_url='/').
    exported_by_id defaults to session['user_id'] when not supplied.
    """
    if exported_by_id is None:
        exported_by_id = session.get('user_id')
    if base_url is None:
        base_url = request.url_root

    db = get_db()
    show = db.execute('SELECT * FROM shows WHERE id = ?', (show_id,)).fetchone()
    all_sched_rows = db.execute(
        'SELECT * FROM schedule_rows WHERE show_id=? ORDER BY sort_order,id', (show_id,)
    ).fetchall()
    meta_rows = db.execute(
        'SELECT field_key, field_value FROM schedule_meta WHERE show_id=?', (show_id,)
    ).fetchall()
    schedule_meta = {r['field_key']: r['field_value'] for r in meta_rows}
    adv_rows = db.execute(
        'SELECT field_key, field_value FROM advance_data WHERE show_id=?', (show_id,)
    ).fetchall()
    advance_data = {r['field_key']: r['field_value'] for r in adv_rows}
    performances = [dict(p) for p in db.execute(
        'SELECT * FROM show_performances WHERE show_id=? ORDER BY sort_order, perf_date, perf_time, id', (show_id,)
    ).fetchall()]
    contacts = db.execute('SELECT * FROM contacts ORDER BY name').fetchall()
    contact_map = {c['id']: dict(c) for c in contacts}

    logo_data = get_app_setting('logo_data', '')

    # WiFi always from global settings (not per-show)
    wifi_ssid = get_app_setting('wifi_network', '')
    wifi_pass  = get_app_setting('wifi_password', '')
    wifi_qr_b64 = None
    if wifi_ssid:
        try:
            import qrcode, io, base64
            qr = qrcode.make(f"WIFI:S:{wifi_ssid};T:WPA;P:{wifi_pass};;")
            buf = io.BytesIO()
            qr.save(buf, format='PNG')
            wifi_qr_b64 = base64.b64encode(buf.getvalue()).decode()
        except Exception:
            pass

    # Group schedule rows: prefer day_date when set, else perf_id (legacy).
    rows_by_date = {}
    rows_by_perf = {}
    for row in all_sched_rows:
        rd = dict(row)
        dd = rd.get('day_date')
        if dd:
            rows_by_date.setdefault(str(dd)[:10], []).append(rd)
        else:
            rows_by_perf.setdefault(rd.get('perf_id'), []).append(rd)

    # Build per-day data for the PDF template.  Union of performance dates
    # and the load-in → load-out range so the schedule covers every day of
    # the production.
    def _iso_pdf(v):
        s = str(v) if v is not None else ''
        return s[:10] if s else ''

    li = _iso_pdf(show['load_in_date'])
    lo = _iso_pdf(show['load_out_date'])
    seen_dates = set()
    schedule_days = []
    day_num = 1
    perf_by_date = {}
    for p in performances:
        pd = _iso_pdf(p.get('perf_date'))
        if pd:
            perf_by_date.setdefault(pd, []).append(p)

    ordered_dates = []
    if li and lo:
        try:
            from datetime import date as _d, timedelta as _td
            d1, d2 = _d.fromisoformat(li), _d.fromisoformat(lo)
            if d2 >= d1:
                cur = d1
                while cur <= d2:
                    ordered_dates.append(cur.isoformat())
                    cur += _td(days=1)
        except (ValueError, TypeError):
            pass
    for pd in sorted(perf_by_date.keys()):
        if pd not in ordered_dates:
            ordered_dates.append(pd)
    ordered_dates.sort()

    for pd in ordered_dates:
        perfs_here = perf_by_date.get(pd, [])
        primary = perfs_here[0] if perfs_here else {'perf_date': pd, 'perf_time': ''}
        day_rows = list(rows_by_date.get(pd, []))
        for p in perfs_here:
            day_rows += rows_by_perf.get(p['id'], [])
        if day_num == 1:
            day_rows = rows_by_perf.get(None, []) + day_rows
        schedule_days.append({'perf': primary, 'rows': day_rows, 'day_num': day_num})
        seen_dates.add(pd)
        day_num += 1

    # Performances without a date (rare) fall back to legacy handling
    for p in performances:
        if not _iso_pdf(p.get('perf_date')):
            schedule_days.append({'perf': p,
                                  'rows': rows_by_perf.get(p['id'], []),
                                  'day_num': day_num})
            day_num += 1

    if not schedule_days:  # Fallback: show with no performances & no load-in/out
        schedule_days = [{'perf': {'perf_date': show['show_date'], 'perf_time': show['show_time']},
                          'rows': rows_by_perf.get(None, []), 'day_num': 1}]

    # Crew call times — unique in_times from labor_requests, sorted
    labor_in_times = db.execute(
        "SELECT in_time FROM labor_requests WHERE show_id=? AND in_time != '' ORDER BY in_time",
        (show_id,)
    ).fetchall()
    crew_call_times = sorted(set(r['in_time'] for r in labor_in_times if r['in_time']))

    new_v = (show['schedule_version'] or 0) + 1
    db.execute('UPDATE shows SET schedule_version=? WHERE id=?', (new_v, show_id))
    log_cur = db.execute("""INSERT INTO export_log (show_id, export_type, version, exported_by)
                  VALUES (?, 'schedule', ?, ?)""", (show_id, new_v, exported_by_id))
    log_id = log_cur.lastrowid
    db.commit()
    db.close()

    html = render_template('pdf/schedule_pdf.html',
                           show=show,
                           schedule_days=schedule_days,
                           schedule_meta=schedule_meta,
                           sched_meta_fields=get_schedule_meta_fields(),
                           advance_data=advance_data,
                           contact_map=contact_map,
                           logo_data=logo_data,
                           wifi_ssid=wifi_ssid,
                           wifi_pass=wifi_pass,
                           wifi_qr_b64=wifi_qr_b64,
                           crew_call_times=crew_call_times,
                           version=new_v,
                           layout=pdf_layouts.PdfLayout('schedule', get_app_setting),
                           export_date=datetime.now().strftime('%B %d, %Y at %I:%M %p'))

    # Generate PDF bytes (S3 push is handled by the caller)
    try:
        from weasyprint import HTML as WP_HTML
        pdf_bytes = WP_HTML(string=html, base_url=base_url).write_pdf()
    except Exception as e:
        app.logger.error(f"PDF_GENERATION_FAILED show_id={show_id} type=schedule error={e}")
        pdf_bytes = None

    syslog_logger.info(
        f"PDF_EXPORT show_id={show_id} type=schedule v={new_v} by={exported_by_id}"
    )
    return html, new_v, dict(show), pdf_bytes, log_id


@app.route('/shows/<int:show_id>/export/advance')
@login_required
def export_advance(show_id):
    if not can_access_show(session['user_id'], show_id):
        abort(403)
    get_show_or_404(show_id)
    html, version, show, pdf_bytes, log_id = _build_advance_pdf(show_id)
    safe_name = show['name'].replace(' ', '_').replace('/', '-')
    filename  = f"Advance_{safe_name}_{show.get('show_date','nodate')}_v{version}.pdf"
    if pdf_bytes:
        resp = make_response(pdf_bytes)
        resp.headers['Content-Type'] = 'application/pdf'
        resp.headers['Content-Disposition'] = _safe_content_disposition(filename)
        # Push to S3 in background for archival (public links, export history)
        if s3_storage.is_configured():
            _s3_key = f"exports/{show_id}/advance/v{version}.pdf"
            _pdf = pdf_bytes
            _lid = log_id
            def _push_advance():
                try:
                    s3_storage.upload_file(_s3_key, _pdf, 'application/pdf')
                    with app.app_context():
                        db2 = get_db()
                        db2.execute('UPDATE export_log SET s3_key=? WHERE id=?', (_s3_key, _lid))
                        db2.commit()
                        db2.close()
                except Exception as e:
                    app.logger.error(f"S3 push failed for advance PDF log_id={_lid}: {e}")
                    syslog_logger.error(f"S3_PUSH_FAILED context=advance_export show_id={show_id} log_id={_lid} error={e}")
            threading.Thread(target=_push_advance, daemon=True).start()
        return resp
    # Fallback to HTML if weasyprint failed
    try:
        from weasyprint import HTML
        pdf = HTML(string=html, base_url=request.url_root).write_pdf()
        resp = make_response(pdf)
        resp.headers['Content-Type'] = 'application/pdf'
        resp.headers['Content-Disposition'] = _safe_content_disposition(filename)
        return resp
    except Exception as e:
        app.logger.error(f"PDF_FALLBACK_FAILED show_id={show_id} type=advance error={e}")
        resp = make_response(html)
        resp.headers['Content-Type'] = 'text/html'
        return resp


@app.route('/shows/<int:show_id>/export/schedule')
@login_required
def export_schedule(show_id):
    if not can_access_show(session['user_id'], show_id):
        abort(403)
    get_show_or_404(show_id)
    html, version, show, pdf_bytes, log_id = _build_schedule_pdf(show_id)
    safe_name = show['name'].replace(' ', '_').replace('/', '-')
    filename  = f"Schedule_{safe_name}_{show.get('show_date','nodate')}_v{version}.pdf"
    if pdf_bytes:
        resp = make_response(pdf_bytes)
        resp.headers['Content-Type'] = 'application/pdf'
        resp.headers['Content-Disposition'] = _safe_content_disposition(filename)
        # Push to S3 in background for archival (public links, export history)
        if s3_storage.is_configured():
            _s3_key = f"exports/{show_id}/schedule/v{version}.pdf"
            _pdf = pdf_bytes
            _lid = log_id
            def _push_schedule():
                try:
                    s3_storage.upload_file(_s3_key, _pdf, 'application/pdf')
                    with app.app_context():
                        db2 = get_db()
                        db2.execute('UPDATE export_log SET s3_key=? WHERE id=?', (_s3_key, _lid))
                        db2.commit()
                        db2.close()
                except Exception as e:
                    app.logger.error(f"S3 push failed for schedule PDF log_id={_lid}: {e}")
                    syslog_logger.error(f"S3_PUSH_FAILED context=schedule_export show_id={show_id} log_id={_lid} error={e}")
            threading.Thread(target=_push_schedule, daemon=True).start()
        return resp
    try:
        from weasyprint import HTML
        pdf = HTML(string=html, base_url=request.url_root).write_pdf()
        resp = make_response(pdf)
        resp.headers['Content-Type'] = 'application/pdf'
        resp.headers['Content-Disposition'] = _safe_content_disposition(filename)
        return resp
    except Exception as e:
        app.logger.error(f"PDF_FALLBACK_FAILED show_id={show_id} type=schedule error={e}")
        resp = make_response(html)
        resp.headers['Content-Type'] = 'text/html'
        return resp


@app.route('/shows/<int:show_id>/export/history/<int:log_id>/download')
@login_required
def download_export_history(show_id, log_id):
    if not can_access_show(session['user_id'], show_id):
        abort(403)
    db = get_db()
    row = db.execute(
        'SELECT * FROM export_log WHERE id=? AND show_id=?', (log_id, show_id)
    ).fetchone()
    db.close()
    if not row or (not row['s3_key'] and not row['pdf_data']):
        abort(404)
    filename = f"{row['export_type'].capitalize()}_v{row['version']}.pdf"
    if row['s3_key']:
        try:
            data = s3_storage.download_file(row['s3_key'])
        except Exception as e:
            app.logger.error(f"S3 download failed for export_log {log_id}: {e}")
            abort(503)
    else:
        data = bytes(row['pdf_data'])
    resp = make_response(data)
    resp.headers['Content-Type'] = 'application/pdf'
    resp.headers['Content-Disposition'] = _safe_content_disposition(filename)
    return resp


def _send_pdf_email_async(show_id, pdf_type, triggered_by, exported_by_id):
    """Run _send_pdf_email in a background daemon thread.

    Used by the manual "send email" UI so the user can navigate away
    immediately. Failures during the build / SMTP / audit-log stages are
    surfaced through the email_send_errors panel rather than the original
    HTTP response.
    """
    def _runner():
        with app.app_context():
            try:
                ok, msg, count = _send_pdf_email(
                    show_id, pdf_type, triggered_by,
                    exported_by_id=exported_by_id,
                )
                if ok:
                    try:
                        _adb = get_db()
                        log_audit(
                            _adb, f'PDF_EMAIL_{pdf_type.upper()}', 'show', show_id,
                            show_id=show_id,
                            detail=f'Background email to {count} recipient(s)',
                        )
                        _adb.commit()
                        _adb.close()
                    except Exception as e:
                        app.logger.warning(f'audit log failed for PDF_EMAIL_{pdf_type}: {e}')
                    syslog_logger.info(
                        f'PDF_EMAIL_BG_OK show_id={show_id} type={pdf_type} '
                        f'by={triggered_by} recipients={count}'
                    )
                else:
                    syslog_logger.warning(
                        f'PDF_EMAIL_BG_FAIL show_id={show_id} type={pdf_type} '
                        f'by={triggered_by} msg={msg!r}'
                    )
            except Exception as e:
                app.logger.exception(
                    f'email_pdf background failed show_id={show_id} type={pdf_type}'
                )
                syslog_logger.error(
                    f'PDF_EMAIL_BG_CRASHED show_id={show_id} type={pdf_type} '
                    f'by={triggered_by} error={e!r}'
                )
                # Best-effort: surface the crash in the email errors panel so
                # an admin can see it without grepping syslog.
                try:
                    _log_email_error(
                        ['(unknown)'],
                        f'{pdf_type} for show {show_id}',
                        f'Background send crashed: {e}',
                        pdf_type=pdf_type,
                        show_id=show_id,
                        triggered_by=triggered_by,
                    )
                except Exception:
                    pass

    threading.Thread(
        target=_runner,
        name=f'email-{pdf_type}-{show_id}',
        daemon=True,
    ).start()


@app.route('/shows/<int:show_id>/email/<pdf_type>', methods=['POST'])
@login_required
def email_pdf(show_id, pdf_type):
    """Queue a PDF email for advance, schedule, or postnotes.

    Returns 202 immediately; the actual PDF render + SMTP send runs on a
    background thread. Any failures show up in
    Settings → Email → Email Send Errors.
    """
    if pdf_type not in ('advance', 'schedule', 'postnotes'):
        return jsonify({'success': False, 'error': 'Invalid PDF type.'}), 400
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    get_show_or_404(show_id)
    triggered_by = session.get('username') or session.get('user_display') or 'user'
    exported_by_id = session.get('user_id')

    _send_pdf_email_async(show_id, pdf_type, triggered_by, exported_by_id)

    syslog_logger.info(
        f'PDF_EMAIL_QUEUED show_id={show_id} type={pdf_type} by={triggered_by}'
    )
    return jsonify({
        'success': True,
        'queued':  True,
        'message': ('Email queued. Delivery happens in the background — any '
                    'failures will appear in Settings → Email → Email Send Errors.'),
        'recipients': None,
    }), 202


def _build_postnotes_pdf(show_id, exported_by_id=None, base_url=None):
    """Build the post-show notes PDF. Versions and logs to export_log so the
    same machinery (history download, S3 archival, scheduled email) works
    identically to advance/schedule. Returns (html, version, show_dict,
    pdf_bytes, log_id)."""
    if exported_by_id is None:
        exported_by_id = session.get('user_id')
    if base_url is None:
        base_url = request.url_root

    db = get_db()
    show = db.execute('SELECT * FROM shows WHERE id=?', (show_id,)).fetchone()
    note_rows = db.execute(
        'SELECT field_key, field_value FROM post_show_notes WHERE show_id=?', (show_id,)
    ).fetchall()
    notes_data = {r['field_key']: r['field_value'] for r in note_rows}
    adv_rows = db.execute(
        'SELECT field_key, field_value FROM advance_data WHERE show_id=?', (show_id,)
    ).fetchall()
    advance_data = {r['field_key']: r['field_value'] for r in adv_rows}
    sched_rows = db.execute(
        'SELECT * FROM schedule_rows WHERE show_id=? ORDER BY sort_order,id', (show_id,)
    ).fetchall()
    logo_data = get_app_setting('logo_data', '')

    new_v = (show['postnotes_version'] or 0) + 1
    db.execute('UPDATE shows SET postnotes_version=? WHERE id=?', (new_v, show_id))
    log_cur = db.execute("""INSERT INTO export_log (show_id, export_type, version, exported_by)
                  VALUES (?, 'postnotes', ?, ?)""", (show_id, new_v, exported_by_id))
    log_id = log_cur.lastrowid
    db.commit()
    db.close()

    layout = pdf_layouts.PdfLayout('postnotes', get_app_setting)
    html = render_template('pdf/postnotes_pdf.html',
                           show=show,
                           notes_data=notes_data,
                           advance_data=advance_data,
                           schedule_rows=sched_rows,
                           logo_data=logo_data,
                           version=new_v,
                           layout=layout,
                           export_date=datetime.now().strftime('%B %d, %Y at %I:%M %p'))
    try:
        from weasyprint import HTML as WP_HTML
        pdf_bytes = WP_HTML(string=html, base_url=base_url).write_pdf()
    except Exception as e:
        app.logger.error(f"PDF_GENERATION_FAILED show_id={show_id} type=postnotes error={e}")
        pdf_bytes = None

    syslog_logger.info(
        f"PDF_EXPORT show_id={show_id} type=postnotes v={new_v} by={exported_by_id}"
    )
    return html, new_v, dict(show), pdf_bytes, log_id


@app.route('/shows/<int:show_id>/export/postnotes')
@login_required
def export_postnotes(show_id):
    if not can_access_show(session['user_id'], show_id):
        abort(403)
    get_show_or_404(show_id)
    html, version, show_dict, pdf_bytes, _log_id = _build_postnotes_pdf(show_id)
    safe_name = show_dict['name'].replace(' ', '_').replace('/', '-')
    filename = f"PostNotes_{safe_name}_{show_dict.get('show_date') or 'nodate'}_v{version}.pdf"
    if pdf_bytes:
        resp = make_response(pdf_bytes)
        resp.headers['Content-Type'] = 'application/pdf'
        resp.headers['Content-Disposition'] = _safe_content_disposition(filename)
        return resp
    app.logger.error(f"PDF_FALLBACK_FAILED show_id={show_id} type=postnotes")
    resp = make_response(html)
    resp.headers['Content-Type'] = 'text/html'
    return resp


# ─── Document Viewer Routes ───────────────────────────────────────────────────
# Read-only landing for users with users.is_document_viewer=1. The @before_request
# gate above redirects them here from every other endpoint.

@app.route('/viewer')
@login_required
def viewer_home():
    if not session.get('is_document_viewer'):
        return redirect(url_for('dashboard'))
    user_id = session['user_id']
    is_viewer, venues_allow, doc_types_allow = get_document_viewer_settings(user_id)
    accessible = viewer_accessible_shows(user_id)
    db = get_db()
    if not accessible:
        rows = []
    else:
        ph = ','.join(['?'] * len(accessible))
        rows = db.execute(
            f"SELECT id, name, venue, show_date, show_time "
            f"FROM shows WHERE id IN ({ph}) AND COALESCE(status, 'active') != 'archived' "
            f"ORDER BY show_date IS NULL, show_date, name",
            accessible
        ).fetchall()
    db.close()
    # Group by venue for an easy-to-scan layout
    by_venue = {}
    for r in rows:
        v = (r['venue'] or '— No venue —').strip() or '— No venue —'
        by_venue.setdefault(v, []).append(dict(r))
    grouped = sorted(by_venue.items(), key=lambda kv: kv[0].lower())
    doc_types = [(d, DOCUMENT_TYPE_LABELS[d]) for d in DOCUMENT_TYPES
                 if (not doc_types_allow) or d in doc_types_allow]
    return render_template(
        'viewer_home.html',
        user=get_current_user(),
        grouped=grouped,
        doc_types=doc_types,
        venues_allow=venues_allow,
        doc_types_allow=doc_types_allow,
    )


@app.route('/viewer/shows/<int:show_id>')
@login_required
def viewer_show(show_id):
    if not session.get('is_document_viewer'):
        return redirect(url_for('show_page', show_id=show_id))
    user_id = session['user_id']
    accessible = viewer_accessible_shows(user_id)
    if not accessible or show_id not in accessible:
        abort(403)
    db = get_db()
    show = db.execute(
        "SELECT id, name, venue, show_date, show_time FROM shows WHERE id=?",
        (show_id,)
    ).fetchone()
    db.close()
    if not show:
        abort(404)
    if not viewer_can_see_venue(show['venue']):
        abort(403)
    _, _, doc_types_allow = get_document_viewer_settings(user_id)
    doc_types = [(d, DOCUMENT_TYPE_LABELS[d]) for d in DOCUMENT_TYPES
                 if (not doc_types_allow) or d in doc_types_allow]
    return render_template(
        'viewer_show.html',
        user=get_current_user(),
        show=dict(show),
        doc_types=doc_types,
    )


@app.route('/viewer/shows/<int:show_id>/<doc_type>.pdf')
@login_required
def viewer_export(show_id, doc_type):
    if not session.get('is_document_viewer'):
        return redirect(url_for('show_page', show_id=show_id))
    if doc_type not in DOCUMENT_TYPES:
        abort(404)
    if not viewer_can_see_doc_type(doc_type):
        abort(403)
    user_id = session['user_id']
    accessible = viewer_accessible_shows(user_id)
    if not accessible or show_id not in accessible:
        abort(403)
    db = get_db()
    show = db.execute("SELECT venue FROM shows WHERE id=?", (show_id,)).fetchone()
    db.close()
    if not show:
        abort(404)
    if not viewer_can_see_venue(show['venue']):
        abort(403)
    # Audit trail — viewer document access is a sensitive event worth logging.
    syslog_logger.info(
        f"VIEWER_DOC_ACCESS show_id={show_id} doc_type={doc_type} "
        f"by={session.get('username')} ip={request.remote_addr}"
    )
    # Reuse the existing PDF builders — they already enforce can_access_show
    # via the user's group ACL, which we've already passed above.
    if doc_type == 'advance':
        return export_advance(show_id)
    if doc_type == 'schedule':
        return export_schedule(show_id)
    return export_postnotes(show_id)


# ─── Show Management ──────────────────────────────────────────────────────────

@app.route('/shows/<int:show_id>/archive', methods=['POST'])
@staff_or_admin_required
def archive_show(show_id):
    if session.get('user_role') != 'admin' and not can_access_show(session['user_id'], show_id):
        abort(403)
    db = get_db()
    db.execute("UPDATE shows SET status='archived' WHERE id=?", (show_id,))
    log_audit(db, 'SHOW_ARCHIVE', 'show', show_id, show_id=show_id)
    db.commit(); db.close()
    syslog_logger.info(f"SHOW_ARCHIVE show_id={show_id} by={session.get('username')}")
    flash('Show archived.', 'success')
    return redirect(url_for('dashboard'))


@app.route('/shows/<int:show_id>/restore', methods=['POST'])
@staff_or_admin_required
def restore_show(show_id):
    if session.get('user_role') != 'admin' and not can_access_show(session['user_id'], show_id):
        abort(403)
    db = get_db()
    db.execute("UPDATE shows SET status='active' WHERE id=?", (show_id,))
    log_audit(db, 'SHOW_RESTORE', 'show', show_id, show_id=show_id)
    db.commit(); db.close()
    syslog_logger.info(f"SHOW_RESTORE show_id={show_id} by={session.get('username')}")
    flash('Show restored to active.', 'success')
    return redirect(url_for('dashboard'))


@app.route('/shows/<int:show_id>/delete', methods=['POST'])
@admin_required
def delete_show(show_id):
    db = get_db()
    show = db.execute('SELECT name FROM shows WHERE id=?', (show_id,)).fetchone()
    show_name = show['name'] if show else str(show_id)
    for tbl in ['advance_data', 'schedule_rows', 'schedule_meta',
                'post_show_notes', 'export_log', 'form_history', 'show_group_access']:
        db.execute(f'DELETE FROM {tbl} WHERE show_id=?', (show_id,))
    db.execute('DELETE FROM shows WHERE id=?', (show_id,))
    log_audit(db, 'SHOW_DELETE', 'show', show_id, detail=show_name)
    db.commit(); db.close()
    syslog_logger.info(f"SHOW_DELETE show_id={show_id} by={session.get('username')}")
    flash('Show permanently deleted.', 'success')
    return redirect(url_for('dashboard'))


# ─── Show Access (Groups ↔ Shows) ─────────────────────────────────────────────

@app.route('/shows/<int:show_id>/access/add', methods=['POST'])
@admin_required
def add_show_access(show_id):
    data = request.get_json(force=True) or {}
    group_id = data.get('group_id')
    if not group_id:
        return jsonify({'success': False, 'error': 'group_id required'}), 400
    db = get_db()
    db.execute('INSERT OR IGNORE INTO show_group_access (show_id, group_id) VALUES (?,?)',
               (show_id, group_id))
    log_audit(db, 'SHOW_ACCESS_ADD', 'show', show_id, show_id=show_id,
              detail=f'group_id={group_id}')
    db.commit(); db.close()
    syslog_logger.info(
        f"SHOW_ACCESS_ADD show_id={show_id} group_id={group_id} by={session.get('username')}"
    )
    return jsonify({'success': True})


@app.route('/shows/<int:show_id>/access/remove', methods=['POST'])
@admin_required
def remove_show_access(show_id):
    data = request.get_json(force=True) or {}
    group_id = data.get('group_id')
    if not group_id:
        return jsonify({'success': False, 'error': 'group_id required'}), 400
    db = get_db()
    db.execute('DELETE FROM show_group_access WHERE show_id=? AND group_id=?',
               (show_id, group_id))
    log_audit(db, 'SHOW_ACCESS_REMOVE', 'show', show_id, show_id=show_id,
              detail=f'group_id={group_id}')
    db.commit(); db.close()
    syslog_logger.info(f"SHOW_ACCESS_REMOVE show_id={show_id} group_id={group_id} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/api/shows/<int:show_id>/access')
@admin_required
def get_show_access(show_id):
    db = get_db()
    rows = db.execute("""
        SELECT ug.id, ug.name, ug.group_type
        FROM show_group_access sga
        JOIN user_groups ug ON sga.group_id = ug.id
        WHERE sga.show_id = ?
    """, (show_id,)).fetchall()
    db.close()
    return jsonify([dict(r) for r in rows])


# ─── Audit Log ────────────────────────────────────────────────────────────────

@app.route('/admin/s3-test')
@admin_required
def admin_s3_test():
    """Verify SeaweedFS S3 connectivity by uploading, reading back, and deleting a test object."""
    result = s3_storage.test_connection()
    return jsonify(result)


@app.route('/admin/migrate-files-to-s3', methods=['POST'])
@admin_required
def admin_migrate_files_to_s3():
    """
    One-time migration: move existing BLOB/BYTEA file data from the database
    to SeaweedFS S3.  Safe to re-run — skips rows that already have an s3_key.
    Returns a JSON summary of migrated / failed counts.
    """
    if not s3_storage.is_configured():
        return jsonify({'success': False, 'error': 'SeaweedFS not configured in db_config.ini.'}), 400

    migrated = 0
    failed = 0
    errors = []

    db = get_db()
    try:
        # ── show_attachments ──────────────────────────────────────────────────
        rows = db.execute(
            "SELECT id, show_id, filename, mime_type, file_data FROM show_attachments "
            "WHERE file_data IS NOT NULL AND s3_key IS NULL"
        ).fetchall()
        for row in rows:
            try:
                key = f"attachments/{row['show_id']}/{row['id']}/{row['filename']}"
                s3_storage.upload_file(key, bytes(row['file_data']),
                                       row['mime_type'] or 'application/octet-stream')
                db.execute('UPDATE show_attachments SET s3_key=?, file_data=NULL WHERE id=?',
                           (key, row['id']))
                migrated += 1
            except Exception as e:
                failed += 1
                errors.append(f"attachment id={row['id']}: {e}")
        db.commit()

        # ── export_log ────────────────────────────────────────────────────────
        rows = db.execute(
            "SELECT id, show_id, export_type, version, pdf_data FROM export_log "
            "WHERE pdf_data IS NOT NULL AND s3_key IS NULL"
        ).fetchall()
        for row in rows:
            try:
                key = f"exports/{row['show_id']}/{row['export_type']}/v{row['version']}.pdf"
                s3_storage.upload_file(key, bytes(row['pdf_data']), 'application/pdf')
                db.execute('UPDATE export_log SET s3_key=?, pdf_data=NULL WHERE id=?',
                           (key, row['id']))
                migrated += 1
            except Exception as e:
                failed += 1
                errors.append(f"export_log id={row['id']}: {e}")
        db.commit()

        # ── asset_types photos ────────────────────────────────────────────────
        rows = db.execute(
            "SELECT id, photo, photo_mime FROM asset_types "
            "WHERE photo IS NOT NULL AND photo_s3_key IS NULL"
        ).fetchall()
        for row in rows:
            try:
                key = f"asset-photos/{row['id']}"
                s3_storage.upload_file(key, bytes(row['photo']),
                                       row['photo_mime'] or 'image/jpeg')
                db.execute('UPDATE asset_types SET photo_s3_key=?, photo=NULL WHERE id=?',
                           (key, row['id']))
                migrated += 1
            except Exception as e:
                failed += 1
                errors.append(f"asset_type id={row['id']}: {e}")
        db.commit()

        # ── show_external_rentals ─────────────────────────────────────────────
        rows = db.execute(
            "SELECT id, pdf_data, pdf_filename FROM show_external_rentals "
            "WHERE pdf_data IS NOT NULL AND s3_key IS NULL"
        ).fetchall()
        for row in rows:
            try:
                fname = row['pdf_filename'] or 'rental.pdf'
                key = f"external-rentals/{row['id']}/{fname}"
                s3_storage.upload_file(key, bytes(row['pdf_data']), 'application/pdf')
                db.execute('UPDATE show_external_rentals SET s3_key=?, pdf_data=NULL WHERE id=?',
                           (key, row['id']))
                migrated += 1
            except Exception as e:
                failed += 1
                errors.append(f"external_rental id={row['id']}: {e}")
        db.commit()

    finally:
        db.close()

    return jsonify({'success': failed == 0, 'migrated': migrated, 'failed': failed, 'errors': errors})


@app.route('/admin/audit')
@admin_required
def audit_log_view():
    db = get_db()
    page = max(1, int(request.args.get('page', 1)))
    per_page = 50
    offset = (page - 1) * per_page

    filters = []
    params = []
    if request.args.get('user_id'):
        try:
            filters.append('al.user_id = ?')
            params.append(int(request.args['user_id']))
        except ValueError:
            pass
    if request.args.get('show_id'):
        try:
            filters.append('al.show_id = ?')
            params.append(int(request.args['show_id']))
        except ValueError:
            pass
    if request.args.get('action'):
        filters.append('al.action LIKE ?')
        params.append(f"%{request.args['action'].upper()}%")
    if request.args.get('date_from'):
        filters.append('al.timestamp >= ?')
        params.append(request.args['date_from'])
    if request.args.get('date_to'):
        filters.append('al.timestamp <= ?')
        params.append(request.args['date_to'] + ' 23:59:59')

    where = ('WHERE ' + ' AND '.join(filters)) if filters else ''

    total_row = db.execute(
        f'SELECT COUNT(*) FROM audit_log al {where}', params
    ).fetchone()
    total = list(total_row.values())[0] if total_row else 0

    rows = db.execute(f"""
        SELECT al.id, al.timestamp, al.username, al.action, al.entity_type,
               al.entity_id, al.show_id, al.before_json, al.after_json,
               al.ip_address, al.detail,
               al.undone_at, al.undone_by, al.undone_by_log_id,
               s.name as show_name,
               u.display_name as display_name
        FROM audit_log al
        LEFT JOIN shows s ON al.show_id = s.id
        LEFT JOIN users u ON al.user_id = u.id
        {where}
        ORDER BY al.timestamp DESC
        LIMIT ? OFFSET ?
    """, params + [per_page, offset]).fetchall()

    users = db.execute(
        'SELECT id, username, display_name FROM users ORDER BY username'
    ).fetchall()
    shows = db.execute(
        'SELECT id, name FROM shows ORDER BY name'
    ).fetchall()
    db.close()

    entries = []
    for r in rows:
        entry = dict(r)
        entry['display_name'] = r['display_name'] or r['username']
        ok, _reason = _can_undo_audit_row(r)
        entry['undoable'] = ok
        entries.append(entry)

    return render_template('audit_log.html',
        entries=entries,
        users=users,
        shows=shows,
        total=total,
        page=page,
        per_page=per_page,
        total_pages=max(1, (total + per_page - 1) // per_page),
        filters=request.args,
        user=get_current_user(),
    )


@app.route('/admin/audit/<int:log_id>/undo', methods=['POST'])
@admin_required
def audit_undo(log_id):
    """Reverse a single audit_log entry, if enough state was captured to do so.

    Supports three shapes of action:
      * create  — deletes the row whose id == audit.entity_id
      * update  — restores columns from audit.before_json on that row
      * delete  — re-inserts audit.before_json, preserving the original id
    Writes a new audit row (ACTION_UNDONE) and marks the original as undone.
    """
    db = get_db()
    row = db.execute("""
        SELECT id, action, entity_type, entity_id, show_id,
               before_json, after_json, undone_at
        FROM audit_log WHERE id = ?
    """, (log_id,)).fetchone()
    if not row:
        db.close()
        return jsonify({'error': 'Audit entry not found'}), 404

    ok, kind_or_reason = _can_undo_audit_row(row)
    if not ok:
        db.close()
        return jsonify({'error': kind_or_reason}), 400

    kind = kind_or_reason  # 'create' | 'update' | 'delete'
    table = UNDO_TABLE_MAP[row['entity_type']]
    entity_id = row['entity_id']
    before = json.loads(row['before_json']) if row['before_json'] else None
    after  = json.loads(row['after_json'])  if row['after_json']  else None

    try:
        if kind == 'create':
            # Snapshot the current row before deletion (so the undo is itself undoable)
            before_undo = _snapshot_row(db, table, entity_id)
            db.execute(f'DELETE FROM {table} WHERE id = ?', (entity_id,))
            undo_before, undo_after = before_undo, None

        elif kind == 'update':
            # Capture current row, then restore columns from before_json
            before_undo = _snapshot_row(db, table, entity_id)
            if before_undo is None:
                db.close()
                return jsonify({'error': f'{row["entity_type"]} #{entity_id} no longer exists'}), 409
            # Only restore columns that exist in both before_json and the current row
            cols = [c for c in before.keys() if c in before_undo and c != 'id']
            if not cols:
                db.close()
                return jsonify({'error': 'No matching columns to restore'}), 400
            set_clause = ', '.join(f'{c} = ?' for c in cols)
            values = [before[c] for c in cols] + [entity_id]
            db.execute(f'UPDATE {table} SET {set_clause} WHERE id = ?', values)
            undo_before, undo_after = before_undo, before

        elif kind == 'delete':
            # Re-insert the deleted row from before_json, preserving its id
            cols = list(before.keys())
            placeholders = ', '.join('?' for _ in cols)
            col_list = ', '.join(cols)
            values = [before[c] for c in cols]
            db.execute(f'INSERT INTO {table} ({col_list}) VALUES ({placeholders})', values)
            undo_before, undo_after = None, before

        else:
            db.close()
            return jsonify({'error': f'Unknown undo kind: {kind}'}), 500

        # Write the undo audit row
        undo_cur = db.execute("""
            INSERT INTO audit_log
              (user_id, username, action, entity_type, entity_id,
               show_id, before_json, after_json, ip_address, detail)
            VALUES (?,?,?,?,?,?,?,?,?,?)
        """, (
            session.get('user_id'),
            session.get('username', ''),
            'ACTION_UNDONE',
            row['entity_type'],
            entity_id,
            row['show_id'],
            json.dumps(undo_before) if undo_before is not None else None,
            json.dumps(undo_after)  if undo_after  is not None else None,
            request.remote_addr,
            f'Undid audit #{log_id} ({row["action"]})',
        ))
        new_log_id = getattr(undo_cur, 'lastrowid', None)

        # Mark the original row as undone
        db.execute("""
            UPDATE audit_log
               SET undone_at = CURRENT_TIMESTAMP,
                   undone_by = ?,
                   undone_by_log_id = ?
             WHERE id = ?
        """, (session.get('user_id'), new_log_id, log_id))

        db.commit()
        db.close()
        return jsonify({'success': True, 'undo_log_id': new_log_id})

    except Exception as e:
        try: db.rollback()
        except Exception: pass
        db.close()
        syslog_logger.warning(f"AUDIT_UNDO_FAILED log_id={log_id} err={e}")
        return jsonify({'error': f'Undo failed: {e}'}), 500


# ─── Settings ─────────────────────────────────────────────────────────────────

@app.route('/settings')
@login_required
def settings():
    db = get_db()
    contacts = db.execute('SELECT * FROM contacts ORDER BY department, name').fetchall()
    users_raw = db.execute(
        'SELECT id, username, display_name, email, role, created_at, '
        '       is_readonly, is_scheduler, is_asset_manager, '
        '       is_document_viewer, viewer_venues, viewer_doc_types '
        'FROM users ORDER BY display_name'
    ).fetchall()
    # Decode the viewer JSON columns so the template can use |tojson cleanly
    users = []
    for u in users_raw:
        ud = dict(u)
        ud['viewer_venues_list']    = _decode_json_list(ud.get('viewer_venues'))
        ud['viewer_doc_types_list'] = _decode_json_list(ud.get('viewer_doc_types'))
        users.append(ud)
    groups   = db.execute('SELECT * FROM user_groups ORDER BY name').fetchall()

    # Attach members and shows to each group
    groups_data = []
    for g in groups:
        members = db.execute("""
            SELECT u.id, u.display_name, u.username FROM user_group_members ugm
            JOIN users u ON ugm.user_id = u.id
            WHERE ugm.group_id = ?
        """, (g['id'],)).fetchall()
        shows = db.execute("""
            SELECT s.id, s.name, s.show_date FROM show_group_access sga
            JOIN shows s ON sga.show_id = s.id
            WHERE sga.group_id = ?
            ORDER BY s.show_date DESC
        """, (g['id'],)).fetchall()
        gd = dict(g)
        gd['members'] = [dict(m) for m in members]
        gd['shows'] = [dict(s) for s in shows]
        groups_data.append(gd)

    all_settings = {r['key']: r['value'] for r in
                    db.execute("SELECT key, value FROM app_settings").fetchall()}

    db.close()
    _is_ca = session.get('is_content_admin', False) or session.get('user_role') == 'admin'
    _can_manage_crew = _is_ca or session.get('is_scheduler', False) or session.get('is_labor_scheduler', False)
    form_sections = get_form_fields_for_template() if _is_ca else []
    sched_meta_fields = get_schedule_meta_fields() if _is_ca else []

    db3 = get_db()
    sched_templates = [dict(t) for t in db3.execute(
        'SELECT id, name FROM schedule_templates ORDER BY sort_order, name'
    ).fetchall()] if _is_ca else []

    # Asset categories for the form-section editor (link a section to a category
    # so rented assets in that category appear under it in the advance PDF).
    asset_categories = [dict(c) for c in db3.execute(
        'SELECT id, name FROM asset_categories ORDER BY sort_order, name'
    ).fetchall()] if _is_ca else []

    # Arts groups (touring companies / artists) — global free-text dropdown
    # used by the arts_group_dropdown form field type.
    arts_groups = [dict(r) for r in db3.execute(
        'SELECT id, name, sort_order FROM arts_groups ORDER BY sort_order, name'
    ).fetchall()] if _is_ca else []

    # Job positions data for settings tab — visible to content admins AND schedulers
    position_categories = [dict(c) for c in db3.execute(
        'SELECT * FROM position_categories WHERE is_venue=0 OR is_venue IS NULL ORDER BY sort_order, id'
    ).fetchall()] if _can_manage_crew else []
    positions_raw = db3.execute(
        'SELECT jp.*, pc.name as category_name FROM job_positions jp LEFT JOIN position_categories pc ON jp.category_id = pc.id ORDER BY jp.venue, pc.sort_order, jp.sort_order, jp.id'
    ).fetchall() if _can_manage_crew else []
    job_positions = [dict(p) for p in positions_raw]
    distinct_venues = _get_distinct_venues(db3) if _can_manage_crew else []

    # Crew members with rate level info
    crew_members_list = [dict(m) for m in db3.execute(
        '''SELECT cm.*, prl.name as level_name, prl.hourly_rate as level_rate
           FROM crew_members cm
           LEFT JOIN pay_rate_levels prl ON prl.id = cm.rate_level_id
           ORDER BY cm.sort_order, cm.name'''
    ).fetchall()] if _can_manage_crew else []
    pay_rate_levels = [dict(r) for r in db3.execute(
        'SELECT * FROM pay_rate_levels ORDER BY sort_order, name'
    ).fetchall()] if _can_manage_crew else []
    db3.close()

    db_settings = {
        'db_type':          all_settings.get('db_type', 'sqlite'),
        'pg_host':          all_settings.get('pg_host', 'localhost'),
        'pg_port':          all_settings.get('pg_port', '5432'),
        'pg_dbname':        all_settings.get('pg_dbname', '321theater'),
        'pg_user':          all_settings.get('pg_user', ''),
        'pg_app_schema':    all_settings.get('pg_app_schema', 'theater321'),
        'pg_shared_schema': all_settings.get('pg_shared_schema', 'shared'),
    }
    ai_settings = {
        'ollama_enabled':   all_settings.get('ollama_enabled', '0'),
        'ollama_url':       all_settings.get('ollama_url', 'http://localhost:11434'),
        'ollama_model':     all_settings.get('ollama_model', 'llama3.2'),
        'ai_system_prompt': all_settings.get('ai_system_prompt', ''),
    }
    _is_admin = session.get('user_role') == 'admin'
    smtp_settings = {
        'smtp_host':  all_settings.get('smtp_host', ''),
        'smtp_port':  all_settings.get('smtp_port', '587'),
        'smtp_user':  all_settings.get('smtp_user', ''),
        'smtp_pass':  all_settings.get('smtp_pass', '') if _is_admin else '',
        'smtp_from':  all_settings.get('smtp_from', ''),
        'smtp_tls':   all_settings.get('smtp_tls', '1'),
    }
    email_provider_settings = {
        'email_provider':        all_settings.get('email_provider', 'smtp'),
        'direct_ehlo_hostname':  all_settings.get('direct_ehlo_hostname', ''),
        'direct_display_name':   all_settings.get('direct_display_name', ''),
    }
    pdf_email_settings = {
        'pdf_email_send_hour':       all_settings.get('pdf_email_send_hour', '6'),
        'advance_email_enabled':     all_settings.get('advance_email_enabled', '0'),
        'advance_email_days_before': all_settings.get('advance_email_days_before', '7'),
        'schedule_email_enabled_1':  all_settings.get('schedule_email_enabled_1', '0'),
        'schedule_email_days_1':     all_settings.get('schedule_email_days_1', '10'),
        'schedule_email_enabled_2':  all_settings.get('schedule_email_enabled_2', '0'),
        'schedule_email_days_2':     all_settings.get('schedule_email_days_2', '1'),
    }

    # Strip sensitive keys from syslog_settings for non-admin users
    safe_settings = all_settings if _is_admin else {
        k: v for k, v in all_settings.items()
        if k not in ('smtp_pass', 'pg_password', 'wifi_password')
    }

    # Pending registrations for admin approval panel
    pending_regs = []
    if _is_admin:
        db4 = get_db()
        pending_regs = [dict(r) for r in db4.execute("""
            SELECT id, username, display_name, email, created_at, email_confirmed
            FROM user_pending_registration
            WHERE admin_approved=0
            ORDER BY created_at
        """).fetchall()]
        db4.close()

    return render_template('settings.html',
                           contacts=contacts,
                           users=users,
                           groups=groups_data,
                           form_sections=form_sections,
                           asset_categories=asset_categories,
                           arts_groups=arts_groups,
                           sched_meta_fields=sched_meta_fields,
                           syslog_settings=safe_settings,
                           db_settings=db_settings if _is_admin else {},
                           ai_settings=ai_settings,
                           departments=DEPARTMENTS,
                           is_content_admin=_is_ca,
                           sched_templates=sched_templates,
                           position_categories=position_categories,
                           job_positions=job_positions,
                           distinct_venues=distinct_venues,
                           crew_members_list=crew_members_list,
                           pay_rate_levels=pay_rate_levels,
                           wifi_network=all_settings.get('wifi_network', ''),
                           wifi_password=all_settings.get('wifi_password', ''),
                           upload_max_mb=all_settings.get('upload_max_mb', '20'),
                           logo_data=all_settings.get('logo_data', ''),
                           smtp_settings=smtp_settings,
                           email_provider_settings=email_provider_settings,
                           pdf_email_settings=pdf_email_settings,
                           pending_regs=pending_regs,
                           ai_max_sessions=all_settings.get('ai_max_sessions', '2'),
                           user=get_current_user())


@app.route('/settings/contacts/add', methods=['POST'])
@content_admin_required
def add_contact():
    db = get_db()
    name = request.form.get('name','').strip()
    cur = db.execute("""
        INSERT INTO contacts (name, title, department, phone, email,
                              report_recipient, advance_recipient, production_recipient,
                              postnotes_recipient)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (name,
          request.form.get('title','').strip(),
          request.form.get('department','').strip(),
          request.form.get('phone','').strip(),
          request.form.get('email','').strip(),
          1 if request.form.get('report_recipient') else 0,
          1 if request.form.get('advance_recipient') else 0,
          1 if request.form.get('production_recipient') else 0,
          1 if request.form.get('postnotes_recipient') else 0))
    cid_new = cur.lastrowid
    log_audit_change(db, 'CONTACT_ADD', 'contact', cid_new, detail=name,
                     table='contacts')
    db.commit(); db.close()
    syslog_logger.info(f"CONTACT_ADD id={cid_new} name={name!r} by={session.get('username')}")
    flash('Contact added.', 'success')
    return redirect(url_for('settings') + '#contacts')


@app.route('/settings/contacts/<int:cid>/edit', methods=['POST'])
@content_admin_required
def edit_contact(cid):
    data = request.get_json(force=True) or {}
    db = get_db()
    before = _snapshot_row(db, 'contacts', cid)
    db.execute("""
        UPDATE contacts SET name=?, title=?, department=?, phone=?, email=?,
                            report_recipient=?, advance_recipient=?, production_recipient=?,
                            postnotes_recipient=?
        WHERE id=?
    """, (data.get('name',''), data.get('title',''), data.get('department',''),
          data.get('phone',''), data.get('email',''),
          1 if data.get('report_recipient') else 0,
          1 if data.get('advance_recipient') else 0,
          1 if data.get('production_recipient') else 0,
          1 if data.get('postnotes_recipient') else 0,
          cid))
    after = _snapshot_row(db, 'contacts', cid)
    log_audit(db, 'CONTACT_EDIT', 'contact', cid, detail=data.get('name',''),
              before=before, after=after)
    db.commit(); db.close()
    syslog_logger.info(f"CONTACT_EDIT id={cid} name={data.get('name','')!r} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/contacts/<int:cid>/delete', methods=['POST'])
@content_admin_required
def delete_contact(cid):
    db = get_db()
    before = _snapshot_row(db, 'contacts', cid)
    name = before['name'] if before else str(cid)
    log_audit(db, 'CONTACT_DELETE', 'contact', cid, detail=name, before=before)
    db.execute('DELETE FROM contacts WHERE id=?', (cid,))
    db.commit(); db.close()
    syslog_logger.info(f"CONTACT_DELETE id={cid} by={session.get('username')}")
    return jsonify({'success': True})


_MIN_PASSWORD_LENGTH = 8


def _validate_password(pw):
    """Return None if password is acceptable, else an error string."""
    if not pw or len(pw) < _MIN_PASSWORD_LENGTH:
        return f'Password must be at least {_MIN_PASSWORD_LENGTH} characters.'
    return None


@app.route('/settings/users/add', methods=['POST'])
@admin_required
def add_user():
    username = request.form.get('username','').strip()
    password = request.form.get('password','')
    display  = request.form.get('display_name','').strip() or username
    role     = request.form.get('role','user')
    if not username or not password:
        flash('Username and password are required.', 'error')
        return redirect(url_for('settings') + '#users')
    pw_err = _validate_password(password)
    if pw_err:
        flash(pw_err, 'error')
        return redirect(url_for('settings') + '#users')
    email    = request.form.get('email','').strip()
    is_readonly = 1 if request.form.get('is_readonly') else 0
    db = get_db()
    try:
        cur = db.execute("""INSERT INTO users (username, password_hash, display_name, role, email, is_readonly)
                      VALUES (?, ?, ?, ?, ?, ?)""",
                   (username, generate_password_hash(password), display, role, email, is_readonly))
        log_audit(db, 'USER_CREATE', 'user', cur.lastrowid, detail=f'{username} role={role}')
        db.commit()
        flash(f'User "{username}" created.', 'success')
        syslog_logger.info(f"USER_CREATE username={username} role={role} by={session.get('username')}")
    except sqlite3.IntegrityError:
        flash('Username already exists.', 'error')
    db.close()
    return redirect(url_for('settings') + '#users')


@app.route('/settings/users/<int:uid>/edit', methods=['POST'])
@admin_required
def edit_user(uid):
    data = request.get_json(force=True) or {}
    display_name = (data.get('display_name') or '').strip()
    email = (data.get('email') or '').strip()
    role = data.get('role', 'user')
    is_readonly = 1 if data.get('is_readonly') else 0
    is_scheduler = 1 if data.get('is_scheduler') else 0
    is_asset_manager = 1 if data.get('is_asset_manager') else 0
    is_document_viewer = 1 if data.get('is_document_viewer') else 0
    # Doc viewer implies read-only — they only see read views/PDFs.
    if is_document_viewer:
        is_readonly = 1
    # Sanitize JSON lists; only persist known doc types.
    viewer_venues_list = []
    for v in (data.get('viewer_venues') or []):
        s = str(v).strip()
        if s:
            viewer_venues_list.append(s)
    viewer_doc_types_list = []
    for v in (data.get('viewer_doc_types') or []):
        s = str(v).strip().lower()
        if s in DOCUMENT_TYPES:
            viewer_doc_types_list.append(s)
    viewer_venues_json    = json.dumps(viewer_venues_list)    if viewer_venues_list    else None
    viewer_doc_types_json = json.dumps(viewer_doc_types_list) if viewer_doc_types_list else None

    if role not in ('user', 'staff', 'admin'):
        return jsonify({'success': False, 'error': 'Invalid role'}), 400
    # Guard against admin self-lockout: a doc-viewer cannot reach the settings
    # page (the _viewer_gate redirects them to /viewer), so flipping it on
    # yourself locks you out of the admin UI permanently.
    if is_document_viewer and uid == session.get('user_id'):
        return jsonify({'success': False,
                        'error': "You can't make your own account a document viewer — "
                                 "you'd be locked out of admin."}), 400
    db = get_db()
    row = db.execute(
        'SELECT username, is_document_viewer FROM users WHERE id=?', (uid,)
    ).fetchone()
    if not row:
        db.close()
        return jsonify({'success': False, 'error': 'User not found'}), 404
    was_document_viewer = bool(row.get('is_document_viewer', 0))
    db.execute(
        'UPDATE users SET display_name=?, email=?, role=?, is_readonly=?, '
        '                 is_scheduler=?, is_asset_manager=?, '
        '                 is_document_viewer=?, viewer_venues=?, viewer_doc_types=? '
        'WHERE id=?',
        (display_name or row['username'], email, role, is_readonly,
         is_scheduler, is_asset_manager,
         is_document_viewer, viewer_venues_json, viewer_doc_types_json, uid)
    )
    # If the doc-viewer flag flipped (either direction), invalidate any active
    # sessions for this user so the new permission set takes effect immediately
    # instead of waiting for the next 5-minute role refresh.
    if was_document_viewer != bool(is_document_viewer):
        try:
            db.execute('DELETE FROM app_sessions WHERE user_id=?', (uid,))
        except Exception:
            pass  # Table missing or DB-sessions disabled — fall through
    log_audit(db, 'USER_EDIT', 'user', uid,
              detail=(f'role={role} readonly={is_readonly} scheduler={is_scheduler} '
                      f'asset_mgr={is_asset_manager} doc_viewer={is_document_viewer} '
                      f'by={session.get("username")}'))
    db.commit()
    db.close()
    syslog_logger.info(
        f"USER_EDIT user_id={uid} role={role} readonly={is_readonly} "
        f"scheduler={is_scheduler} asset_mgr={is_asset_manager} "
        f"doc_viewer={is_document_viewer} by={session.get('username')}"
    )
    return jsonify({'success': True})


@app.route('/settings/users/<int:uid>/delete', methods=['POST'])
@admin_required
def delete_user(uid):
    if uid == session['user_id']:
        return jsonify({'success': False, 'error': "You can't delete your own account."})
    db = get_db()
    row = db.execute('SELECT username FROM users WHERE id=?', (uid,)).fetchone()
    log_audit(db, 'USER_DELETE', 'user', uid, detail=row['username'] if row else str(uid))
    db.execute('DELETE FROM users WHERE id=?', (uid,))
    db.commit(); db.close()
    syslog_logger.info(f"USER_DELETE user_id={uid} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/users/<int:uid>/reset_password', methods=['POST'])
@admin_required
def admin_reset_user_password(uid):
    data = request.get_json(force=True) or {}
    pw = data.get('password','')
    if not pw:
        return jsonify({'success': False, 'error': 'Password required.'})
    pw_err = _validate_password(pw)
    if pw_err:
        return jsonify({'success': False, 'error': pw_err})
    db = get_db()
    db.execute('UPDATE users SET password_hash=? WHERE id=?', (generate_password_hash(pw), uid))
    log_audit(db, 'USER_PASSWORD_RESET', 'user', uid, detail=f'reset by {session.get("username")}')
    db.commit(); db.close()
    syslog_logger.info(f"PASSWORD_CHANGE user_id={uid} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/account/theme', methods=['POST'])
@login_required
def set_theme():
    data = request.get_json(force=True) or {}
    theme = data.get('theme', 'dark')
    if theme not in ('dark', 'light'):
        theme = 'dark'
    db = get_db()
    db.execute('UPDATE users SET theme=? WHERE id=?', (theme, session['user_id']))
    db.commit()
    db.close()
    session['theme'] = theme
    return jsonify({'success': True})


@app.route('/account/change_password', methods=['POST'])
@login_required
def change_own_password():
    data = request.get_json(force=True) or {}
    current = data.get('current_password','')
    new_pw  = data.get('new_password','')
    db = get_db()
    user = db.execute('SELECT * FROM users WHERE id=?', (session['user_id'],)).fetchone()
    if not check_password_hash(user['password_hash'], current):
        db.close()
        return jsonify({'success': False, 'error': 'Current password incorrect.'})
    pw_err = _validate_password(new_pw)
    if pw_err:
        db.close()
        return jsonify({'success': False, 'error': pw_err})
    db.execute('UPDATE users SET password_hash=? WHERE id=?',
               (generate_password_hash(new_pw), session['user_id']))
    db.commit(); db.close()
    syslog_logger.info(f"PASSWORD_CHANGE user_id={session['user_id']} by={session.get('username')} (self)")
    return jsonify({'success': True})


# ─── User Groups Management ───────────────────────────────────────────────────

@app.route('/settings/groups/add', methods=['POST'])
@admin_required
def add_group():
    data = request.get_json(force=True) or {}
    name = data.get('name','').strip()
    group_type = data.get('group_type','all_access')
    desc = data.get('description','').strip()
    if not name:
        return jsonify({'success': False, 'error': 'Name required.'}), 400
    db = get_db()
    try:
        cur = db.execute(
            'INSERT INTO user_groups (name, group_type, description) VALUES (?,?,?)',
            (name, group_type, desc)
        )
        gid = cur.lastrowid
        log_audit(db, 'GROUP_CREATE', 'group', gid, detail=name)
        db.commit()
        syslog_logger.info(f"GROUP_CREATE name={name} by={session.get('username')}")
        return jsonify({'success': True, 'id': gid})
    except sqlite3.IntegrityError:
        return jsonify({'success': False, 'error': 'Group name already exists.'}), 400
    finally:
        db.close()


@app.route('/settings/groups/<int:gid>/edit', methods=['POST'])
@admin_required
def edit_group(gid):
    data = request.get_json(force=True) or {}
    db = get_db()
    db.execute("""
        UPDATE user_groups SET name=?, group_type=?, description=? WHERE id=?
    """, (data.get('name',''), data.get('group_type','all_access'),
          data.get('description',''), gid))
    log_audit(db, 'GROUP_EDIT', 'group', gid, detail=data.get('name',''))
    db.commit(); db.close()
    syslog_logger.info(f"GROUP_EDIT id={gid} name={data.get('name','')!r} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/groups/<int:gid>/delete', methods=['POST'])
@admin_required
def delete_group(gid):
    db = get_db()
    row = db.execute('SELECT name FROM user_groups WHERE id=?', (gid,)).fetchone()
    log_audit(db, 'GROUP_DELETE', 'group', gid, detail=row['name'] if row else str(gid))
    db.execute('DELETE FROM user_groups WHERE id=?', (gid,))
    db.commit(); db.close()
    syslog_logger.info(f"GROUP_DELETE id={gid} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/groups/<int:gid>/members/add', methods=['POST'])
@admin_required
def add_group_member(gid):
    data = request.get_json(force=True) or {}
    uid = data.get('user_id')
    if not uid:
        return jsonify({'success': False, 'error': 'user_id required'}), 400
    db = get_db()
    db.execute('INSERT OR IGNORE INTO user_group_members (user_id, group_id) VALUES (?,?)',
               (uid, gid))
    log_audit(db, 'GROUP_MEMBER_ADD', 'group', gid, detail=f'user_id={uid}')

    # Refresh the affected user's is_restricted in their session (best effort)
    # Session is server-side cookie; we can't update other sessions directly.
    # User will see updated access on next login.

    db.commit()
    db.close()
    syslog_logger.info(
        f"GROUP_ASSIGN user_id={uid} group_id={gid} by={session.get('username')}"
    )
    return jsonify({'success': True})


@app.route('/settings/groups/<int:gid>/members/remove', methods=['POST'])
@admin_required
def remove_group_member(gid):
    data = request.get_json(force=True) or {}
    uid = data.get('user_id')
    if not uid:
        return jsonify({'success': False, 'error': 'user_id required'}), 400
    db = get_db()
    db.execute('DELETE FROM user_group_members WHERE user_id=? AND group_id=?', (uid, gid))
    log_audit(db, 'GROUP_MEMBER_REMOVE', 'group', gid, detail=f'user_id={uid}')
    db.commit(); db.close()
    syslog_logger.info(
        f"GROUP_REMOVE user_id={uid} group_id={gid} by={session.get('username')}"
    )
    return jsonify({'success': True})


@app.route('/api/groups')
@login_required
def api_groups():
    db = get_db()
    groups = db.execute('SELECT * FROM user_groups ORDER BY name').fetchall()
    db.close()
    return jsonify([dict(g) for g in groups])


# ─── Form Field Editor ────────────────────────────────────────────────────────

@app.route('/settings/form-fields')
@content_admin_required
def form_fields_settings():
    form_sections = get_form_fields_for_template()
    db = get_db()
    users = db.execute(
        'SELECT id, username, display_name, role, created_at FROM users ORDER BY display_name'
    ).fetchall()
    contacts = db.execute('SELECT * FROM contacts ORDER BY department, name').fetchall()
    groups = db.execute('SELECT * FROM user_groups ORDER BY name').fetchall()

    groups_data = []
    for g in groups:
        members = db.execute("""
            SELECT u.id, u.display_name, u.username FROM user_group_members ugm
            JOIN users u ON ugm.user_id = u.id
            WHERE ugm.group_id = ?
        """, (g['id'],)).fetchall()
        shows = db.execute("""
            SELECT s.id, s.name, s.show_date FROM show_group_access sga
            JOIN shows s ON sga.show_id = s.id
            WHERE sga.group_id = ?
            ORDER BY s.show_date DESC
        """, (g['id'],)).fetchall()
        gd = dict(g)
        gd['members'] = [dict(m) for m in members]
        gd['shows'] = [dict(s) for s in shows]
        groups_data.append(gd)

    all_settings = {r['key']: r['value'] for r in
                    db.execute("SELECT key, value FROM app_settings").fetchall()}
    db.close()

    _is_ca = session.get('is_content_admin', False) or session.get('user_role') == 'admin'

    db4 = get_db()
    sched_templates2 = [dict(t) for t in db4.execute(
        'SELECT id, name FROM schedule_templates ORDER BY sort_order, name'
    ).fetchall()] if _is_ca else []
    db4.close()

    return render_template('settings.html',
                           contacts=contacts,
                           users=users,
                           groups=groups_data,
                           form_sections=form_sections,
                           sched_meta_fields=get_schedule_meta_fields(),
                           syslog_settings=all_settings,
                           departments=DEPARTMENTS,
                           active_tab='fields',
                           is_content_admin=_is_ca,
                           sched_templates=sched_templates2,
                           wifi_network=all_settings.get('wifi_network', ''),
                           wifi_password=all_settings.get('wifi_password', ''),
                           upload_max_mb=all_settings.get('upload_max_mb', '20'),
                           logo_data=all_settings.get('logo_data', ''),
                           db_settings=all_settings,
                           ai_settings=all_settings,
                           user=get_current_user())


@app.route('/settings/form-fields/add', methods=['POST'])
@content_admin_required
def add_form_field():
    data = request.get_json(force=True) or {}
    section_id = data.get('section_id')
    field_key  = data.get('field_key','').strip().lower().replace(' ','_')
    label      = data.get('label','').strip()
    if not section_id or not field_key or not label:
        return jsonify({'success': False, 'error': 'section_id, field_key, and label required.'}), 400

    options = data.get('options', [])
    options_json = json.dumps(options) if options else None

    db = get_db()
    # Put it at the end of the section
    max_order = db.execute(
        'SELECT MAX(sort_order) FROM form_fields WHERE section_id=?', (section_id,)
    ).fetchone()[0] or 0
    try:
        cur = db.execute("""
            INSERT INTO form_fields
            (section_id, field_key, label, field_type, sort_order,
             options_json, contact_dept, conditional_show_when,
             help_text, placeholder, width_hint, is_notes_field, ai_hint,
             display_as, allow_multi, auto_select_visible, hide_from_pdf, upload_button_only)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
        """, (section_id, field_key, label,
              data.get('field_type','text'), max_order + 10,
              options_json,
              data.get('contact_dept'),
              data.get('conditional_show_when'),
              data.get('help_text'),
              data.get('placeholder',''),
              data.get('width_hint','full'),
              1 if data.get('is_notes_field') else 0,
              data.get('ai_hint') or None,
              data.get('display_as') or None,
              1 if data.get('allow_multi') else 0,
              1 if data.get('auto_select_visible') else 0,
              1 if data.get('hide_from_pdf') else 0,
              1 if data.get('upload_button_only') else 0))
        fid = cur.lastrowid
        log_audit_change(db, 'FIELD_ADD', 'form_field', fid, detail=field_key,
                         table='form_fields')
        db.commit()
        syslog_logger.info(f"FIELD_ADD key={field_key} by={session.get('username')}")
        return jsonify({'success': True, 'id': fid})
    except (sqlite3.IntegrityError, DBIntegrityError):
        return jsonify({'success': False, 'error': f'field_key "{field_key}" already exists.'}), 400
    finally:
        db.close()


@app.route('/settings/form-fields/<int:fid>/edit', methods=['POST'])
@content_admin_required
def edit_form_field(fid):
    data = request.get_json(force=True) or {}
    options = data.get('options', [])
    options_json = json.dumps(options) if options else None
    db = get_db()
    before = _snapshot_row(db, 'form_fields', fid)
    db.execute("""
        UPDATE form_fields SET
            section_id=?, label=?, field_type=?,
            options_json=?, contact_dept=?, conditional_show_when=?,
            help_text=?, placeholder=?, width_hint=?, is_notes_field=?, ai_hint=?,
            display_as=?, allow_multi=?, auto_select_visible=?, hide_from_pdf=?, upload_button_only=?
        WHERE id=?
    """, (data.get('section_id'), data.get('label',''),
          data.get('field_type','text'), options_json,
          data.get('contact_dept'), data.get('conditional_show_when'),
          data.get('help_text'), data.get('placeholder',''),
          data.get('width_hint','full'),
          1 if data.get('is_notes_field') else 0,
          data.get('ai_hint') or None,
          data.get('display_as') or None,
          1 if data.get('allow_multi') else 0,
          1 if data.get('auto_select_visible') else 0,
          1 if data.get('hide_from_pdf') else 0,
          1 if data.get('upload_button_only') else 0,
          fid))
    after = _snapshot_row(db, 'form_fields', fid)
    log_audit(db, 'FIELD_EDIT', 'form_field', fid, detail=data.get('label',''),
              before=before, after=after)
    db.commit(); db.close()
    syslog_logger.info(f"FIELD_EDIT id={fid} label={data.get('label','')!r} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/form-fields/<int:fid>/delete', methods=['POST'])
@content_admin_required
def delete_form_field(fid):
    db = get_db()
    before = _snapshot_row(db, 'form_fields', fid)
    log_audit(db, 'FIELD_DELETE', 'form_field', fid, before=before)
    db.execute('DELETE FROM form_fields WHERE id=?', (fid,))
    db.commit(); db.close()
    syslog_logger.info(f"FIELD_DELETE id={fid} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/form-fields/reorder', methods=['POST'])
@content_admin_required
def reorder_form_fields():
    data = request.get_json(force=True) or {}
    field_ids = data.get('field_ids', [])
    db = get_db()
    for i, fid in enumerate(field_ids):
        db.execute('UPDATE form_fields SET sort_order=? WHERE id=?', (i * 10, fid))
    db.commit(); db.close()
    return jsonify({'success': True})


@app.route('/settings/form-sections/add', methods=['POST'])
@content_admin_required
def add_form_section():
    data = request.get_json(force=True) or {}
    section_key = data.get('section_key','').strip().lower().replace(' ','_')
    label = data.get('label','').strip()
    if not section_key or not label:
        return jsonify({'success': False, 'error': 'section_key and label required.'}), 400
    db = get_db()
    max_order = db.execute('SELECT MAX(sort_order) FROM form_sections').fetchone()[0] or 0
    asset_cat_raw = data.get('asset_category_id')
    asset_cat_id = int(asset_cat_raw) if str(asset_cat_raw or '').strip().isdigit() else None
    try:
        cur = db.execute("""
            INSERT INTO form_sections (section_key, label, sort_order, collapsible, icon, default_open, asset_category_id)
            VALUES (?,?,?,?,?,?,?)
        """, (section_key, label, max_order + 10,
              1 if data.get('collapsible', True) else 0,
              data.get('icon', '◈'),
              0 if str(data.get('default_open', '1')) == '0' else 1,
              asset_cat_id))
        sid = cur.lastrowid
        log_audit_change(db, 'SECTION_ADD', 'form_section', sid, detail=label,
                         table='form_sections')
        db.commit()
        syslog_logger.info(f"SECTION_ADD id={sid} label={label!r} by={session.get('username')}")
        return jsonify({'success': True, 'id': sid})
    except sqlite3.IntegrityError:
        return jsonify({'success': False, 'error': f'section_key "{section_key}" already exists.'}), 400
    finally:
        db.close()


@app.route('/settings/form-sections/<int:sid>/edit', methods=['POST'])
@content_admin_required
def edit_form_section(sid):
    data = request.get_json(force=True) or {}
    db = get_db()
    before = _snapshot_row(db, 'form_sections', sid)
    asset_cat_raw = data.get('asset_category_id')
    asset_cat_id = int(asset_cat_raw) if str(asset_cat_raw or '').strip().isdigit() else None
    db.execute("""
        UPDATE form_sections SET label=?, collapsible=?, icon=?, default_open=?, asset_category_id=? WHERE id=?
    """, (data.get('label',''),
          1 if data.get('collapsible', True) else 0,
          data.get('icon','◈'),
          0 if str(data.get('default_open', '1')) == '0' else 1,
          asset_cat_id,
          sid))
    after = _snapshot_row(db, 'form_sections', sid)
    log_audit(db, 'SECTION_EDIT', 'form_section', sid, detail=data.get('label',''),
              before=before, after=after)
    db.commit(); db.close()
    syslog_logger.info(f"SECTION_EDIT id={sid} label={data.get('label','')!r} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/form-sections/<int:sid>/delete', methods=['POST'])
@content_admin_required
def delete_form_section(sid):
    db = get_db()
    before = _snapshot_row(db, 'form_sections', sid)
    log_audit(db, 'SECTION_DELETE', 'form_section', sid, before=before)
    db.execute('DELETE FROM form_sections WHERE id=?', (sid,))
    db.commit(); db.close()
    syslog_logger.info(f"SECTION_DELETE id={sid} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/form-sections/reorder', methods=['POST'])
@content_admin_required
def reorder_form_sections():
    data = request.get_json(force=True) or {}
    section_ids = data.get('section_ids', [])
    db = get_db()
    for i, sid in enumerate(section_ids):
        db.execute('UPDATE form_sections SET sort_order=? WHERE id=?', (i * 10, sid))
    db.commit(); db.close()
    return jsonify({'success': True})


# ─── PDF Layout Designer (admin) ──────────────────────────────────────────────
#
# Admin-only editor that lets content admins reorder, hide, relabel, and font-
# size sections (and per-section fields) for the 5 PDF exports. Layouts are
# stored as JSON in app_settings under `pdf_layout_<type>`. See pdf_layouts.py
# for the catalog (what's configurable) and PdfLayout (template-side helper).
# Phase rollout: only `advance` is enabled in v1; other types return 404.

PDF_DESIGNER_ENABLED_TYPES = {'advance', 'postnotes', 'asset_invoice', 'post_show_invoice', 'schedule'}


def _resolve_advance_form_sections():
    """Helper: returns form_sections (or [] on failure) for advance catalog."""
    try:
        return get_form_fields_for_template()
    except Exception as e:
        app.logger.warning(f'pdf_designer: form_sections fetch failed: {e}')
        return []


def _pdf_designer_catalog(pdf_type):
    if pdf_type == 'advance':
        return pdf_layouts.get_catalog('advance', form_sections=_resolve_advance_form_sections())
    return pdf_layouts.get_catalog(pdf_type)


@app.route('/admin/pdf-designer')
@content_admin_required
def pdf_designer_page():
    db = get_db()
    recent_shows = db.execute(
        "SELECT id, name, show_date FROM shows "
        "WHERE status != 'archived' "
        "ORDER BY show_date DESC, id DESC LIMIT 10"
    ).fetchall()
    db.close()
    return render_template('pdf_designer.html',
                           pdf_types=list(pdf_layouts.PDF_TYPES),
                           enabled_types=sorted(PDF_DESIGNER_ENABLED_TYPES),
                           recent_shows=[dict(r) for r in recent_shows],
                           font_size_choices=list(pdf_layouts.FONT_SIZE_CHOICES),
                           user=get_current_user())


@app.route('/admin/pdf-designer/<pdf_type>/catalog.json')
@content_admin_required
def pdf_designer_catalog(pdf_type):
    if pdf_type not in pdf_layouts.PDF_TYPES:
        abort(404)
    if pdf_type not in PDF_DESIGNER_ENABLED_TYPES:
        return jsonify({'error': 'pdf type not enabled yet'}), 404
    return jsonify({'pdf_type': pdf_type, 'catalog': _pdf_designer_catalog(pdf_type)})


@app.route('/admin/pdf-designer/<pdf_type>/layout.json', methods=['GET'])
@content_admin_required
def pdf_designer_layout_get(pdf_type):
    if pdf_type not in pdf_layouts.PDF_TYPES:
        abort(404)
    if pdf_type not in PDF_DESIGNER_ENABLED_TYPES:
        return jsonify({'error': 'pdf type not enabled yet'}), 404
    catalog = _pdf_designer_catalog(pdf_type)
    raw = get_app_setting(f'pdf_layout_{pdf_type}', '')
    layout = pdf_layouts._parse_or_default(raw, catalog, pdf_type)
    return jsonify({'pdf_type': pdf_type, 'layout': layout, 'catalog': catalog})


@app.route('/admin/pdf-designer/<pdf_type>/layout.json', methods=['POST'])
@content_admin_required
def pdf_designer_layout_save(pdf_type):
    if pdf_type not in pdf_layouts.PDF_TYPES:
        abort(404)
    if pdf_type not in PDF_DESIGNER_ENABLED_TYPES:
        return jsonify({'error': 'pdf type not enabled yet'}), 404
    payload = request.get_json(force=True, silent=True)
    if payload is None:
        return jsonify({'error': 'invalid JSON body'}), 400
    catalog = _pdf_designer_catalog(pdf_type)
    cleaned, err = pdf_layouts.validate_payload(payload, catalog, pdf_type)
    if err:
        return jsonify({'error': err}), 400
    db = get_db()
    key = f'pdf_layout_{pdf_type}'
    before_raw = get_app_setting(key, '')
    try:
        before_obj = json.loads(before_raw) if before_raw else None
    except Exception:
        before_obj = None
    new_value = json.dumps(cleaned)
    db.execute('INSERT OR REPLACE INTO app_settings (key, value) VALUES (?, ?)',
               (key, new_value))
    log_audit_change(db, 'LAYOUT_EDIT', 'pdf_layout', None,
                     detail=pdf_type, before=before_obj, after=cleaned)
    db.commit(); db.close()
    return jsonify({'success': True, 'layout': cleaned})


@app.route('/admin/pdf-designer/<pdf_type>/reset', methods=['POST'])
@content_admin_required
def pdf_designer_layout_reset(pdf_type):
    if pdf_type not in pdf_layouts.PDF_TYPES:
        abort(404)
    if pdf_type not in PDF_DESIGNER_ENABLED_TYPES:
        return jsonify({'error': 'pdf type not enabled yet'}), 404
    db = get_db()
    key = f'pdf_layout_{pdf_type}'
    before_raw = get_app_setting(key, '')
    try:
        before_obj = json.loads(before_raw) if before_raw else None
    except Exception:
        before_obj = None
    db.execute('DELETE FROM app_settings WHERE key=?', (key,))
    log_audit_change(db, 'LAYOUT_RESET', 'pdf_layout', None,
                     detail=pdf_type, before=before_obj, after=None)
    db.commit(); db.close()
    return jsonify({'success': True})


@app.route('/admin/pdf-designer/<pdf_type>/preview.pdf')
@content_admin_required
def pdf_designer_preview(pdf_type):
    if pdf_type not in pdf_layouts.PDF_TYPES:
        abort(404)
    if pdf_type not in PDF_DESIGNER_ENABLED_TYPES:
        return ('PDF type not enabled yet', 404)
    try:
        show_id = int(request.args.get('show_id', '0'))
    except (TypeError, ValueError):
        show_id = 0
    if show_id <= 0:
        return ('show_id query param required', 400)

    db = get_db()
    show = db.execute('SELECT id FROM shows WHERE id=?', (show_id,)).fetchone()
    db.close()
    if not show:
        return ('show not found', 404)

    # Tuple-returning builders (advance, schedule, postnotes): unpack 5-tuple.
    tuple_builders = {
        'advance': _build_advance_pdf,
        'postnotes': _build_postnotes_pdf,
        'schedule': _build_schedule_pdf,
    }
    # Response-returning route handlers (invoices): call directly, extract .data.
    response_builders = {
        'asset_invoice': show_asset_invoice,
        'post_show_invoice': show_post_invoice,
    }

    pdf_bytes = None
    try:
        if pdf_type in tuple_builders:
            _html, _ver, _show, pdf_bytes, _log_id = tuple_builders[pdf_type](show_id)
        elif pdf_type in response_builders:
            sub_resp = response_builders[pdf_type](show_id)
            if hasattr(sub_resp, 'data'):
                pdf_bytes = sub_resp.data
            elif isinstance(sub_resp, (bytes, bytearray)):
                pdf_bytes = bytes(sub_resp)
            else:
                return ('builder returned unsupported type', 500)
        else:
            return ('builder not wired', 404)
    except Exception as e:
        app.logger.error(f'pdf_designer preview {pdf_type} show={show_id}: {e}')
        return (f'preview failed: {e}', 500)

    if not pdf_bytes:
        return ('empty PDF', 500)
    resp = make_response(pdf_bytes)
    resp.headers['Content-Type'] = 'application/pdf'
    resp.headers['Content-Disposition'] = f'inline; filename="{pdf_type}_preview.pdf"'
    resp.headers['Cache-Control'] = 'no-store'
    return resp


@app.route('/api/form-fields')
@login_required
def api_form_fields():
    return jsonify(get_form_fields_for_template())


# ─── Schedule Meta Field Editor ───────────────────────────────────────────────

@app.route('/api/schedule-meta-fields')
@login_required
def api_schedule_meta_fields():
    return jsonify(get_schedule_meta_fields())


@app.route('/settings/schedule-meta-fields/add', methods=['POST'])
@content_admin_required
def add_sched_meta_field():
    data = request.get_json(force=True) or {}
    field_key = data.get('field_key', '').strip().lower().replace(' ', '_')
    label = data.get('label', '').strip()
    if not field_key or not label:
        return jsonify({'success': False, 'error': 'field_key and label required.'}), 400
    db = get_db()
    max_order = db.execute('SELECT MAX(sort_order) FROM schedule_meta_fields').fetchone()[0] or 0
    try:
        cur = db.execute("""
            INSERT INTO schedule_meta_fields
              (field_key, label, field_type, advance_field_key, sort_order, width_hint, show_in_contacts)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (field_key, label,
              data.get('field_type', 'text'),
              data.get('advance_field_key', '').strip() or None,
              max_order + 10,
              data.get('width_hint', 'half'),
              1 if data.get('show_in_contacts') else 0))
        fid = cur.lastrowid
        db.commit()
        return jsonify({'success': True, 'id': fid})
    except sqlite3.IntegrityError:
        return jsonify({'success': False, 'error': f'field_key "{field_key}" already exists.'}), 400
    finally:
        db.close()


@app.route('/settings/schedule-meta-fields/<int:fid>/edit', methods=['POST'])
@content_admin_required
def edit_sched_meta_field(fid):
    data = request.get_json(force=True) or {}
    db = get_db()
    db.execute("""
        UPDATE schedule_meta_fields
        SET label=?, field_type=?, advance_field_key=?, width_hint=?, show_in_contacts=?
        WHERE id=?
    """, (data.get('label', ''),
          data.get('field_type', 'text'),
          data.get('advance_field_key', '').strip() or None,
          data.get('width_hint', 'half'),
          1 if data.get('show_in_contacts') else 0,
          fid))
    db.commit(); db.close()
    return jsonify({'success': True})


@app.route('/settings/schedule-meta-fields/<int:fid>/delete', methods=['POST'])
@content_admin_required
def delete_sched_meta_field(fid):
    db = get_db()
    db.execute('DELETE FROM schedule_meta_fields WHERE id=?', (fid,))
    db.commit(); db.close()
    return jsonify({'success': True})


@app.route('/settings/schedule-meta-fields/reorder', methods=['POST'])
@content_admin_required
def reorder_sched_meta_fields():
    data = request.get_json(force=True) or {}
    field_ids = data.get('field_ids', [])
    db = get_db()
    for i, fid in enumerate(field_ids):
        db.execute('UPDATE schedule_meta_fields SET sort_order=? WHERE id=?', (i * 10, fid))
    db.commit(); db.close()
    return jsonify({'success': True})


# ─── Syslog Settings ──────────────────────────────────────────────────────────

_last_port_change = 0  # timestamp of last port change — rate limiter

@app.route('/settings/server', methods=['POST'])
@admin_required
def save_server_settings():
    global _last_port_change
    import time as _time_mod
    now = _time_mod.time()
    if now - _last_port_change < 30:
        return jsonify({'success': False, 'error': 'Port was changed recently. Wait 30 seconds.'}), 429
    data = request.get_json(force=True) or {}
    port_str = str(data.get('app_port', '5400')).strip()
    try:
        port_val = int(port_str)
        if not (1024 <= port_val <= 65535):
            return jsonify({'success': False, 'error': 'Port must be between 1024 and 65535.'}), 400
    except ValueError:
        return jsonify({'success': False, 'error': 'Invalid port number.'}), 400

    db = get_db()
    db.execute('INSERT OR REPLACE INTO app_settings (key, value) VALUES (?,?)',
               ('app_port', str(port_val)))
    log_audit(db, 'SETTINGS_CHANGE', 'setting', None, detail=f'app_port={port_val}')
    db.commit(); db.close()
    _last_port_change = now
    syslog_logger.info(f"SETTINGS_CHANGE key=app_port value={port_val} by={session.get('username')}")

    # Check if running under the showadvance systemd service
    svc_active = False
    try:
        result = subprocess.run(
            ['systemctl', 'is-active', 'showadvance'],
            capture_output=True, text=True, timeout=3
        )
        svc_active = result.stdout.strip() == 'active'
    except Exception:
        pass

    if svc_active:
        # Schedule restart after response is sent (1 s delay so the JSON
        # response reaches the browser before the process is killed)
        def _do_restart():
            import time as _time
            _time.sleep(1.0)
            try:
                subprocess.run(
                    ['sudo', 'systemctl', 'restart', 'showadvance'],
                    timeout=10
                )
            except Exception as exc:
                app.logger.error(f'Service restart failed: {exc}')
        threading.Thread(target=_do_restart, daemon=True).start()

    return jsonify({
        'success': True,
        'new_port': port_val,
        'restarting': svc_active,
        'message': (
            f'Port changed to {port_val}. Service is restarting...'
            if svc_active else
            f'Port set to {port_val}. Restart the service to apply.'
        )
    })


@app.route('/settings/syslog', methods=['POST'])
@admin_required
def save_syslog_settings():
    data = request.get_json(force=True) or {}
    # Validate syslog host — block metadata/link-local to prevent exfiltration
    syslog_host = data.get('syslog_host', '')
    if syslog_host and _is_blocked_host(syslog_host):
        return jsonify({'success': False, 'error': 'Invalid syslog host.'}), 400
    # Validate port range
    try:
        syslog_port = int(data.get('syslog_port', 514))
        if not (1 <= syslog_port <= 65535):
            raise ValueError
    except (ValueError, TypeError):
        return jsonify({'success': False, 'error': 'Invalid syslog port.'}), 400
    # Validate facility against known values
    valid_facilities = [f'LOG_LOCAL{i}' for i in range(8)] + [
        'LOG_USER', 'LOG_DAEMON', 'LOG_SYSLOG', 'LOG_AUTH']
    if data.get('syslog_facility') and data['syslog_facility'] not in valid_facilities:
        return jsonify({'success': False, 'error': 'Invalid syslog facility.'}), 400
    db = get_db()
    for key in ('syslog_host', 'syslog_port', 'syslog_facility', 'syslog_enabled'):
        if key in data:
            db.execute('INSERT OR REPLACE INTO app_settings (key, value) VALUES (?,?)',
                       (key, str(data[key])))
    log_audit(db, 'SETTINGS_CHANGE', 'setting', None, detail='syslog')
    db.commit(); db.close()
    reload_syslog_handler()
    syslog_logger.info(f"SETTINGS_CHANGE key=syslog by={session.get('username')}")
    return jsonify({'success': True})


# ─── Backup Management ────────────────────────────────────────────────────────

@app.route('/settings/backups')
@admin_required
def backup_status():
    result = {'hourly': [], 'daily': [], 'db_type': db_adapter.read_db_settings(DATABASE).get('db_type', 'sqlite')}
    for kind in ('hourly', 'daily'):
        d = os.path.join(BACKUP_DIR, kind)
        if os.path.isdir(d):
            files = sorted(
                [f for f in os.listdir(d) if f.endswith('.db') or f.endswith('.sql.gz')],
                reverse=True
            )
            result[kind] = [{
                'filename': f,
                'size_kb': round(os.path.getsize(os.path.join(d, f)) / 1024, 1),
                'mtime': datetime.fromtimestamp(
                    os.path.getmtime(os.path.join(d, f))
                ).strftime('%Y-%m-%d %H:%M')
            } for f in files[:10]]
    return jsonify(result)


@app.route('/settings/backups/run', methods=['POST'])
@admin_required
def manual_backup():
    try:
        run_hourly_backup()
        return jsonify({'success': True, 'message': 'Backup created successfully.'})
    except Exception as e:
        app.logger.error(f'Backup failed: {e}')
        return jsonify({'success': False, 'error': 'Backup failed. Check server logs.'}), 500


@app.route('/settings/backups/download/<kind>/<filename>')
@admin_required
def download_backup(kind, filename):
    if kind not in ('hourly', 'daily'):
        abort(404)
    if not re.match(r'^advance_[\d_]+\.(db|sql\.gz)$', filename):
        abort(404)
    path = os.path.join(BACKUP_DIR, kind, filename)
    if not os.path.isfile(path):
        abort(404)
    return send_file(path, as_attachment=True, download_name=filename)


# ─── API ──────────────────────────────────────────────────────────────────────

_gs_rate_limit = limiter.limit("200 per minute") if (_limiter_available and limiter) else (lambda f: f)


@app.route('/api/search')
@login_required
@_gs_rate_limit
def global_search():
    """Universal search across shows, contacts, asset types, and asset items."""
    q = (request.args.get('q') or '').strip()
    if len(q) < 2 or len(q) > 255:
        return jsonify([])

    db = get_db()
    results = []
    like = f'%{q}%'
    is_admin = session.get('user_role') == 'admin'

    # ── Shows ────────────────────────────────────────────────────────────────
    accessible = get_accessible_shows(session['user_id'])  # None=all, []=none, list=ids
    if accessible != []:
        if accessible is None:
            show_rows = db.execute("""
                SELECT id, name, show_date, venue, performance_company, status
                FROM shows
                WHERE name LIKE ? OR venue LIKE ? OR performance_company LIKE ? OR show_date LIKE ?
                ORDER BY show_date DESC LIMIT 6
            """, (like, like, like, like)).fetchall()
        else:
            placeholders = ','.join('?' * len(accessible))
            show_rows = db.execute(f"""
                SELECT id, name, show_date, venue, performance_company, status
                FROM shows
                WHERE id IN ({placeholders})
                  AND (name LIKE ? OR venue LIKE ? OR performance_company LIKE ? OR show_date LIKE ?)
                ORDER BY show_date DESC LIMIT 6
            """, (*accessible, like, like, like, like)).fetchall()
        for r in show_rows:
            sub_parts = [p for p in [r['show_date'], r['venue'], r['performance_company']] if p]
            results.append({
                'type': 'show',
                'icon': '🎭',
                'label': r['name'],
                'sub': '  ·  '.join(sub_parts),
                'url': f"/shows/{r['id']}",
                'status': r['status'],
            })

    # ── Contacts ────────────────────────────────────────────────────────────
    contact_rows = db.execute("""
        SELECT id, name, title, department, email, phone
        FROM contacts
        WHERE name LIKE ? OR department LIKE ? OR email LIKE ? OR phone LIKE ? OR title LIKE ?
        ORDER BY department, name LIMIT 5
    """, (like, like, like, like, like)).fetchall()
    for r in contact_rows:
        sub_parts = [p for p in [r['department'], r['title'], r['email']] if p]
        results.append({
            'type': 'contact',
            'icon': '👤',
            'label': r['name'],
            'sub': '  ·  '.join(sub_parts),
            'url': None,  # contacts don't have their own page; sub-label carries the info
        })

    # ── Asset Types (admin only) ─────────────────────────────────────────────
    if is_admin:
        type_rows = db.execute("""
            SELECT at.id, at.name, at.manufacturer, at.model, ac.name as cat_name,
                   at.storage_location, at.is_retired
            FROM asset_types at
            JOIN asset_categories ac ON ac.id = at.category_id
            WHERE at.name LIKE ? OR at.manufacturer LIKE ? OR at.model LIKE ?
            ORDER BY at.is_retired, at.name LIMIT 5
        """, (like, like, like)).fetchall()
        for r in type_rows:
            label_parts = [p for p in [r['manufacturer'], r['model']] if p]
            sub_parts = [r['cat_name']] + ([r['storage_location']] if r['storage_location'] else [])
            results.append({
                'type': 'asset_type',
                'icon': '◈',
                'label': r['name'] + (f" — {' '.join(label_parts)}" if label_parts else ''),
                'sub': '  ·  '.join(sub_parts) + ('  ·  RETIRED' if r['is_retired'] else ''),
                'url': '/assets',
                'retired': bool(r['is_retired']),
            })

        # ── Asset Items / Barcodes (leading-zero tolerant) ──────────────────
        # Strip leading zeros from stored barcodes and compare with stripped query
        norm_q = q.lstrip('0') or '0'
        item_rows = db.execute("""
            SELECT ai.id, ai.barcode, ai.status, ai.condition,
                   at.name as type_name, at.id as type_id, ac.name as cat_name
            FROM asset_items ai
            JOIN asset_types at ON at.id = ai.asset_type_id
            JOIN asset_categories ac ON ac.id = at.category_id
            WHERE ai.barcode LIKE ?
               OR ltrim(ai.barcode, '0') = ?
               OR ai.barcode = ?
            ORDER BY ai.status, ai.id LIMIT 5
        """, (like, norm_q, q)).fetchall()
        for r in item_rows:
            results.append({
                'type': 'asset_item',
                'icon': '🔖',
                'label': f"Unit #{r['id']}" + (f" — {r['barcode']}" if r['barcode'] else ''),
                'sub': f"{r['type_name']}  ·  {r['cat_name']}  ·  {r['status']}",
                'url': '/assets',
                'status': r['status'],
            })

    db.close()
    return jsonify(results)


@app.route('/api/contacts')
@login_required
def api_contacts():
    db = get_db()
    contacts = db.execute('SELECT * FROM contacts ORDER BY department, name').fetchall()
    db.close()
    return jsonify([dict(c) for c in contacts])


@app.route('/api/users')
@admin_required
def api_users():
    db = get_db()
    users = db.execute(
        'SELECT id, username, display_name, role FROM users ORDER BY display_name'
    ).fetchall()
    db.close()
    return jsonify([dict(u) for u in users])


@app.route('/api/shows')
@admin_required
def api_shows():
    db = get_db()
    shows = db.execute(
        "SELECT id, name, show_date, status FROM shows ORDER BY show_date DESC"
    ).fetchall()
    db.close()
    return jsonify([dict(s) for s in shows])


# ─── API Time ─────────────────────────────────────────────────────────────────

@app.route('/api/time')
@login_required
def api_time():
    return jsonify({
        'utc': datetime.utcnow().isoformat(),
        'local': datetime.now().isoformat()
    })


# ─── God Mode (Admin) ─────────────────────────────────────────────────────────

@app.route('/api/admin/god-mode')
@admin_required
def api_god_mode():
    db = get_db()
    sessions = db.execute("""
        SELECT u.display_name, u.username, acs.tab, acs.last_seen,
               s.name as show_name, s.id as show_id
        FROM active_sessions acs
        JOIN users u ON acs.user_id = u.id
        JOIN shows s ON acs.show_id = s.id
        WHERE acs.last_seen > datetime('now', '-5 minutes')
        ORDER BY acs.last_seen DESC
    """).fetchall()
    users = db.execute("""
        SELECT id, display_name, username, role, last_login
        FROM users ORDER BY display_name
    """).fetchall()
    db.close()
    return jsonify({
        'sessions': [{
            'user': r['display_name'] or r['username'],
            'show': r['show_name'],
            'show_id': r['show_id'],
            'tab': r['tab'],
            'last_seen': r['last_seen'],
        } for r in sessions],
        'users': [{
            'id': u['id'],
            'name': u['display_name'] or u['username'],
            'username': u['username'],
            'role': u['role'],
            'last_login': u['last_login'] or '—',
        } for u in users],
    })


# ─── File Manager (Admin) ─────────────────────────────────────────────────────

@app.route('/api/admin/files')
@admin_required
def api_file_manager():
    db = get_db()
    files = []

    for r in db.execute("""
        SELECT sa.id, sa.filename, sa.mime_type, sa.file_size, sa.created_at,
               s.id as show_id, COALESCE(s.name, 'Deleted Show') as show_name,
               u.display_name, u.username
        FROM show_attachments sa
        LEFT JOIN shows s ON sa.show_id = s.id
        LEFT JOIN users u ON sa.uploaded_by = u.id
    """).fetchall():
        files.append({
            'id':           r['id'],
            'file_type':    'attachment',
            'show_id':      r['show_id'],
            'show_name':    r['show_name'],
            'filename':     r['filename'],
            'mime_type':    r['mime_type'],
            'file_size':    r['file_size'] or 0,
            'created_at':   r['created_at'],
            'uploader':     r['display_name'] or r['username'] or 'Unknown',
            'download_url': f"/shows/{r['show_id']}/attachments/{r['id']}/download" if r['show_id'] else None,
            'delete_url':   f"/shows/{r['show_id']}/attachments/{r['id']}/delete" if r['show_id'] else None,
        })

    for r in db.execute("""
        SELECT el.id, el.export_type, el.version, el.exported_at, el.s3_key,
               COALESCE(NULLIF(el.filename,''), el.export_type || '_v' || CAST(el.version AS TEXT) || '.pdf') as filename,
               CASE WHEN el.pdf_data IS NOT NULL THEN LENGTH(el.pdf_data) ELSE 0 END as file_size,
               s.id as show_id, COALESCE(s.name, 'Deleted Show') as show_name,
               u.display_name, u.username
        FROM export_log el
        LEFT JOIN shows s ON el.show_id = s.id
        LEFT JOIN users u ON el.exported_by = u.id
        WHERE el.pdf_data IS NOT NULL OR el.s3_key IS NOT NULL
    """).fetchall():
        files.append({
            'id':           r['id'],
            'file_type':    'export',
            'show_id':      r['show_id'],
            'show_name':    r['show_name'],
            'filename':     r['filename'],
            'mime_type':    'application/pdf',
            'file_size':    r['file_size'] or 0,
            'created_at':   r['exported_at'],
            'uploader':     r['display_name'] or r['username'] or 'Unknown',
            'download_url': f"/shows/{r['show_id']}/export/history/{r['id']}/download" if r['show_id'] else None,
            'delete_url':   None,
        })

    for r in db.execute("""
        SELECT er.id, er.pdf_filename, er.s3_key, er.created_at,
               CASE WHEN er.pdf_data IS NOT NULL THEN LENGTH(er.pdf_data) ELSE 0 END as file_size,
               s.id as show_id, COALESCE(s.name, 'Deleted Show') as show_name
        FROM show_external_rentals er
        LEFT JOIN shows s ON er.show_id = s.id
        WHERE er.pdf_data IS NOT NULL OR er.s3_key IS NOT NULL
    """).fetchall():
        files.append({
            'id':           r['id'],
            'file_type':    'rental',
            'show_id':      r['show_id'],
            'show_name':    r['show_name'],
            'filename':     r['pdf_filename'] or 'rental.pdf',
            'mime_type':    'application/pdf',
            'file_size':    r['file_size'] or 0,
            'created_at':   r['created_at'],
            'uploader':     '—',
            'download_url': f"/shows/{r['show_id']}/external-rentals/{r['id']}/pdf" if r['show_id'] else None,
            'delete_url':   None,
        })

    files.sort(key=lambda f: f['created_at'] or '', reverse=True)
    total_bytes = sum(f['file_size'] for f in files)
    db.close()
    return jsonify({'files': files, 'total_bytes': total_bytes})


# ─── Schedule Templates ───────────────────────────────────────────────────────

@app.route('/api/schedule-templates')
@login_required
def api_schedule_templates():
    db = get_db()
    templates = [dict(t) for t in db.execute(
        'SELECT * FROM schedule_templates ORDER BY sort_order, name'
    ).fetchall()]
    for t in templates:
        t['rows'] = [dict(r) for r in db.execute(
            'SELECT * FROM schedule_template_rows WHERE template_id=? ORDER BY sort_order',
            (t['id'],)
        ).fetchall()]
    db.close()
    return jsonify(templates)


@app.route('/api/schedule-templates/<int:tid>')
@login_required
def api_schedule_template(tid):
    db = get_db()
    t = db.execute('SELECT * FROM schedule_templates WHERE id=?', (tid,)).fetchone()
    if not t:
        db.close(); return jsonify({'error': 'Not found'}), 404
    rows = [dict(r) for r in db.execute(
        'SELECT * FROM schedule_template_rows WHERE template_id=? ORDER BY sort_order', (tid,)
    ).fetchall()]
    db.close()
    return jsonify({'id': t['id'], 'name': t['name'], 'rows': rows})


@app.route('/settings/schedule-templates/add', methods=['POST'])
@content_admin_required
def add_schedule_template():
    data = request.get_json(force=True) or {}
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'success': False, 'error': 'Name required.'}), 400
    db = get_db()
    max_o = db.execute('SELECT MAX(sort_order) FROM schedule_templates').fetchone()[0] or 0
    cur = db.execute('INSERT INTO schedule_templates (name, sort_order) VALUES (?,?)',
                     (name, max_o + 10))
    tid = cur.lastrowid
    for i, row in enumerate(data.get('rows', [])):
        db.execute("""INSERT INTO schedule_template_rows
                      (template_id, sort_order, start_time, end_time, description, notes)
                      VALUES (?,?,?,?,?,?)""",
                   (tid, i, row.get('start_time',''), row.get('end_time',''),
                    row.get('description',''), row.get('notes','')))
    log_audit(db, 'TEMPLATE_ADD', 'schedule_template', tid, detail=name)
    db.commit(); db.close()
    syslog_logger.info(f"TEMPLATE_ADD id={tid} name={name!r} by={session.get('username')}")
    return jsonify({'success': True, 'id': tid})


@app.route('/settings/schedule-templates/<int:tid>/edit', methods=['POST'])
@content_admin_required
def edit_schedule_template(tid):
    data = request.get_json(force=True) or {}
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'success': False, 'error': 'Name required.'}), 400
    db = get_db()
    db.execute('UPDATE schedule_templates SET name=? WHERE id=?', (name, tid))
    db.execute('DELETE FROM schedule_template_rows WHERE template_id=?', (tid,))
    for i, row in enumerate(data.get('rows', [])):
        db.execute("""INSERT INTO schedule_template_rows
                      (template_id, sort_order, start_time, end_time, description, notes)
                      VALUES (?,?,?,?,?,?)""",
                   (tid, i, row.get('start_time',''), row.get('end_time',''),
                    row.get('description',''), row.get('notes','')))
    log_audit(db, 'TEMPLATE_EDIT', 'schedule_template', tid, detail=name)
    db.commit(); db.close()
    syslog_logger.info(f"TEMPLATE_EDIT id={tid} name={name!r} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/schedule-templates/<int:tid>/delete', methods=['POST'])
@content_admin_required
def delete_schedule_template(tid):
    db = get_db()
    row = db.execute('SELECT name FROM schedule_templates WHERE id=?', (tid,)).fetchone()
    log_audit(db, 'TEMPLATE_DELETE', 'schedule_template', tid,
              detail=row['name'] if row else str(tid))
    db.execute('DELETE FROM schedule_templates WHERE id=?', (tid,))
    db.commit(); db.close()
    syslog_logger.info(f"TEMPLATE_DELETE id={tid} by={session.get('username')}")
    return jsonify({'success': True})


# ─── WiFi / Logo Settings ─────────────────────────────────────────────────────

@app.route('/settings/wifi', methods=['POST'])
@admin_required
def save_wifi_settings():
    data = request.get_json(force=True) or {}
    db = get_db()
    for key in ('wifi_network', 'wifi_password'):
        if key in data:
            db.execute('INSERT OR REPLACE INTO app_settings (key, value) VALUES (?,?)',
                       (key, data[key]))
    log_audit(db, 'SETTINGS_CHANGE', 'setting', None, detail='wifi')
    db.commit(); db.close()
    syslog_logger.info(f"SETTINGS_CHANGE detail=wifi by={session.get('username')}")
    return jsonify({'success': True})


def _get_distinct_venues(db):
    """Return sorted distinct venue names from the venue field options, shows, and advance_data."""
    venues = set()

    # Primary source: options defined on the 'venue' form field (dropdown options list)
    ff = db.execute(
        "SELECT options_json FROM form_fields WHERE field_key='venue' AND options_json IS NOT NULL"
    ).fetchone()
    if ff and ff['options_json']:
        try:
            opts = json.loads(ff['options_json'])
            for o in opts:
                v = (o or '').strip()
                if v and v != '—' and v != '-':
                    venues.add(v)
        except Exception:
            pass

    # Secondary sources: values actually saved on shows
    for row in db.execute(
        "SELECT DISTINCT venue FROM shows WHERE venue IS NOT NULL AND TRIM(venue) != ''"
    ).fetchall():
        venues.add(row[0].strip())

    for row in db.execute(
        "SELECT DISTINCT field_value FROM advance_data WHERE field_key='venue' AND field_value IS NOT NULL AND TRIM(field_value) != ''"
    ).fetchall():
        venues.add(row[0].strip())

    return sorted(venues)


@app.route('/settings/venues', methods=['GET', 'POST'])
@login_required
def venues_settings():
    db = get_db()
    venues = _get_distinct_venues(db)
    db.close()
    return jsonify({'venues': venues})


@app.route('/settings/arts-groups', methods=['GET'])
@login_required
def arts_groups_list():
    db = get_db()
    rows = db.execute(
        'SELECT id, name, sort_order FROM arts_groups ORDER BY sort_order, name'
    ).fetchall()
    db.close()
    return jsonify({'arts_groups': [dict(r) for r in rows]})


@app.route('/settings/arts-groups/add', methods=['POST'])
@content_admin_required
def arts_groups_add():
    data = request.get_json(force=True) or {}
    name = (data.get('name') or '').strip()
    if not name:
        return jsonify({'success': False, 'error': 'Name required.'}), 400
    db = get_db()
    max_order = db.execute('SELECT MAX(sort_order) FROM arts_groups').fetchone()[0] or 0
    try:
        cur = db.execute(
            'INSERT INTO arts_groups (name, sort_order) VALUES (?, ?)',
            (name, max_order + 10)
        )
        gid = cur.lastrowid
        log_audit(db, 'ARTS_GROUP_ADD', 'arts_group', gid, detail=name)
        db.commit()
        syslog_logger.info(f"ARTS_GROUP_ADD id={gid} name={name!r} by={session.get('username')}")
        return jsonify({'success': True, 'id': gid, 'name': name})
    except sqlite3.IntegrityError:
        return jsonify({'success': False, 'error': f'"{name}" already exists.'}), 400
    finally:
        db.close()


@app.route('/settings/arts-groups/<int:gid>/edit', methods=['POST'])
@content_admin_required
def arts_groups_edit(gid):
    data = request.get_json(force=True) or {}
    name = (data.get('name') or '').strip()
    if not name:
        return jsonify({'success': False, 'error': 'Name required.'}), 400
    db = get_db()
    before = _snapshot_row(db, 'arts_groups', gid)
    try:
        db.execute('UPDATE arts_groups SET name=? WHERE id=?', (name, gid))
        after = _snapshot_row(db, 'arts_groups', gid)
        log_audit(db, 'ARTS_GROUP_EDIT', 'arts_group', gid, detail=name,
                  before=before, after=after)
        db.commit()
        syslog_logger.info(f"ARTS_GROUP_EDIT id={gid} name={name!r} by={session.get('username')}")
        return jsonify({'success': True})
    except sqlite3.IntegrityError:
        return jsonify({'success': False, 'error': f'"{name}" already exists.'}), 400
    finally:
        db.close()


@app.route('/settings/arts-groups/<int:gid>/delete', methods=['POST'])
@content_admin_required
def arts_groups_delete(gid):
    db = get_db()
    before = _snapshot_row(db, 'arts_groups', gid)
    log_audit(db, 'ARTS_GROUP_DELETE', 'arts_group', gid, before=before)
    db.execute('DELETE FROM arts_groups WHERE id=?', (gid,))
    db.commit(); db.close()
    syslog_logger.info(f"ARTS_GROUP_DELETE id={gid} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/logo', methods=['POST'])
@admin_required
def save_logo():
    f = request.files.get('logo')
    if f and f.filename:
        import base64
        data = f.read()
        if len(data) > 2 * 1024 * 1024:
            return jsonify({'success': False, 'error': 'Logo too large (max 2 MB).'}), 413
        mime = f.content_type or 'image/png'
        # Only allow safe image MIME types (reject SVG — can contain JavaScript)
        _allowed_logo_mimes = ('image/png', 'image/jpeg', 'image/gif', 'image/webp')
        if mime not in _allowed_logo_mimes:
            return jsonify({'success': False,
                            'error': f'Unsupported image type. Allowed: PNG, JPEG, GIF, WebP.'}), 400
        b64 = base64.b64encode(data).decode()
        logo_data = f'data:{mime};base64,{b64}'
    else:
        data_uri = (request.get_json(force=True) or {}).get('logo_data', '')
        logo_data = data_uri

    db = get_db()
    db.execute('INSERT OR REPLACE INTO app_settings (key, value) VALUES (?,?)',
               ('logo_data', logo_data))
    log_audit(db, 'SETTINGS_CHANGE', 'setting', None, detail='logo_upload')
    db.commit(); db.close()
    syslog_logger.info(f"SETTINGS_CHANGE detail=logo_upload by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/logo/delete', methods=['POST'])
@admin_required
def delete_logo():
    db = get_db()
    db.execute("INSERT OR REPLACE INTO app_settings (key, value) VALUES ('logo_data', '')")
    log_audit(db, 'SETTINGS_CHANGE', 'setting', None, detail='logo_delete')
    db.commit(); db.close()
    syslog_logger.info(f"SETTINGS_CHANGE detail=logo_delete by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/upload-size', methods=['POST'])
@admin_required
def save_upload_size():
    data = request.get_json(force=True) or {}
    try:
        mb = int(data.get('upload_max_mb', 20))
        mb = max(1, min(mb, 500))
    except (ValueError, TypeError):
        mb = 20
    db = get_db()
    db.execute('INSERT OR REPLACE INTO app_settings (key, value) VALUES (?,?)',
               ('upload_max_mb', str(mb)))
    log_audit(db, 'SETTINGS_CHANGE', 'setting', None, detail=f'upload_max_mb={mb}')
    db.commit(); db.close()
    syslog_logger.info(f"SETTINGS_CHANGE detail=upload_max_mb={mb} by={session.get('username')}")
    return jsonify({'success': True, 'upload_max_mb': mb})


# ─── Public Show Page ─────────────────────────────────────────────────────────

@app.route('/public')
def public_shows():
    db = get_db()
    shows = db.execute("""
        SELECT id, name, show_date, show_time, venue, advance_version, schedule_version
        FROM shows WHERE status='active'
        ORDER BY show_date ASC NULLS LAST
    """).fetchall()
    db.close()
    return render_template('public.html', shows=shows)


@app.route('/public/shows/<int:show_id>/advance')
def public_advance_pdf(show_id):
    db = get_db()
    row = db.execute("""
        SELECT s3_key, pdf_data FROM export_log
        WHERE show_id=? AND export_type='advance'
        ORDER BY exported_at DESC LIMIT 1
    """, (show_id,)).fetchone()
    show = db.execute('SELECT * FROM shows WHERE id=? AND status="active"', (show_id,)).fetchone()
    db.close()
    if not show:
        abort(404)
    if row and row['s3_key']:
        try:
            data = s3_storage.download_file(row['s3_key'])
        except Exception as e:
            app.logger.error(f"S3 download failed for public advance PDF show_id={show_id}: {e}")
            abort(503)
        resp = make_response(data)
        resp.headers['Content-Type'] = 'application/pdf'
        resp.headers['Content-Disposition'] = f'inline; filename="Advance_{show_id}.pdf"'
        return resp
    if row and row['pdf_data']:
        resp = make_response(bytes(row['pdf_data']))
        resp.headers['Content-Type'] = 'application/pdf'
        resp.headers['Content-Disposition'] = f'inline; filename="Advance_{show_id}.pdf"'
        return resp
    abort(404)


@app.route('/public/shows/<int:show_id>/schedule')
def public_schedule_pdf(show_id):
    db = get_db()
    row = db.execute("""
        SELECT s3_key, pdf_data FROM export_log
        WHERE show_id=? AND export_type='schedule'
        ORDER BY exported_at DESC LIMIT 1
    """, (show_id,)).fetchone()
    show = db.execute('SELECT * FROM shows WHERE id=? AND status="active"', (show_id,)).fetchone()
    db.close()
    if not show:
        abort(404)
    if row and row['s3_key']:
        try:
            data = s3_storage.download_file(row['s3_key'])
        except Exception as e:
            app.logger.error(f"S3 download failed for public schedule PDF show_id={show_id}: {e}")
            abort(503)
        resp = make_response(data)
        resp.headers['Content-Type'] = 'application/pdf'
        resp.headers['Content-Disposition'] = f'inline; filename="Schedule_{show_id}.pdf"'
        return resp
    if row and row['pdf_data']:
        resp = make_response(bytes(row['pdf_data']))
        resp.headers['Content-Type'] = 'application/pdf'
        resp.headers['Content-Disposition'] = f'inline; filename="Schedule_{show_id}.pdf"'
        return resp
    abort(404)


# ─── Field Key Availability Check ─────────────────────────────────────────────

@app.route('/settings/form-fields/check-key')
@content_admin_required
def check_field_key():
    key = request.args.get('key', '').strip().lower().replace(' ', '_')
    exclude_id = request.args.get('exclude_id', type=int)
    if not key:
        return jsonify({'available': False, 'conflict': None})
    db = get_db()
    if exclude_id:
        row = db.execute(
            'SELECT id, label FROM form_fields WHERE field_key=? AND id!=?', (key, exclude_id)
        ).fetchone()
    else:
        row = db.execute(
            'SELECT id, label FROM form_fields WHERE field_key=?', (key,)
        ).fetchone()
    db.close()
    if row:
        return jsonify({'available': False, 'conflict': row['label']})
    return jsonify({'available': True})


# ─── Database Settings ─────────────────────────────────────────────────────────

@app.route('/settings/database', methods=['POST'])
@admin_required
def save_database_settings():
    data = request.get_json(force=True) or {}
    db_type = data.get('db_type', 'sqlite')

    # Only db_type is stored in the database. PG credentials live in db_config.ini.
    # Write to SQLite bootstrap directly so it works even when active DB is PostgreSQL.
    _sqlite_conn = sqlite3.connect(DATABASE)
    _sqlite_conn.execute(
        'INSERT OR REPLACE INTO app_settings (key, value) VALUES (?,?)', ('db_type', db_type)
    )
    _sqlite_conn.commit(); _sqlite_conn.close()

    db_adapter.clear_settings_cache()
    syslog_logger.info(f"SETTINGS_CHANGE key=database db_type={db_type} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/database/test', methods=['POST'])
@admin_required
def test_database_connection():
    data = request.get_json(force=True) or {}
    db_type = data.get('db_type', 'sqlite')

    if db_type == 'sqlite':
        if os.path.exists(DATABASE):
            return jsonify({'success': True, 'message': 'SQLite database found and accessible.'})
        return jsonify({'success': False, 'message': 'SQLite database not found. Run init_db.py first.'})

    if db_type == 'postgres':
        # Credentials come from db_config.ini, not the request
        settings = db_adapter.read_db_settings(DATABASE)
        if not settings.get('pg_host'):
            return jsonify({'success': False, 'message': 'db_config.ini not found or missing [postgresql] section. See db_config.ini.example.'})
        ok, err = db_adapter.test_postgres_connection(
            host=settings.get('pg_host', 'localhost'),
            port=settings.get('pg_port', 5432),
            dbname=settings.get('pg_dbname', '321theater'),
            user=settings.get('pg_user', ''),
            password=settings.get('pg_password', ''),
            app_schema=settings.get('pg_app_schema', 'theater321'),
            shared_schema=settings.get('pg_shared_schema', 'shared'),
        )
        if ok:
            return jsonify({'success': True, 'message': 'Connected to PostgreSQL successfully.'})
        app.logger.warning(f'PostgreSQL test failed: {err}')
        return jsonify({'success': False, 'message': err or 'PostgreSQL connection failed.'})

    return jsonify({'success': False, 'message': 'Unknown database type.'})


@app.route('/settings/database/migrate', methods=['POST'])
@admin_required
def migrate_database():
    """Migrate data from SQLite to PostgreSQL. Safe to run multiple times."""
    from init_db import migrate_sqlite_to_postgres

    settings = db_adapter.read_db_settings(DATABASE)
    if settings.get('db_type') != 'postgres':
        return jsonify({'success': False, 'error': 'Database type must be PostgreSQL to migrate.'}), 400

    try:
        stats = migrate_sqlite_to_postgres(DATABASE, settings)
    except Exception as e:
        app.logger.error(f'Database migration failed: {e}')
        return jsonify({'success': False, 'error': 'Migration failed. Check server logs.'}), 500

    if 'error' in stats:
        return jsonify({'success': False, 'error': stats['error']}), 500

    total_copied = sum(v.get('copied', 0) for v in stats.values() if isinstance(v, dict))
    total_skipped = sum(v.get('skipped', 0) for v in stats.values() if isinstance(v, dict))

    syslog_logger.info(f"DB_MIGRATE copied={total_copied} skipped={total_skipped} by={session.get('username')}")
    return jsonify({
        'success': True,
        'stats': stats,
        'total_copied': total_copied,
        'total_skipped': total_skipped,
    })


# ─── AI / Ollama Settings ──────────────────────────────────────────────────────

def _is_blocked_host(hostname):
    """Return True if hostname resolves to a cloud metadata or link-local address."""
    import ipaddress
    if not hostname:
        return True
    if hostname in ('169.254.169.254', 'metadata.google.internal'):
        return True
    try:
        ip = ipaddress.ip_address(hostname)
        if ip.is_link_local or ip.is_reserved:
            return True
        # Block metadata IP range (169.254.x.x)
        if ip.is_private and str(ip).startswith('169.254.'):
            return True
    except ValueError:
        pass  # DNS name, not an IP literal — not blocked
    return False


def _validate_ollama_url(url):
    """Validate that an Ollama URL is safe (no SSRF to internal/metadata endpoints)."""
    from urllib.parse import urlparse
    import ipaddress
    try:
        parsed = urlparse(url)
        if parsed.scheme not in ('http', 'https'):
            return False
        hostname = parsed.hostname or ''
        # Allow localhost / 127.x (typical Ollama install)
        if hostname in ('localhost', '127.0.0.1', '::1') or hostname.startswith('127.'):
            return True
        # Block cloud metadata, link-local, and private ranges
        try:
            ip = ipaddress.ip_address(hostname)
            if ip.is_private or ip.is_link_local or ip.is_loopback or ip.is_reserved:
                return False
        except ValueError:
            pass  # hostname is a DNS name, not an IP — allow it
        if _is_blocked_host(hostname):
            return False
        return True
    except Exception:
        return False


@app.route('/settings/ai', methods=['POST'])
@admin_required
def save_ai_settings():
    data = request.get_json(force=True) or {}
    db = get_db()
    for key in ('ollama_enabled', 'ollama_url', 'ollama_model', 'ai_max_sessions', 'ai_system_prompt'):
        if key in data:
            db.execute('INSERT OR REPLACE INTO app_settings (key, value) VALUES (?,?)',
                       (key, str(data[key])))
    db.commit(); db.close()
    syslog_logger.info(f"SETTINGS_CHANGE key=ai by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/ai/test', methods=['POST'])
@admin_required
def test_ai_connection():
    data = request.get_json(force=True) or {}
    url = data.get('ollama_url', 'http://localhost:11434').rstrip('/')
    model = data.get('ollama_model', 'llama3.2')

    # SSRF protection — only allow http(s) to non-internal hosts (except localhost for Ollama)
    if not _validate_ollama_url(url):
        return jsonify({'success': False, 'message': 'Invalid Ollama URL. Only http/https to localhost or non-internal hosts allowed.'})

    import urllib.request
    import urllib.error
    try:
        req = urllib.request.Request(f'{url}/api/tags', method='GET')
        with urllib.request.urlopen(req, timeout=5) as resp:
            body = json.loads(resp.read())
            models = [m['name'] for m in body.get('models', [])]
            if model in models or any(m.startswith(model.split(':')[0]) for m in models):
                return jsonify({'success': True, 'message': f'Connected. Model "{model}" available.', 'models': models})
            return jsonify({'success': True, 'message': f'Connected, but model "{model}" not found. Available: {", ".join(models[:5])}', 'models': models})
    except urllib.error.URLError as e:
        app.logger.warning(f'Ollama connection failed: {e}')
        return jsonify({'success': False, 'message': 'Cannot reach Ollama. Check URL and ensure Ollama is running.'})
    except Exception as e:
        app.logger.warning(f'Ollama test error: {e}')
        return jsonify({'success': False, 'message': 'Ollama connection test failed.'})


# ─── SMTP / PDF Email Settings ────────────────────────────────────────────────

@app.route('/settings/smtp', methods=['POST'])
@admin_required
def save_smtp_settings():
    data = request.get_json(force=True) or {}
    db = get_db()
    for key in ('smtp_host', 'smtp_port', 'smtp_user', 'smtp_pass',
                'smtp_from', 'smtp_tls'):
        if key in data:
            db.execute('INSERT OR REPLACE INTO app_settings (key, value) VALUES (?,?)',
                       (key, str(data[key])))
    log_audit(db, 'SETTINGS_CHANGE', 'setting', None,
              after={k: v for k, v in data.items() if 'pass' not in k},
              detail='smtp_settings')
    db.commit(); db.close()
    syslog_logger.info(f"SETTINGS_CHANGE key=smtp by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/smtp/test', methods=['POST'])
@admin_required
def test_smtp_connection():
    import smtplib
    data = request.get_json(force=True) or {}
    host     = data.get('smtp_host', '')
    port     = int(data.get('smtp_port') or 587)
    user     = data.get('smtp_user', '')
    password = data.get('smtp_pass', '')
    use_tls  = data.get('smtp_tls', '1') not in ('0', 'false', 'False', '')
    from_addr = data.get('smtp_from') or user
    to_addr   = data.get('test_to') or user
    if not host:
        return jsonify({'success': False, 'message': 'SMTP host is required.'})
    # Block SSRF — prevent connecting to cloud metadata or link-local addresses
    if _is_blocked_host(host):
        return jsonify({'success': False, 'message': 'Invalid SMTP host.'})
    # Validate email addresses to prevent header injection / open relay abuse
    import re
    _email_re = re.compile(r'^[^@\s\r\n]+@[^@\s\r\n]+\.[^@\s\r\n]+$')
    if from_addr and not _email_re.match(from_addr):
        return jsonify({'success': False, 'message': 'Invalid from address.'})
    if to_addr and not _email_re.match(to_addr):
        return jsonify({'success': False, 'message': 'Invalid test recipient address.'})
    try:
        if use_tls:
            server = smtplib.SMTP(host, port, timeout=10)
            server.ehlo(); server.starttls(); server.ehlo()
        else:
            server = smtplib.SMTP_SSL(host, port, timeout=10)
        if user and password:
            server.login(user, password)
        if to_addr:
            from email.mime.text import MIMEText
            msg = MIMEText('3·2·1→Theater SMTP test — connection successful.')
            msg['Subject'] = '3·2·1→Theater SMTP Test'
            msg['From'] = from_addr
            msg['To']   = to_addr
            server.sendmail(from_addr, [to_addr], msg.as_string())
            server.quit()
            syslog_logger.info(f"SMTP_TEST to={to_addr} by={session.get('username')}")
            return jsonify({'success': True, 'message': f'Test email sent to {to_addr}.'})
        server.quit()
        return jsonify({'success': True, 'message': 'Connected successfully (no test email sent).'})
    except Exception as e:
        app.logger.warning(f'SMTP test failed: {e}')
        return jsonify({'success': False, 'message': 'SMTP connection failed. Check host, port, and credentials.'})


@app.route('/settings/email-provider', methods=['POST'])
@admin_required
def save_email_provider_settings():
    data = request.get_json(force=True) or {}
    db = get_db()
    provider = data.get('email_provider', 'smtp')
    if provider not in ('smtp', 'direct'):
        provider = 'smtp'
    db.execute('INSERT OR REPLACE INTO app_settings (key, value) VALUES (?,?)',
               ('email_provider', provider))
    # Save direct send settings
    for key in ('smtp_from', 'direct_ehlo_hostname', 'direct_display_name'):
        if key in data:
            db.execute('INSERT OR REPLACE INTO app_settings (key, value) VALUES (?,?)',
                       (key, str(data[key])))
    log_audit(db, 'SETTINGS_CHANGE', 'setting', None,
              after={'email_provider': provider}, detail='email_provider')
    db.commit(); db.close()
    syslog_logger.info(f"SETTINGS_CHANGE key=email_provider value={provider} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/email/test', methods=['POST'])
@admin_required
def test_email_provider():
    """Send a test email through the currently configured email provider."""
    data = request.get_json(force=True) or {}
    to_addr = data.get('test_to', '').strip()
    if not to_addr:
        return jsonify({'success': False, 'message': 'Test recipient address is required.'})
    success, message = _send_email(
        subject='3·2·1→THEATER Email Test',
        recipients=[to_addr],
        body_text='This is a test email from 3·2·1→THEATER to verify your email configuration.',
        error_context={'pdf_type': 'test', 'triggered_by': session.get('username')},
    )
    return jsonify({'success': success, 'message': message})


# ─── Email Error Log ──────────────────────────────────────────────────────────

@app.route('/settings/email-errors', methods=['GET'])
@admin_required
def email_errors_list():
    """Return paginated email send failures.

    Query: ?status=unresolved|all (default unresolved), ?limit=N (default 200)."""
    status = (request.args.get('status') or 'unresolved').lower()
    try:
        limit = max(1, min(int(request.args.get('limit') or 200), 1000))
    except ValueError:
        limit = 200
    db = get_db()
    sql = """
        SELECT e.id, e.sent_at, e.recipient, e.subject, e.error_msg, e.smtp_code,
               e.pdf_type, e.show_id, e.triggered_by, e.resolved, e.resolved_at,
               s.name AS show_name,
               u.username AS resolved_by_name
        FROM email_send_errors e
        LEFT JOIN shows s ON s.id = e.show_id
        LEFT JOIN users u ON u.id = e.resolved_by
    """
    params = []
    if status == 'unresolved':
        sql += ' WHERE COALESCE(e.resolved, 0) = 0'
    sql += ' ORDER BY e.sent_at DESC, e.id DESC LIMIT ?'
    params.append(limit)
    rows = db.execute(sql, params).fetchall()
    counts = db.execute("""
        SELECT
            SUM(CASE WHEN COALESCE(resolved,0)=0 THEN 1 ELSE 0 END) AS unresolved,
            COUNT(*) AS total
        FROM email_send_errors
    """).fetchone()
    db.close()
    return jsonify({
        'errors':     [_normalize_row_dates(dict(r)) for r in rows],
        'unresolved': (counts['unresolved'] if counts else 0) or 0,
        'total':      (counts['total']      if counts else 0) or 0,
    })


@app.route('/settings/email-errors/<int:eid>/resolve', methods=['POST'])
@admin_required
def email_errors_resolve(eid):
    """Mark a single error as reviewed/handled."""
    db = get_db()
    row = db.execute('SELECT id FROM email_send_errors WHERE id=?', (eid,)).fetchone()
    if not row:
        db.close()
        return jsonify({'success': False, 'error': 'Not found.'}), 404
    db.execute(
        'UPDATE email_send_errors SET resolved=1, resolved_at=CURRENT_TIMESTAMP, resolved_by=? WHERE id=?',
        (session['user_id'], eid),
    )
    db.commit()
    db.close()
    return jsonify({'success': True})


@app.route('/settings/email-errors/resolve-all', methods=['POST'])
@admin_required
def email_errors_resolve_all():
    """Mark every currently-unresolved error as reviewed in one click."""
    db = get_db()
    cur = db.execute(
        'UPDATE email_send_errors SET resolved=1, resolved_at=CURRENT_TIMESTAMP, resolved_by=? '
        'WHERE COALESCE(resolved,0)=0',
        (session['user_id'],),
    )
    affected = cur.rowcount or 0
    db.commit()
    db.close()
    syslog_logger.info(f'EMAIL_ERRORS_RESOLVE_ALL count={affected} by={session.get("username")}')
    return jsonify({'success': True, 'count': affected})


@app.route('/settings/email-errors/clear', methods=['POST'])
@admin_required
def email_errors_clear():
    """Permanently delete all rows. Use sparingly — review first.

    Body: {scope: 'resolved' (default) | 'all'}.
    'resolved' is the safe default — only removes rows already marked
    handled. Pass 'all' to wipe everything (e.g. before re-running tests).
    """
    data = request.get_json(silent=True) or {}
    scope = (data.get('scope') or 'resolved').lower()
    db = get_db()
    if scope == 'all':
        cur = db.execute('DELETE FROM email_send_errors')
    else:
        cur = db.execute('DELETE FROM email_send_errors WHERE COALESCE(resolved,0)=1')
    affected = cur.rowcount or 0
    db.commit()
    db.close()
    syslog_logger.info(f'EMAIL_ERRORS_CLEAR scope={scope} count={affected} by={session.get("username")}')
    return jsonify({'success': True, 'count': affected, 'scope': scope})


@app.route('/settings/pdf-emails', methods=['POST'])
@admin_required
def save_pdf_email_settings():
    data = request.get_json(force=True) or {}
    db = get_db()
    keys = ('pdf_email_send_hour',
            'advance_email_enabled',   'advance_email_days_before',
            'schedule_email_enabled_1','schedule_email_days_1',
            'schedule_email_enabled_2','schedule_email_days_2')
    for key in keys:
        if key in data:
            db.execute('INSERT OR REPLACE INTO app_settings (key, value) VALUES (?,?)',
                       (key, str(data[key])))
    log_audit(db, 'SETTINGS_CHANGE', 'setting', None, after=data, detail='pdf_email_settings')
    db.commit(); db.close()
    syslog_logger.info(f"SETTINGS_CHANGE key=pdf_emails by={session.get('username')}")
    return jsonify({'success': True})


# ─── Cluster (multi-server heartbeat & leader election) ──────────────────────

@app.route('/api/cluster/peers', methods=['GET'])
@admin_required
def api_cluster_peers():
    """Return cluster status (live peers, current leader, this instance)."""
    return jsonify(get_cluster_status())


@app.route('/settings/cluster', methods=['POST'])
@admin_required
def save_cluster_settings():
    """Save cluster heartbeat settings. Restarts heartbeat thread if needed."""
    data = request.get_json(force=True) or {}
    valid_force = {'auto', 'always', 'never'}
    keys = ('cluster_heartbeat_enabled',
            'cluster_heartbeat_interval_sec',
            'cluster_peer_timeout_sec',
            'cluster_force_leader')
    db = get_db()
    for key in keys:
        if key not in data:
            continue
        val = str(data[key])
        if key == 'cluster_force_leader' and val not in valid_force:
            val = 'auto'
        if key in ('cluster_heartbeat_interval_sec', 'cluster_peer_timeout_sec'):
            try:
                ival = max(2, int(val))
                val = str(ival)
            except ValueError:
                continue
        if key == 'cluster_heartbeat_enabled':
            val = '1' if val in ('1', 'true', 'on') else '0'
        db.execute('INSERT OR REPLACE INTO app_settings (key, value) VALUES (?,?)',
                   (key, val))
    log_audit(db, 'SETTINGS_CHANGE', 'setting', None, after=data, detail='cluster_settings')
    db.commit(); db.close()
    db_adapter.clear_settings_cache()
    # Invalidate leader cache so changes take effect on next call
    with _cluster_lock:
        _leader_cache['at'] = 0.0
    # Restart heartbeat thread so a freshly-enabled cluster starts pinging
    # immediately (or a freshly-disabled one will see the flag on next loop tick)
    stop_cluster_heartbeat()
    if get_app_setting('cluster_heartbeat_enabled', '1') in ('1', 'true'):
        start_cluster_heartbeat()
    syslog_logger.info(f"SETTINGS_CHANGE key=cluster by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/contacts/<int:cid>/recipient', methods=['POST'])
@admin_required
def toggle_contact_recipient(cid):
    data = request.get_json(force=True) or {}
    email_type = data.get('email_type', 'report')
    val = 1 if data.get('recipient') else 0
    allowed_cols = {'report': 'report_recipient', 'advance': 'advance_recipient',
                    'production': 'production_recipient',
                    'postnotes': 'postnotes_recipient'}
    col = allowed_cols.get(email_type, 'report_recipient')
    db = get_db()
    db.execute(f'UPDATE contacts SET {col}=? WHERE id=?', (val, cid))
    log_audit(db, 'CONTACT_RECIPIENT_TOGGLE', 'contact', cid,
              detail=f"type={email_type} recipient={'yes' if val else 'no'}")
    db.commit(); db.close()
    syslog_logger.info(f"CONTACT_RECIPIENT_TOGGLE id={cid} type={email_type} recipient={'yes' if val else 'no'} by={session.get('username')}")
    return jsonify({'success': True})


# ─── AI Document Extraction ────────────────────────────────────────────────────

@app.route('/shows/<int:show_id>/ai-extract', methods=['POST'])
@login_required
def ai_extract(show_id):
    """Extract form field values from an uploaded document using Ollama."""
    ai_sid, slot_error = _claim_ai_session(show_id)
    if slot_error:
        return jsonify({'success': False, 'error': slot_error}), 429
    try:
        return _ai_extract_impl(show_id)
    except Exception as e:
        app.logger.exception("ai_extract unhandled error")
        return jsonify({'success': False, 'error': f'Server error: {e}'}), 500
    finally:
        _release_ai_session(ai_sid)

def _ai_extract_impl(show_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403

    # Check Ollama is enabled
    ollama_enabled = get_app_setting('ollama_enabled', '0')
    if ollama_enabled != '1':
        return jsonify({'success': False, 'error': 'AI extraction is not enabled. Enable it in Settings → AI.'}), 400

    ollama_url = get_app_setting('ollama_url', 'http://localhost:11434').rstrip('/')
    ollama_model = get_app_setting('ollama_model', 'llama3.2')

    # Get document text
    doc_text = ''
    attachment_id = request.form.get('attachment_id', type=int)
    uploaded_file = request.files.get('document')

    if attachment_id:
        db = get_db()
        row = db.execute(
            'SELECT file_data, mime_type, filename FROM show_attachments WHERE id=? AND show_id=?',
            (attachment_id, show_id)
        ).fetchone()
        db.close()
        if not row:
            return jsonify({'success': False, 'error': 'Attachment not found.'}), 404
        file_bytes = bytes(row['file_data'])
        mime = row['mime_type']
        fname = row['filename']
    elif uploaded_file and uploaded_file.filename:
        file_bytes = uploaded_file.read()
        mime = uploaded_file.content_type or 'application/octet-stream'
        fname = uploaded_file.filename
    else:
        return jsonify({'success': False, 'error': 'No document provided.'}), 400

    # Extract text from document
    try:
        if mime == 'application/pdf' or fname.lower().endswith('.pdf'):
            try:
                import pdfplumber
                from io import BytesIO as _BytesIO
                with pdfplumber.open(_BytesIO(file_bytes)) as pdf:
                    doc_text = '\n'.join(
                        page.extract_text() or '' for page in pdf.pages
                    )
            except ImportError:
                return jsonify({'success': False, 'error': 'pdfplumber not installed. Run: pip install pdfplumber'}), 500
        elif fname.lower().endswith('.docx') or mime in (
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',):
            try:
                import docx as _docx
                from io import BytesIO as _BytesIO
                document = _docx.Document(_BytesIO(file_bytes))
                doc_text = '\n'.join(p.text for p in document.paragraphs if p.text.strip())
            except ImportError:
                return jsonify({'success': False, 'error': 'python-docx not installed. Run: pip install python-docx'}), 500
        elif fname.lower().endswith('.doc'):
            return jsonify({'success': False, 'error': 'Legacy .doc format is not supported. Please save as .docx and re-upload.'}), 400
        elif fname.lower().endswith('.rtf'):
            try:
                from striprtf.striprtf import rtf_to_text as _rtf_to_text
                doc_text = _rtf_to_text(file_bytes.decode('utf-8', errors='replace'))
            except ImportError:
                return jsonify({'success': False, 'error': 'striprtf not installed. Run: pip install striprtf'}), 500
        elif fname.lower().endswith(('.xlsx', '.xls')) or mime in (
                'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                'application/vnd.ms-excel'):
            try:
                from io import BytesIO as _BytesIO
                if fname.lower().endswith('.xls'):
                    import xlrd as _xlrd
                    wb = _xlrd.open_workbook(file_contents=file_bytes)
                    rows = []
                    for sheet in wb.sheets():
                        for i in range(sheet.nrows):
                            rows.append('\t'.join(str(sheet.cell_value(i, j)) for j in range(sheet.ncols)))
                    doc_text = '\n'.join(rows)
                else:
                    import openpyxl as _openpyxl
                    wb = _openpyxl.load_workbook(_BytesIO(file_bytes), data_only=True)
                    rows = []
                    for sheet in wb.worksheets:
                        for row in sheet.iter_rows(values_only=True):
                            line = '\t'.join(str(c) if c is not None else '' for c in row)
                            if line.strip():
                                rows.append(line)
                    doc_text = '\n'.join(rows)
            except ImportError as ie:
                pkg = 'xlrd' if '.xls' in fname.lower() and not fname.lower().endswith('.xlsx') else 'openpyxl'
                return jsonify({'success': False, 'error': f'{pkg} not installed. Run: pip install {pkg}'}), 500
        else:
            # Plain text (.txt, .csv and others)
            doc_text = file_bytes.decode('utf-8', errors='replace')
    except Exception as e:
        return jsonify({'success': False, 'error': f'Could not extract text: {e}'}), 500

    if not doc_text.strip():
        return jsonify({'success': False, 'error': 'No readable text found in document.'}), 400

    # Load form fields with ai_hint
    db = get_db()
    field_rows = db.execute(
        'SELECT field_key, label, ai_hint FROM form_fields WHERE ai_hint IS NOT NULL AND ai_hint != \'\''
    ).fetchall()
    db.close()

    if not field_rows:
        return jsonify({'success': False, 'error': 'No fields have AI hints configured. Add hints in Settings → Form Fields.'}), 400

    field_map = {r['field_key']: {'label': r['label'], 'hint': r['ai_hint']} for r in field_rows}
    field_schema = {k: v['hint'] for k, v in field_map.items()}

    # Build prompt — custom system prompt prefix overrides the default if set
    _default_system = (
        'You are extracting information from a document to populate a show advance form for a live event. '
        'The document may be a technical rider, production spec sheet, artist contract, or similar. '
        'Return ONLY valid JSON — no explanation, no markdown, just the JSON object. '
        'Use null for fields where the information is not found in the document.'
    )
    _custom_system = get_app_setting('ai_system_prompt', '').strip()
    system_prefix = _custom_system if _custom_system else _default_system
    prompt = (
        f'{system_prefix}\n\n'
        f'Fields to extract (key: description): {json.dumps(field_schema)}\n\n'
        f'Document text:\n{doc_text[:8000]}'
    )

    # Call Ollama (with SSRF validation)
    if not _validate_ollama_url(ollama_url):
        return jsonify({'success': False, 'error': 'Invalid Ollama URL.'}), 400
    try:
        import urllib.request as _urlreq
        payload = json.dumps({
            'model': ollama_model,
            'messages': [{'role': 'user', 'content': prompt}],
            'format': 'json',
            'stream': True,  # Stream tokens; timeout applies per-chunk, not total
        }).encode()
        req = _urlreq.Request(
            f'{ollama_url}/api/chat',
            data=payload,
            headers={'Content-Type': 'application/json'},
            method='POST',
        )
        raw_content = ''
        with _urlreq.urlopen(req, timeout=60) as resp:
            for line in resp:
                line = line.strip()
                if not line:
                    continue
                chunk = json.loads(line)
                raw_content += chunk.get('message', {}).get('content', '')
                if chunk.get('done'):
                    break
    except Exception as e:
        app.logger.error(f'Ollama request failed: {e}')
        return jsonify({'success': False, 'error': 'AI extraction request failed. Check Ollama connection.'}), 500

    # Parse JSON response
    try:
        # Strip markdown code fences if model added them
        clean = raw_content.strip()
        if clean.startswith('```'):
            clean = '\n'.join(clean.split('\n')[1:])
            if clean.endswith('```'):
                clean = clean[:-3]
        extracted = json.loads(clean)
    except (json.JSONDecodeError, ValueError):
        app.logger.warning(f'AI returned invalid JSON: {raw_content[:200]}')
        return jsonify({'success': False, 'error': 'AI returned an unparseable response. Try again or use a different model.'}), 500

    # Build suggestions (only non-null values)
    suggestions = {}
    for field_key, value in extracted.items():
        if value is not None and str(value).strip() and field_key in field_map:
            suggestions[field_key] = {
                'value': str(value).strip(),
                'label': field_map[field_key]['label'],
            }

    syslog_logger.info(
        f"AI_EXTRACT show_id={show_id} document={fname} "
        f"fields_found={len(suggestions)} model={ollama_model} "
        f"by={session.get('username')}"
    )
    return jsonify({
        'success': True,
        'suggestions': suggestions,
        'document': fname,
        'model': ollama_model,
        'field_count': len(field_rows),
    })


# ─── Job Positions & Position Categories ─────────────────────────────────────

@app.route('/api/job-positions')
@login_required
def api_job_positions():
    db = get_db()
    rows = db.execute("""
        SELECT jp.id, jp.category_id, pc.name as category_name, jp.name, jp.venue, jp.sort_order
        FROM job_positions jp
        LEFT JOIN position_categories pc ON jp.category_id = pc.id
        ORDER BY pc.sort_order, jp.sort_order, jp.id
    """).fetchall()
    db.close()
    return jsonify([dict(r) for r in rows])


@app.route('/api/position-categories')
@login_required
def api_position_categories():
    db = get_db()
    cats = db.execute('SELECT * FROM position_categories ORDER BY sort_order, id').fetchall()
    db.close()
    return jsonify([dict(c) for c in cats])


@app.route('/settings/position-categories/add', methods=['POST'])
@scheduler_required
def add_position_category():
    data = request.get_json(force=True) or {}
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'success': False, 'error': 'Name is required.'}), 400
    db = get_db()
    max_order = db.execute('SELECT MAX(sort_order) FROM position_categories').fetchone()[0] or 0
    cur = db.execute(
        'INSERT INTO position_categories (name, sort_order) VALUES (?, ?)',
        (name, max_order + 10)
    )
    cid = cur.lastrowid
    log_audit_change(db, 'POSITION_CATEGORY_ADD', 'position_category', cid, detail=name,
                     table='position_categories')
    db.commit()
    db.close()
    syslog_logger.info(f"POSITION_CATEGORY_ADD id={cid} name={name!r} by={session.get('username')}")
    return jsonify({'success': True, 'id': cid, 'name': name})


@app.route('/settings/position-categories/<int:cid>/edit', methods=['POST'])
@scheduler_required
def edit_position_category(cid):
    data = request.get_json(force=True) or {}
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'success': False, 'error': 'Name is required.'}), 400
    db = get_db()
    before = _snapshot_row(db, 'position_categories', cid)
    db.execute('UPDATE position_categories SET name=? WHERE id=?', (name, cid))
    after = _snapshot_row(db, 'position_categories', cid)
    log_audit(db, 'POSITION_CATEGORY_EDIT', 'position_category', cid, detail=name,
              before=before, after=after)
    db.commit()
    db.close()
    syslog_logger.info(f"POSITION_CATEGORY_EDIT id={cid} name={name!r} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/position-categories/<int:cid>/delete', methods=['POST'])
@scheduler_required
def delete_position_category(cid):
    db = get_db()
    # Null out category_id on positions in this category
    before = _snapshot_row(db, 'position_categories', cid)
    db.execute('UPDATE job_positions SET category_id=NULL WHERE category_id=?', (cid,))
    db.execute('DELETE FROM position_categories WHERE id=?', (cid,))
    log_audit(db, 'POSITION_CATEGORY_DELETE', 'position_category', cid,
              detail=before['name'] if before else str(cid), before=before)
    db.commit()
    db.close()
    syslog_logger.info(f"POSITION_CATEGORY_DELETE id={cid} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/job-positions/add', methods=['POST'])
@scheduler_required
def add_job_position():
    data = request.get_json(force=True) or {}
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'success': False, 'error': 'Name is required.'}), 400
    category_id = data.get('category_id') or None
    venue = (data.get('venue') or '').strip() or None
    db = get_db()
    max_order = db.execute('SELECT MAX(sort_order) FROM job_positions').fetchone()[0] or 0
    override_rate = data.get('override_rate')
    override_rate = float(override_rate) if override_rate not in (None, '') else None
    cur = db.execute(
        'INSERT INTO job_positions (category_id, name, venue, override_rate, sort_order) VALUES (?, ?, ?, ?, ?)',
        (category_id, name, venue, override_rate, max_order + 10)
    )
    pid = cur.lastrowid
    log_audit_change(db, 'JOB_POSITION_ADD', 'job_position', pid, detail=name,
                     table='job_positions')
    db.commit()
    db.close()
    syslog_logger.info(f"JOB_POSITION_ADD id={pid} name={name!r} category_id={category_id} by={session.get('username')}")
    return jsonify({'success': True, 'id': pid, 'name': name})


@app.route('/settings/job-positions/<int:pid>/edit', methods=['POST'])
@scheduler_required
def edit_job_position(pid):
    data = request.get_json(force=True) or {}
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'success': False, 'error': 'Name is required.'}), 400
    category_id = data.get('category_id') or None
    venue = (data.get('venue') or '').strip() or None
    override_rate = data.get('override_rate')
    override_rate = float(override_rate) if override_rate not in (None, '') else None
    db = get_db()
    before = _snapshot_row(db, 'job_positions', pid)
    db.execute(
        'UPDATE job_positions SET name=?, category_id=?, venue=?, override_rate=? WHERE id=?',
        (name, category_id, venue, override_rate, pid)
    )
    after = _snapshot_row(db, 'job_positions', pid)
    log_audit(db, 'JOB_POSITION_EDIT', 'job_position', pid, detail=name,
              before=before, after=after)
    db.commit()
    db.close()
    syslog_logger.info(f"JOB_POSITION_EDIT id={pid} name={name!r} category_id={category_id} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/job-positions/<int:pid>/delete', methods=['POST'])
@scheduler_required
def delete_job_position(pid):
    db = get_db()
    before = _snapshot_row(db, 'job_positions', pid)
    log_audit(db, 'JOB_POSITION_DELETE', 'job_position', pid,
              detail=before['name'] if before else str(pid), before=before)
    db.execute('DELETE FROM job_positions WHERE id=?', (pid,))
    db.commit()
    db.close()
    syslog_logger.info(f"JOB_POSITION_DELETE id={pid} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/job-positions/reorder', methods=['POST'])
@scheduler_required
def reorder_job_positions():
    data = request.get_json(force=True) or {}
    position_ids = data.get('position_ids', [])
    db = get_db()
    for i, pid in enumerate(position_ids):
        db.execute('UPDATE job_positions SET sort_order=? WHERE id=?', (i * 10, pid))
    db.commit()
    db.close()
    return jsonify({'success': True})


# ─── Labor Requests (per show) ────────────────────────────────────────────────

@app.route('/shows/<int:show_id>/labor-requests', methods=['GET'])
@login_required
def get_labor_requests(show_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    db = get_db()
    rows = db.execute("""
        SELECT lr.*, jp.name as position_name,
               cm.name as scheduled_crew_name
        FROM labor_requests lr
        LEFT JOIN job_positions jp ON lr.position_id = jp.id
        LEFT JOIN crew_members cm ON lr.scheduled_crew_member_id = cm.id
        WHERE lr.show_id = ?
        ORDER BY lr.sort_order, lr.id
    """, (show_id,)).fetchall()
    db.close()
    return jsonify([_normalize_row_dates(dict(r)) for r in rows])


@app.route('/shows/<int:show_id>/labor-notes', methods=['PUT'])
@login_required
def save_show_labor_notes(show_id):
    """Save show-level notes that travel with every labor request — used by
    requesters and schedulers to share split-shift instructions, task lists,
    and other nuance across the staffing flow."""
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    if session.get('is_restricted') or session.get('is_readonly'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403
    data = request.get_json(force=True) or {}
    notes = (data.get('labor_notes') or '').strip()
    db = get_db()
    db.execute('UPDATE shows SET labor_notes=? WHERE id=?', (notes, show_id))
    log_audit(db, 'SHOW_LABOR_NOTES_EDIT', 'show', show_id, show_id=show_id)
    db.commit()
    db.close()
    return jsonify({'success': True})


@app.route('/shows/<int:show_id>/labor-requests', methods=['POST'])
@login_required
def add_labor_request(show_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    if session.get('is_restricted'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403
    data = request.get_json(force=True) or {}
    db = get_db()
    max_order = db.execute(
        'SELECT MAX(sort_order) FROM labor_requests WHERE show_id=?', (show_id,)
    ).fetchone()[0] or 0
    cur = db.execute("""
        INSERT INTO labor_requests (show_id, position_id, work_date, in_time, out_time,
                                    break_start, break_end, break2_start, break2_end,
                                    requested_name, notes, sort_order)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (show_id,
          data.get('position_id') or None,
          data.get('work_date') or None,
          _normalize_perf_time(data.get('in_time', '')),
          _normalize_perf_time(data.get('out_time', '')),
          _normalize_perf_time(data.get('break_start', '')),
          _normalize_perf_time(data.get('break_end', '')),
          _normalize_perf_time(data.get('break2_start', '')),
          _normalize_perf_time(data.get('break2_end', '')),
          data.get('requested_name', ''),
          (data.get('notes') or '').strip(),
          max_order + 10))
    rid = cur.lastrowid
    log_audit(db, 'LABOR_REQUEST_ADD', 'labor_request', rid, show_id=show_id,
              detail=data.get('requested_name', '') or f'position_id={data.get("position_id")}')
    db.commit()
    db.close()
    syslog_logger.info(f"LABOR_REQUEST_ADD show_id={show_id} id={rid} by={session.get('username')}")
    return jsonify({'success': True, 'id': rid})


@app.route('/shows/<int:show_id>/labor-requests/<int:rid>', methods=['PUT'])
@login_required
def update_labor_request(show_id, rid):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    if session.get('is_restricted'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403
    data = request.get_json(force=True) or {}
    db = get_db()
    db.execute("""
        UPDATE labor_requests
        SET position_id=?, work_date=?, in_time=?, out_time=?,
            break_start=?, break_end=?, break2_start=?, break2_end=?,
            requested_name=?, notes=?
        WHERE id=? AND show_id=?
    """, (data.get('position_id') or None,
          data.get('work_date') or None,
          _normalize_perf_time(data.get('in_time', '')),
          _normalize_perf_time(data.get('out_time', '')),
          _normalize_perf_time(data.get('break_start', '')),
          _normalize_perf_time(data.get('break_end', '')),
          _normalize_perf_time(data.get('break2_start', '')),
          _normalize_perf_time(data.get('break2_end', '')),
          data.get('requested_name', ''),
          (data.get('notes') or '').strip(),
          rid, show_id))
    log_audit(db, 'LABOR_REQUEST_EDIT', 'labor_request', rid, show_id=show_id)
    db.commit()
    db.close()
    syslog_logger.info(f"LABOR_REQUEST_EDIT show_id={show_id} id={rid} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/shows/<int:show_id>/labor-requests/<int:rid>', methods=['DELETE'])
@login_required
def delete_labor_request(show_id, rid):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    if session.get('is_restricted'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403
    db = get_db()
    db.execute('DELETE FROM labor_requests WHERE id=? AND show_id=?', (rid, show_id))
    log_audit(db, 'LABOR_REQUEST_DELETE', 'labor_request', rid, show_id=show_id)
    db.commit()
    db.close()
    syslog_logger.info(f"LABOR_REQUEST_DELETE show_id={show_id} id={rid} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/shows/<int:show_id>/labor-requests/reorder', methods=['POST'])
@login_required
def reorder_labor_requests(show_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    if session.get('is_restricted'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403
    data = request.get_json(force=True) or {}
    request_ids = data.get('request_ids', [])
    db = get_db()
    for i, rid in enumerate(request_ids):
        db.execute(
            'UPDATE labor_requests SET sort_order=? WHERE id=? AND show_id=?',
            (i * 10, rid, show_id)
        )
    db.commit()
    db.close()
    return jsonify({'success': True})


# ─── Labor Scheduler ─────────────────────────────────────────────────────────

def _calc_labor_cost_for_show(db, show_id):
    """Return labor line items and total cost for a show."""
    rows = db.execute("""
        SELECT lr.id, lr.work_date, lr.in_time, lr.out_time,
               lr.break_start, lr.break_end, lr.break2_start, lr.break2_end,
               lr.is_scheduled, lr.scheduled_crew_member_id,
               jp.name as position_name, jp.override_rate,
               cm.name as tech_name,
               prl.hourly_rate as level_rate, prl.name as level_name
        FROM labor_requests lr
        LEFT JOIN job_positions jp ON jp.id = lr.position_id
        LEFT JOIN crew_members cm ON cm.id = lr.scheduled_crew_member_id
        LEFT JOIN pay_rate_levels prl ON prl.id = cm.rate_level_id
        WHERE lr.show_id = ?
        ORDER BY lr.work_date, lr.sort_order
    """, (show_id,)).fetchall()

    lines = []
    total = 0.0
    for r in rows:
        hours = _calc_hours(r['in_time'], r['out_time'],
                            r['break_start'], r['break_end'],
                            r['break2_start'], r['break2_end'])
        rate = r['override_rate'] if r['override_rate'] is not None else (r['level_rate'] or 0)
        cost = round(hours * rate, 2)
        total += cost
        lines.append({
            'id': r['id'],
            'work_date': r['work_date'],
            'position_name': r['position_name'] or '',
            'tech_name': r['tech_name'] or r['scheduled_crew_member_id'] or 'Unassigned',
            'in_time': r['in_time'],
            'out_time': r['out_time'],
            'hours': round(hours, 2),
            'hourly_rate': rate,
            'line_total': cost,
            'level_name': r['level_name'] or '',
            'is_scheduled': bool(r['is_scheduled']),
        })
    return lines, round(total, 2)


def _calc_hours(in_time, out_time, break_start=None, break_end=None,
                break2_start=None, break2_end=None):
    """Calculate hours between in/out times, minus up to two unpaid breaks.

    Accepts HH:MM or HHMM format on every field — values are normalised
    via _normalize_perf_time before parsing. Either break can be omitted.
    """
    in_time  = _normalize_perf_time(in_time)
    out_time = _normalize_perf_time(out_time)
    if not in_time or not out_time:
        return 0.0
    try:
        from datetime import datetime as _dt
        fmt = '%H:%M'
        t_in  = _dt.strptime(in_time[:5],  fmt)
        t_out = _dt.strptime(out_time[:5], fmt)
        if t_out <= t_in:
            t_out = t_out.replace(day=t_out.day + 1)
        hours = (t_out - t_in).total_seconds() / 3600
        for bs_raw, be_raw in (
            (break_start,  break_end),
            (break2_start, break2_end),
        ):
            bs = _normalize_perf_time(bs_raw)
            be = _normalize_perf_time(be_raw)
            if bs and be:
                t_bs = _dt.strptime(bs[:5], fmt)
                t_be = _dt.strptime(be[:5], fmt)
                if t_be > t_bs:
                    hours -= (t_be - t_bs).total_seconds() / 3600
        return max(0.0, hours)
    except Exception:
        return 0.0


@app.route('/shows/<int:show_id>/labor-cost')
@login_required
def show_labor_cost(show_id):
    """Return labor cost breakdown for a show."""
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'error': 'Access denied'}), 403
    db = get_db()
    lines, total = _calc_labor_cost_for_show(db, show_id)
    db.close()
    return jsonify({'lines': lines, 'total': total})


@app.route('/labor-scheduler')
@scheduler_required
def labor_scheduler():
    """Cross-show labor scheduler view."""
    return render_template('labor_scheduler.html', user=get_current_user())


@app.route('/labor-overview')
@login_required
def labor_overview():
    """Read-only week-at-a-time view of every scheduled labor line — both
    show labor and overhead/project crew. Anyone logged in can view; nothing
    here is editable."""
    from datetime import date, datetime, timedelta

    raw_start = (request.args.get('start') or '').strip()
    try:
        anchor = date.fromisoformat(raw_start) if raw_start else date.today()
    except ValueError:
        anchor = date.today()
    # Snap to Monday of the anchor's week (weekday(): Mon=0..Sun=6)
    week_start = anchor - timedelta(days=anchor.weekday())
    week_end   = week_start + timedelta(days=6)
    prev_week  = (week_start - timedelta(days=7)).isoformat()
    next_week  = (week_start + timedelta(days=7)).isoformat()
    today_iso  = date.today().isoformat()

    db = get_db()

    # ── Show labor — joined with the show, position, and the scheduled tech.
    show_rows = db.execute("""
        SELECT
            COALESCE(lr.work_date, s.show_date) AS work_date,
            s.id   AS show_id,
            s.name AS show_name,
            s.venue AS venue,
            jp.name AS position_name,
            lr.in_time, lr.out_time,
            lr.break_start, lr.break_end, lr.break2_start, lr.break2_end,
            lr.is_scheduled,
            cm.name AS scheduled_tech,
            lr.requested_name AS requested_tech,
            s.labor_notes AS group_notes,
            lr.sort_order AS sort_order,
            ad.field_value AS pm_name
        FROM labor_requests lr
        JOIN shows s ON s.id = lr.show_id
        LEFT JOIN job_positions jp ON jp.id = lr.position_id
        LEFT JOIN crew_members cm ON cm.id = lr.scheduled_crew_member_id
        LEFT JOIN advance_data ad
               ON ad.show_id = s.id AND ad.field_key = 'production_manager'
        WHERE COALESCE(s.status, 'active') != 'archived'
          AND COALESCE(lr.work_date, s.show_date) BETWEEN ? AND ?
        ORDER BY work_date, s.name, lr.sort_order, lr.id
    """, (week_start.isoformat(), week_end.isoformat())).fetchall()

    # ── Overhead / project labor — group provides the contact name; project
    # provides a fallback contact and a colour for visual grouping.
    oh_rows = db.execute("""
        SELECT
            COALESCE(r.work_date, g.work_date) AS work_date,
            g.id AS group_id,
            COALESCE(p.name, g.name, 'General') AS project_name,
            p.client_name AS client_name,
            p.color AS project_color,
            jp.name AS position_name,
            r.in_time, r.out_time,
            r.break_start, r.break_end, r.break2_start, r.break2_end,
            r.is_scheduled,
            cm.name AS scheduled_tech,
            r.requested_name AS requested_tech,
            COALESCE(NULLIF(g.project_notes, ''), p.project_notes, '') AS group_notes,
            r.sort_order AS sort_order,
            COALESCE(NULLIF(g.contact_name, ''), p.contact_name) AS contact_name
        FROM overhead_labor_requests r
        JOIN overhead_labor_groups g ON g.id = r.group_id
        LEFT JOIN overhead_projects p ON p.id = g.project_id
        LEFT JOIN job_positions jp ON jp.id = r.position_id
        LEFT JOIN crew_members cm ON cm.id = r.scheduled_crew_member_id
        WHERE COALESCE(r.work_date, g.work_date) BETWEEN ? AND ?
        ORDER BY work_date, project_name, r.sort_order, r.id
    """, (week_start.isoformat(), week_end.isoformat())).fetchall()
    db.close()

    def _iso(v):
        if v is None: return ''
        return str(v)[:10] if not isinstance(v, str) else v[:10]

    # Build a dict keyed by ISO date so the template can render every day
    # of the week, even ones with no entries.
    by_day = {(week_start + timedelta(days=i)).isoformat(): [] for i in range(7)}

    for r in show_rows:
        d = _iso(r['work_date'])
        if d not in by_day:
            continue
        by_day[d].append({
            'kind': 'show',
            'group_key': f"show:{r['show_id']}",
            'name': r['show_name'] or '—',
            'venue': r['venue'] or '',
            'position': r['position_name'] or '—',
            'in_time': r['in_time'] or '',
            'break_start': r['break_start'] or '',
            'break_end': r['break_end'] or '',
            'break2_start': r['break2_start'] or '',
            'break2_end': r['break2_end'] or '',
            'out_time': r['out_time'] or '',
            'tech': r['scheduled_tech'] or (r['requested_tech'] or ''),
            'is_scheduled': bool(r['is_scheduled']),
            'pm': r['pm_name'] or '',
            'group_notes': (r['group_notes'] or '').strip() if r['group_notes'] else '',
            'color': '',
            'show_id': r['show_id'],
        })

    for r in oh_rows:
        d = _iso(r['work_date'])
        if d not in by_day:
            continue
        by_day[d].append({
            'kind': 'overhead',
            'group_key': f"oh:{r['group_id']}",
            'name': r['project_name'] or 'General',
            'venue': r['client_name'] or '',
            'position': r['position_name'] or '—',
            'in_time': r['in_time'] or '',
            'break_start': r['break_start'] or '',
            'break_end': r['break_end'] or '',
            'break2_start': r['break2_start'] or '',
            'break2_end': r['break2_end'] or '',
            'out_time': r['out_time'] or '',
            'tech': r['scheduled_tech'] or (r['requested_tech'] or ''),
            'is_scheduled': bool(r['is_scheduled']),
            'pm': r['contact_name'] or '',
            'group_notes': (r['group_notes'] or '').strip() if r['group_notes'] else '',
            'color': r['project_color'] or '',
            'group_id': r['group_id'],
        })

    # Build the per-day list in calendar order with day labels
    day_names = ['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday']
    days = []
    for i in range(7):
        d = (week_start + timedelta(days=i))
        d_iso = d.isoformat()
        rows = by_day[d_iso]
        # Group rows by show / project so notes render once per group
        rows.sort(key=lambda r: (r['name'].lower(), (r['in_time'] or '99:99')))
        seen_groups = set()
        for row in rows:
            key = row['group_key']
            row['is_group_first'] = key not in seen_groups
            seen_groups.add(key)
        days.append({
            'name': day_names[i],
            'date': d,
            'iso': d_iso,
            'pretty': d.strftime('%A, %B %-d, %Y'),
            'rows': rows,
            'count': len(rows),
        })

    return render_template('labor_overview.html',
                           user=get_current_user(),
                           days=days,
                           week_start=week_start,
                           week_end=week_end,
                           prev_week=prev_week,
                           next_week=next_week,
                           today_iso=today_iso)


@app.route('/api/labor-scheduler', methods=['GET'])
@scheduler_required
def api_labor_scheduler_list():
    """Return labor requests whose work_date falls in [from, to] (inclusive),
    grouped by show on the client. Rows with NULL work_date fall back to
    the show's show_date so legacy rows still appear."""
    date_from = (request.args.get('from') or '').strip()
    date_to   = (request.args.get('to') or '').strip()
    if not date_from or not date_to:
        return jsonify({'error': 'from and to dates required'}), 400

    db = get_db()
    accessible = get_accessible_shows(session['user_id'])

    sql = """
        SELECT lr.id, lr.show_id, lr.position_id, lr.work_date,
               lr.in_time, lr.out_time,
               lr.break_start, lr.break_end, lr.break2_start, lr.break2_end,
               lr.requested_name, lr.notes, lr.is_scheduled,
               lr.scheduled_crew_member_id, lr.sort_order,
               jp.name as position_name,
               pc.name as category_name,
               pc.sort_order as category_sort,
               cm.name as scheduled_crew_name,
               s.name as show_name,
               s.venue as show_venue,
               s.show_date as show_date,
               s.status as show_status,
               s.labor_notes as show_labor_notes
        FROM labor_requests lr
        JOIN shows s ON lr.show_id = s.id
        LEFT JOIN job_positions jp ON lr.position_id = jp.id
        LEFT JOIN position_categories pc ON jp.category_id = pc.id
        LEFT JOIN crew_members cm ON lr.scheduled_crew_member_id = cm.id
        WHERE s.status != 'archived'
          AND COALESCE(lr.work_date, s.show_date) BETWEEN ? AND ?
    """
    params = [date_from, date_to]
    if accessible is not None:
        if not accessible:
            db.close()
            return jsonify({'shows': []})
        placeholders = ','.join(['?'] * len(accessible))
        sql += f' AND lr.show_id IN ({placeholders})'
        params.extend(accessible)
    sql += ' ORDER BY COALESCE(lr.work_date, s.show_date), s.name, pc.sort_order, jp.sort_order, lr.sort_order, lr.id'

    rows = db.execute(sql, params).fetchall()

    # Optional: include shows even when they have zero labor entries — used by
    # the scheduler when adding a show / picking an existing show that doesn't
    # yet have any labor lines.
    include_raw = (request.args.get('include_show_ids') or '').strip()
    extra_ids = []
    if include_raw:
        for tok in include_raw.split(','):
            tok = tok.strip()
            if tok.isdigit():
                extra_ids.append(int(tok))

    # Group by show
    shows = {}
    order = []
    for r in rows:
        rd = _normalize_row_dates(dict(r))
        sid = rd['show_id']
        if sid not in shows:
            shows[sid] = {
                'show_id': sid,
                'show_name': rd['show_name'],
                'show_venue': rd['show_venue'],
                'show_date': rd['show_date'],
                'show_status': rd['show_status'],
                'show_labor_notes': rd.get('show_labor_notes') or '',
                'requests': [],
            }
            order.append(sid)
        shows[sid]['requests'].append(rd)

    # Pull in any forced-include shows that didn't surface above because they
    # have zero labor rows. They become empty sections in the scheduler so the
    # user can add the first labor line.
    if extra_ids:
        missing = [sid for sid in extra_ids if sid not in shows]
        if missing:
            if accessible is not None:
                missing = [sid for sid in missing if sid in accessible]
            if missing:
                ph = ','.join(['?'] * len(missing))
                extra_rows = db.execute(
                    f"SELECT id, name, venue, show_date, status, labor_notes "
                    f"FROM shows WHERE id IN ({ph}) AND status != 'archived'",
                    missing
                ).fetchall()
                for er in extra_rows:
                    erd = _normalize_row_dates(dict(er))
                    sid = erd['id']
                    shows[sid] = {
                        'show_id':          sid,
                        'show_name':        erd['name'],
                        'show_venue':       erd['venue'],
                        'show_date':        erd['show_date'],
                        'show_status':      erd['status'],
                        'show_labor_notes': erd.get('labor_notes') or '',
                        'requests':         [],
                    }
                    order.insert(0, sid)  # surface at top so they're easy to find

    # ── Overhead & Project Crew labor (not tied to any show) ─────────────────
    # Pulled in here so the labor scheduler sees a single unified to-do list.
    oh_rows = db.execute("""
        SELECT r.id, r.group_id, r.work_date, r.position_id,
               r.in_time, r.out_time,
               r.break_start, r.break_end, r.break2_start, r.break2_end,
               r.requested_name, r.notes, r.is_scheduled,
               r.scheduled_crew_member_id, r.sort_order,
               jp.name AS position_name,
               pc.name AS category_name,
               pc.sort_order AS category_sort,
               cm.name AS scheduled_crew_name,
               g.name AS group_name,
               COALESCE(p.name, g.name) AS project_name,
               p.client_name AS project_client,
               COALESCE(NULLIF(g.project_notes, ''), p.project_notes, '') AS group_notes
        FROM overhead_labor_requests r
        JOIN overhead_labor_groups g ON g.id = r.group_id
        LEFT JOIN overhead_projects p ON p.id = g.project_id
        LEFT JOIN job_positions jp ON jp.id = r.position_id
        LEFT JOIN position_categories pc ON pc.id = jp.category_id
        LEFT JOIN crew_members cm ON cm.id = r.scheduled_crew_member_id
        WHERE r.work_date BETWEEN ? AND ?
        ORDER BY r.work_date, g.sort_order, g.id, pc.sort_order, jp.sort_order,
                 r.sort_order, r.id
    """, (date_from, date_to)).fetchall()
    db.close()

    overhead_groups = {}
    oh_order = []
    for r in oh_rows:
        rd = _normalize_row_dates(dict(r))
        gid = rd['group_id']
        if gid not in overhead_groups:
            overhead_groups[gid] = {
                'group_id':       gid,
                'work_date':      rd['work_date'],
                'project_name':   rd['project_name'],
                'project_client': rd['project_client'],
                'group_name':     rd['group_name'],
                'group_notes':    rd.get('group_notes') or '',
                'requests':       [],
            }
            oh_order.append(gid)
        overhead_groups[gid]['requests'].append(rd)

    return jsonify({
        'shows': [shows[sid] for sid in order],
        'overhead_groups': [overhead_groups[gid] for gid in oh_order],
    })


@app.route('/api/labor-scheduler/<int:rid>', methods=['PUT'])
@scheduler_required
def api_labor_scheduler_update(rid):
    """Update only the scheduling fields on a labor request."""
    if session.get('is_readonly'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403
    data = request.get_json(force=True) or {}
    db = get_db()
    row = db.execute(
        'SELECT show_id FROM labor_requests WHERE id=?', (rid,)
    ).fetchone()
    if not row:
        db.close()
        return jsonify({'success': False, 'error': 'Not found.'}), 404
    show_id = row['show_id']
    if not can_access_show(session['user_id'], show_id):
        db.close()
        return jsonify({'success': False, 'error': 'Access denied.'}), 403

    updates = []
    params = []
    detail_parts = []
    if 'is_scheduled' in data:
        updates.append('is_scheduled=?')
        params.append(1 if data.get('is_scheduled') else 0)
        detail_parts.append(f"is_scheduled={1 if data.get('is_scheduled') else 0}")
    if 'scheduled_crew_member_id' in data:
        cmid = data.get('scheduled_crew_member_id')
        cmid = int(cmid) if cmid else None
        updates.append('scheduled_crew_member_id=?')
        params.append(cmid)
        detail_parts.append(f"crew_id={cmid}")
    for field in ('in_time', 'out_time', 'break_start', 'break_end',
                  'break2_start', 'break2_end'):
        if field in data:
            updates.append(f'{field}=?')
            params.append(_normalize_perf_time((data[field] or '').strip()))
            detail_parts.append(f"{field}={data[field]}")
    if 'notes' in data:
        updates.append('notes=?')
        params.append((data['notes'] or '').strip())
        detail_parts.append('notes=updated')

    if not updates:
        db.close()
        return jsonify({'success': False, 'error': 'No changes.'}), 400

    updates.append('scheduled_by=?')
    params.append(session['user_id'])
    updates.append('scheduled_at=CURRENT_TIMESTAMP')

    params.append(rid)
    db.execute(
        f"UPDATE labor_requests SET {', '.join(updates)} WHERE id=?",
        params,
    )
    log_audit(db, 'LABOR_SCHEDULED', 'labor_request', rid, show_id=show_id,
              detail='; '.join(detail_parts))
    db.commit()

    # Return the refreshed row so the client can render (esp. scheduled_crew_name)
    refreshed = db.execute("""
        SELECT lr.id, lr.is_scheduled, lr.scheduled_crew_member_id,
               cm.name as scheduled_crew_name
        FROM labor_requests lr
        LEFT JOIN crew_members cm ON lr.scheduled_crew_member_id = cm.id
        WHERE lr.id = ?
    """, (rid,)).fetchone()
    db.close()
    syslog_logger.info(
        f"LABOR_SCHEDULED id={rid} show_id={show_id} by={session.get('username')}"
    )
    return jsonify({'success': True, 'row': _normalize_row_dates(dict(refreshed)) if refreshed else None})


@app.route('/api/labor-scheduler/<int:rid>', methods=['DELETE'])
@scheduler_required
def api_labor_scheduler_delete(rid):
    """Delete a labor request from the scheduler view."""
    db = get_db()
    row = db.execute('SELECT show_id FROM labor_requests WHERE id=?', (rid,)).fetchone()
    if not row:
        db.close()
        return jsonify({'success': False, 'error': 'Not found.'}), 404
    if not can_access_show(session['user_id'], row['show_id']):
        db.close()
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    db.execute('DELETE FROM labor_requests WHERE id=?', (rid,))
    log_audit(db, 'LABOR_REQUEST_DELETE', 'labor_request', rid, show_id=row['show_id'])
    db.commit()
    db.close()
    syslog_logger.info(f"LABOR_REQUEST_DELETE (scheduler) id={rid} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/api/labor-scheduler/add', methods=['POST'])
@scheduler_required
def api_labor_scheduler_add():
    """Add a new labor request from the scheduler view."""
    data = request.get_json(force=True) or {}
    show_id = data.get('show_id')
    if not show_id:
        return jsonify({'success': False, 'error': 'show_id required'}), 400
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'success': False, 'error': 'Access denied.'}), 403
    db = get_db()
    max_order = db.execute(
        'SELECT MAX(sort_order) FROM labor_requests WHERE show_id=?', (show_id,)
    ).fetchone()[0] or 0
    cur = db.execute("""
        INSERT INTO labor_requests (show_id, position_id, work_date, in_time, out_time,
                                    break_start, break_end, break2_start, break2_end,
                                    requested_name, sort_order)
        VALUES (?,?,?,?,?,?,?,?,?,?,?)
    """, (show_id,
          data.get('position_id') or None,
          data.get('work_date') or None,
          _normalize_perf_time(data.get('in_time', '')),
          _normalize_perf_time(data.get('out_time', '')),
          _normalize_perf_time(data.get('break_start', '')),
          _normalize_perf_time(data.get('break_end', '')),
          _normalize_perf_time(data.get('break2_start', '')),
          _normalize_perf_time(data.get('break2_end', '')),
          data.get('requested_name', ''), max_order + 10))
    rid = cur.lastrowid
    log_audit(db, 'LABOR_REQUEST_ADD', 'labor_request', rid, show_id=show_id, detail='via scheduler')
    db.commit()
    row = db.execute("""
        SELECT lr.id, lr.show_id, lr.position_id, lr.work_date, lr.in_time, lr.out_time,
               lr.break_start, lr.break_end, lr.break2_start, lr.break2_end,
               lr.requested_name, lr.notes, lr.is_scheduled,
               lr.scheduled_crew_member_id, lr.sort_order,
               jp.name as position_name, pc.name as category_name,
               cm.name as scheduled_crew_name
        FROM labor_requests lr
        LEFT JOIN job_positions jp ON lr.position_id = jp.id
        LEFT JOIN position_categories pc ON jp.category_id = pc.id
        LEFT JOIN crew_members cm ON lr.scheduled_crew_member_id = cm.id
        WHERE lr.id=?
    """, (rid,)).fetchone()
    db.close()
    syslog_logger.info(f"LABOR_REQUEST_ADD (scheduler) show_id={show_id} id={rid} by={session.get('username')}")
    return jsonify({'success': True, 'row': _normalize_row_dates(dict(row)) if row else {'id': rid}})


@app.route('/api/labor-scheduler/create-show', methods=['POST'])
@scheduler_required
def api_labor_scheduler_create_show():
    """Create a barebones show from the labor scheduler. The PM is expected
    to come back later via the normal advance flow and fill in the rest."""
    if session.get('is_readonly'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403
    data = request.get_json(force=True) or {}
    name = (data.get('name') or '').strip()
    if not name:
        return jsonify({'success': False, 'error': 'Show name is required.'}), 400
    if len(name) > 200:
        return jsonify({'success': False, 'error': 'Show name is too long.'}), 400
    show_date = (data.get('show_date') or '').strip() or None
    show_time = _normalize_perf_time(data.get('show_time', ''))
    venue     = (data.get('venue') or '').strip() or "Judson's Live"
    if len(venue) > 120:
        return jsonify({'success': False, 'error': 'Venue name is too long.'}), 400

    db = get_db()
    cur = db.execute("""
        INSERT INTO shows (name, show_date, show_time, venue, created_by)
        VALUES (?, ?, ?, ?, ?)
    """, (name, show_date, show_time, venue, session['user_id']))
    show_id = cur.lastrowid

    for key, val in [('show_name', name), ('show_date', show_date or ''),
                     ('show_time', show_time), ('venue', venue)]:
        if val:
            db.execute(
                "INSERT OR REPLACE INTO advance_data (show_id, field_key, field_value) "
                "VALUES (?, ?, ?)",
                (show_id, key, val)
            )

    if show_date:
        db.execute(
            "INSERT INTO show_performances (show_id, perf_date, perf_time, sort_order) "
            "VALUES (?, ?, ?, 0)",
            (show_id, show_date, show_time)
        )

    log_audit(db, 'SHOW_CREATE', 'show', show_id, show_id=show_id,
              after={'name': name, 'show_date': show_date, 'venue': venue,
                     'via': 'labor_scheduler'})
    db.commit()
    db.close()
    syslog_logger.info(
        f"SHOW_CREATE (scheduler) show_id={show_id} name={name!r} venue={venue!r} "
        f"by={session.get('username')}"
    )
    return jsonify({
        'success':   True,
        'show_id':   show_id,
        'show_name': name,
        'show_date': show_date or '',
        'show_time': show_time,
        'venue':     venue,
    })


@app.route('/api/labor-scheduler/shows-without-labor', methods=['GET'])
@scheduler_required
def api_labor_scheduler_shows_without_labor():
    """List active shows that have zero labor_requests rows so the scheduler
    can pull one up and start adding labor. Optional 'from'/'to' ISO dates
    constrain by show_date (inclusive)."""
    date_from = (request.args.get('from') or '').strip()
    date_to   = (request.args.get('to') or '').strip()

    db = get_db()
    accessible = get_accessible_shows(session['user_id'])

    sql = """
        SELECT s.id, s.name, s.show_date, s.show_time, s.venue
        FROM shows s
        LEFT JOIN labor_requests lr ON lr.show_id = s.id
        WHERE COALESCE(s.status, 'active') != 'archived'
          AND lr.id IS NULL
    """
    params = []
    if date_from and date_to:
        # Either show_date in range, OR no show_date set at all (so brand-new
        # shows that haven't picked a date yet still appear).
        sql += " AND (s.show_date IS NULL OR s.show_date BETWEEN ? AND ?)"
        params.extend([date_from, date_to])
    if accessible is not None:
        if not accessible:
            db.close()
            return jsonify({'shows': []})
        ph = ','.join(['?'] * len(accessible))
        sql += f' AND s.id IN ({ph})'
        params.extend(accessible)
    sql += ' ORDER BY s.show_date IS NULL, s.show_date, s.name'

    rows = db.execute(sql, params).fetchall()
    db.close()
    return jsonify({'shows': [_normalize_row_dates(dict(r)) for r in rows]})


# ─── Overhead & Project Crew ─────────────────────────────────────────────────
#
# Standalone labor scheduling that isn't tied to a show — running list of days,
# each with one or more sub-groups (per project / contact) and labor requests
# beneath them. Recurring templates can auto-generate requests for repeating
# weekly needs (e.g. 2 overhead techs Mon–Fri).
#
# Access: any logged-in user may view; read-only users cannot mutate.

def _overhead_write_check():
    """Return None if the caller may mutate, otherwise a (json, status) tuple."""
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': 'Not signed in.'}), 403
    if session.get('is_readonly'):
        return jsonify({'success': False, 'error': 'Read-only access.'}), 403
    return None


def _parse_iso_date(s):
    """Parse 'YYYY-MM-DD' (or empty) → datetime.date or None. Never raises."""
    if not s:
        return None
    try:
        return datetime.strptime(str(s)[:10], '%Y-%m-%d').date()
    except Exception:
        return None


# ── Shared overhead helpers ──────────────────────────────────────────────────

# SELECT used in three places — keep it here so the column list / joins stay
# in sync between create/update endpoints and the list view.
_OVERHEAD_REQUEST_JOIN_SELECT = """
    SELECT r.*,
           jp.name AS position_name,
           pc.name AS category_name,
           cm.name AS scheduled_crew_name,
           prl.name AS scheduled_level_name,
           prl.hourly_rate AS scheduled_level_rate
    FROM overhead_labor_requests r
    LEFT JOIN job_positions jp ON jp.id = r.position_id
    LEFT JOIN position_categories pc ON pc.id = jp.category_id
    LEFT JOIN crew_members cm ON cm.id = r.scheduled_crew_member_id
    LEFT JOIN pay_rate_levels prl ON prl.id = cm.rate_level_id
"""

_OVERHEAD_GROUP_JOIN_SELECT = """
    SELECT g.*,
           p.name AS project_name,
           p.client_name,
           p.color,
           p.contact_name  AS project_contact_name,
           p.contact_email AS project_contact_email,
           p.contact_phone AS project_contact_phone,
           p.project_notes AS project_default_notes
    FROM overhead_labor_groups g
    LEFT JOIN overhead_projects p ON p.id = g.project_id
"""


def _fetch_overhead_request(db, rid):
    """Fetch a single labor request row joined to position + crew + rate level."""
    return db.execute(_OVERHEAD_REQUEST_JOIN_SELECT + ' WHERE r.id=?', (rid,)).fetchone()


def _fetch_overhead_group(db, gid):
    """Fetch a single sub-group row joined to its project (if any)."""
    return db.execute(_OVERHEAD_GROUP_JOIN_SELECT + ' WHERE g.id=?', (gid,)).fetchone()


def _annotate_request_metrics(req_dict, *, rate_keys=('scheduled_level_rate',)):
    """Attach hours_planned / hours_actual / effective_rate / cost_* fields
    in-place. Snapshot rate is used when present so historical cost is stable.

    `rate_keys` lists the dict keys to fall back to when `pay_rate_snapshot`
    is None (different SELECT aliases use different names)."""
    h_p = _calc_hours(
        req_dict.get('in_time'), req_dict.get('out_time'),
        req_dict.get('break_start'), req_dict.get('break_end'),
        req_dict.get('break2_start'), req_dict.get('break2_end'))
    h_a = _calc_hours(
        req_dict.get('actual_in_time'), req_dict.get('actual_out_time'),
        req_dict.get('actual_break_start'), req_dict.get('actual_break_end'),
        req_dict.get('actual_break2_start'), req_dict.get('actual_break2_end'))
    rate = req_dict.get('pay_rate_snapshot')
    if rate is None:
        for k in rate_keys:
            if req_dict.get(k) is not None:
                rate = req_dict.get(k)
                break
    rate = float(rate or 0)
    req_dict['hours_planned'] = round(h_p, 4)
    req_dict['hours_actual']  = round(h_a, 4)
    req_dict['effective_rate'] = rate
    req_dict['cost_planned'] = round(h_p * rate, 2)
    req_dict['cost_actual']  = round(h_a * rate, 2)
    return req_dict


def _parse_dow_csv(csv):
    """Parse a CSV of weekday digits (0=Sun..6=Sat) → sorted set of valid ints."""
    out = set()
    for piece in (csv or '').split(','):
        piece = piece.strip()
        if not piece:
            continue
        try:
            n = int(piece)
        except ValueError:
            continue
        if 0 <= n <= 6:
            out.add(n)
    return out


def _max_sort_order(db, table, where_sql='', where_params=()):
    """Return next sort_order value (current max + 10). Centralizes the pattern."""
    sql = f'SELECT COALESCE(MAX(sort_order), 0) FROM {table}'
    if where_sql:
        sql += ' WHERE ' + where_sql
    return (db.execute(sql, where_params).fetchone()[0] or 0) + 10


def _overhead_log(action, **kw):
    """Single syslog entry point for overhead crew mutations.
    Keeps the line format consistent with the rest of the app
    (LABOR_REQUEST_ADD show_id=… id=… by=…)."""
    parts = [f'{k}={v}' for k, v in kw.items() if v is not None]
    parts.append(f"by={session.get('username')}")
    syslog_logger.info(f"{action} {' '.join(parts)}")


@app.route('/overhead-crew')
@login_required
def overhead_crew_page():
    """Render the Overhead & Project Crew running schedule."""
    return render_template(
        'overhead_crew.html',
        user=get_current_user(),
        can_edit=not session.get('is_readonly'),
        can_schedule=_can_schedule_labor(),
    )


@app.route('/api/overhead-crew/list', methods=['GET'])
@login_required
def api_overhead_list():
    """Return all groups + requests whose work_date falls in [from, to] (inclusive),
    grouped by date then group, sorted closest-date first."""
    date_from = (request.args.get('from') or '').strip()
    date_to   = (request.args.get('to') or '').strip()
    if not date_from or not date_to:
        return jsonify({'error': 'from and to dates required'}), 400

    db = get_db()
    groups = db.execute(
        _OVERHEAD_GROUP_JOIN_SELECT +
        ' WHERE g.work_date BETWEEN ? AND ? ORDER BY g.work_date, g.sort_order, g.id',
        (date_from, date_to),
    ).fetchall()

    requests_rows = db.execute(
        _OVERHEAD_REQUEST_JOIN_SELECT +
        ' WHERE r.work_date BETWEEN ? AND ? ORDER BY r.work_date, r.group_id, r.sort_order, r.id',
        (date_from, date_to),
    ).fetchall()
    db.close()

    # Index requests by group id, attaching pre-computed hours/cost so the
    # client doesn't have to mirror the time-math logic.
    by_group = {}
    for r in requests_rows:
        rd = _annotate_request_metrics(_normalize_row_dates(dict(r)))
        by_group.setdefault(rd['group_id'], []).append(rd)

    # Build days[] → groups[] → requests[]
    days = {}
    day_order = []
    for g in groups:
        gd = _normalize_row_dates(dict(g))
        wd = gd['work_date']
        if wd not in days:
            days[wd] = {'work_date': wd, 'groups': []}
            day_order.append(wd)
        gd['requests'] = by_group.get(gd['id'], [])
        # Display name = group's own override if non-empty, else the project's name
        gd['display_name'] = (
            gd.get('name')
            if (gd.get('name') and gd.get('name') != 'General')
            else (gd.get('project_name') or gd.get('name') or 'General')
        )
        days[wd]['groups'].append(gd)

    return jsonify({
        'days': [days[d] for d in sorted(day_order)],
        'from': date_from,
        'to':   date_to,
    })


# ── Projects (sub-group catalog) ─────────────────────────────────────────────

@app.route('/api/overhead-crew/projects', methods=['GET'])
@login_required
def api_overhead_projects_list():
    include_archived = request.args.get('include_archived') == '1'
    db = get_db()
    sql = """
        SELECT p.*,
               (SELECT COUNT(*) FROM overhead_labor_groups g WHERE g.project_id = p.id) AS group_count,
               (SELECT MIN(g.work_date) FROM overhead_labor_groups g WHERE g.project_id = p.id) AS first_date,
               (SELECT MAX(g.work_date) FROM overhead_labor_groups g WHERE g.project_id = p.id) AS last_date
        FROM overhead_projects p
    """
    if not include_archived:
        sql += ' WHERE COALESCE(p.archived, 0) = 0'
    sql += ' ORDER BY p.archived ASC, p.sort_order, p.name'
    rows = db.execute(sql).fetchall()
    db.close()
    return jsonify([_normalize_row_dates(dict(r)) for r in rows])


def _project_payload(data):
    return {
        'name':          (data.get('name') or '').strip(),
        'description':   (data.get('description') or '').strip(),
        'client_name':   (data.get('client_name') or '').strip(),
        'billing_code':  (data.get('billing_code') or '').strip(),
        'contact_name':  (data.get('contact_name') or '').strip(),
        'contact_email': (data.get('contact_email') or '').strip(),
        'contact_phone': (data.get('contact_phone') or '').strip(),
        'project_notes': data.get('project_notes') or '',
        'color':         (data.get('color') or '').strip(),
        'archived':      1 if data.get('archived') else 0,
    }


def _find_or_create_project(db, name, *, defaults=None):
    """Look up a project by name (case-insensitive). If missing, create it with
    optional default contact / notes. Returns the project id."""
    name = (name or '').strip()
    if not name:
        return None
    row = db.execute(
        'SELECT id FROM overhead_projects WHERE LOWER(name) = LOWER(?)',
        (name,)
    ).fetchone()
    if row:
        return row['id']
    d = defaults or {}
    cur = db.execute("""
        INSERT INTO overhead_projects
            (name, description, client_name, billing_code,
             contact_name, contact_email, contact_phone, project_notes, color,
             archived, sort_order, created_by)
        VALUES (?,?,?,?,?,?,?,?,?,0,?,?)
    """, (
        name,
        (d.get('description') or '').strip(),
        (d.get('client_name') or '').strip(),
        (d.get('billing_code') or '').strip(),
        (d.get('contact_name') or '').strip(),
        (d.get('contact_email') or '').strip(),
        (d.get('contact_phone') or '').strip(),
        d.get('project_notes') or '',
        (d.get('color') or '').strip(),
        _max_sort_order(db, 'overhead_projects'),
        session.get('user_id'),
    ))
    pid = cur.lastrowid
    log_audit(db, 'OVERHEAD_PROJECT_ADD', 'overhead_project', pid,
              detail=f'auto-created via group: {name}')
    _overhead_log('OVERHEAD_PROJECT_ADD', id=pid, name=name, source='auto')
    return pid


@app.route('/api/overhead-crew/projects', methods=['POST'])
@login_required
def api_overhead_projects_add():
    block = _overhead_write_check()
    if block: return block
    data = request.get_json(force=True) or {}
    payload = _project_payload(data)
    if not payload['name']:
        return jsonify({'success': False, 'error': 'Project name is required.'}), 400
    db = get_db()
    existing = db.execute(
        'SELECT id FROM overhead_projects WHERE LOWER(name) = LOWER(?)', (payload['name'],)
    ).fetchone()
    if existing:
        db.close()
        return jsonify({'success': False, 'error': 'A project with that name already exists.'}), 409
    cur = db.execute("""
        INSERT INTO overhead_projects
            (name, description, client_name, billing_code,
             contact_name, contact_email, contact_phone, project_notes, color,
             archived, sort_order, created_by)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?)
    """, (
        payload['name'], payload['description'], payload['client_name'],
        payload['billing_code'], payload['contact_name'], payload['contact_email'],
        payload['contact_phone'], payload['project_notes'], payload['color'],
        payload['archived'], _max_sort_order(db, 'overhead_projects'),
        session.get('user_id'),
    ))
    pid = cur.lastrowid
    log_audit(db, 'OVERHEAD_PROJECT_ADD', 'overhead_project', pid, detail=payload['name'])
    db.commit()
    row = db.execute('SELECT * FROM overhead_projects WHERE id=?', (pid,)).fetchone()
    db.close()
    _overhead_log('OVERHEAD_PROJECT_ADD', id=pid, name=payload['name'])
    return jsonify({'success': True, 'project': _normalize_row_dates(dict(row)) if row else {'id': pid}})


@app.route('/api/overhead-crew/projects/<int:pid>', methods=['PUT'])
@login_required
def api_overhead_projects_update(pid):
    block = _overhead_write_check()
    if block: return block
    data = request.get_json(force=True) or {}
    payload = _project_payload(data)
    if not payload['name']:
        return jsonify({'success': False, 'error': 'Project name is required.'}), 400
    db = get_db()
    row = db.execute('SELECT id FROM overhead_projects WHERE id=?', (pid,)).fetchone()
    if not row:
        db.close()
        return jsonify({'success': False, 'error': 'Not found.'}), 404
    # Detect a name collision against a *different* project
    dup = db.execute(
        'SELECT id FROM overhead_projects WHERE LOWER(name) = LOWER(?) AND id <> ?',
        (payload['name'], pid)
    ).fetchone()
    if dup:
        db.close()
        return jsonify({'success': False, 'error': 'Another project already uses that name.'}), 409
    set_clause = ', '.join(f'{k}=?' for k in payload.keys())
    db.execute(
        f'UPDATE overhead_projects SET {set_clause} WHERE id=?',
        list(payload.values()) + [pid],
    )
    log_audit(db, 'OVERHEAD_PROJECT_EDIT', 'overhead_project', pid, detail=payload['name'])
    db.commit()
    refreshed = db.execute('SELECT * FROM overhead_projects WHERE id=?', (pid,)).fetchone()
    db.close()
    _overhead_log('OVERHEAD_PROJECT_EDIT', id=pid, name=payload['name'])
    return jsonify({'success': True, 'project': _normalize_row_dates(dict(refreshed)) if refreshed else None})


@app.route('/api/overhead-crew/projects/<int:pid>', methods=['DELETE'])
@login_required
def api_overhead_projects_delete(pid):
    """Archive a project (soft delete) by default. Hard delete only if no groups
    reference it. Existing groups keep their override name/contact when unlinked."""
    block = _overhead_write_check()
    if block: return block
    hard = request.args.get('hard') == '1'
    db = get_db()
    row = db.execute('SELECT * FROM overhead_projects WHERE id=?', (pid,)).fetchone()
    if not row:
        db.close()
        return jsonify({'success': False, 'error': 'Not found.'}), 404
    in_use = db.execute(
        'SELECT COUNT(*) FROM overhead_labor_groups WHERE project_id=?', (pid,)
    ).fetchone()[0] or 0
    if hard and not in_use:
        db.execute('DELETE FROM overhead_projects WHERE id=?', (pid,))
        log_audit(db, 'OVERHEAD_PROJECT_DELETE', 'overhead_project', pid, detail=row['name'])
        action = 'OVERHEAD_PROJECT_DELETE'
    else:
        db.execute('UPDATE overhead_projects SET archived=1 WHERE id=?', (pid,))
        log_audit(db, 'OVERHEAD_PROJECT_ARCHIVE', 'overhead_project', pid,
                  detail=f"{row['name']} (in_use={in_use})")
        action = 'OVERHEAD_PROJECT_ARCHIVE'
    db.commit()
    db.close()
    _overhead_log(action, id=pid, name=row['name'], in_use=in_use)
    return jsonify({'success': True, 'archived_only': bool(in_use) or not hard})


@app.route('/api/overhead-crew/project-stats', methods=['GET'])
@login_required
def api_overhead_project_stats():
    """Return per-project stats (range + lifetime) for the date range and overall.

    Each project's `range_*` numbers cover the [from, to] window if provided;
    `lifetime_*` numbers are computed across every overhead labor request ever
    logged against that project's groups.
    """
    date_from = (request.args.get('from') or '').strip() or None
    date_to   = (request.args.get('to')   or '').strip() or None

    db = get_db()
    projects = db.execute(
        'SELECT * FROM overhead_projects ORDER BY archived, sort_order, name'
    ).fetchall()

    # Pull ALL requests joined to their group → project. We do the per-project
    # aggregation in Python so we can compute hours via _calc_hours (which deals
    # with break math + over-midnight wrap) consistently with the rest of the app.
    rows = db.execute("""
        SELECT r.id, r.work_date, r.in_time, r.out_time,
               r.break_start, r.break_end, r.break2_start, r.break2_end,
               r.actual_in_time, r.actual_out_time,
               r.actual_break_start, r.actual_break_end,
               r.actual_break2_start, r.actual_break2_end,
               r.is_scheduled, r.pay_rate_snapshot,
               g.project_id, g.name AS group_name,
               prl.hourly_rate AS level_rate
        FROM overhead_labor_requests r
        JOIN overhead_labor_groups g ON g.id = r.group_id
        LEFT JOIN crew_members cm ON cm.id = r.scheduled_crew_member_id
        LEFT JOIN pay_rate_levels prl ON prl.id = cm.rate_level_id
    """).fetchall()
    db.close()

    def _empty_bucket():
        return {
            'lifetime_hours_planned': 0.0, 'lifetime_hours_actual':  0.0,
            'lifetime_cost_planned':  0.0, 'lifetime_cost_actual':   0.0,
            'lifetime_lines': 0,
            'lifetime_first_date': None, 'lifetime_last_date': None,
            'range_hours_planned': 0.0, 'range_hours_actual': 0.0,
            'range_cost_planned': 0.0, 'range_cost_actual': 0.0,
            'range_lines': 0,
        }

    by_pid = {}
    # Track unmatched (no project) lines under a synthetic project_id = None bucket
    for r in rows:
        rd = _annotate_request_metrics(
            _normalize_row_dates(dict(r)),
            rate_keys=('level_rate',),
        )
        pid = rd['project_id']
        rec = by_pid.setdefault(pid, _empty_bucket())
        wd = rd.get('work_date') or ''
        h_p, h_a = rd['hours_planned'], rd['hours_actual']
        cp, ca   = rd['cost_planned'],  rd['cost_actual']

        rec['lifetime_hours_planned'] += h_p
        rec['lifetime_hours_actual']  += h_a
        rec['lifetime_cost_planned']  += cp
        rec['lifetime_cost_actual']   += ca
        rec['lifetime_lines']         += 1
        if wd:
            if rec['lifetime_first_date'] is None or wd < rec['lifetime_first_date']:
                rec['lifetime_first_date'] = wd
            if rec['lifetime_last_date'] is None or wd > rec['lifetime_last_date']:
                rec['lifetime_last_date'] = wd

        in_range = bool(wd)
        if in_range and date_from and wd < date_from: in_range = False
        if in_range and date_to   and wd > date_to:   in_range = False
        if in_range:
            rec['range_hours_planned'] += h_p
            rec['range_hours_actual']  += h_a
            rec['range_cost_planned']  += cp
            rec['range_cost_actual']   += ca
            rec['range_lines']         += 1

    _NUM_KEYS = (
        'lifetime_hours_planned', 'lifetime_hours_actual',
        'lifetime_cost_planned',  'lifetime_cost_actual',
        'range_hours_planned',    'range_hours_actual',
        'range_cost_planned',     'range_cost_actual',
    )
    def _round_bucket(b):
        for k in _NUM_KEYS:
            b[k] = round(b[k], 2)
        return b

    out = []
    for p in projects:
        pd = _normalize_row_dates(dict(p))
        bucket = by_pid.get(pd['id']) or _empty_bucket()
        out.append({**pd, **_round_bucket(bucket)})

    # Also include un-projected lines (project_id=NULL) as a synthetic row
    if None in by_pid:
        out.append({
            'id': None, 'name': '(No project)', 'archived': 0,
            **_round_bucket(by_pid[None]),
        })

    return jsonify({'from': date_from, 'to': date_to, 'projects': out})


@app.route('/api/overhead-crew/groups', methods=['POST'])
@login_required
def api_overhead_group_add():
    """Create a sub-group under a date.

    Body: {
      work_date,
      project_id?:     int      (existing project — preferred when known)
      project_name?:   string   (auto-creates if a project with that name doesn't exist)
      name?:           string   (per-group display override; falls back to project name)
      contact_*?:      strings  (per-group override)
      project_notes?:  string   (per-group override)
    }

    If neither project_id nor project_name is given the group is "ungrouped" —
    it still appears under its date but won't roll up to a project total.
    """
    block = _overhead_write_check()
    if block: return block
    data = request.get_json(force=True) or {}
    wd = _parse_iso_date(data.get('work_date'))
    if not wd:
        return jsonify({'success': False, 'error': 'Valid work_date is required.'}), 400
    db = get_db()

    project_id = data.get('project_id') or None
    if project_id:
        try:
            project_id = int(project_id)
        except (TypeError, ValueError):
            project_id = None
    if not project_id and (data.get('project_name') or '').strip():
        project_id = _find_or_create_project(
            db, data.get('project_name'),
            defaults={
                'contact_name':  data.get('contact_name'),
                'contact_email': data.get('contact_email'),
                'contact_phone': data.get('contact_phone'),
                'project_notes': data.get('project_notes'),
            },
        )

    # Pull project defaults so the group display can resolve a sensible name
    proj = None
    if project_id:
        proj = db.execute('SELECT * FROM overhead_projects WHERE id=?', (project_id,)).fetchone()

    # Sub-group display name: prefer explicit override, then project name, then 'General'
    name = (data.get('name') or '').strip()
    if not name:
        name = (proj['name'] if proj else (data.get('project_name') or '').strip()) or 'General'

    cur = db.execute("""
        INSERT INTO overhead_labor_groups
            (work_date, project_id, name, contact_name, contact_email, contact_phone,
             project_notes, sort_order, created_by)
        VALUES (?,?,?,?,?,?,?,?,?)
    """, (
        wd.isoformat(),
        project_id,
        name,
        (data.get('contact_name') or '').strip(),
        (data.get('contact_email') or '').strip(),
        (data.get('contact_phone') or '').strip(),
        data.get('project_notes') or '',
        _max_sort_order(db, 'overhead_labor_groups', 'work_date=?', (wd.isoformat(),)),
        session.get('user_id'),
    ))
    gid = cur.lastrowid
    log_audit(db, 'OVERHEAD_GROUP_ADD', 'overhead_labor_group', gid,
              detail=f"{wd.isoformat()} · project={project_id} · {name}")
    db.commit()
    row = _fetch_overhead_group(db, gid)
    db.close()
    out = _normalize_row_dates(dict(row)) if row else {'id': gid}
    out['requests'] = []
    _overhead_log('OVERHEAD_GROUP_ADD',
                  id=gid, date=wd.isoformat(), project_id=project_id, name=name)
    return jsonify({'success': True, 'group': out})


@app.route('/api/overhead-crew/groups/<int:gid>', methods=['PUT'])
@login_required
def api_overhead_group_update(gid):
    """Update sub-group fields.

    Body: any of {work_date, project_id, project_name, name, contact_*, project_notes}.
    Passing project_id=null clears the project link. Passing project_name auto-
    creates the project if necessary (same upsert as the create endpoint).
    """
    block = _overhead_write_check()
    if block: return block
    data = request.get_json(force=True) or {}
    db = get_db()
    row = db.execute('SELECT * FROM overhead_labor_groups WHERE id=?', (gid,)).fetchone()
    if not row:
        db.close()
        return jsonify({'success': False, 'error': 'Not found.'}), 404

    updates, params = [], []
    if 'work_date' in data:
        wd = _parse_iso_date(data.get('work_date'))
        if not wd:
            db.close()
            return jsonify({'success': False, 'error': 'Invalid work_date.'}), 400
        updates.append('work_date=?'); params.append(wd.isoformat())
        # Cascade work_date to all requests under this group so the running list
        # stays consistent if the group's date is changed.
        db.execute(
            'UPDATE overhead_labor_requests SET work_date=? WHERE group_id=?',
            (wd.isoformat(), gid)
        )

    # Project linking — accept either an explicit id or a name (to upsert)
    if 'project_id' in data:
        pid = data.get('project_id')
        if pid in (None, '', 0):
            updates.append('project_id=?'); params.append(None)
        else:
            try:
                updates.append('project_id=?'); params.append(int(pid))
            except (TypeError, ValueError):
                pass
    elif 'project_name' in data:
        nm = (data.get('project_name') or '').strip()
        if nm:
            new_pid = _find_or_create_project(db, nm)
            updates.append('project_id=?'); params.append(new_pid)
        else:
            updates.append('project_id=?'); params.append(None)

    for key in ('name', 'contact_name', 'contact_email', 'contact_phone', 'project_notes'):
        if key in data:
            updates.append(f'{key}=?')
            params.append((data.get(key) or '').strip() if key != 'project_notes' else (data.get(key) or ''))
    if not updates:
        db.close()
        return jsonify({'success': False, 'error': 'No changes.'}), 400
    params.append(gid)
    db.execute(f"UPDATE overhead_labor_groups SET {', '.join(updates)} WHERE id=?", params)
    log_audit(db, 'OVERHEAD_GROUP_EDIT', 'overhead_labor_group', gid)
    db.commit()
    refreshed = _fetch_overhead_group(db, gid)
    db.close()
    _overhead_log('OVERHEAD_GROUP_EDIT', id=gid)
    return jsonify({'success': True, 'group': _normalize_row_dates(dict(refreshed)) if refreshed else None})


@app.route('/api/overhead-crew/groups/<int:gid>', methods=['DELETE'])
@login_required
def api_overhead_group_delete(gid):
    """Delete a sub-group and all its requests (ON DELETE CASCADE)."""
    block = _overhead_write_check()
    if block: return block
    db = get_db()
    row = db.execute('SELECT * FROM overhead_labor_groups WHERE id=?', (gid,)).fetchone()
    if not row:
        db.close()
        return jsonify({'success': False, 'error': 'Not found.'}), 404
    db.execute('DELETE FROM overhead_labor_requests WHERE group_id=?', (gid,))
    db.execute('DELETE FROM overhead_labor_groups WHERE id=?', (gid,))
    log_audit(db, 'OVERHEAD_GROUP_DELETE', 'overhead_labor_group', gid,
              detail=f"{row['work_date']} · {row['name']}")
    db.commit()
    db.close()
    _overhead_log('OVERHEAD_GROUP_DELETE', id=gid, date=row['work_date'], name=row['name'])
    return jsonify({'success': True})


@app.route('/api/overhead-crew/requests', methods=['POST'])
@login_required
def api_overhead_request_add():
    """Add a labor line to a sub-group."""
    block = _overhead_write_check()
    if block: return block
    data = request.get_json(force=True) or {}
    gid = data.get('group_id')
    if not gid:
        return jsonify({'success': False, 'error': 'group_id required'}), 400
    db = get_db()
    grp = db.execute('SELECT id, work_date FROM overhead_labor_groups WHERE id=?', (gid,)).fetchone()
    if not grp:
        db.close()
        return jsonify({'success': False, 'error': 'Group not found.'}), 404
    cur = db.execute("""
        INSERT INTO overhead_labor_requests
            (group_id, work_date, position_id, in_time, out_time, break_start, break_end,
             break2_start, break2_end, requested_name, sort_order, created_by)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?)
    """, (
        gid,
        grp['work_date'],
        data.get('position_id') or None,
        _normalize_perf_time((data.get('in_time') or '').strip()),
        _normalize_perf_time((data.get('out_time') or '').strip()),
        _normalize_perf_time((data.get('break_start') or '').strip()),
        _normalize_perf_time((data.get('break_end') or '').strip()),
        _normalize_perf_time((data.get('break2_start') or '').strip()),
        _normalize_perf_time((data.get('break2_end') or '').strip()),
        (data.get('requested_name') or '').strip(),
        _max_sort_order(db, 'overhead_labor_requests', 'group_id=?', (gid,)),
        session.get('user_id'),
    ))
    rid = cur.lastrowid
    log_audit(db, 'OVERHEAD_REQUEST_ADD', 'overhead_labor_request', rid,
              detail=f"group={gid}")
    db.commit()
    row = _fetch_overhead_request(db, rid)
    db.close()
    out = _annotate_request_metrics(_normalize_row_dates(dict(row))) if row else {'id': rid}
    _overhead_log('OVERHEAD_REQUEST_ADD', id=rid, group_id=gid)
    return jsonify({'success': True, 'request': out})


@app.route('/api/overhead-crew/requests/<int:rid>', methods=['PUT'])
@login_required
def api_overhead_request_update(rid):
    """Update fields on a labor request line. Any of position_id, in/out/break times,
    requested_name, is_scheduled, scheduled_crew_member_id."""
    block = _overhead_write_check()
    if block: return block
    data = request.get_json(force=True) or {}
    db = get_db()
    row = db.execute('SELECT * FROM overhead_labor_requests WHERE id=?', (rid,)).fetchone()
    if not row:
        db.close()
        return jsonify({'success': False, 'error': 'Not found.'}), 404

    # Scheduling fields (is_scheduled / scheduled_crew_member_id) are only
    # writable by users with scheduler permission — anyone else is editing
    # request details, not making the schedule.
    if ('is_scheduled' in data or 'scheduled_crew_member_id' in data) and not _can_schedule_labor():
        db.close()
        return jsonify({
            'success': False,
            'error': 'Only labor schedulers can change the SCHED checkbox or assign technicians.',
        }), 403

    updates, params, detail_parts = [], [], []
    if 'position_id' in data:
        pid = data.get('position_id')
        pid = int(pid) if pid else None
        updates.append('position_id=?'); params.append(pid)
        detail_parts.append(f'pos={pid}')
    for f in ('in_time', 'out_time',
              'break_start', 'break_end', 'break2_start', 'break2_end',
              'requested_name',
              'actual_in_time', 'actual_out_time',
              'actual_break_start', 'actual_break_end',
              'actual_break2_start', 'actual_break2_end',
              'notes'):
        if f in data:
            val = (data.get(f) or '').strip()
            if f.endswith('_time') or f.startswith('break_') or f.startswith('break2_') \
                    or f.startswith('actual_break_') or f.startswith('actual_break2_'):
                val = _normalize_perf_time(val)
            updates.append(f'{f}=?'); params.append(val)
    if 'is_scheduled' in data:
        updates.append('is_scheduled=?'); params.append(1 if data.get('is_scheduled') else 0)
        detail_parts.append(f"sched={1 if data.get('is_scheduled') else 0}")
    if 'scheduled_crew_member_id' in data:
        cmid = data.get('scheduled_crew_member_id')
        cmid = int(cmid) if cmid else None
        updates.append('scheduled_crew_member_id=?'); params.append(cmid)
        updates.append('scheduled_by=?'); params.append(session['user_id'])
        updates.append('scheduled_at=CURRENT_TIMESTAMP')
        detail_parts.append(f'crew={cmid}')

        # Snapshot the crew member's current rate level so future rate edits
        # don't rewrite the cost history. Cleared when the scheduling is removed.
        if cmid:
            level = db.execute("""
                SELECT prl.id, prl.hourly_rate
                FROM crew_members cm
                LEFT JOIN pay_rate_levels prl ON prl.id = cm.rate_level_id
                WHERE cm.id=?
            """, (cmid,)).fetchone()
            if level:
                updates.append('pay_rate_snapshot=?');           params.append(level['hourly_rate'])
                updates.append('pay_rate_level_id_snapshot=?');  params.append(level['id'])
        else:
            updates.append('pay_rate_snapshot=?');          params.append(None)
            updates.append('pay_rate_level_id_snapshot=?'); params.append(None)

    if not updates:
        db.close()
        return jsonify({'success': False, 'error': 'No changes.'}), 400

    params.append(rid)
    db.execute(f"UPDATE overhead_labor_requests SET {', '.join(updates)} WHERE id=?", params)
    log_audit(db, 'OVERHEAD_REQUEST_EDIT', 'overhead_labor_request', rid,
              detail='; '.join(detail_parts))
    db.commit()
    refreshed = _fetch_overhead_request(db, rid)
    db.close()
    out = _annotate_request_metrics(_normalize_row_dates(dict(refreshed))) if refreshed else None
    _overhead_log('OVERHEAD_REQUEST_EDIT', id=rid)
    return jsonify({'success': True, 'request': out})


@app.route('/api/overhead-crew/requests/<int:rid>', methods=['DELETE'])
@login_required
def api_overhead_request_delete(rid):
    """Delete a single labor line."""
    block = _overhead_write_check()
    if block: return block
    db = get_db()
    row = db.execute('SELECT id FROM overhead_labor_requests WHERE id=?', (rid,)).fetchone()
    if not row:
        db.close()
        return jsonify({'success': False, 'error': 'Not found.'}), 404
    db.execute('DELETE FROM overhead_labor_requests WHERE id=?', (rid,))
    log_audit(db, 'OVERHEAD_REQUEST_DELETE', 'overhead_labor_request', rid)
    db.commit()
    db.close()
    _overhead_log('OVERHEAD_REQUEST_DELETE', id=rid)
    return jsonify({'success': True})


# ── Recurring templates ──────────────────────────────────────────────────────

_OVERHEAD_TEMPLATE_JOIN_SELECT = """
    SELECT t.*, jp.name AS position_name
    FROM overhead_labor_templates t
    LEFT JOIN job_positions jp ON jp.id = t.position_id
"""


def _template_row_to_dict(r):
    d = _normalize_row_dates(dict(r))
    d['days_of_week'] = sorted(_parse_dow_csv(d.get('days_of_week')))
    return d


@app.route('/api/overhead-crew/templates', methods=['GET'])
@login_required
def api_overhead_templates_list():
    db = get_db()
    rows = db.execute(
        _OVERHEAD_TEMPLATE_JOIN_SELECT + ' ORDER BY t.is_active DESC, t.sort_order, t.id'
    ).fetchall()
    db.close()
    return jsonify([_template_row_to_dict(r) for r in rows])


def _prune_template_requests(db, tid, *, today=None):
    """Remove any future (work_date >= today) labor requests stamped with this
    template_id that no longer fit the template's current
    start_date / end_date / days_of_week / is_active values.

    Past requests are preserved as historical records — only upcoming ones get
    cleaned up so the running schedule stays in sync when a template is
    shortened, has weekdays removed, or is deactivated.

    Returns the number of rows deleted. Does NOT commit — caller does that.
    """
    today = today or date.today()
    t = db.execute(
        'SELECT id, start_date, end_date, days_of_week, is_active '
        'FROM overhead_labor_templates WHERE id=?',
        (tid,),
    ).fetchone()
    if not t:
        return 0

    is_active = bool(t['is_active'])
    wanted_days = _parse_dow_csv(t['days_of_week'])
    t_start = _parse_iso_date(t['start_date']) if t['start_date'] else None
    t_end   = _parse_iso_date(t['end_date'])   if t['end_date']   else None

    # Pull every future template-stamped request and keep the ones that no
    # longer fit. Doing this in Python keeps the weekday math identical to
    # the generator (Python weekday Mon=0..Sun=6 → JS getDay Sun=0..Sat=6).
    rows = db.execute(
        'SELECT id, work_date FROM overhead_labor_requests '
        'WHERE template_id=? AND work_date >= ?',
        (tid, today.isoformat()),
    ).fetchall()
    to_delete = []
    for r in rows:
        wd = _parse_iso_date(r['work_date'])
        if not wd:
            continue
        # Inactive template → drop every future stamped row
        if not is_active:
            to_delete.append(r['id']); continue
        if t_start and wd < t_start: to_delete.append(r['id']); continue
        if t_end   and wd > t_end:   to_delete.append(r['id']); continue
        # Empty wanted_days set means no weekday is allowed → drop everything
        if not wanted_days:
            to_delete.append(r['id']); continue
        js_w = (wd.weekday() + 1) % 7
        if js_w not in wanted_days:
            to_delete.append(r['id'])

    if not to_delete:
        return 0

    placeholders = ','.join(['?'] * len(to_delete))
    db.execute(
        f'DELETE FROM overhead_labor_requests WHERE id IN ({placeholders})',
        to_delete,
    )
    return len(to_delete)


def _template_payload_from_request(data):
    """Normalize an inbound payload into the column tuple/values we INSERT/UPDATE.
    Returns (col_assignments_dict)."""
    days = data.get('days_of_week')
    if isinstance(days, list):
        # Same validation as _parse_dow_csv: keep only 0..6
        days_csv = ','.join(str(n) for n in sorted(_parse_dow_csv(','.join(str(d) for d in days))))
    else:
        days_csv = ','.join(str(n) for n in sorted(_parse_dow_csv(data.get('days_of_week') or '')))
    return {
        'name':                   (data.get('name') or '').strip(),
        'position_id':            (int(data['position_id']) if data.get('position_id') else None),
        'quantity':               max(1, int(data.get('quantity') or 1)),
        'days_of_week':           days_csv,
        'start_date':             (_parse_iso_date(data.get('start_date')).isoformat() if data.get('start_date') else None),
        'end_date':               (_parse_iso_date(data.get('end_date')).isoformat() if data.get('end_date') else None),
        'in_time':                _normalize_perf_time((data.get('in_time') or '').strip()),
        'out_time':               _normalize_perf_time((data.get('out_time') or '').strip()),
        'break_start':            _normalize_perf_time((data.get('break_start') or '').strip()),
        'break_end':              _normalize_perf_time((data.get('break_end') or '').strip()),
        'break2_start':           _normalize_perf_time((data.get('break2_start') or '').strip()),
        'break2_end':             _normalize_perf_time((data.get('break2_end') or '').strip()),
        'default_group_name':     (data.get('default_group_name') or 'Overhead').strip() or 'Overhead',
        'default_contact_name':   (data.get('default_contact_name') or '').strip(),
        'default_contact_email':  (data.get('default_contact_email') or '').strip(),
        'default_contact_phone':  (data.get('default_contact_phone') or '').strip(),
        'default_project_notes':  data.get('default_project_notes') or '',
        'is_active':              1 if data.get('is_active', True) else 0,
    }


@app.route('/api/overhead-crew/templates', methods=['POST'])
@login_required
def api_overhead_templates_add():
    block = _overhead_write_check()
    if block: return block
    data = request.get_json(force=True) or {}
    payload = _template_payload_from_request(data)
    if not payload['name']:
        return jsonify({'success': False, 'error': 'Template name is required.'}), 400
    db = get_db()
    cols = list(payload.keys()) + ['sort_order']
    vals = list(payload.values()) + [_max_sort_order(db, 'overhead_labor_templates')]
    placeholders = ','.join(['?'] * len(cols))
    cur = db.execute(
        f"INSERT INTO overhead_labor_templates ({','.join(cols)}) VALUES ({placeholders})",
        vals,
    )
    tid = cur.lastrowid
    log_audit(db, 'OVERHEAD_TEMPLATE_ADD', 'overhead_labor_template', tid,
              detail=payload['name'])
    db.commit()
    row = db.execute(_OVERHEAD_TEMPLATE_JOIN_SELECT + ' WHERE t.id=?', (tid,)).fetchone()
    db.close()
    _overhead_log('OVERHEAD_TEMPLATE_ADD', id=tid, name=payload['name'])
    return jsonify({'success': True, 'template': _template_row_to_dict(row) if row else {'id': tid}})


@app.route('/api/overhead-crew/templates/<int:tid>', methods=['PUT'])
@login_required
def api_overhead_templates_update(tid):
    block = _overhead_write_check()
    if block: return block
    data = request.get_json(force=True) or {}
    payload = _template_payload_from_request(data)
    if not payload['name']:
        return jsonify({'success': False, 'error': 'Template name is required.'}), 400
    db = get_db()
    row = db.execute('SELECT id FROM overhead_labor_templates WHERE id=?', (tid,)).fetchone()
    if not row:
        db.close()
        return jsonify({'success': False, 'error': 'Not found.'}), 404
    set_clause = ', '.join(f'{k}=?' for k in payload.keys())
    db.execute(
        f'UPDATE overhead_labor_templates SET {set_clause} WHERE id=?',
        list(payload.values()) + [tid],
    )
    # Sync future requests with the (possibly narrowed) template — drops any
    # template-stamped lines that now fall outside the date range / weekdays,
    # so shortening a template removes its leftover lines automatically.
    pruned = _prune_template_requests(db, tid)
    log_audit(db, 'OVERHEAD_TEMPLATE_EDIT', 'overhead_labor_template', tid,
              detail=f"{payload['name']} (pruned={pruned})" if pruned else payload['name'])
    db.commit()
    refreshed = db.execute(_OVERHEAD_TEMPLATE_JOIN_SELECT + ' WHERE t.id=?', (tid,)).fetchone()
    db.close()
    _overhead_log('OVERHEAD_TEMPLATE_EDIT', id=tid, name=payload['name'], pruned=pruned)
    return jsonify({
        'success': True,
        'template': _template_row_to_dict(refreshed) if refreshed else None,
        'pruned': pruned,
    })


@app.route('/api/overhead-crew/templates/<int:tid>', methods=['DELETE'])
@login_required
def api_overhead_templates_delete(tid):
    block = _overhead_write_check()
    if block: return block
    db = get_db()
    row = db.execute('SELECT * FROM overhead_labor_templates WHERE id=?', (tid,)).fetchone()
    if not row:
        db.close()
        return jsonify({'success': False, 'error': 'Not found.'}), 404
    db.execute('DELETE FROM overhead_labor_templates WHERE id=?', (tid,))
    log_audit(db, 'OVERHEAD_TEMPLATE_DELETE', 'overhead_labor_template', tid,
              detail=row['name'])
    db.commit()
    db.close()
    _overhead_log('OVERHEAD_TEMPLATE_DELETE', id=tid, name=row['name'])
    return jsonify({'success': True})


@app.route('/api/overhead-crew/templates/generate', methods=['POST'])
@login_required
def api_overhead_templates_generate():
    """Generate labor requests from active templates for every date in [from, to]
    that matches a template's days_of_week. Idempotent: skips creating duplicate
    request lines that already exist for the same (date, group, position, in_time, out_time).

    Body: {from: 'YYYY-MM-DD', to: 'YYYY-MM-DD', template_id?: int}
    If template_id is omitted, all active templates are applied.
    """
    block = _overhead_write_check()
    if block: return block
    data = request.get_json(force=True) or {}
    df = _parse_iso_date(data.get('from'))
    dt = _parse_iso_date(data.get('to'))
    if not df or not dt:
        return jsonify({'success': False, 'error': 'Valid from + to dates required.'}), 400
    if dt < df:
        return jsonify({'success': False, 'error': 'to must be on or after from.'}), 400

    only_id = data.get('template_id')
    db = get_db()
    sql = "SELECT * FROM overhead_labor_templates WHERE is_active=1"
    params = []
    if only_id:
        sql += " AND id=?"
        params.append(int(only_id))
    templates = db.execute(sql, params).fetchall()

    requests_created = 0
    groups_created   = 0
    skipped          = 0

    # Cache project upserts per (run, template) so multiple matching dates re-use
    # the same project_id without a DB round-trip per day.
    project_id_cache = {}

    for t in templates:
        wanted = _parse_dow_csv(t['days_of_week'])
        if not wanted:
            continue

        t_start = _parse_iso_date(t['start_date']) if t['start_date'] else None
        t_end   = _parse_iso_date(t['end_date'])   if t['end_date']   else None

        # Auto-link the template's default group name to a project (creates one
        # on first generate so lifetime stats roll up correctly).
        gname = t['default_group_name'] or 'Overhead'
        if gname not in project_id_cache:
            project_id_cache[gname] = _find_or_create_project(
                db, gname,
                defaults={
                    'contact_name':  t['default_contact_name'],
                    'contact_email': t['default_contact_email'],
                    'contact_phone': t['default_contact_phone'],
                    'project_notes': t['default_project_notes'],
                },
            )
        project_id = project_id_cache[gname]

        # Build an equality predicate for position_id that works on both engines
        # (PostgreSQL rejects `col IS ?` when ? binds to NULL).
        pos_id = t['position_id']
        if pos_id is None:
            pos_pred  = 'position_id IS NULL'
            pos_param = []
        else:
            pos_pred  = 'position_id = ?'
            pos_param = [pos_id]

        cur_date = df
        last_generated = _parse_iso_date(t['last_generated_through']) if t['last_generated_through'] else None
        while cur_date <= dt:
            # Python weekday: Mon=0..Sun=6 → JS getDay: Sun=0..Sat=6
            js_w = (cur_date.weekday() + 1) % 7
            in_range = (
                (t_start is None or cur_date >= t_start) and
                (t_end   is None or cur_date <= t_end)
            )
            if js_w in wanted and in_range:
                iso = cur_date.isoformat()

                # Find or create the target group (same date + group name).
                grp = db.execute(
                    "SELECT id FROM overhead_labor_groups WHERE work_date=? AND name=? LIMIT 1",
                    (iso, gname),
                ).fetchone()
                if not grp:
                    cur_g = db.execute("""
                        INSERT INTO overhead_labor_groups
                            (work_date, project_id, name, contact_name, contact_email,
                             contact_phone, project_notes, sort_order, created_by)
                        VALUES (?,?,?,?,?,?,?,?,?)
                    """, (
                        iso, project_id, gname,
                        t['default_contact_name'] or '',
                        t['default_contact_email'] or '',
                        t['default_contact_phone'] or '',
                        t['default_project_notes'] or '',
                        _max_sort_order(db, 'overhead_labor_groups', 'work_date=?', (iso,)),
                        session.get('user_id'),
                    ))
                    gid = cur_g.lastrowid
                    groups_created += 1
                else:
                    gid = grp['id']

                # Idempotency: count matching lines once, only insert (qty - existing)
                # additional ones. Matching = same position + in/out within this group.
                existing = db.execute(
                    f"""SELECT COUNT(*) FROM overhead_labor_requests
                        WHERE group_id=? AND {pos_pred} AND in_time=? AND out_time=?""",
                    [gid] + pos_param + [t['in_time'] or '', t['out_time'] or ''],
                ).fetchone()[0] or 0
                qty = max(1, int(t['quantity'] or 1))
                to_insert = max(0, qty - existing)
                if to_insert == 0:
                    skipped += 1
                else:
                    next_sort = _max_sort_order(db, 'overhead_labor_requests', 'group_id=?', (gid,))
                    for i in range(to_insert):
                        _t_dict = dict(t)
                        db.execute("""
                            INSERT INTO overhead_labor_requests
                                (group_id, template_id, work_date, position_id,
                                 in_time, out_time,
                                 break_start, break_end, break2_start, break2_end,
                                 requested_name, sort_order, created_by)
                            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
                        """, (
                            gid, _t_dict['id'], iso, _t_dict['position_id'],
                            _t_dict.get('in_time') or '', _t_dict.get('out_time') or '',
                            _t_dict.get('break_start') or '', _t_dict.get('break_end') or '',
                            _t_dict.get('break2_start') or '', _t_dict.get('break2_end') or '',
                            '', next_sort + (i * 10),
                            session.get('user_id'),
                        ))
                        requests_created += 1
            cur_date = cur_date + timedelta(days=1)

        # Bump last_generated_through if we extended coverage
        if last_generated is None or dt > last_generated:
            db.execute(
                'UPDATE overhead_labor_templates SET last_generated_through=? WHERE id=?',
                (dt.isoformat(), t['id'])
            )

    log_audit(db, 'OVERHEAD_TEMPLATE_GENERATE', 'overhead_labor_template',
              detail=f"{df.isoformat()}→{dt.isoformat()} created={requests_created} groups={groups_created}")
    db.commit()
    db.close()
    _overhead_log('OVERHEAD_GENERATE',
                  range_from=df.isoformat(), range_to=dt.isoformat(),
                  template_id=only_id,
                  requests_created=requests_created,
                  groups_created=groups_created,
                  skipped=skipped)
    return jsonify({
        'success': True,
        'requests_created': requests_created,
        'groups_created':   groups_created,
        'skipped':          skipped,
    })


# ─── Crew Members ────────────────────────────────────────────────────────────

@app.route('/crew')
@login_required
def crew_tracker():
    db = get_db()
    categories = db.execute(
        'SELECT * FROM position_categories ORDER BY sort_order, id'
    ).fetchall()

    # Build categories with their positions
    positions = db.execute(
        'SELECT * FROM job_positions ORDER BY sort_order, id'
    ).fetchall()
    pos_by_cat = {}
    all_positions = []
    for p in positions:
        pos_by_cat.setdefault(p['category_id'], []).append(dict(p))
        all_positions.append(dict(p))

    cats_with_positions = []
    for c in categories:
        c_dict = dict(c)
        c_dict['positions'] = pos_by_cat.get(c['id'], [])
        cats_with_positions.append(c_dict)
    # Uncategorized positions
    uncategorized = pos_by_cat.get(None, [])

    # Crew members
    members = db.execute(
        'SELECT * FROM crew_members ORDER BY sort_order, name'
    ).fetchall()

    # Qualifications — build set of (crew_member_id, position_id)
    quals = db.execute('SELECT crew_member_id, position_id FROM crew_qualifications').fetchall()
    qual_set = {(q['crew_member_id'], q['position_id']) for q in quals}

    # Build member rows with qual flags
    member_rows = []
    for m in members:
        m_dict = dict(m)
        m_dict['qualifications'] = [q[1] for q in qual_set if q[0] == m['id']]
        member_rows.append(m_dict)

    db.close()
    can_edit = (
        session.get('user_role') == 'admin'
        or session.get('is_scheduler')
        or session.get('is_labor_scheduler')
        or session.get('is_content_admin')
    )
    return render_template('crew_tracker.html',
                           categories=cats_with_positions,
                           uncategorized_positions=uncategorized,
                           all_positions=all_positions,
                           members=member_rows,
                           can_edit=can_edit,
                           user=get_current_user())


@app.route('/api/pay-rate-levels')
@login_required
def api_pay_rate_levels():
    db = get_db()
    levels = db.execute('SELECT * FROM pay_rate_levels ORDER BY sort_order, name').fetchall()
    db.close()
    return jsonify([dict(l) for l in levels])


@app.route('/settings/pay-rate-levels/add', methods=['POST'])
@admin_required
def add_pay_rate_level():
    data = request.get_json(force=True) or {}
    name = data.get('name', '').strip()
    rate = float(data.get('hourly_rate') or 0)
    if not name:
        return jsonify({'success': False, 'error': 'Name is required.'}), 400
    db = get_db()
    max_order = db.execute('SELECT COALESCE(MAX(sort_order),0) FROM pay_rate_levels').fetchone()[0]
    cur = db.execute('INSERT INTO pay_rate_levels (name, hourly_rate, sort_order) VALUES (?,?,?)',
                     (name, rate, max_order + 10))
    lid = cur.lastrowid
    log_audit_change(db, 'PAY_LEVEL_ADD', 'pay_rate_level', lid,
                     detail=f'{name} ${rate}/hr', table='pay_rate_levels')
    db.commit(); db.close()
    syslog_logger.info(f"PAY_LEVEL_ADD id={lid} name={name!r} rate={rate} by={session.get('username')}")
    return jsonify({'success': True, 'id': lid, 'name': name, 'hourly_rate': rate})


@app.route('/settings/pay-rate-levels/<int:lid>/edit', methods=['POST'])
@admin_required
def edit_pay_rate_level(lid):
    data = request.get_json(force=True) or {}
    name = data.get('name', '').strip()
    rate = float(data.get('hourly_rate') or 0)
    if not name:
        return jsonify({'success': False, 'error': 'Name is required.'}), 400
    db = get_db()
    before = _snapshot_row(db, 'pay_rate_levels', lid)
    db.execute('UPDATE pay_rate_levels SET name=?, hourly_rate=? WHERE id=?', (name, rate, lid))
    after = _snapshot_row(db, 'pay_rate_levels', lid)
    log_audit(db, 'PAY_LEVEL_EDIT', 'pay_rate_level', lid, detail=f'{name} ${rate}/hr',
              before=before, after=after)
    db.commit(); db.close()
    syslog_logger.info(f"PAY_LEVEL_EDIT id={lid} name={name!r} rate={rate} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/pay-rate-levels/<int:lid>/delete', methods=['POST'])
@admin_required
def delete_pay_rate_level(lid):
    db = get_db()
    before = _snapshot_row(db, 'pay_rate_levels', lid)
    db.execute('DELETE FROM pay_rate_levels WHERE id=?', (lid,))
    log_audit(db, 'PAY_LEVEL_DELETE', 'pay_rate_level', lid, before=before)
    db.commit(); db.close()
    syslog_logger.info(f"PAY_LEVEL_DELETE id={lid} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/api/crew-members')
@login_required
def api_crew_members():
    db = get_db()
    members = db.execute("""
        SELECT cm.*, prl.name as level_name, prl.hourly_rate as level_rate
        FROM crew_members cm
        LEFT JOIN pay_rate_levels prl ON prl.id = cm.rate_level_id
        ORDER BY cm.sort_order, cm.name
    """).fetchall()
    quals = db.execute('SELECT crew_member_id, position_id FROM crew_qualifications').fetchall()
    qual_map = {}
    for q in quals:
        qual_map.setdefault(q['crew_member_id'], []).append(q['position_id'])

    result = []
    for m in members:
        m_dict = dict(m)
        m_dict['qualifications'] = qual_map.get(m['id'], [])
        result.append(m_dict)
    db.close()
    return jsonify(result)


@app.route('/settings/crew-members/add', methods=['POST'])
@scheduler_required
def add_crew_member():
    data = request.get_json(force=True) or {}
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'success': False, 'error': 'Name is required.'}), 400
    rate_level_id = data.get('rate_level_id') or None
    db = get_db()
    max_order = db.execute('SELECT MAX(sort_order) FROM crew_members').fetchone()[0] or 0
    cur = db.execute(
        'INSERT INTO crew_members (name, rate_level_id, sort_order) VALUES (?, ?, ?)',
        (name, rate_level_id, max_order + 10)
    )
    mid = cur.lastrowid
    log_audit_change(db, 'CREW_MEMBER_ADD', 'crew_member', mid, detail=name,
                     table='crew_members')
    db.commit()
    level_row = db.execute(
        'SELECT name, hourly_rate FROM pay_rate_levels WHERE id=?', (rate_level_id,)
    ).fetchone() if rate_level_id else None
    db.close()
    syslog_logger.info(f"TECHNICIAN_ADD id={mid} name={name!r} by={session.get('username')}")
    return jsonify({
        'success': True, 'id': mid, 'name': name,
        'rate_level_id': rate_level_id,
        'level_name': level_row['name'] if level_row else None,
        'level_rate': level_row['hourly_rate'] if level_row else None,
    })


@app.route('/settings/crew-members/<int:mid>/edit', methods=['POST'])
@scheduler_required
def edit_crew_member(mid):
    data = request.get_json(force=True) or {}
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'success': False, 'error': 'Name is required.'}), 400
    db = get_db()
    rate_level_id = data.get('rate_level_id') or None
    before = _snapshot_row(db, 'crew_members', mid)
    db.execute('UPDATE crew_members SET name=?, rate_level_id=? WHERE id=?', (name, rate_level_id, mid))
    after = _snapshot_row(db, 'crew_members', mid)
    log_audit(db, 'CREW_MEMBER_EDIT', 'crew_member', mid, detail=name,
              before=before, after=after)
    db.commit()
    level_row = db.execute(
        'SELECT name, hourly_rate FROM pay_rate_levels WHERE id=?', (rate_level_id,)
    ).fetchone() if rate_level_id else None
    db.close()
    syslog_logger.info(f"TECHNICIAN_EDIT id={mid} name={name!r} by={session.get('username')}")
    return jsonify({
        'success': True,
        'level_name': level_row['name'] if level_row else None,
        'level_rate': level_row['hourly_rate'] if level_row else None,
    })


@app.route('/settings/crew-members/<int:mid>/delete', methods=['POST'])
@scheduler_required
def delete_crew_member(mid):
    db = get_db()
    before = _snapshot_row(db, 'crew_members', mid)
    db.execute('DELETE FROM crew_members WHERE id=?', (mid,))
    log_audit(db, 'CREW_MEMBER_DELETE', 'crew_member', mid, before=before)
    db.commit()
    db.close()
    syslog_logger.info(f"TECHNICIAN_DELETE id={mid} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/crew-members/reorder', methods=['POST'])
@scheduler_required
def reorder_crew_members():
    data = request.get_json(force=True) or {}
    member_ids = data.get('member_ids', [])
    db = get_db()
    for i, mid in enumerate(member_ids):
        db.execute('UPDATE crew_members SET sort_order=? WHERE id=?', (i * 10, mid))
    db.commit()
    db.close()
    return jsonify({'success': True})


@app.route('/api/crew-qualifications/toggle', methods=['POST'])
@scheduler_required
def toggle_crew_qualification():
    data = request.get_json(force=True) or {}
    crew_member_id = data.get('crew_member_id')
    position_id = data.get('position_id')
    if not crew_member_id or not position_id:
        return jsonify({'success': False, 'error': 'crew_member_id and position_id required.'}), 400
    db = get_db()
    existing = db.execute(
        'SELECT 1 FROM crew_qualifications WHERE crew_member_id=? AND position_id=?',
        (crew_member_id, position_id)
    ).fetchone()
    if existing:
        db.execute(
            'DELETE FROM crew_qualifications WHERE crew_member_id=? AND position_id=?',
            (crew_member_id, position_id)
        )
        has = False
    else:
        db.execute(
            'INSERT INTO crew_qualifications (crew_member_id, position_id) VALUES (?, ?)',
            (crew_member_id, position_id)
        )
        has = True
    action = 'QUAL_ADD' if has else 'QUAL_REMOVE'
    log_audit(db, f'CREW_{action}', 'crew_qualification', crew_member_id,
              detail=f'position_id={position_id}')
    db.commit()
    db.close()
    syslog_logger.info(f"TECHNICIAN_{action} crew_member_id={crew_member_id} position_id={position_id} by={session.get('username')}")
    return jsonify({'success': True, 'has': has})


# ─── Asset Manager — Warehouse Locations ──────────────────────────────────────

@app.route('/settings/warehouse-locations', methods=['GET'])
@admin_required
def warehouse_locations_list():
    db = get_db()
    rows = db.execute('SELECT * FROM warehouse_locations ORDER BY sort_order, name').fetchall()
    db.close()
    return jsonify([dict(r) for r in rows])


@app.route('/settings/warehouse-locations', methods=['POST'])
@admin_required
def warehouse_location_add():
    data = request.get_json() or {}
    name = (data.get('name') or '').strip()
    if not name:
        return jsonify({'error': 'Name required'}), 400
    db = get_db()
    try:
        max_order = db.execute('SELECT COALESCE(MAX(sort_order),0) FROM warehouse_locations').fetchone()[0]
        db.execute('INSERT INTO warehouse_locations (name, sort_order) VALUES (?,?)', (name, max_order + 1))
        db.commit()
        row = db.execute('SELECT * FROM warehouse_locations WHERE name=?', (name,)).fetchone()
        log_audit_change(db, 'WAREHOUSE_LOC_ADD', 'warehouse_location', row['id'],
                         detail=name, table='warehouse_locations')
        db.commit()
        syslog_logger.info(f"WAREHOUSE_LOC_ADD name={name} by={session.get('username')}")
        result = dict(row)
        db.close()
        return jsonify(result), 201
    except DBIntegrityError:
        db.close()
        return jsonify({'error': 'Location name already exists'}), 409


@app.route('/settings/warehouse-locations/<int:loc_id>', methods=['PUT'])
@admin_required
def warehouse_location_edit(loc_id):
    data = request.get_json() or {}
    name = (data.get('name') or '').strip()
    if not name:
        return jsonify({'error': 'Name required'}), 400
    db = get_db()
    try:
        before = _snapshot_row(db, 'warehouse_locations', loc_id)
        db.execute('UPDATE warehouse_locations SET name=? WHERE id=?', (name, loc_id))
        db.commit()
        after = _snapshot_row(db, 'warehouse_locations', loc_id)
        log_audit(db, 'WAREHOUSE_LOC_EDIT', 'warehouse_location', loc_id,
                  detail=name, before=before, after=after)
        db.commit()
        db.close()
        return jsonify({'success': True})
    except DBIntegrityError:
        db.close()
        return jsonify({'error': 'Location name already exists'}), 409


@app.route('/settings/warehouse-locations/<int:loc_id>', methods=['DELETE'])
@admin_required
def warehouse_location_delete(loc_id):
    db = get_db()
    before = _snapshot_row(db, 'warehouse_locations', loc_id)
    row = db.execute('SELECT name FROM warehouse_locations WHERE id=?', (loc_id,)).fetchone()
    db.execute('DELETE FROM warehouse_locations WHERE id=?', (loc_id,))
    db.commit()
    log_audit(db, 'WAREHOUSE_LOC_DELETE', 'warehouse_location', loc_id,
              detail=row['name'] if row else str(loc_id),
              before=before)
    db.commit()
    db.close()
    return jsonify({'success': True})


# ─── Asset Manager — Categories ───────────────────────────────────────────────

@app.route('/settings/asset-categories', methods=['GET'])
@asset_manager_required
def asset_categories_list():
    db = get_db()
    rows = db.execute('SELECT * FROM asset_categories ORDER BY sort_order, name').fetchall()
    db.close()
    return jsonify([dict(r) for r in rows])


@app.route('/settings/asset-categories', methods=['POST'])
@asset_manager_required
def asset_category_add():
    data = request.get_json() or {}
    name = (data.get('name') or '').strip()
    if not name:
        return jsonify({'error': 'Name required'}), 400
    db = get_db()
    max_order = db.execute('SELECT COALESCE(MAX(sort_order),0) FROM asset_categories').fetchone()[0]
    db.execute('INSERT INTO asset_categories (name, sort_order) VALUES (?,?)', (name, max_order + 1))
    db.commit()
    row = db.execute('SELECT * FROM asset_categories WHERE name=? ORDER BY id DESC LIMIT 1', (name,)).fetchone()
    log_audit_change(db, 'ASSET_CATEGORY_ADD', 'asset_category', row['id'],
                     detail=name, table='asset_categories')
    db.commit()
    syslog_logger.info(f"ASSET_CATEGORY_ADD name={name} by={session.get('username')}")
    result = dict(row)
    db.close()
    return jsonify(result), 201


@app.route('/settings/asset-categories/<int:cat_id>', methods=['PUT'])
@asset_manager_required
def asset_category_edit(cat_id):
    data = request.get_json() or {}
    name = (data.get('name') or '').strip()
    if not name:
        return jsonify({'error': 'Name required'}), 400
    db = get_db()
    before = _snapshot_row(db, 'asset_categories', cat_id)
    db.execute('UPDATE asset_categories SET name=? WHERE id=?', (name, cat_id))
    db.commit()
    after = _snapshot_row(db, 'asset_categories', cat_id)
    log_audit(db, 'ASSET_CATEGORY_EDIT', 'asset_category', cat_id,
              detail=name, before=before, after=after)
    db.commit()
    db.close()
    return jsonify({'success': True})


@app.route('/settings/asset-categories/<int:cat_id>', methods=['DELETE'])
@asset_manager_required
def asset_category_delete(cat_id):
    db = get_db()
    # Block deletion if any types (including retired) exist — preserves history
    type_count = db.execute(
        'SELECT COUNT(*) FROM asset_types WHERE category_id=?', (cat_id,)
    ).fetchone()[0]
    if type_count > 0:
        db.close()
        return jsonify({'error': f'Cannot delete: this category still has {type_count} item type(s). Retire all types first.'}), 400
    before = _snapshot_row(db, 'asset_categories', cat_id)
    row = db.execute('SELECT name FROM asset_categories WHERE id=?', (cat_id,)).fetchone()
    db.execute('DELETE FROM asset_categories WHERE id=?', (cat_id,))
    db.commit()
    log_audit(db, 'ASSET_CATEGORY_DELETE', 'asset_category', cat_id,
              detail=row['name'] if row else str(cat_id),
              before=before)
    db.commit()
    db.close()
    return jsonify({'success': True})


# ─── Asset Manager — Types ────────────────────────────────────────────────────

@app.route('/api/asset-types', methods=['GET'])
@login_required
def asset_types_api():
    """Return active (non-retired) asset types for search/browse."""
    db = get_db()
    rows = db.execute("""
        SELECT at.*, ac.name as category_name,
               pt.name as parent_name
        FROM asset_types at
        JOIN asset_categories ac ON ac.id = at.category_id
        LEFT JOIN asset_types pt ON pt.id = at.parent_type_id
        WHERE at.is_retired = 0
        ORDER BY ac.sort_order, ac.name, at.sort_order, at.name
    """).fetchall()
    db.close()
    result = []
    for r in rows:
        d = dict(r)
        d.pop('photo', None)  # Don't send blob over API
        result.append(d)
    return jsonify(result)


@app.route('/settings/asset-types', methods=['GET'])
@asset_manager_required
def asset_types_admin_list():
    db = get_db()
    show_retired = request.args.get('show_retired') == '1'
    where = '' if show_retired else 'WHERE at.is_retired = 0'
    rows = db.execute(f"""
        SELECT at.*, ac.name as category_name,
               pt.name as parent_name,
               (SELECT COUNT(*) FROM asset_items ai WHERE ai.asset_type_id = at.id) as item_count,
               (SELECT COUNT(*) FROM asset_items ai WHERE ai.asset_type_id = at.id AND ai.status = 'retired') as retired_item_count
        FROM asset_types at
        JOIN asset_categories ac ON ac.id = at.category_id
        LEFT JOIN asset_types pt ON pt.id = at.parent_type_id
        {where}
        ORDER BY ac.sort_order, ac.name, at.sort_order, at.name
    """).fetchall()
    db.close()
    result = []
    for r in rows:
        d = dict(r)
        d.pop('photo', None)
        result.append(d)
    return jsonify(result)


@app.route('/settings/asset-types', methods=['POST'])
@asset_manager_required
def asset_type_add():
    data = request.get_json() or {}
    name = (data.get('name') or '').strip()
    category_id = data.get('category_id')
    if not name or not category_id:
        return jsonify({'error': 'Name and category required'}), 400
    db = get_db()
    max_order = db.execute('SELECT COALESCE(MAX(sort_order),0) FROM asset_types WHERE category_id=?',
                           (category_id,)).fetchone()[0]
    db.execute("""
        INSERT INTO asset_types
          (category_id, parent_type_id, name, manufacturer, model,
           storage_location, rental_cost, weekly_rate, reserve_count, is_consumable, track_quantity,
           supplier_name, supplier_contact, is_system, is_package, hide_from_pm, sort_order)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (
        category_id,
        data.get('parent_type_id') or None,
        name,
        (data.get('manufacturer') or '').strip(),
        (data.get('model') or '').strip(),
        (data.get('storage_location') or '').strip(),
        float(data.get('rental_cost') or 0),
        float(data.get('weekly_rate') or 0),
        int(data.get('reserve_count') or 0),
        1 if data.get('is_consumable') else 0,
        1 if data.get('track_quantity', True) else 0,
        (data.get('supplier_name') or '').strip(),
        (data.get('supplier_contact') or '').strip(),
        1 if data.get('is_system') else 0,
        1 if data.get('is_package') else 0,
        1 if data.get('hide_from_pm') else 0,
        max_order + 1,
    ))
    db.commit()
    row = db.execute('SELECT * FROM asset_types ORDER BY id DESC LIMIT 1').fetchone()
    log_audit(db, 'ASSET_TYPE_ADD', 'asset_type', row['id'], detail=name)
    db.commit()
    syslog_logger.info(f"ASSET_TYPE_ADD name={name} category_id={category_id} by={session.get('username')}")
    result = dict(row)
    result.pop('photo', None)
    db.close()
    return jsonify(result), 201


@app.route('/settings/asset-types/<int:type_id>', methods=['PUT'])
@asset_manager_required
def asset_type_edit(type_id):
    data = request.get_json() or {}
    name = (data.get('name') or '').strip()
    if not name:
        return jsonify({'error': 'Name required'}), 400
    db = get_db()
    db.execute("""
        UPDATE asset_types SET
          name=?, manufacturer=?, model=?, storage_location=?,
          rental_cost=?, weekly_rate=?, reserve_count=?, is_consumable=?, track_quantity=?,
          supplier_name=?, supplier_contact=?,
          category_id=?, parent_type_id=?, is_system=?, is_package=?, hide_from_pm=?
        WHERE id=?
    """, (
        name,
        (data.get('manufacturer') or '').strip(),
        (data.get('model') or '').strip(),
        (data.get('storage_location') or '').strip(),
        float(data.get('rental_cost') or 0),
        float(data.get('weekly_rate') or 0),
        int(data.get('reserve_count') or 0),
        1 if data.get('is_consumable') else 0,
        1 if data.get('track_quantity', True) else 0,
        (data.get('supplier_name') or '').strip(),
        (data.get('supplier_contact') or '').strip(),
        data.get('category_id'),
        data.get('parent_type_id') or None,
        1 if data.get('is_system') else 0,
        1 if data.get('is_package') else 0,
        1 if data.get('hide_from_pm') else 0,
        type_id,
    ))
    db.commit()
    log_audit(db, 'ASSET_TYPE_EDIT', 'asset_type', type_id, detail=name)
    db.commit()
    db.close()
    return jsonify({'success': True})


@app.route('/settings/asset-types/<int:type_id>', methods=['DELETE'])
@asset_manager_required
def asset_type_delete(type_id):
    """Retire an asset type (soft delete) — history is preserved."""
    db = get_db()
    row = db.execute('SELECT name FROM asset_types WHERE id=?', (type_id,)).fetchone()
    db.execute("""
        UPDATE asset_types SET is_retired=1, retired_at=CURRENT_TIMESTAMP WHERE id=?
    """, (type_id,))
    # Retire all active items under this type too
    db.execute("""
        UPDATE asset_items SET status='retired' WHERE asset_type_id=? AND status='available'
    """, (type_id,))
    db.commit()
    log_audit(db, 'ASSET_TYPE_RETIRE', 'asset_type', type_id,
              detail=row['name'] if row else str(type_id))
    db.commit()
    syslog_logger.info(f"ASSET_TYPE_RETIRE type_id={type_id} by={session.get('username')}")
    db.close()
    return jsonify({'success': True})


@app.route('/settings/asset-types/<int:type_id>/photo', methods=['POST'])
@asset_manager_required
def asset_type_photo_upload(type_id):
    f = request.files.get('photo')
    if not f:
        return jsonify({'error': 'No file'}), 400
    mime = f.mimetype or 'image/jpeg'
    data = f.read()
    db = get_db()
    if s3_storage.is_configured():
        try:
            s3_key = f"asset-photos/{type_id}"
            s3_storage.upload_file(s3_key, data, mime)
            db.execute('UPDATE asset_types SET photo=NULL, photo_s3_key=?, photo_mime=? WHERE id=?',
                       (s3_key, mime, type_id))
        except Exception as e:
            app.logger.warning(f"S3 upload failed for asset photo type_id={type_id}, falling back to DB: {e}")
            syslog_logger.warning(f"S3_UPLOAD_FAILED table=asset_types id={type_id} error={e}")
            db.execute('UPDATE asset_types SET photo=?, photo_s3_key=NULL, photo_mime=? WHERE id=?',
                       (data, mime, type_id))
    else:
        db.execute('UPDATE asset_types SET photo=?, photo_s3_key=NULL, photo_mime=? WHERE id=?',
                   (data, mime, type_id))
    db.commit()
    log_audit(db, 'ASSET_TYPE_PHOTO', 'asset_type', type_id)
    db.commit()
    db.close()
    syslog_logger.info(f"ASSET_TYPE_PHOTO type_id={type_id} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/asset-types/<int:type_id>/photo', methods=['DELETE'])
@asset_manager_required
def asset_type_photo_delete(type_id):
    db = get_db()
    row = db.execute('SELECT photo_s3_key FROM asset_types WHERE id=?', (type_id,)).fetchone()
    if row and row['photo_s3_key']:
        try:
            s3_storage.delete_file(row['photo_s3_key'])
        except Exception as e:
            app.logger.error(f"S3 delete failed for asset photo type_id={type_id}: {e}")
            syslog_logger.error(f"S3_DELETE_FAILED table=asset_types id={type_id} error={e}")
    db.execute("UPDATE asset_types SET photo=NULL, photo_s3_key=NULL, photo_mime='' WHERE id=?", (type_id,))
    db.commit()
    db.close()
    return jsonify({'success': True})


@app.route('/asset-types/<int:type_id>/photo')
@login_required
def asset_type_photo(type_id):
    db = get_db()
    row = db.execute('SELECT photo, photo_mime, photo_s3_key FROM asset_types WHERE id=?', (type_id,)).fetchone()
    db.close()
    if not row or (not row['photo_s3_key'] and not row['photo']):
        abort(404)
    if row['photo_s3_key']:
        try:
            data = s3_storage.download_file(row['photo_s3_key'])
        except Exception as e:
            app.logger.error(f"S3 download failed for asset photo type_id={type_id}: {e}")
            abort(503)
    else:
        data = bytes(row['photo'])
    resp = make_response(data)
    resp.headers['Content-Type'] = row['photo_mime'] or 'image/jpeg'
    resp.headers['Cache-Control'] = 'max-age=86400'
    return resp


# ─── Asset Manager — Used-in (reverse membership lookup) ─────────────────────

@app.route('/settings/asset-types/<int:type_id>/used-in')
@login_required
def asset_type_used_in(type_id):
    """Return system/package types that include this type as a component."""
    db = get_db()
    rows = db.execute("""
        SELECT at.id, at.name, at.is_system, at.is_package
        FROM asset_type_system_members m
        JOIN asset_types at ON at.id = m.system_type_id
        WHERE m.component_type_id = ?
        ORDER BY at.name
    """, (type_id,)).fetchall()
    db.close()
    return jsonify([dict(r) for r in rows])


# ─── Asset Manager — System/Package Members ──────────────────────────────────

@app.route('/settings/asset-types/<int:type_id>/members', methods=['GET'])
@asset_manager_required
def asset_type_members_list(type_id):
    db = get_db()
    try:
        # Per-system count (requires system_type_id column — added in migration)
        rows = db.execute("""
            SELECT at.id, at.name, at.manufacturer, at.model, at.is_system, at.is_package,
                   ac.name as category_name,
                   (SELECT COUNT(*) FROM asset_items ai
                    WHERE ai.asset_type_id = at.id
                      AND ai.system_type_id = m.system_type_id
                      AND ai.status != 'retired') as unit_count
            FROM asset_type_system_members m
            JOIN asset_types at ON at.id = m.component_type_id
            JOIN asset_categories ac ON ac.id = at.category_id
            WHERE m.system_type_id = ?
            ORDER BY m.sort_order, at.name
        """, (type_id,)).fetchall()
    except Exception:
        # Fallback: column not yet migrated — show total unit count
        rows = db.execute("""
            SELECT at.id, at.name, at.manufacturer, at.model, at.is_system, at.is_package,
                   ac.name as category_name,
                   (SELECT COUNT(*) FROM asset_items ai
                    WHERE ai.asset_type_id = at.id AND ai.status != 'retired') as unit_count
            FROM asset_type_system_members m
            JOIN asset_types at ON at.id = m.component_type_id
            JOIN asset_categories ac ON ac.id = at.category_id
            WHERE m.system_type_id = ?
            ORDER BY m.sort_order, at.name
        """, (type_id,)).fetchall()
    db.close()
    return jsonify([dict(r) for r in rows])


@app.route('/settings/asset-types/<int:type_id>/members', methods=['POST'])
@asset_manager_required
def asset_type_member_add(type_id):
    data = request.get_json() or {}
    component_id = data.get('component_type_id')
    if not component_id:
        return jsonify({'error': 'component_type_id required'}), 400
    db = get_db()
    try:
        db.execute("""
            INSERT OR IGNORE INTO asset_type_system_members (system_type_id, component_type_id)
            VALUES (?, ?)
        """, (type_id, component_id))
        db.commit()
    except Exception as e:
        db.close()
        return jsonify({'error': str(e)}), 400
    log_audit(db, 'ASSET_MEMBER_ADD', 'asset_type', type_id, detail=f'component={component_id}')
    db.commit()
    db.close()
    return jsonify({'success': True}), 201


@app.route('/settings/asset-types/<int:type_id>/members/<int:component_id>', methods=['DELETE'])
@asset_manager_required
def asset_type_member_remove(type_id, component_id):
    db = get_db()
    db.execute("""
        DELETE FROM asset_type_system_members
        WHERE system_type_id = ? AND component_type_id = ?
    """, (type_id, component_id))
    db.commit()
    log_audit(db, 'ASSET_MEMBER_REMOVE', 'asset_type', type_id, detail=f'component={component_id}')
    db.commit()
    db.close()
    return jsonify({'success': True})


# ─── Asset Manager — Items ────────────────────────────────────────────────────

@app.route('/settings/asset-types/<int:type_id>/items', methods=['GET'])
@asset_manager_required
def asset_items_list(type_id):
    db = get_db()
    show_retired = request.args.get('show_retired') == '1'
    status_filter = '' if show_retired else "AND ai.status != 'retired'"
    rows = db.execute(f"""
        SELECT ai.*,
               COALESCE(am.status, 'available') as maint_status,
               am.reason as maint_reason,
               am.notes as maint_notes,
               am.id as maint_id
        FROM asset_items ai
        LEFT JOIN asset_maintenance am ON am.asset_item_id = ai.id AND am.status = 'in_progress'
        WHERE ai.asset_type_id = ? {status_filter}
        ORDER BY ai.status, ai.sort_order, ai.id
    """, (type_id,)).fetchall()
    db.close()
    return jsonify([dict(r) for r in rows])


@app.route('/settings/asset-types/<int:type_id>/items', methods=['POST'])
@asset_manager_required
def asset_item_add(type_id):
    data = request.get_json() or {}
    count = int(data.get('count') or 1)
    barcode = (data.get('barcode') or '').strip()
    db = get_db()
    max_order = db.execute('SELECT COALESCE(MAX(sort_order),0) FROM asset_items WHERE asset_type_id=?',
                           (type_id,)).fetchone()[0]
    added_ids = []
    for i in range(count):
        bc = barcode if count == 1 else ''
        db.execute('INSERT INTO asset_items (asset_type_id, barcode, status, sort_order) VALUES (?,?,?,?)',
                   (type_id, bc, 'available', max_order + i + 1))
        db.commit()
        row = db.execute('SELECT * FROM asset_items ORDER BY id DESC LIMIT 1').fetchone()
        added_ids.append(row['id'])
    log_audit(db, 'ASSET_ITEM_ADD', 'asset_item', type_id, detail=f'count={count}')
    db.commit()
    syslog_logger.info(f"ASSET_ITEM_ADD type_id={type_id} count={count} by={session.get('username')}")
    db.close()
    return jsonify({'success': True, 'added': len(added_ids)}), 201


@app.route('/settings/asset-items/<int:item_id>', methods=['PUT'])
@asset_manager_required
def asset_item_edit(item_id):
    data = request.get_json() or {}
    db = get_db()
    def _int_or_none(v):
        try: return int(v) if v not in (None, '', 'null') else None
        except (ValueError, TypeError): return None
    def _float_or_none(v):
        try: return float(v) if v not in (None, '', 'null') else None
        except (ValueError, TypeError): return None
    valid_conditions = {'excellent', 'good', 'fair', 'poor', 'retired'}
    condition = data.get('condition', 'good')
    if condition not in valid_conditions:
        condition = 'good'
    db.execute("""
        UPDATE asset_items SET
          barcode=?, condition=?, year_purchased=?, purchase_value=?,
          depreciation_years=?, warranty_expires=?,
          depreciation_start_date=?, replacement_cost=?, is_container=?
        WHERE id=?
    """, (
        (data.get('barcode') or '').strip(),
        condition,
        _int_or_none(data.get('year_purchased')),
        _float_or_none(data.get('purchase_value')),
        _int_or_none(data.get('depreciation_years')),
        (data.get('warranty_expires') or '').strip() or None,
        (data.get('depreciation_start_date') or '').strip() or None,
        _float_or_none(data.get('replacement_cost')),
        1 if data.get('is_container') else 0,
        item_id,
    ))
    db.commit()
    log_audit(db, 'ASSET_ITEM_EDIT', 'asset_item', item_id)
    db.commit()
    db.close()
    return jsonify({'success': True})


@app.route('/settings/asset-items/<int:item_id>/contents', methods=['GET'])
@asset_manager_required
def asset_item_contents(item_id):
    """List items contained within a container item."""
    db = get_db()
    rows = db.execute("""
        SELECT ai.*, at.name as type_name
        FROM asset_items ai
        JOIN asset_types at ON at.id = ai.asset_type_id
        WHERE ai.container_item_id = ?
        ORDER BY at.name, ai.barcode
    """, (item_id,)).fetchall()
    db.close()
    return jsonify([dict(r) for r in rows])


@app.route('/settings/asset-items/<int:item_id>/set-container', methods=['POST'])
@asset_manager_required
def asset_item_set_container(item_id):
    """Assign an item to a container (or clear its container)."""
    data = request.get_json() or {}
    container_item_id = data.get('container_item_id') or None
    db = get_db()
    # Prevent an item from being its own container
    if container_item_id and int(container_item_id) == item_id:
        db.close()
        return jsonify({'error': 'An item cannot contain itself'}), 400
    db.execute('UPDATE asset_items SET container_item_id=? WHERE id=?', (container_item_id, item_id))
    db.commit()
    log_audit(db, 'ASSET_ITEM_CONTAINER', 'asset_item', item_id,
              detail=f'container_item_id={container_item_id}')
    db.commit()
    db.close()
    return jsonify({'success': True})


@app.route('/settings/asset-items/<int:item_id>/logs', methods=['GET'])
@asset_manager_required
def asset_item_logs_list(item_id):
    db = get_db()
    rows = db.execute("""
        SELECT al.*, u.display_name as author_name
        FROM asset_logs al
        LEFT JOIN users u ON u.id = al.user_id
        WHERE al.asset_item_id = ?
        ORDER BY al.log_date DESC, al.created_at DESC
    """, (item_id,)).fetchall()
    db.close()
    return jsonify([dict(r) for r in rows])


@app.route('/settings/asset-items/<int:item_id>/logs', methods=['POST'])
@asset_manager_required
def asset_item_log_add(item_id):
    data = request.get_json() or {}
    body = (data.get('body') or '').strip()
    if not body:
        return jsonify({'error': 'Entry body required'}), 400
    log_type = data.get('log_type', 'note')
    if log_type not in ('note', 'damage', 'service', 'usage'):
        log_type = 'note'
    import re as _re
    from datetime import date as _date
    log_date = (data.get('log_date') or '').strip()
    if not log_date:
        log_date = _date.today().isoformat()
    elif not _re.match(r'^\d{4}-\d{2}-\d{2}$', log_date):
        return jsonify({'error': 'Invalid log_date format; use YYYY-MM-DD'}), 400
    db = get_db()
    db.execute("""
        INSERT INTO asset_logs (asset_item_id, user_id, log_date, log_type, body)
        VALUES (?,?,?,?,?)
    """, (item_id, session['user_id'], log_date, log_type, body))
    db.commit()
    row = db.execute("""
        SELECT al.*, u.display_name as author_name
        FROM asset_logs al LEFT JOIN users u ON u.id = al.user_id
        WHERE al.id = last_insert_rowid()
    """).fetchone()
    db.close()
    syslog_logger.info(f"ASSET_LOG_ADD item_id={item_id} log_type={log_type} by={session.get('username')}")
    return jsonify(dict(row)), 201


@app.route('/settings/asset-logs/<int:log_id>', methods=['DELETE'])
@asset_manager_required
def asset_log_delete(log_id):
    db = get_db()
    db.execute('DELETE FROM asset_logs WHERE id=?', (log_id,))
    db.commit()
    db.close()
    syslog_logger.info(f"ASSET_LOG_DELETE log_id={log_id} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/settings/asset-items/<int:item_id>', methods=['DELETE'])
@asset_manager_required
def asset_item_delete(item_id):
    """Retire an asset item (soft delete) — history is preserved."""
    db = get_db()
    db.execute("UPDATE asset_items SET status='retired' WHERE id=?", (item_id,))
    db.commit()
    log_audit(db, 'ASSET_ITEM_RETIRE', 'asset_item', item_id)
    db.commit()
    syslog_logger.info(f"ASSET_ITEM_RETIRE item_id={item_id} by={session.get('username')}")
    db.close()
    return jsonify({'success': True})


@app.route('/settings/asset-items/<int:item_id>/maintenance', methods=['POST'])
@asset_manager_required
def asset_item_maintenance_start(item_id):
    data = request.get_json() or {}
    reason = (data.get('reason') or '').strip()
    notes = (data.get('notes') or '').strip()
    db = get_db()
    # Close any open maintenance records first
    db.execute("UPDATE asset_maintenance SET status='resolved', resolved_at=CURRENT_TIMESTAMP WHERE asset_item_id=? AND status='in_progress'",
               (item_id,))
    db.execute("""
        INSERT INTO asset_maintenance (asset_item_id, removed_by, reason, notes, status)
        VALUES (?,?,?,?,'in_progress')
    """, (item_id, session['user_id'], reason, notes))
    db.execute("UPDATE asset_items SET status='maintenance' WHERE id=?", (item_id,))
    db.commit()
    log_audit(db, 'ASSET_MAINT_START', 'asset_item', item_id, detail=reason)
    db.commit()
    syslog_logger.info(f"ASSET_MAINT_START item_id={item_id} reason={reason} by={session.get('username')}")
    db.close()
    return jsonify({'success': True})


@app.route('/settings/asset-items/<int:item_id>/maintenance/resolve', methods=['POST'])
@asset_manager_required
def asset_item_maintenance_resolve(item_id):
    data = request.get_json() or {}
    notes = (data.get('notes') or '').strip()
    db = get_db()
    db.execute("""
        UPDATE asset_maintenance
        SET status='resolved', resolved_at=CURRENT_TIMESTAMP, notes=COALESCE(NULLIF(?,''), notes)
        WHERE asset_item_id=? AND status='in_progress'
    """, (notes, item_id))
    db.execute("UPDATE asset_items SET status='available' WHERE id=?", (item_id,))
    db.commit()
    log_audit(db, 'ASSET_MAINT_RESOLVE', 'asset_item', item_id)
    db.commit()
    syslog_logger.info(f"ASSET_MAINT_RESOLVE item_id={item_id} by={session.get('username')}")
    db.close()
    return jsonify({'success': True})


# ─── Asset Manager — Availability ─────────────────────────────────────────────

def _get_asset_availability(db, asset_type_id, start_date=None, end_date=None):
    """
    Return dict with:
      total_items   — count of all items for this type
      reserve_count — held back as spares
      in_maintenance — items currently in maintenance
      available     — items available for the date range (may be negative)
      shows         — list of shows requesting this asset with quantities
    """
    type_row = db.execute('SELECT * FROM asset_types WHERE id=?', (asset_type_id,)).fetchone()
    if not type_row:
        return None

    # System/package types have no individually tracked units — treat as always available
    if type_row['is_system'] or type_row['is_package']:
        return {'unlimited': True, 'kit': True}

    total_items = db.execute(
        "SELECT COUNT(*) FROM asset_items WHERE asset_type_id=? AND status != 'retired'",
        (asset_type_id,)
    ).fetchone()[0]

    in_maintenance = db.execute(
        "SELECT COUNT(*) FROM asset_items WHERE asset_type_id=? AND status='maintenance'",
        (asset_type_id,)
    ).fetchone()[0]

    reserve_count = type_row['reserve_count'] or 0

    # For consumables with unlimited stock
    if type_row['is_consumable'] and not type_row['track_quantity']:
        return {
            'total_items': None,
            'reserve_count': reserve_count,
            'in_maintenance': 0,
            'available': None,
            'shows': [],
            'unlimited': True,
        }

    # Shows requesting this asset (optionally filtered by date range overlap)
    params = [asset_type_id]
    date_filter = ''
    if start_date and end_date:
        date_filter = ' AND (sa.rental_end >= ? AND sa.rental_start <= ?)'
        params.extend([start_date, end_date])

    shows = db.execute(f"""
        SELECT sa.id, sa.show_id, sa.quantity, sa.rental_start, sa.rental_end,
               sa.is_hidden, s.name as show_name
        FROM show_assets sa
        JOIN shows s ON s.id = sa.show_id
        WHERE sa.asset_type_id = ?{date_filter}
        ORDER BY sa.rental_start
    """, params).fetchall()

    total_reserved = sum(r['quantity'] for r in shows)
    available = total_items - in_maintenance - reserve_count - total_reserved

    return {
        'total_items': total_items,
        'reserve_count': reserve_count,
        'in_maintenance': in_maintenance,
        'total_reserved': total_reserved,
        'available': available,
        'shows': [dict(r) for r in shows],
        'unlimited': False,
    }


@app.route('/api/asset-types/<int:type_id>/availability')
@login_required
def asset_type_availability(type_id):
    start = request.args.get('start')
    end = request.args.get('end')
    db = get_db()
    result = _get_asset_availability(db, type_id, start, end)
    db.close()
    if result is None:
        abort(404)
    return jsonify(result)


@app.route('/api/assets/availability')
def assets_availability_bulk():
    """Return availability summary for all asset types, plus by-show data.

    Accessible without login for public dashboards; access control applied
    to the by-show section only when a user is logged in.
    """
    date_from = request.args.get('from')
    date_to   = request.args.get('to')
    db = get_db()
    type_ids = [r['id'] for r in db.execute('SELECT id FROM asset_types WHERE is_retired=0').fetchall()]
    by_type = {}
    for tid in type_ids:
        info = _get_asset_availability(db, tid, date_from, date_to)
        if info:
            by_type[tid] = {
                'total':       info.get('total_items'),
                'maintenance': info.get('in_maintenance', 0),
                'reserved':    info.get('reserve_count', 0),
                'available':   info.get('available'),
            }

    # By-show summary (for 'by_show' layout) — requires login for access control
    if not session.get('user_id'):
        db.close()
        return jsonify({'by_type': by_type, 'by_show': []})

    accessible_ids = get_accessible_shows(session['user_id'])  # None = all, [] = none, list = specific
    params = []
    where  = []
    if accessible_ids is not None and len(accessible_ids) > 0:
        placeholders = ','.join('?' * len(accessible_ids))
        where.append(f's.id IN ({placeholders})')
        params.extend(accessible_ids)
    elif accessible_ids is not None and len(accessible_ids) == 0:
        db.close()
        return jsonify({'by_type': by_type, 'by_show': []})
    if date_from: where.append("COALESCE(s.show_date,'9999-12-31') >= ?"); params.append(date_from)
    if date_to:   where.append("COALESCE(s.show_date,'0001-01-01') <= ?"); params.append(date_to)
    where_sql = ('WHERE ' + ' AND '.join(where)) if where else ''
    shows_raw = db.execute(f"""
        SELECT s.id, s.name, s.show_date FROM shows s {where_sql}
        ORDER BY s.show_date
    """, params).fetchall()
    by_show = []
    for sr in shows_raw:
        assets = db.execute("""
            SELECT sa.quantity, sa.locked_price, sa.rental_start, sa.rental_end,
                   at.name as type_name, at.manufacturer,
                   ac.name as category_name
            FROM show_assets sa
            JOIN asset_types at ON at.id = sa.asset_type_id
            JOIN asset_categories ac ON ac.id = at.category_id
            WHERE sa.show_id = ? AND sa.is_hidden = 0
            ORDER BY ac.name, at.name
        """, (sr['id'],)).fetchall()
        if assets:
            by_show.append({
                'id':       sr['id'],
                'name':     sr['name'],
                'show_date': str(sr['show_date']) if sr['show_date'] else None,
                'assets':   [{k: (str(v) if hasattr(v, 'isoformat') else v) for k, v in dict(a).items()} for a in assets],
            })
    db.close()
    return jsonify({'by_type': by_type, 'by_show': by_show})


# ─── Asset Manager — Show Assets (per-show tab) ───────────────────────────────

# ─── Asset Approval — helpers ─────────────────────────────────────────────────

def _get_asset_notification_recipients(db, exclude_user_id=None):
    """Return [(display_name, email)] for users who can manage assets.

    Mirrors the access rules of asset_manager_required: admins, staff,
    members of an 'admin_group' user group, and users with the
    is_asset_manager flag. Skips users without an email and the optional
    exclude_user_id (typically the actor that triggered the event, so they
    don't receive an email about their own action).
    """
    rows = db.execute("""
        SELECT DISTINCT u.id, u.username, u.display_name, u.email
        FROM users u
        LEFT JOIN user_group_members ugm ON ugm.user_id = u.id
        LEFT JOIN user_groups ug ON ug.id = ugm.group_id
        WHERE u.email IS NOT NULL AND u.email != ''
          AND (
            u.role IN ('admin', 'staff')
            OR u.is_asset_manager = 1
            OR ug.group_type = 'admin_group'
          )
    """).fetchall()
    out = []
    seen = set()
    for r in rows:
        if exclude_user_id is not None and r['id'] == exclude_user_id:
            continue
        if r['email'] in seen:
            continue
        seen.add(r['email'])
        out.append(((r['display_name'] or r['username']), r['email']))
    return out


def _notify_asset_recipients(db, subject, body_text, exclude_user_id=None):
    """Send a system notification to all asset-permission holders.

    Failures are swallowed so that the triggering action still succeeds; each
    recipient send is independent.
    """
    try:
        recipients = _get_asset_notification_recipients(db, exclude_user_id)
    except Exception as exc:
        app.logger.error(f'asset notify recipient lookup failed: {exc}')
        return
    for _name, _email in recipients:
        try:
            _send_simple_email(_email, subject, body_text)
        except Exception as exc:
            app.logger.error(f'asset notify send to {_email} failed: {exc}')


def _reset_asset_approval(db, show_id, reason):
    """If the show's assets were approved, flip back to unapproved and log it.

    Called from every write path that changes a show's gear so an advance
    update silently re-queues the show for approval.
    """
    row = db.execute(
        'SELECT s.assets_approved, s.name FROM shows s WHERE s.id=?', (show_id,)
    ).fetchone()
    if not row or not row['assets_approved']:
        return
    db.execute(
        'UPDATE shows SET assets_approved=0, assets_approved_by=NULL, '
        'assets_approved_at=NULL WHERE id=?',
        (show_id,),
    )
    log_audit(db, 'ASSET_APPROVAL_RESET', 'show', show_id, show_id=show_id,
              detail=f'reason={reason}')
    show_name = row['name'] or f'Show #{show_id}'
    actor = session.get('display_name') or session.get('username') or 'Someone'
    try:
        approvals_url = url_for('asset_approvals', _external=True)
    except Exception:
        approvals_url = ''
    body = (
        f'Show "{show_name}" is now waiting for asset approval again.\n\n'
        f'Trigger: {reason}\n'
        f'Changed by: {actor}\n\n'
        + (f'Review pending approvals:\n{approvals_url}\n' if approvals_url else '')
    )
    _notify_asset_recipients(
        db,
        f'3·2·1→THEATER: Show Waiting for Asset Approval — {show_name}',
        body,
        exclude_user_id=session.get('user_id'),
    )


@app.route('/shows/<int:show_id>/assets', methods=['GET'])
@login_required
def show_assets_list(show_id):
    if not can_access_show(session['user_id'], show_id):
        return jsonify({'error': 'Access denied'}), 403
    db = get_db()
    user_is_restricted = session.get('is_restricted', False)
    user_is_admin = session.get('user_role') == 'admin'

    rows = db.execute("""
        SELECT sa.*, at.name as type_name, at.is_consumable, at.manufacturer, at.model,
               at.rental_cost as current_price,
               ac.name as category_name
        FROM show_assets sa
        JOIN asset_types at ON at.id = sa.asset_type_id
        JOIN asset_categories ac ON ac.id = at.category_id
        WHERE sa.show_id = ?
          AND (? = 1 OR sa.is_hidden = 0)
        ORDER BY ac.name, at.name, sa.created_at
    """, (show_id, 1 if user_is_admin else 0)).fetchall()

    # External rentals
    ext_rows = db.execute("""
        SELECT * FROM show_external_rentals WHERE show_id=? ORDER BY sort_order, id
    """, (show_id,)).fetchall()

    # Approval state
    appr = db.execute("""
        SELECT s.assets_approved, s.assets_approved_at,
               u.display_name AS approver_name, u.username AS approver_username
        FROM shows s LEFT JOIN users u ON u.id = s.assets_approved_by
        WHERE s.id = ?
    """, (show_id,)).fetchone()
    approval = {
        'approved': bool(appr['assets_approved']) if appr else False,
        'approved_at': appr['assets_approved_at'] if appr else None,
        'approver_name': (appr['approver_name'] or appr['approver_username']) if appr else None,
    }

    db.close()
    return jsonify({
        'assets': [dict(r) for r in rows],
        'external_rentals': [{k: v for k, v in dict(r).items() if k != 'pdf_data'} for r in ext_rows],
        'approval': approval,
    })


@app.route('/shows/<int:show_id>/assets', methods=['POST'])
@show_advance_editor_required
def show_asset_add(show_id):
    data = request.get_json() or {}
    asset_type_id = data.get('asset_type_id')
    quantity = int(data.get('quantity') or 1)
    if not asset_type_id or quantity < 1:
        return jsonify({'error': 'asset_type_id and quantity required'}), 400

    db = get_db()

    # Get show dates for rental period defaults
    show = db.execute('SELECT * FROM shows WHERE id=?', (show_id,)).fetchone()
    perfs = db.execute(
        'SELECT perf_date FROM show_performances WHERE show_id=? ORDER BY perf_date', (show_id,)
    ).fetchall()
    # Prefer load-in/out dates over first/last performance dates
    default_start = show['load_in_date'] or (perfs[0]['perf_date'] if perfs else show['show_date'])
    default_end   = show['load_out_date'] or (perfs[-1]['perf_date'] if perfs else show['show_date'])
    rental_start = data.get('rental_start') or default_start
    rental_end   = data.get('rental_end')   or default_end

    # Smart rate: weekly if weekly_rate set and rental >= 7 days, else daily × days
    type_row = db.execute('SELECT rental_cost, weekly_rate, hide_from_pm FROM asset_types WHERE id=?', (asset_type_id,)).fetchone()
    if data.get('locked_price') is not None:
        locked_price = float(data['locked_price'])
    elif type_row:
        daily_rate  = float(type_row['rental_cost'] or 0)
        weekly_rate = float(type_row['weekly_rate'] or 0)
        try:
            d_start = date.fromisoformat(str(rental_start)) if rental_start else None
            d_end   = date.fromisoformat(str(rental_end))   if rental_end   else None
            days = max(1, (d_end - d_start).days + 1) if d_start and d_end else 1
        except (ValueError, TypeError):
            days = 1
        if weekly_rate > 0 and days >= 7:
            locked_price = weekly_rate * math.ceil(days / 7)
        else:
            locked_price = daily_rate * days
    else:
        locked_price = 0.0

    # Default is_hidden from the asset type's hide_from_pm flag
    is_hidden = 1 if (type_row and type_row['hide_from_pm']) else 0

    db.execute("""
        INSERT INTO show_assets
          (show_id, asset_type_id, quantity, rental_start, rental_end, locked_price, is_hidden, notes, added_by)
        VALUES (?,?,?,?,?,?,?,?,?)
    """, (show_id, asset_type_id, quantity, rental_start, rental_end, locked_price,
          is_hidden, (data.get('notes') or '').strip(), session['user_id']))
    db.commit()
    row = db.execute('SELECT * FROM show_assets ORDER BY id DESC LIMIT 1').fetchone()
    log_audit(db, 'ASSET_ADDED_TO_SHOW', 'show_asset', row['id'], show_id=show_id,
              detail=f'type_id={asset_type_id} qty={quantity}')
    _reset_asset_approval(db, show_id, 'asset_added')
    db.commit()
    syslog_logger.info(f"ASSET_ADDED_TO_SHOW show_id={show_id} type_id={asset_type_id} qty={quantity} by={session.get('username')}")
    result = dict(row)
    # Check for over-allocation and notify admins
    try:
        _avail = _get_asset_availability(db, asset_type_id, rental_start, rental_end)
        if _avail and not _avail.get('unlimited') and _avail.get('available') is not None and _avail['available'] < 0:
            _trow = db.execute('SELECT name FROM asset_types WHERE id=?', (asset_type_id,)).fetchone()
            _type_name = _trow['name'] if _trow else f'Type #{asset_type_id}'
            _show_name = show['name'] if show else f'Show #{show_id}'
            _notify_asset_recipients(
                db,
                f'3\u00b72\u00b71\u2192THEATER: Asset Over-Allocated \u2014 {_type_name}',
                f'Asset "{_type_name}" is now over-allocated for show "{_show_name}".\n\n'
                f'Current availability: {_avail["available"]} (negative = over-allocated)\n'
                f'Total units: {_avail.get("total_items","?")}, In maintenance: {_avail.get("in_maintenance",0)}, '
                f'Reserved spares: {_avail.get("reserve_count",0)}\n\n'
                f'Review the show\'s assets tab for details.',
                exclude_user_id=session.get('user_id'),
            )
    except Exception:
        pass
    db.close()
    return jsonify(result), 201


@app.route('/shows/<int:show_id>/assets/<int:sa_id>', methods=['PUT'])
@show_advance_editor_required
def show_asset_edit(show_id, sa_id):
    data = request.get_json() or {}
    db = get_db()
    existing = db.execute(
        'SELECT * FROM show_assets WHERE id=? AND show_id=?',
        (sa_id, show_id),
    ).fetchone()
    if not existing:
        db.close()
        return jsonify({'error': 'Not found'}), 404

    if 'quantity' in data:
        try:
            quantity = max(1, int(data.get('quantity') or 1))
        except (TypeError, ValueError):
            db.close()
            return jsonify({'error': 'Invalid quantity'}), 400
    else:
        quantity = existing['quantity']

    rental_start = data['rental_start'] if 'rental_start' in data else existing['rental_start']
    rental_end   = data['rental_end']   if 'rental_end'   in data else existing['rental_end']
    is_hidden    = (1 if data.get('is_hidden') else 0) if 'is_hidden' in data else existing['is_hidden']
    notes        = (data.get('notes') or '').strip()  if 'notes'      in data else existing['notes']

    db.execute("""
        UPDATE show_assets SET quantity=?, rental_start=?, rental_end=?,
               is_hidden=?, notes=?
        WHERE id=? AND show_id=?
    """, (quantity, rental_start, rental_end, is_hidden, notes, sa_id, show_id))
    db.commit()
    log_audit(db, 'ASSET_SHOW_EDIT', 'show_asset', sa_id, show_id=show_id)
    _reset_asset_approval(db, show_id, 'asset_edited')
    db.commit()
    db.close()
    return jsonify({'success': True})


@app.route('/shows/<int:show_id>/assets/<int:sa_id>', methods=['DELETE'])
@show_advance_editor_required
def show_asset_remove(show_id, sa_id):
    db = get_db()
    row = db.execute('SELECT * FROM show_assets WHERE id=? AND show_id=?', (sa_id, show_id)).fetchone()
    if not row:
        db.close()
        return jsonify({'error': 'Not found'}), 404
    db.execute('DELETE FROM show_assets WHERE id=?', (sa_id,))
    db.commit()
    log_audit(db, 'ASSET_REMOVED_FROM_SHOW', 'show_asset', sa_id, show_id=show_id,
              detail=f'type_id={row["asset_type_id"]}')
    _reset_asset_approval(db, show_id, 'asset_removed')
    db.commit()
    syslog_logger.info(f"ASSET_REMOVED_FROM_SHOW show_id={show_id} sa_id={sa_id} by={session.get('username')}")
    db.close()
    return jsonify({'success': True})


@app.route('/shows/<int:show_id>/assets/<int:sa_id>/toggle-hidden', methods=['POST'])
@show_advance_editor_required
def show_asset_toggle_hidden(show_id, sa_id):
    db = get_db()
    row = db.execute('SELECT is_hidden FROM show_assets WHERE id=? AND show_id=?', (sa_id, show_id)).fetchone()
    if not row:
        db.close()
        return jsonify({'error': 'Not found'}), 404
    new_val = 0 if row['is_hidden'] else 1
    db.execute('UPDATE show_assets SET is_hidden=? WHERE id=?', (new_val, sa_id))
    db.commit()
    log_audit(db, 'ASSET_HIDE_TOGGLE', 'show_asset', sa_id, show_id=show_id,
              detail=f'hidden={new_val}')
    _reset_asset_approval(db, show_id, 'asset_visibility_toggled')
    db.commit()
    db.close()
    return jsonify({'success': True, 'is_hidden': new_val})


# ─── Asset Manager — External Rentals ─────────────────────────────────────────

@app.route('/shows/<int:show_id>/external-rentals', methods=['POST'])
@show_advance_editor_required
def external_rental_add(show_id):
    db = get_db()
    description = (request.form.get('description') or '').strip()
    cost = float(request.form.get('cost') or 0)
    if not description:
        return jsonify({'error': 'Description required'}), 400
    pdf_bytes = None
    pdf_filename = ''
    f = request.files.get('pdf')
    if f:
        pdf_bytes = f.read()
        pdf_filename = secure_filename(f.filename)
    max_order = db.execute('SELECT COALESCE(MAX(sort_order),0) FROM show_external_rentals WHERE show_id=?',
                           (show_id,)).fetchone()[0]
    db.execute("""
        INSERT INTO show_external_rentals (show_id, description, cost, pdf_data, pdf_filename, sort_order)
        VALUES (?,?,?,?,?,?)
    """, (show_id, description, cost, None, pdf_filename, max_order + 1))
    db.commit()
    row = db.execute('SELECT * FROM show_external_rentals ORDER BY id DESC LIMIT 1').fetchone()
    er_id = row['id']
    # Upload PDF to S3 if provided
    if pdf_bytes:
        if s3_storage.is_configured():
            try:
                s3_key = f"external-rentals/{er_id}/{pdf_filename}"
                s3_storage.upload_file(s3_key, pdf_bytes, 'application/pdf')
                db.execute('UPDATE show_external_rentals SET s3_key=? WHERE id=?', (s3_key, er_id))
            except Exception as e:
                app.logger.warning(f"S3 upload failed for external rental {er_id}, falling back to DB: {e}")
                syslog_logger.warning(f"S3_UPLOAD_FAILED table=show_external_rentals id={er_id} show_id={show_id} error={e}")
                db.execute('UPDATE show_external_rentals SET pdf_data=? WHERE id=?', (pdf_bytes, er_id))
        else:
            db.execute('UPDATE show_external_rentals SET pdf_data=? WHERE id=?', (pdf_bytes, er_id))
        db.commit()
        row = db.execute('SELECT * FROM show_external_rentals WHERE id=?', (er_id,)).fetchone()
    log_audit(db, 'EXTERNAL_RENTAL_ADD', 'show_external_rental', er_id, show_id=show_id,
              detail=description)
    _reset_asset_approval(db, show_id, 'external_rental_added')
    db.commit()
    result = {k: v for k, v in dict(row).items() if k not in ('pdf_data', 's3_key')}
    db.close()
    syslog_logger.info(f"EXTERNAL_RENTAL_ADD show_id={show_id} er_id={er_id} desc={description!r} by={session.get('username')}")
    return jsonify(result), 201


@app.route('/shows/<int:show_id>/external-rentals/<int:er_id>', methods=['POST'])
@show_advance_editor_required
def external_rental_update(show_id, er_id):
    """Update an existing external rental — description, cost, and optionally
    replace the attached PDF. Sent as multipart/form-data so a new PDF can be
    uploaded; if no file is attached the existing PDF is preserved."""
    db = get_db()
    row = db.execute('SELECT * FROM show_external_rentals WHERE id=? AND show_id=?',
                     (er_id, show_id)).fetchone()
    if not row:
        db.close()
        return jsonify({'error': 'Not found'}), 404

    description = (request.form.get('description') or '').strip()
    if not description:
        db.close()
        return jsonify({'error': 'Description required'}), 400
    try:
        cost = float(request.form.get('cost') or 0)
    except (TypeError, ValueError):
        db.close()
        return jsonify({'error': 'Invalid cost'}), 400

    db.execute(
        'UPDATE show_external_rentals SET description=?, cost=? WHERE id=?',
        (description, cost, er_id),
    )

    f = request.files.get('pdf')
    if f and f.filename:
        new_bytes = f.read()
        new_filename = secure_filename(f.filename)
        old_s3_key = row['s3_key']
        if old_s3_key:
            try:
                s3_storage.delete_file(old_s3_key)
            except Exception as e:
                app.logger.warning(f"S3 delete failed for old external rental PDF {er_id}: {e}")
        if s3_storage.is_configured():
            try:
                s3_key = f"external-rentals/{er_id}/{new_filename}"
                s3_storage.upload_file(s3_key, new_bytes, 'application/pdf')
                db.execute(
                    'UPDATE show_external_rentals SET s3_key=?, pdf_filename=?, pdf_data=NULL WHERE id=?',
                    (s3_key, new_filename, er_id),
                )
            except Exception as e:
                app.logger.warning(f"S3 upload failed for external rental {er_id}, falling back to DB: {e}")
                db.execute(
                    'UPDATE show_external_rentals SET pdf_data=?, pdf_filename=?, s3_key=NULL WHERE id=?',
                    (new_bytes, new_filename, er_id),
                )
        else:
            db.execute(
                'UPDATE show_external_rentals SET pdf_data=?, pdf_filename=?, s3_key=NULL WHERE id=?',
                (new_bytes, new_filename, er_id),
            )

    log_audit(db, 'EXTERNAL_RENTAL_UPDATE', 'show_external_rental', er_id, show_id=show_id,
              detail=description)
    _reset_asset_approval(db, show_id, 'external_rental_updated')
    db.commit()
    updated = db.execute('SELECT * FROM show_external_rentals WHERE id=?', (er_id,)).fetchone()
    result = {k: v for k, v in dict(updated).items() if k not in ('pdf_data', 's3_key')}
    db.close()
    syslog_logger.info(
        f"EXTERNAL_RENTAL_UPDATE show_id={show_id} er_id={er_id} desc={description!r} "
        f"by={session.get('username')}"
    )
    return jsonify(result)


@app.route('/shows/<int:show_id>/external-rentals/<int:er_id>', methods=['DELETE'])
@show_advance_editor_required
def external_rental_delete(show_id, er_id):
    db = get_db()
    row = db.execute('SELECT s3_key FROM show_external_rentals WHERE id=? AND show_id=?',
                     (er_id, show_id)).fetchone()
    if row and row['s3_key']:
        try:
            s3_storage.delete_file(row['s3_key'])
        except Exception as e:
            app.logger.error(f"S3 delete failed for external rental {er_id}: {e}")
            syslog_logger.error(f"S3_DELETE_FAILED table=show_external_rentals id={er_id} show_id={show_id} error={e}")
    db.execute('DELETE FROM show_external_rentals WHERE id=? AND show_id=?', (er_id, show_id))
    db.commit()
    log_audit(db, 'EXTERNAL_RENTAL_DELETE', 'show_external_rental', er_id, show_id=show_id)
    _reset_asset_approval(db, show_id, 'external_rental_removed')
    db.commit()
    db.close()
    syslog_logger.info(f"EXTERNAL_RENTAL_DELETE show_id={show_id} er_id={er_id} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/shows/<int:show_id>/external-rentals/<int:er_id>/pdf')
@login_required
def external_rental_pdf(show_id, er_id):
    if not can_access_show(session['user_id'], show_id):
        abort(403)
    db = get_db()
    row = db.execute('SELECT * FROM show_external_rentals WHERE id=? AND show_id=?',
                     (er_id, show_id)).fetchone()
    db.close()
    if not row or (not row['s3_key'] and not row['pdf_data']):
        abort(404)
    if row['s3_key']:
        try:
            data = s3_storage.download_file(row['s3_key'])
        except Exception as e:
            app.logger.error(f"S3 download failed for external rental PDF {er_id}: {e}")
            abort(503)
    else:
        data = bytes(row['pdf_data'])
    resp = make_response(data)
    resp.headers['Content-Type'] = 'application/pdf'
    resp.headers['Content-Disposition'] = _safe_content_disposition(row['pdf_filename'] or 'rental.pdf')
    return resp


# ─── Asset Manager — Asset Page (admin view) ──────────────────────────────────

# ─── Asset Manager — Approvals ────────────────────────────────────────────────

@app.route('/assets/approvals')
@asset_manager_required
def asset_approvals():
    """Rental approval sub-page: shows whose load-in falls in a rolling
    (or custom) window, with their requested assets and approval state."""
    from datetime import date, timedelta
    try:
        start_str = (request.args.get('start') or '').strip()
        end_str   = (request.args.get('end')   or '').strip()
        start = date.fromisoformat(start_str) if start_str else date.today()
        end   = date.fromisoformat(end_str)   if end_str   else start + timedelta(days=21)
        if end < start:
            end = start
    except ValueError:
        start = date.today()
        end   = start + timedelta(days=21)

    db = get_db()
    # Shows whose load-in date (fallback: show_date, then earliest performance)
    # falls inside the window.  "All shows" means no status filter beyond active.
    shows = db.execute("""
        SELECT s.id, s.name, s.venue, s.show_date, s.show_time,
               s.load_in_date, s.load_in_time, s.load_out_date, s.load_out_time,
               s.assets_approved, s.assets_approved_at,
               u.display_name AS approver_name, u.username AS approver_username,
               COALESCE(s.load_in_date, s.show_date,
                        (SELECT MIN(perf_date) FROM show_performances sp
                          WHERE sp.show_id = s.id)) AS effective_date
        FROM shows s
        LEFT JOIN users u ON u.id = s.assets_approved_by
        WHERE COALESCE(s.status, 'active') = 'active'
        ORDER BY effective_date NULLS LAST, s.id
    """).fetchall()
    shows = [dict(r) for r in shows]
    # Filter to the date window (effective_date inside [start, end])
    def _in_range(v):
        if not v:
            return False
        try:
            return start <= date.fromisoformat(str(v)[:10]) <= end
        except ValueError:
            return False
    shows = [s for s in shows if _in_range(s.get('effective_date'))]

    # Aggregate per-show asset + external-rental totals, counts, and rows.
    for s in shows:
        assets = db.execute("""
            SELECT sa.*, at.name AS type_name, at.manufacturer, at.model,
                   ac.name AS category_name,
                   at.rental_cost AS catalog_daily_rate,
                   at.weekly_rate AS catalog_weekly_rate
            FROM show_assets sa
            JOIN asset_types at ON at.id = sa.asset_type_id
            JOIN asset_categories ac ON ac.id = at.category_id
            WHERE sa.show_id = ? AND sa.is_hidden = 0
            ORDER BY ac.name, at.name, sa.created_at
        """, (s['id'],)).fetchall()
        ext = db.execute("""
            SELECT id, description, cost, pdf_filename,
                   (pdf_data IS NOT NULL OR s3_key IS NOT NULL) AS has_pdf
            FROM show_external_rentals WHERE show_id=? ORDER BY sort_order, id
        """, (s['id'],)).fetchall()
        s['assets']           = [dict(r) for r in assets]
        s['external_rentals'] = [dict(r) for r in ext]
        s['assets_total']     = sum(float(r['locked_price'] or 0) * int(r['quantity'] or 1)
                                    for r in assets)
        s['externals_total']  = sum(float(r['cost'] or 0) for r in ext)
        s['total']            = s['assets_total'] + s['externals_total']
        s['has_any']          = bool(assets) or bool(ext)

    db.close()
    return render_template(
        'asset_approvals.html',
        shows=shows,
        range_start=start.isoformat(),
        range_end=end.isoformat(),
        user=get_current_user(),
    )


@app.route('/shows/<int:show_id>/assets/approve', methods=['POST'])
@asset_manager_required
def show_assets_approve(show_id):
    db = get_db()
    row = db.execute('SELECT assets_approved, name FROM shows WHERE id=?', (show_id,)).fetchone()
    if not row:
        db.close()
        return jsonify({'error': 'Show not found'}), 404
    db.execute("""
        UPDATE shows SET assets_approved=1, assets_approved_by=?,
               assets_approved_at=CURRENT_TIMESTAMP WHERE id=?
    """, (session['user_id'], show_id))
    log_audit(db, 'ASSET_APPROVAL_GRANTED', 'show', show_id, show_id=show_id)
    db.commit()
    me = db.execute('SELECT display_name, username FROM users WHERE id=?',
                    (session['user_id'],)).fetchone()
    approver_name = (me['display_name'] if me else '') or (me['username'] if me else '')
    approved_at = db.execute('SELECT assets_approved_at FROM shows WHERE id=?',
                             (show_id,)).fetchone()['assets_approved_at']
    show_name = row['name'] or f'Show #{show_id}'
    try:
        approvals_url = url_for('asset_approvals', _external=True)
    except Exception:
        approvals_url = ''
    body = (
        f'Show "{show_name}" has been marked as asset-approved.\n\n'
        f'Approved by: {approver_name or "(unknown)"}\n'
        f'Approved at: {approved_at}\n\n'
        + (f'View approvals:\n{approvals_url}\n' if approvals_url else '')
    )
    _notify_asset_recipients(
        db,
        f'3·2·1→THEATER: Show Assets Approved — {show_name}',
        body,
        exclude_user_id=session.get('user_id'),
    )
    db.close()
    syslog_logger.info(f"ASSET_APPROVAL_GRANTED show_id={show_id} by={session.get('username')}")
    return jsonify({'success': True,
                    'approver_name': approver_name,
                    'approved_at': approved_at})


@app.route('/shows/<int:show_id>/assets/unapprove', methods=['POST'])
@asset_manager_required
def show_assets_unapprove(show_id):
    db = get_db()
    row = db.execute('SELECT assets_approved FROM shows WHERE id=?', (show_id,)).fetchone()
    if not row:
        db.close()
        return jsonify({'error': 'Show not found'}), 404
    db.execute("""
        UPDATE shows SET assets_approved=0, assets_approved_by=NULL,
               assets_approved_at=NULL WHERE id=?
    """, (show_id,))
    log_audit(db, 'ASSET_APPROVAL_REVOKED', 'show', show_id, show_id=show_id)
    db.commit()
    db.close()
    syslog_logger.info(f"ASSET_APPROVAL_REVOKED show_id={show_id} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/shows/<int:show_id>/assets/<int:sa_id>/price', methods=['PUT'])
@asset_manager_required
def show_asset_price_override(show_id, sa_id):
    """Approver-only price override — edits the per-show locked_price
    without touching the catalog or other shows, and does NOT reset the
    show's approval state."""
    data = request.get_json() or {}
    try:
        new_price = float(data.get('locked_price'))
    except (TypeError, ValueError):
        return jsonify({'error': 'locked_price required'}), 400
    db = get_db()
    row = db.execute('SELECT locked_price FROM show_assets WHERE id=? AND show_id=?',
                     (sa_id, show_id)).fetchone()
    if not row:
        db.close()
        return jsonify({'error': 'Not found'}), 404
    old_price = float(row['locked_price'] or 0)
    db.execute('UPDATE show_assets SET locked_price=? WHERE id=?',
               (new_price, sa_id))
    log_audit(db, 'ASSET_PRICE_OVERRIDE', 'show_asset', sa_id, show_id=show_id,
              detail=f'old={old_price:.2f} new={new_price:.2f}')
    db.commit()
    db.close()
    syslog_logger.info(
        f"ASSET_PRICE_OVERRIDE show_id={show_id} sa_id={sa_id} "
        f"old={old_price} new={new_price} by={session.get('username')}"
    )
    return jsonify({'success': True, 'locked_price': new_price})


@app.route('/assets')
@asset_manager_required
def assets_admin():
    db = get_db()
    categories = db.execute('SELECT * FROM asset_categories ORDER BY sort_order, name').fetchall()
    locations = db.execute('SELECT * FROM warehouse_locations ORDER BY sort_order, name').fetchall()
    db.close()
    return render_template('assets.html',
                           categories=[dict(c) for c in categories],
                           locations=[dict(l) for l in locations],
                           user=get_current_user())


@app.route('/assets/retired')
@asset_manager_required
def assets_retired():
    db = get_db()
    # Retired types with their items and log counts
    types = db.execute("""
        SELECT at.*, ac.name as category_name,
               (SELECT COUNT(*) FROM asset_items ai WHERE ai.asset_type_id = at.id) as total_items,
               (SELECT COUNT(*) FROM asset_items ai WHERE ai.asset_type_id = at.id AND ai.status='retired') as retired_items
        FROM asset_types at
        JOIN asset_categories ac ON ac.id = at.category_id
        WHERE at.is_retired = 1
        ORDER BY at.retired_at DESC
    """).fetchall()
    # Also standalone retired items (type is still active, but item was individually retired)
    standalone = db.execute("""
        SELECT ai.*, at.name as type_name, at.manufacturer, at.model,
               ac.name as category_name,
               (SELECT COUNT(*) FROM asset_logs al WHERE al.asset_item_id = ai.id) as log_count
        FROM asset_items ai
        JOIN asset_types at ON at.id = ai.asset_type_id
        JOIN asset_categories ac ON ac.id = at.category_id
        WHERE ai.status = 'retired' AND at.is_retired = 0
        ORDER BY ai.created_at DESC
    """).fetchall()
    db.close()
    return render_template('asset_retired.html',
                           retired_types=[dict(t) for t in types],
                           standalone_items=[dict(s) for s in standalone],
                           user=get_current_user())


def _make_watermark_pdf(text):
    """Render a single transparent letter page with a subtle text watermark
    pinned to the bottom-left, used to brand extra-doc pages so the
    source is identifiable when printed. Returns PDF bytes or None."""
    try:
        from weasyprint import HTML as WP_HTML
        from markupsafe import escape
        safe = str(escape(text))
        html = (
            "<!doctype html><html><head><style>"
            "@page { size: letter; margin: 0; }"
            "html, body { margin: 0; padding: 0; height: 100%; }"
            ".wm { position: absolute; bottom: 0.18in; left: 0.4in; right: 0.4in; "
            "      font-family: Arial, Helvetica, sans-serif; font-size: 6pt; "
            "      color: #b0b0b0; letter-spacing: 0.04em; text-align: left; }"
            "</style></head><body>"
            f"<div class=\"wm\">{safe}</div>"
            "</body></html>"
        )
        return WP_HTML(string=html).write_pdf()
    except Exception as e:
        app.logger.warning(f'Watermark generation failed: {e}')
        return None


def _merge_pdfs(base_pdf_bytes, extra_pdfs, extras_watermark=None):
    """Append extra PDF byte-strings to base_pdf_bytes using pdfrw. Returns merged bytes.
    When extras_watermark is set, overlay that text on each extra page so printed
    copies retain a hint of their origin (the base PDF has its own footer)."""
    try:
        from pdfrw import PdfReader, PdfWriter, PageMerge
        from io import BytesIO
        writer = PdfWriter()
        writer.addpages(PdfReader(fdata=base_pdf_bytes).pages)

        watermark_page = None
        if extras_watermark:
            wm_bytes = _make_watermark_pdf(extras_watermark)
            if wm_bytes:
                try:
                    watermark_page = PdfReader(fdata=wm_bytes).pages[0]
                except Exception as e:
                    app.logger.warning(f'Watermark read failed: {e}')

        for extra in extra_pdfs:
            if not extra:
                continue
            try:
                pages = PdfReader(fdata=extra).pages
                for p in pages:
                    if watermark_page is not None:
                        try:
                            PageMerge(p).add(watermark_page).render()
                        except Exception as e:
                            app.logger.warning(f'Watermark overlay failed on a page: {e}')
                    writer.addpage(p)
            except Exception as e:
                app.logger.warning(f'PDF merge: skipping page set due to error: {e}')
        buf = BytesIO()
        writer.write(buf)
        return buf.getvalue()
    except Exception as e:
        app.logger.error(f'PDF merge failed: {e}')
        return base_pdf_bytes


def _fetch_external_rental_pdfs(db, show_id):
    """Return list of PDF byte-strings for all external rentals that have attached PDFs."""
    rows = db.execute(
        'SELECT id, s3_key, pdf_data, pdf_filename FROM show_external_rentals WHERE show_id=? ORDER BY sort_order',
        (show_id,)
    ).fetchall()
    result = []
    for row in rows:
        if not row['pdf_filename']:
            continue  # no PDF was ever attached to this rental
        if row['s3_key']:
            try:
                result.append(s3_storage.download_file(row['s3_key']))
            except Exception as e:
                app.logger.error(
                    f'PDF merge: S3 download failed for external_rental id={row["id"]} '
                    f'key={row["s3_key"]!r}: {e}'
                )
        elif row['pdf_data']:
            result.append(bytes(row['pdf_data']))
        else:
            app.logger.warning(
                f'PDF merge: external_rental id={row["id"]} has pdf_filename={row["pdf_filename"]!r} '
                f'but no s3_key and no pdf_data — PDF was lost (run S3 migration or re-upload)'
            )
    return result




@app.route('/shows/<int:show_id>/assets/invoice.pdf')
@login_required
def show_asset_invoice(show_id):
    """Generate a PDF invoice for all show assets and external rentals."""
    if not can_access_show(session['user_id'], show_id):
        abort(403)
    db = get_db()
    show = db.execute('SELECT * FROM shows WHERE id=?', (show_id,)).fetchone()
    if not show:
        db.close()
        abort(404)

    assets = db.execute("""
        SELECT sa.quantity, sa.locked_price, sa.rental_start, sa.rental_end, sa.notes,
               at.name as type_name, at.manufacturer, at.model,
               ac.name as category_name
        FROM show_assets sa
        JOIN asset_types at ON at.id = sa.asset_type_id
        JOIN asset_categories ac ON ac.id = at.category_id
        WHERE sa.show_id = ? AND sa.is_hidden = 0
        ORDER BY ac.sort_order, at.name
    """, (show_id,)).fetchall()

    external_rentals = db.execute("""
        SELECT description, cost, pdf_filename
        FROM show_external_rentals
        WHERE show_id = ?
        ORDER BY sort_order
    """, (show_id,)).fetchall()

    perf_company_row = db.execute(
        "SELECT field_value FROM advance_data WHERE show_id=? AND field_key='performance_company'",
        (show_id,)
    ).fetchone()
    performance_company = perf_company_row['field_value'] if perf_company_row else ''

    assets_list = [dict(a) for a in assets]
    ext_list    = [dict(e) for e in external_rentals]
    assets_subtotal  = sum((a['locked_price'] or 0) * a['quantity'] for a in assets_list)
    external_subtotal = sum(e['cost'] or 0 for e in ext_list)
    grand_total = assets_subtotal + external_subtotal

    html_str = render_template(
        'pdf/asset_invoice_pdf.html',
        show=dict(show),
        assets=assets_list,
        external_rentals=ext_list,
        assets_subtotal=assets_subtotal,
        external_subtotal=external_subtotal,
        grand_total=grand_total,
        performance_company=performance_company,
        layout=pdf_layouts.PdfLayout('asset_invoice', get_app_setting),
        generated_date=date.today().isoformat(),
    )

    try:
        from weasyprint import HTML as WP_HTML
        pdf_bytes = WP_HTML(string=html_str, base_url=request.host_url).write_pdf()
    except Exception as e:
        app.logger.error(f'WeasyPrint invoice error: {e}')
        return f'PDF generation failed: {e}', 500

    # Append any uploaded external rental PDFs
    er_pdfs = _fetch_external_rental_pdfs(db, show_id)
    if er_pdfs:
        pdf_bytes = _merge_pdfs(pdf_bytes, er_pdfs)
    db.close()

    safe_name = secure_filename(show['name'] or f'show_{show_id}')
    resp = make_response(pdf_bytes)
    resp.headers['Content-Type'] = 'application/pdf'
    resp.headers['Content-Disposition'] = _safe_content_disposition(f'{safe_name}_invoice.pdf')
    return resp


# ─── Post-Show Combined Invoice PDF ──────────────────────────────────────────

@app.route('/shows/<int:show_id>/post-show-invoice.pdf')
@login_required
def show_post_invoice(show_id):
    """Generate a combined PDF invoice for show assets, external rentals, and labor."""
    if not can_access_show(session['user_id'], show_id):
        abort(403)
    db = get_db()
    show = db.execute('SELECT * FROM shows WHERE id=?', (show_id,)).fetchone()
    if not show:
        db.close()
        abort(404)

    assets = db.execute("""
        SELECT sa.quantity, sa.locked_price, sa.rental_start, sa.rental_end, sa.notes,
               at.name as type_name, at.manufacturer, at.model,
               ac.name as category_name
        FROM show_assets sa
        JOIN asset_types at ON at.id = sa.asset_type_id
        JOIN asset_categories ac ON ac.id = at.category_id
        WHERE sa.show_id = ? AND sa.is_hidden = 0
        ORDER BY ac.sort_order, at.name
    """, (show_id,)).fetchall()

    external_rentals = db.execute("""
        SELECT description, cost, pdf_filename
        FROM show_external_rentals
        WHERE show_id = ?
        ORDER BY sort_order
    """, (show_id,)).fetchall()

    perf_company_row = db.execute(
        "SELECT field_value FROM advance_data WHERE show_id=? AND field_key='performance_company'",
        (show_id,)
    ).fetchone()
    performance_company = perf_company_row['field_value'] if perf_company_row else ''

    assets_list = [dict(a) for a in assets]
    ext_list    = [dict(e) for e in external_rentals]
    assets_subtotal   = sum((a['locked_price'] or 0) * a['quantity'] for a in assets_list)
    external_subtotal = sum(e['cost'] or 0 for e in ext_list)

    labor_lines, labor_total = _calc_labor_cost_for_show(db, show_id)
    er_pdfs = _fetch_external_rental_pdfs(db, show_id)
    db.close()

    grand_total = assets_subtotal + external_subtotal + labor_total

    html_str = render_template(
        'pdf/post_show_invoice_pdf.html',
        show=dict(show),
        assets=assets_list,
        external_rentals=ext_list,
        assets_subtotal=assets_subtotal,
        external_subtotal=external_subtotal,
        labor_lines=labor_lines,
        labor_total=labor_total,
        grand_total=grand_total,
        performance_company=performance_company,
        layout=pdf_layouts.PdfLayout('post_show_invoice', get_app_setting),
        generated_date=date.today().isoformat(),
    )

    try:
        from weasyprint import HTML as WP_HTML
        pdf_bytes = WP_HTML(string=html_str, base_url=request.host_url).write_pdf()
    except Exception as e:
        app.logger.error(f'WeasyPrint post-invoice error: {e}')
        return f'PDF generation failed: {e}', 500

    if er_pdfs:
        pdf_bytes = _merge_pdfs(pdf_bytes, er_pdfs)

    safe_name = secure_filename(show['name'] or f'show_{show_id}')
    resp = make_response(pdf_bytes)
    resp.headers['Content-Type'] = 'application/pdf'
    resp.headers['Content-Disposition'] = _safe_content_disposition(f'{safe_name}_post_show_invoice.pdf')
    return resp


# ─── In-App Updates ───────────────────────────────────────────────────────────

_update_state = {'running': False, 'log': [], 'phase': 'idle', 'error': None}
_update_lock  = threading.Lock()

def _update_log(msg):
    ts = datetime.now().strftime('%H:%M:%S')
    entry = f'[{ts}] {msg}'
    _update_state['log'].append(entry)
    app.logger.info(f'[updater] {msg}')

def _detect_service_name():
    """Auto-detect the systemd service this process is running under."""
    pid = os.getpid()
    try:
        r = subprocess.run(['systemctl', 'status', str(pid)],
                           capture_output=True, text=True, timeout=5)
        for line in r.stdout.split('\n'):
            m = re.search(r'(\S+\.service)', line)
            if m:
                return m.group(1)
    except Exception:
        pass
    # Try common names
    for name in ['showadvance', 'showadvance.service', '321theater', '321theater.service',
                 'gunicorn', 'gunicorn.service']:
        try:
            r = subprocess.run(['systemctl', 'is-active', name],
                               capture_output=True, text=True, timeout=2)
            if r.stdout.strip() == 'active':
                return name
        except Exception:
            pass
    return None

def _run_update(service_name):
    """Background thread: git pull + archive + restart + rollback on failure."""
    import glob as _glob
    update_archive = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                  'backups', f'pre_update_{datetime.now().strftime("%Y%m%d_%H%M%S")}')
    archived_files = []

    try:
        _update_state['phase'] = 'checking'
        _update_log('Fetching latest commits from remote…')
        r = subprocess.run(['git', 'fetch', 'origin'], capture_output=True, text=True, timeout=30,
                           cwd=os.path.dirname(os.path.abspath(__file__)))
        if r.returncode != 0:
            raise RuntimeError(f'git fetch failed: {r.stderr.strip()}')
        _update_log('Fetch complete.')

        # Get list of files that will change
        r = subprocess.run(['git', 'diff', '--name-only', 'HEAD', 'origin/HEAD'],
                           capture_output=True, text=True, timeout=10,
                           cwd=os.path.dirname(os.path.abspath(__file__)))
        changed = [f.strip() for f in r.stdout.split('\n') if f.strip()]
        if not changed:
            _update_log('Already up to date — no changes to apply.')
            _update_state['phase'] = 'done'
            return

        _update_log(f'Files to update: {", ".join(changed)}')

        # Archive changed files
        _update_state['phase'] = 'archiving'
        _update_log(f'Archiving {len(changed)} files to {update_archive}…')
        os.makedirs(update_archive, exist_ok=True)
        base = os.path.dirname(os.path.abspath(__file__))
        for rel in changed:
            src = os.path.join(base, rel)
            if os.path.exists(src):
                dst = os.path.join(update_archive, rel)
                os.makedirs(os.path.dirname(dst), exist_ok=True)
                shutil.copy2(src, dst)
                archived_files.append((src, dst))
        _update_log(f'Archived {len(archived_files)} files.')

        # Pull
        _update_state['phase'] = 'pulling'
        _update_log('Running git pull…')
        r = subprocess.run(['git', 'pull', '--ff-only'],
                           capture_output=True, text=True, timeout=60,
                           cwd=os.path.dirname(os.path.abspath(__file__)))
        if r.returncode != 0:
            raise RuntimeError(f'git pull failed: {r.stderr.strip() or r.stdout.strip()}')
        _update_log(r.stdout.strip() or 'Pull successful.')

        # Run DB migration
        _update_log('Running database migrations…')
        r = subprocess.run(['python', 'init_db.py', '--migrate'],
                           capture_output=True, text=True, timeout=60,
                           cwd=os.path.dirname(os.path.abspath(__file__)))
        _update_log(r.stdout.strip() or 'Migrations complete.')

        # Restart service
        if service_name:
            _update_state['phase'] = 'restarting'
            _update_log(f'Restarting service {service_name}…')
            r = subprocess.run(['systemctl', 'restart', service_name],
                               capture_output=True, text=True, timeout=30)
            if r.returncode != 0:
                raise RuntimeError(f'Service restart failed: {r.stderr.strip()}')
            # Wait and check it came back
            import time as _time
            _time.sleep(3)
            r2 = subprocess.run(['systemctl', 'is-active', service_name],
                                capture_output=True, text=True, timeout=5)
            if r2.stdout.strip() != 'active':
                raise RuntimeError(f'Service not active after restart: {r2.stdout.strip()}')
            _update_log(f'Service {service_name} is active.')
        else:
            _update_log('No systemd service detected — please restart the app manually.')

        _update_state['phase'] = 'done'
        _update_log('Update complete!')
        syslog_logger.info('APP_UPDATE applied successfully')

    except Exception as exc:
        _update_state['error'] = str(exc)
        _update_log(f'ERROR: {exc}')
        _update_log('Attempting rollback…')
        _update_state['phase'] = 'rolling_back'
        try:
            for src, bak in archived_files:
                shutil.copy2(bak, src)
                _update_log(f'  Restored {os.path.basename(src)}')
            _update_log('Rollback complete.')
            if service_name:
                _update_log(f'Restarting {service_name} after rollback…')
                subprocess.run(['systemctl', 'restart', service_name], timeout=30)
                _update_log('Restart issued.')
        except Exception as rb_exc:
            _update_log(f'Rollback failed: {rb_exc}')
        _update_state['phase'] = 'failed'
        syslog_logger.error(f'APP_UPDATE failed: {exc}')


@app.route('/settings/update/check')
@admin_required
def update_check():
    """Check whether remote has updates without applying them."""
    try:
        r = subprocess.run(['git', 'fetch', 'origin'], capture_output=True, text=True, timeout=20,
                           cwd=os.path.dirname(os.path.abspath(__file__)))
        r2 = subprocess.run(['git', 'log', 'HEAD..origin/HEAD', '--oneline'],
                            capture_output=True, text=True, timeout=10,
                            cwd=os.path.dirname(os.path.abspath(__file__)))
        commits = [l.strip() for l in r2.stdout.split('\n') if l.strip()]
        r3 = subprocess.run(['git', 'diff', '--name-only', 'HEAD', 'origin/HEAD'],
                            capture_output=True, text=True, timeout=10,
                            cwd=os.path.dirname(os.path.abspath(__file__)))
        files = [f.strip() for f in r3.stdout.split('\n') if f.strip()]
        return jsonify({'available': bool(commits), 'commits': commits, 'files': files})
    except Exception as e:
        return jsonify({'available': False, 'error': str(e)})


@app.route('/settings/update/apply', methods=['POST'])
@admin_required
def update_apply():
    """Start the update process in a background thread."""
    with _update_lock:
        if _update_state['running']:
            return jsonify({'error': 'Update already in progress'}), 409
        _update_state.update({'running': True, 'log': [], 'phase': 'starting', 'error': None})
    data = request.get_json(force=True) or {}
    service_name = data.get('service_name') or _detect_service_name()
    log_audit(get_db(), 'APP_UPDATE_START', 'system', None,
              detail=f'service={service_name} by={session.get("username")}')
    try:
        get_db().commit()
        get_db().close()
    except Exception:
        pass
    t = threading.Thread(target=_run_update, args=(service_name,), daemon=True)
    t.start()
    syslog_logger.info(f'APP_UPDATE_START service={service_name} by={session.get("username")}')
    return jsonify({'success': True, 'service': service_name})


@app.route('/api/update/status')
@admin_required
def update_status_api():
    return jsonify({
        'phase':   _update_state['phase'],
        'running': _update_state['running'],
        'log':     _update_state['log'],
        'error':   _update_state['error'],
    })


@app.route('/settings/update/detect-service')
@admin_required
def detect_service():
    return jsonify({'service': _detect_service_name()})


# ─── User Registration & Recovery ─────────────────────────────────────────────

def _send_simple_email(to_addr, subject, body_text, body_html=None):
    """Send a plain (single-recipient) email via the canonical _send_email
    helper defined at the top of this file."""
    try:
        ok, msg = _send_email(
            subject=subject,
            recipients=[to_addr],
            body_text=body_text,
            body_html=body_html,
        )
        if not ok:
            app.logger.error(f'Email send failed to {to_addr}: {msg}')
        return ok
    except Exception as e:
        app.logger.error(f'Email send failed to {to_addr}: {e}')
        return False


def _register_route():
    if session.get('user_id'):
        return redirect(url_for('dashboard'))
    error = None
    if request.method == 'POST':
        username = request.form.get('username', '').strip().lower()
        display_name = request.form.get('display_name', '').strip()
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')
        confirm  = request.form.get('confirm_password', '')
        score    = int(request.form.get('captcha_score', 0))

        if score < 1:
            error = 'Please complete the mini-game challenge first.'
        elif not username or not email or not password:
            error = 'All fields are required.'
        elif password != confirm:
            error = 'Passwords do not match.'
        elif len(password) < 8:
            error = 'Password must be at least 8 characters.'
        elif not re.match(r'^[a-z0-9_.-]{3,32}$', username):
            error = 'Username: 3-32 characters, letters/numbers/._- only.'
        else:
            db = get_db()
            existing = db.execute('SELECT id FROM users WHERE username=?', (username,)).fetchone()
            if existing:
                error = 'That username or email is not available.'
            else:
                existing_pending = db.execute(
                    'SELECT id FROM user_pending_registration WHERE username=?', (username,)).fetchone()
                if existing_pending:
                    error = 'That username or email is not available.'
                else:
                    token = secrets.token_urlsafe(32)
                    expires = datetime.utcnow() + timedelta(hours=24)
                    pw_hash = generate_password_hash(password)
                    try:
                        db.execute("""
                            INSERT INTO user_pending_registration
                              (username, display_name, email, password_hash, confirm_token, token_expires)
                            VALUES (?,?,?,?,?,?)
                        """, (username, display_name, email, pw_hash, token, expires.isoformat()))
                        db.commit()
                        confirm_url = url_for('confirm_email', token=token, _external=True)
                        _send_simple_email(
                            email,
                            '3·2·1→THEATER: Confirm Your Email',
                            f'Click the link to confirm your email address:\n{confirm_url}\n\nThis link expires in 24 hours.',
                            f'<p>Click the link below to confirm your email address:</p>'
                            f'<p><a href="{confirm_url}">{confirm_url}</a></p>'
                            f'<p>This link expires in 24 hours.</p>'
                        )
                        syslog_logger.info(f'REGISTER_PENDING username={username} email={email}')
                        # Notify admins of new pending registration
                        try:
                            _adb = get_db()
                            _admins = _adb.execute(
                                "SELECT email FROM users WHERE role='admin' AND email != '' AND email IS NOT NULL"
                            ).fetchall()
                            _adb.close()
                            _settings_url = url_for('settings', _external=True) + '#registrations'
                            for _adm in _admins:
                                _send_simple_email(
                                    _adm['email'],
                                    '3\u00b72\u00b71\u2192THEATER: New Registration Pending',
                                    f'A new account registration is awaiting your approval.\n\n'
                                    f'Username: {username}\nDisplay name: {display_name or "(none)"}\nEmail: {email}\n\n'
                                    f'Review and approve at:\n{_settings_url}',
                                )
                        except Exception:
                            pass
                        db.close()
                        return render_template('register.html',
                                               success='Registration submitted! Check your email to confirm, then wait for admin approval.',
                                               user=None)
                    except Exception as exc:
                        app.logger.error(f'Registration error: {exc}')
                        error = 'Registration failed due to a server error. Please try again.'
                    finally:
                        try:
                            db.close()
                        except Exception:
                            pass
    return render_template('register.html', error=error, user=None)


if _limiter_available and limiter:
    @app.route('/register', methods=['GET', 'POST'])
    @limiter.limit("10 per minute", methods=["POST"])
    def register():
        return _register_route()
else:
    @app.route('/register', methods=['GET', 'POST'])
    def register():
        return _register_route()


@app.route('/confirm-email/<token>')
def confirm_email(token):
    db = get_db()
    reg = db.execute(
        'SELECT * FROM user_pending_registration WHERE confirm_token=?', (token,)
    ).fetchone()
    if not reg:
        db.close()
        return render_template('register.html', error='Invalid or expired confirmation link.', user=None)
    if datetime.fromisoformat(reg['token_expires']) < datetime.utcnow():
        db.execute('DELETE FROM user_pending_registration WHERE id=?', (reg['id'],))
        db.commit()
        db.close()
        return render_template('register.html', error='Confirmation link expired. Please register again.', user=None)
    db.execute('UPDATE user_pending_registration SET email_confirmed=1 WHERE id=?', (reg['id'],))
    db.commit()
    syslog_logger.info(f'EMAIL_CONFIRMED username={reg["username"]}')
    db.close()
    return render_template('register.html',
                           success='Email confirmed! Your account is now awaiting admin approval.',
                           user=None)


@app.route('/settings/pending-registrations')
@admin_required
def pending_registrations():
    db = get_db()
    rows = db.execute("""
        SELECT * FROM user_pending_registration
        WHERE admin_approved=0
        ORDER BY created_at
    """).fetchall()
    db.close()
    return jsonify([dict(r) for r in rows])


@app.route('/settings/pending-registrations/<int:reg_id>/approve', methods=['POST'])
@admin_required
def approve_registration(reg_id):
    data = request.get_json(force=True) or {}
    role = data.get('role', 'user')
    db = get_db()
    reg = db.execute('SELECT * FROM user_pending_registration WHERE id=?', (reg_id,)).fetchone()
    if not reg:
        db.close()
        return jsonify({'error': 'Not found'}), 404
    try:
        db.execute("""
            INSERT INTO users (username, display_name, email, password_hash, role, email_confirmed)
            VALUES (?,?,?,?,?,1)
        """, (reg['username'], reg['display_name'] or reg['username'],
              reg['email'], reg['password_hash'], role))
        db.commit()
        uid = db.execute('SELECT id FROM users WHERE username=?', (reg['username'],)).fetchone()['id']
        db.execute('DELETE FROM user_pending_registration WHERE id=?', (reg_id,))
        db.commit()
        log_audit(db, 'USER_APPROVED', 'user', uid,
                  detail=f'username={reg["username"]} role={role} approved_by={session.get("username")}')
        db.commit()
        _send_simple_email(
            reg['email'],
            '3·2·1→THEATER: Account Approved',
            f'Your account "{reg["username"]}" has been approved. You can now log in.',
        )
        syslog_logger.info(f'USER_APPROVED username={reg["username"]} role={role} by={session.get("username")}')
        db.close()
        return jsonify({'success': True})
    except DBIntegrityError:
        db.close()
        return jsonify({'error': 'Username already exists'}), 409


@app.route('/settings/pending-registrations/<int:reg_id>/deny', methods=['POST'])
@admin_required
def deny_registration(reg_id):
    db = get_db()
    reg = db.execute('SELECT * FROM user_pending_registration WHERE id=?', (reg_id,)).fetchone()
    if reg:
        db.execute('DELETE FROM user_pending_registration WHERE id=?', (reg_id,))
        db.commit()
        _send_simple_email(
            reg['email'],
            '3·2·1→THEATER: Registration Declined',
            f'Your registration request for "{reg["username"]}" was not approved.',
        )
        syslog_logger.info(f'USER_DENIED username={reg["username"]} by={session.get("username")}')
    db.close()
    return jsonify({'success': True})


def _forgot_password_route():
    if session.get('user_id'):
        return redirect(url_for('dashboard'))
    sent = False
    error = None
    if request.method == 'POST':
        identifier = request.form.get('identifier', '').strip()
        score = int(request.form.get('captcha_score', 0))
        if score < 1:
            error = 'Please complete the mini-game challenge first.'
        elif not identifier:
            error = 'Enter your username or email.'
        else:
            db = get_db()
            user = db.execute(
                'SELECT * FROM users WHERE username=? OR email=?', (identifier, identifier)
            ).fetchone()
            # Always show success even if user not found (security best practice)
            if user and user.get('email'):
                token = secrets.token_urlsafe(48)
                expires = datetime.utcnow() + timedelta(hours=2)
                # Invalidate old tokens
                db.execute('UPDATE password_reset_tokens SET used=1 WHERE user_id=? AND used=0',
                           (user['id'],))
                db.execute("""
                    INSERT INTO password_reset_tokens (user_id, token, expires_at)
                    VALUES (?,?,?)
                """, (user['id'], token, expires.isoformat()))
                db.commit()
                reset_url = url_for('reset_password', token=token, _external=True)
                _send_simple_email(
                    user['email'],
                    '3·2·1→THEATER: Password Reset',
                    f'Click the link below to reset your password (expires in 2 hours):\n{reset_url}\n\n'
                    f'If you did not request this, ignore this email.',
                    f'<p><a href="{reset_url}">{reset_url}</a></p>'
                    f'<p>This link expires in 2 hours. If you did not request this, ignore this email.</p>'
                )
                syslog_logger.info(f'PASSWORD_RESET_REQUEST user={user["username"]}')
            db.close()
            sent = True
    return render_template('forgot_password.html', sent=sent, error=error, user=None)


if _limiter_available and limiter:
    @app.route('/forgot-password', methods=['GET', 'POST'])
    @limiter.limit("5 per minute", methods=["POST"])
    def forgot_password():
        return _forgot_password_route()
else:
    @app.route('/forgot-password', methods=['GET', 'POST'])
    def forgot_password():
        return _forgot_password_route()


@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    if session.get('user_id'):
        return redirect(url_for('dashboard'))
    db = get_db()
    rec = db.execute(
        'SELECT * FROM password_reset_tokens WHERE token=? AND used=0', (token,)
    ).fetchone()
    if not rec or datetime.fromisoformat(rec['expires_at']) < datetime.utcnow():
        db.close()
        return render_template('forgot_password.html',
                               error='This reset link has expired or already been used.', user=None)
    error = None
    if request.method == 'POST':
        password = request.form.get('password', '')
        confirm  = request.form.get('confirm_password', '')
        if len(password) < 8:
            error = 'Password must be at least 8 characters.'
        elif password != confirm:
            error = 'Passwords do not match.'
        else:
            pw_hash = generate_password_hash(password)
            db.execute('UPDATE users SET password_hash=? WHERE id=?', (pw_hash, rec['user_id']))
            db.execute('UPDATE password_reset_tokens SET used=1 WHERE token=?', (token,))
            db.commit()
            user_row = db.execute('SELECT username FROM users WHERE id=?', (rec['user_id'],)).fetchone()
            log_audit(db, 'PASSWORD_RESET_COMPLETE', 'user', rec['user_id'])
            db.commit()
            syslog_logger.info(f'PASSWORD_RESET_COMPLETE user={user_row["username"] if user_row else rec["user_id"]}')
            db.close()
            flash('Password reset successfully. You can now log in.', 'success')
            return redirect(url_for('login'))
    db.close()
    return render_template('forgot_password.html', token=token, reset_mode=True, error=error, user=None)


# ─── Site-Wide Messaging ───────────────────────────────────────────────────────

def get_active_messages(user_id=None, msg_type=None):
    """Return active, non-dismissed, non-expired messages."""
    db = get_db()
    now = datetime.utcnow().isoformat()
    rows = db.execute("""
        SELECT m.*,
               CASE WHEN d.user_id IS NOT NULL THEN 1 ELSE 0 END as dismissed
        FROM site_messages m
        LEFT JOIN site_message_dismissals d ON d.message_id = m.id AND d.user_id = ?
        WHERE m.is_active = 1
          AND (m.expires_at IS NULL OR m.expires_at > ?)
          AND (m.scheduled_for IS NULL OR m.scheduled_for <= ?)
          AND (? IS NULL OR m.msg_type = ?)
        ORDER BY m.created_at DESC
    """, (user_id or 0, now, now, msg_type, msg_type)).fetchall()
    db.close()
    return [dict(r) for r in rows if not r['dismissed'] or r['dismissible_by'] == 'admin']


@app.route('/api/messages')
@login_required
def get_messages_api():
    msg_type = request.args.get('type')
    msgs = get_active_messages(session['user_id'], msg_type)
    # Filter out already dismissed for users
    result = [m for m in msgs if not m['dismissed']]
    return jsonify(result)


@app.route('/api/messages/<int:msg_id>/dismiss', methods=['POST'])
@login_required
def dismiss_message(msg_id):
    db = get_db()
    msg = db.execute('SELECT * FROM site_messages WHERE id=?', (msg_id,)).fetchone()
    if not msg:
        db.close()
        return jsonify({'error': 'Not found'}), 404
    if msg['dismissible_by'] == 'admin' and session.get('user_role') != 'admin':
        db.close()
        return jsonify({'error': 'Only admins can dismiss this message'}), 403
    try:
        db.execute('INSERT OR IGNORE INTO site_message_dismissals (message_id, user_id) VALUES (?,?)',
                   (msg_id, session['user_id']))
        db.commit()
    except Exception:
        pass
    db.close()
    return jsonify({'success': True})


@app.route('/settings/messages', methods=['GET'])
@admin_required
def messages_list():
    db = get_db()
    rows = db.execute("""
        SELECT m.*, u.display_name as author
        FROM site_messages m
        LEFT JOIN users u ON u.id = m.created_by
        ORDER BY m.created_at DESC
    """).fetchall()
    db.close()
    return jsonify([dict(r) for r in rows])


@app.route('/settings/messages', methods=['POST'])
@admin_required
def message_create():
    data = request.get_json(force=True) or {}
    title = (data.get('title') or '').strip()
    body_html = _sanitize_html((data.get('body_html') or '').strip())
    if not title:
        return jsonify({'error': 'Title required'}), 400
    db = get_db()
    db.execute("""
        INSERT INTO site_messages
          (title, body_html, msg_type, dismissible_by, expires_at, scheduled_for,
           is_active, show_on_login, created_by)
        VALUES (?,?,?,?,?,?,?,?,?)
    """, (
        title, body_html,
        data.get('msg_type', 'motd'),
        data.get('dismissible_by', 'user'),
        data.get('expires_at') or None,
        data.get('scheduled_for') or None,
        1 if data.get('is_active', True) else 0,
        1 if data.get('show_on_login') else 0,
        session['user_id'],
    ))
    db.commit()
    row = db.execute('SELECT * FROM site_messages ORDER BY id DESC LIMIT 1').fetchone()
    log_audit(db, 'MESSAGE_CREATE', 'site_message', row['id'], detail=title)
    db.commit()
    syslog_logger.info(f'MESSAGE_CREATE title="{title}" type={data.get("msg_type","motd")} by={session.get("username")}')
    result = dict(row)
    db.close()
    return jsonify(result), 201


@app.route('/settings/messages/<int:msg_id>', methods=['PUT'])
@admin_required
def message_edit(msg_id):
    data = request.get_json(force=True) or {}
    db = get_db()
    db.execute("""
        UPDATE site_messages SET
          title=?, body_html=?, msg_type=?, dismissible_by=?,
          expires_at=?, scheduled_for=?, is_active=?, show_on_login=?
        WHERE id=?
    """, (
        (data.get('title') or '').strip(),
        _sanitize_html((data.get('body_html') or '').strip()),
        data.get('msg_type', 'motd'),
        data.get('dismissible_by', 'user'),
        data.get('expires_at') or None,
        data.get('scheduled_for') or None,
        1 if data.get('is_active', True) else 0,
        1 if data.get('show_on_login') else 0,
        msg_id,
    ))
    db.commit()
    log_audit(db, 'MESSAGE_EDIT', 'site_message', msg_id, detail=data.get('title', ''))
    db.commit()
    db.close()
    return jsonify({'success': True})


@app.route('/settings/messages/<int:msg_id>', methods=['DELETE'])
@admin_required
def message_delete(msg_id):
    db = get_db()
    db.execute('DELETE FROM site_messages WHERE id=?', (msg_id,))
    db.commit()
    log_audit(db, 'MESSAGE_DELETE', 'site_message', msg_id)
    db.commit()
    db.close()
    return jsonify({'success': True})


@app.route('/settings/messages/<int:msg_id>/dismiss-all', methods=['POST'])
@admin_required
def message_dismiss_all(msg_id):
    """Admin globally deactivates (removes) a message for everyone."""
    db = get_db()
    db.execute('UPDATE site_messages SET is_active=0 WHERE id=?', (msg_id,))
    db.commit()
    log_audit(db, 'MESSAGE_DISMISS_ALL', 'site_message', msg_id)
    db.commit()
    db.close()
    return jsonify({'success': True})


# ─── AI Session Management ─────────────────────────────────────────────────────

def _get_ai_slot_limit():
    return int(get_app_setting('ai_max_sessions', '2'))

def _count_active_ai_sessions():
    """Count running AI sessions, pruning stale ones (>5 min) first."""
    db = get_db()
    cutoff = (datetime.utcnow() - timedelta(minutes=5)).isoformat()
    db.execute("""
        UPDATE ai_sessions SET status='timeout', ended_at=CURRENT_TIMESTAMP
        WHERE status='running' AND started_at < ?
    """, (cutoff,))
    db.commit()
    count = db.execute("SELECT COUNT(*) FROM ai_sessions WHERE status='running'").fetchone()[0]
    db.close()
    return count

def _claim_ai_session(show_id):
    """Reserve a slot. Returns (session_id, None) or (None, error)."""
    db = get_db()
    limit = _get_ai_slot_limit()
    count = _count_active_ai_sessions()
    if count >= limit:
        db.close()
        return None, f'All {limit} AI processing slots are busy. Please try again in a moment.'
    db.execute("""
        INSERT INTO ai_sessions (user_id, show_id, status)
        VALUES (?,?,'running')
    """, (session.get('user_id'), show_id))
    db.commit()
    sid = db.execute('SELECT id FROM ai_sessions ORDER BY id DESC LIMIT 1').fetchone()['id']
    db.close()
    return sid, None

def _release_ai_session(session_id):
    if not session_id:
        return
    db = get_db()
    db.execute("""
        UPDATE ai_sessions SET status='done', ended_at=CURRENT_TIMESTAMP WHERE id=?
    """, (session_id,))
    db.commit()
    db.close()


@app.route('/api/ai/slots')
@login_required
def ai_slots_status():
    """Return current AI slot availability for dynamic UI."""
    limit = _get_ai_slot_limit()
    count = _count_active_ai_sessions()
    return jsonify({
        'limit': limit,
        'active': count,
        'available': max(0, limit - count),
        'busy': count >= limit,
    })


# ─── Asset Availability Dashboard ─────────────────────────────────────────────

@app.route('/api/dashboard/shows-calendar')
def api_dashboard_shows_calendar():
    """Return shows per day for a date range (for calendar widget). Public-safe."""
    from datetime import date as _date, timedelta
    date_from = request.args.get('from', '')
    date_to   = request.args.get('to', '')
    db = get_db()
    accessible = get_accessible_shows(session['user_id']) if session.get('user_id') else None
    params = []
    where_parts = ["s.status != 'archived'"]
    if date_from:
        where_parts.append("COALESCE(s.show_date, s.load_in_date, s.load_out_date) >= ?")
        params.append(date_from)
    if date_to:
        where_parts.append("COALESCE(s.show_date, s.load_in_date, s.load_out_date) <= ?")
        params.append(date_to)
    if accessible is not None:
        if not accessible:
            db.close()
            return jsonify({'days': []})
        placeholders = ','.join('?' * len(accessible))
        where_parts.append(f's.id IN ({placeholders})')
        params.extend(accessible)
    where_sql = 'WHERE ' + ' AND '.join(where_parts)
    rows = db.execute(f"""
        SELECT s.id, s.name, s.show_date, s.load_in_date, s.load_out_date, s.venue, s.status
        FROM shows s {where_sql}
        ORDER BY COALESCE(s.show_date, s.load_in_date), s.name
    """, params).fetchall()
    db.close()

    # Build per-day buckets: a show appears on its load_in, show, and load_out dates
    day_map = {}
    for r in rows:
        dates_for_show = set()
        if r['load_in_date']:  dates_for_show.add(str(r['load_in_date']))
        if r['show_date']:     dates_for_show.add(str(r['show_date']))
        if r['load_out_date']: dates_for_show.add(str(r['load_out_date']))
        li = str(r['load_in_date']) if r['load_in_date'] else None
        sd = str(r['show_date'])    if r['show_date']    else None
        lo = str(r['load_out_date'])if r['load_out_date']else None
        for d in dates_for_show:
            if date_from and d < date_from: continue
            if date_to   and d > date_to:   continue
            if d not in day_map: day_map[d] = []
            day_map[d].append({'id': r['id'], 'name': r['name'], 'venue': r['venue'],
                               'load_in_date': li, 'show_date': sd, 'load_out_date': lo})

    # Fill in all calendar days even if empty
    days = []
    if date_from and date_to:
        try:
            cur = _date.fromisoformat(date_from)
            end = _date.fromisoformat(date_to)
            while cur <= end:
                ds = cur.isoformat()
                days.append({'date': ds, 'shows': day_map.get(ds, []),
                             'show_count': len(day_map.get(ds, []))})
                cur += timedelta(days=1)
        except ValueError:
            pass
    else:
        for d in sorted(day_map):
            days.append({'date': d, 'shows': day_map[d], 'show_count': len(day_map[d])})

    return jsonify({'days': days})


@app.route('/api/dashboard/skills-summary')
def api_dashboard_skills_summary():
    """Return technician skill coverage per position."""
    db = get_db()
    total_crew = db.execute('SELECT COUNT(*) FROM crew_members').fetchone()[0]
    cats = db.execute("""
        SELECT pc.id, pc.name
        FROM position_categories pc
        ORDER BY pc.sort_order, pc.id
    """).fetchall()
    result = []
    for cat in cats:
        positions = db.execute("""
            SELECT jp.id, jp.name, jp.venue,
                   COUNT(cq.crew_member_id) as qualified_count
            FROM job_positions jp
            LEFT JOIN crew_qualifications cq ON cq.position_id = jp.id
            WHERE jp.category_id = ?
            GROUP BY jp.id
            ORDER BY jp.sort_order, jp.id
        """, (cat['id'],)).fetchall()
        if positions:
            result.append({
                'id': cat['id'],
                'name': cat['name'],
                'positions': [{'id': p['id'], 'name': p['name'], 'venue': p['venue'],
                               'qualified': p['qualified_count'],
                               'unqualified': max(0, total_crew - p['qualified_count'])}
                              for p in positions]
            })
    # Uncategorized positions
    uncategorized = db.execute("""
        SELECT jp.id, jp.name, jp.venue,
               COUNT(cq.crew_member_id) as qualified_count
        FROM job_positions jp
        LEFT JOIN crew_qualifications cq ON cq.position_id = jp.id
        WHERE jp.category_id IS NULL
        GROUP BY jp.id
        ORDER BY jp.sort_order, jp.id
    """).fetchall()
    if uncategorized:
        result.append({
            'id': None, 'name': 'Other',
            'positions': [{'id': p['id'], 'name': p['name'], 'venue': p['venue'],
                           'qualified': p['qualified_count'],
                           'unqualified': max(0, total_crew - p['qualified_count'])}
                          for p in uncategorized]
        })
    db.close()
    return jsonify({'total_crew': total_crew, 'categories': result})


@app.route('/api/dashboard/asset-calendar')
def api_dashboard_asset_calendar():
    """Per-day availability for one asset type — used by the asset calendar widget."""
    from datetime import date as _date, timedelta
    type_id   = request.args.get('type_id', type=int)
    date_from = request.args.get('from')
    date_to   = request.args.get('to')
    if not type_id or not date_from or not date_to:
        return jsonify({'error': 'type_id, from, to required'}), 400

    db = get_db()
    row = db.execute(
        'SELECT at.*, ac.name AS category_name FROM asset_types at '
        'JOIN asset_categories ac ON ac.id = at.category_id WHERE at.id = ?',
        (type_id,)
    ).fetchone()
    if not row:
        db.close()
        return jsonify({'error': 'Not found'}), 404

    unlimited = bool(row['is_system'] or row['is_package'] or
                     (row['is_consumable'] and not row['track_quantity']))
    total = in_maint = reserve = 0

    if not unlimited:
        total = db.execute(
            "SELECT COUNT(*) FROM asset_items WHERE asset_type_id=? AND status!='retired'",
            (type_id,)
        ).fetchone()[0]
        in_maint = db.execute(
            "SELECT COUNT(*) FROM asset_items WHERE asset_type_id=? AND status='maintenance'",
            (type_id,)
        ).fetchone()[0]
        reserve = row['reserve_count'] or 0

    reservations = [] if unlimited else db.execute(
        'SELECT sa.quantity, sa.rental_start, sa.rental_end, sa.show_id, s.name AS show_name '
        'FROM show_assets sa JOIN shows s ON s.id = sa.show_id '
        'WHERE sa.asset_type_id=? AND sa.rental_end>=? AND sa.rental_start<=?',
        (type_id, date_from, date_to)
    ).fetchall()
    db.close()

    try:
        cur = _date.fromisoformat(date_from)
        end = _date.fromisoformat(date_to)
    except ValueError:
        return jsonify({'error': 'Invalid date'}), 400

    days = []
    while cur <= end:
        ds = cur.isoformat()
        if unlimited:
            days.append({'date': ds, 'available': None, 'reserved': 0, 'shows': []})
        else:
            day_shows = [
                {'id': r['show_id'], 'name': r['show_name'], 'quantity': r['quantity']}
                for r in reservations
                if r['rental_start'] and r['rental_end']
                and str(r['rental_start']) <= ds <= str(r['rental_end'])
            ]
            day_rsv = sum(s['quantity'] for s in day_shows)
            days.append({'date': ds, 'available': total - in_maint - reserve - day_rsv, 'reserved': day_rsv, 'shows': day_shows})
        cur += timedelta(days=1)

    return jsonify({
        'type_id':       type_id,
        'type_name':     row['name'],
        'category_name': row['category_name'],
        'total':         total,
        'maintenance':   in_maint,
        'reserve':       reserve,
        'unlimited':     unlimited,
        'days':          days,
    })

@app.route('/dashboards')
@login_required
def dashboards_list():
    db = get_db()
    rows = db.execute("""
        SELECT d.*, u.display_name as owner_name
        FROM asset_dashboards d
        JOIN users u ON u.id = d.user_id
        WHERE d.user_id = ? OR d.is_public = 1
        ORDER BY d.user_id = ? DESC, d.name
    """, (session['user_id'], session['user_id'])).fetchall()
    db.close()
    dashboards = []
    for r in rows:
        d = dict(r)
        d['config'] = json.loads(d.get('config_json') or '{}')
        dashboards.append(d)
    return render_template('dashboards.html',
                           dashboards=dashboards,
                           user=get_current_user())


@app.route('/dashboards/new', methods=['POST'])
@login_required
def dashboard_create():
    data = request.get_json(force=True) or {}
    name = (data.get('name') or 'My Dashboard').strip()
    slug = secrets.token_urlsafe(12) if data.get('is_public') else None
    db = get_db()
    db.execute("""
        INSERT INTO asset_dashboards (user_id, name, is_public, public_slug, layout, config_json)
        VALUES (?,?,?,?,?,?)
    """, (session['user_id'], name,
          1 if data.get('is_public') else 0,
          slug,
          data.get('layout', 'combined'),
          json.dumps(data.get('config', {}))))
    db.commit()
    row = db.execute('SELECT * FROM asset_dashboards ORDER BY id DESC LIMIT 1').fetchone()
    db.close()
    return jsonify(dict(row)), 201


@app.route('/dashboards/<int:dash_id>')
@login_required
def dashboard_view(dash_id):
    db = get_db()
    d = db.execute('SELECT * FROM asset_dashboards WHERE id=?', (dash_id,)).fetchone()
    if not d:
        db.close()
        abort(404)
    if d['user_id'] != session['user_id'] and not d['is_public']:
        if session.get('user_role') != 'admin':
            db.close()
            abort(403)
    cats = db.execute('SELECT * FROM asset_categories ORDER BY sort_order, name').fetchall()
    types = db.execute("""
        SELECT at.*, ac.name as category_name
        FROM asset_types at
        JOIN asset_categories ac ON ac.id = at.category_id
        ORDER BY ac.sort_order, at.sort_order, at.name
    """).fetchall()
    db.close()
    config = json.loads(d['config_json'] or '{}')
    return render_template('dashboard_view.html',
                           dash=dict(d),
                           categories=[dict(c) for c in cats],
                           asset_types=[{k: v for k, v in dict(t).items() if k != 'photo'} for t in types],
                           config=config,
                           user=get_current_user())


@app.route('/dashboards/<int:dash_id>', methods=['PUT'])
@login_required
def dashboard_edit(dash_id):
    db = get_db()
    d = db.execute('SELECT * FROM asset_dashboards WHERE id=?', (dash_id,)).fetchone()
    if not d or d['user_id'] != session['user_id']:
        db.close()
        abort(403)
    data = request.get_json(force=True) or {}
    slug = d['public_slug']
    if data.get('is_public') and not slug:
        slug = secrets.token_urlsafe(12)
    db.execute("""
        UPDATE asset_dashboards SET name=?, is_public=?, public_slug=?,
               layout=?, config_json=?, updated_at=CURRENT_TIMESTAMP
        WHERE id=?
    """, (
        (data.get('name') or 'My Dashboard').strip(),
        1 if data.get('is_public') else 0,
        slug if data.get('is_public') else None,
        data.get('layout', 'combined'),
        json.dumps(data.get('config', {})),
        dash_id,
    ))
    db.commit()
    db.close()
    return jsonify({'success': True, 'slug': slug})


@app.route('/dashboards/<int:dash_id>', methods=['DELETE'])
@login_required
def dashboard_delete(dash_id):
    db = get_db()
    d = db.execute('SELECT user_id FROM asset_dashboards WHERE id=?', (dash_id,)).fetchone()
    if not d or (d['user_id'] != session['user_id'] and session.get('user_role') != 'admin'):
        db.close()
        abort(403)
    db.execute('DELETE FROM asset_dashboards WHERE id=?', (dash_id,))
    db.commit()
    db.close()
    return jsonify({'success': True})


@app.route('/api/admin/dashboards/<int:dash_id>/make-private', methods=['POST'])
@admin_required
def api_admin_dashboard_make_private(dash_id):
    db = get_db()
    db.execute('UPDATE asset_dashboards SET is_public=0, public_slug=NULL WHERE id=?', (dash_id,))
    db.commit()
    db.close()
    syslog_logger.info(f"DASHBOARD_MADE_PRIVATE id={dash_id} by={session.get('username')}")
    return jsonify({'success': True})


@app.route('/api/admin/dashboards')
@admin_required
def api_admin_dashboards():
    """Admin: list all dashboards with owner and public URL info."""
    db = get_db()
    rows = db.execute("""
        SELECT d.id, d.name, d.is_public, d.public_slug, d.layout,
               d.created_at, d.updated_at, u.display_name as owner_name, u.username
        FROM asset_dashboards d
        JOIN users u ON u.id = d.user_id
        ORDER BY d.is_public DESC, d.name
    """).fetchall()
    db.close()
    host = request.host_url.rstrip('/')
    result = []
    for r in rows:
        row = dict(r)
        row['public_url'] = f"{host}/d/{r['public_slug']}" if r['is_public'] and r['public_slug'] else None
        result.append(row)
    return jsonify(result)


@app.route('/d/<slug>')
def public_dashboard(slug):
    """Public dashboard — no login required."""
    db = get_db()
    d = db.execute(
        'SELECT * FROM asset_dashboards WHERE public_slug=? AND is_public=1', (slug,)
    ).fetchone()
    if not d:
        db.close()
        abort(404)
    cats = db.execute('SELECT * FROM asset_categories ORDER BY sort_order, name').fetchall()
    types = db.execute("""
        SELECT at.*, ac.name as category_name
        FROM asset_types at
        JOIN asset_categories ac ON ac.id = at.category_id
        ORDER BY ac.sort_order, at.sort_order, at.name
    """).fetchall()
    db.close()
    config = json.loads(d['config_json'] or '{}')
    return render_template('dashboard_view.html',
                           dash=dict(d),
                           categories=[dict(c) for c in cats],
                           asset_types=[{k: v for k, v in dict(t).items() if k != 'photo'} for t in types],
                           config=config,
                           public=True,
                           user=None)


# ─── Asset Reports ─────────────────────────────────────────────────────────────

@app.route('/reports/assets')
@asset_manager_required
def asset_reports():
    db = get_db()
    companies = db.execute("""
        SELECT DISTINCT ad.field_value as company
        FROM advance_data ad
        WHERE ad.field_key = 'performance_company' AND ad.field_value != ''
        ORDER BY ad.field_value
    """).fetchall()
    venues = db.execute("""
        SELECT DISTINCT venue FROM shows WHERE venue != '' ORDER BY venue
    """).fetchall()
    asset_categories = db.execute(
        'SELECT id, name FROM asset_categories ORDER BY sort_order, name'
    ).fetchall()
    asset_types = db.execute(
        'SELECT id, name, category_id FROM asset_types WHERE is_retired=0 ORDER BY name'
    ).fetchall()
    db.close()
    return render_template('asset_reports.html',
                           companies=[r['company'] for r in companies],
                           venues=[r['venue'] for r in venues],
                           asset_categories=[dict(r) for r in asset_categories],
                           asset_types=[dict(r) for r in asset_types],
                           user=get_current_user())


@app.route('/api/reports/assets')
@asset_manager_required
def asset_reports_data():
    company   = request.args.get('company', '')
    venue     = request.args.get('venue', '')
    asset_type_id = request.args.get('asset_type_id', '')
    asset_category_id = request.args.get('asset_category_id', '')
    date_from = request.args.get('from', '')
    date_to   = request.args.get('to', '')
    db = get_db()

    params = []
    where = []

    if company:
        where.append("""
            s.id IN (
                SELECT show_id FROM advance_data
                WHERE field_key='performance_company' AND field_value=?
            )
        """)
        params.append(company)

    if venue:
        where.append("s.venue = ?")
        params.append(venue)

    if asset_type_id:
        where.append("sa.asset_type_id = ?")
        params.append(int(asset_type_id))

    if asset_category_id:
        where.append("at.category_id = ?")
        params.append(int(asset_category_id))

    if date_from:
        where.append("COALESCE(s.show_date, '9999-12-31') >= ?")
        params.append(date_from)
    if date_to:
        where.append("COALESCE(s.show_date, '0001-01-01') <= ?")
        params.append(date_to)

    where_sql = ('WHERE ' + ' AND '.join(where)) if where else ''

    rows = db.execute(f"""
        SELECT sa.id, sa.quantity, sa.locked_price, sa.rental_start, sa.rental_end,
               at.id as asset_type_id, at.name as type_name, at.manufacturer, at.model,
               ac.id as category_id, ac.name as category_name,
               s.id as show_id, s.name as show_name, s.show_date, s.venue,
               (sa.quantity * sa.locked_price) as line_total,
               (SELECT field_value FROM advance_data
                WHERE show_id=s.id AND field_key='performance_company') as performance_company
        FROM show_assets sa
        JOIN asset_types at ON at.id = sa.asset_type_id
        JOIN asset_categories ac ON ac.id = at.category_id
        JOIN shows s ON s.id = sa.show_id
        {where_sql}
        ORDER BY s.show_date DESC, ac.name, at.name
    """, params).fetchall()

    total_revenue = sum(r['line_total'] or 0 for r in rows)
    db.close()
    return jsonify({
        'rows': [dict(r) for r in rows],
        'total_revenue': total_revenue,
        'count': len(rows),
    })


# ─── Error Handlers ───────────────────────────────────────────────────────────

@app.errorhandler(403)
def forbidden(e):
    return render_template('error.html', code=403,
                           message="You don't have permission to do that.",
                           user=get_current_user()), 403


@app.errorhandler(404)
def not_found(e):
    return render_template('error.html', code=404, message="Page not found.",
                           user=get_current_user()), 404


def _wants_json_response():
    """Return True when the current request is XHR / fetch / JSON.

    Used by the error handlers so that an unexpected exception inside an API
    route gives the caller a parseable JSON body instead of an HTML error page
    (which is what triggers `Unexpected token '<'` in the browser console)."""
    if request.is_json:
        return True
    if request.path.startswith('/api/'):
        return True
    accept = (request.headers.get('Accept') or '').lower()
    if 'application/json' in accept:
        return True
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return True
    return False


@app.errorhandler(500)
def internal_error(e):
    app.logger.exception("500 Internal Server Error")
    if _wants_json_response():
        return jsonify({
            'success': False,
            'error': 'Internal server error.',
            'detail': str(getattr(e, 'original_exception', e) or e),
        }), 500
    return render_template('error.html', code=500,
                           message="An unexpected server error occurred.",
                           user=get_current_user()), 500


# ─── Run ──────────────────────────────────────────────────────────────────────

# Initialize syslog at import time (for Gunicorn)
if os.path.exists(DATABASE):
    reload_syslog_handler()
    # Ensure backup dirs exist (fixes PermissionError if dirs were missing)
    try:
        _ensure_backup_dirs()
    except Exception:
        pass
    # Auto-run DB migrations on startup (idempotent — safe to run every time)
    try:
        from init_db import migrate_db, migrate_db_postgres
        migrate_db()
        migrate_db_postgres()
    except Exception as _mig_err:
        print(f"[startup] Migration warning: {_mig_err}")

# Start backup scheduler (guarded against Flask reloader double-start)
_scheduler = None
if not (os.environ.get('WERKZEUG_RUN_MAIN') == 'false'):
    _scheduler = start_scheduler()
    if _scheduler:
        atexit.register(lambda: _scheduler.shutdown(wait=False))
    # Cluster heartbeat for multi-server leader election
    if get_app_setting('cluster_heartbeat_enabled', '1') in ('1', 'true'):
        start_cluster_heartbeat()

if __name__ == '__main__':
    if not os.path.exists(DATABASE):
        print("Database not found. Run: python init_db.py")
        run_port = 5400
    else:
        try:
            _db = get_db()
            _row = _db.execute(
                "SELECT value FROM app_settings WHERE key='app_port'"
            ).fetchone()
            _db.close()
            run_port = int(_row['value']) if _row else 5400
        except Exception:
            run_port = 5400
    app.run(debug=os.environ.get('FLASK_DEBUG', '0') == '1', port=run_port)
